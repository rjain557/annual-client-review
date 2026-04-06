"""
Generate branded AAVA Word documents:
- 2025 Annual Review
- 2026 Q1 Review
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
MED_BLUE = RGBColor(0x2E, 0x75, 0xB6)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

DARK_BLUE_HEX = "1F4E79"
ALT_ROW_HEX = "F2F7FB"
WHITE_HEX = "FFFFFF"
YELLOW_HEX = "FFF3CD"

BASE = r"c:\vscode\annual-client-review\annual-client-review"
OUT_2025 = os.path.join(BASE, "clients", "aava", "2025", "AAVA - 2025 Annual Review.docx")
OUT_2026Q1 = os.path.join(BASE, "clients", "aava", "2026", "AAVA - 2026 Q1 Review.docx")

L = WD_ALIGN_PARAGRAPH.LEFT
R = WD_ALIGN_PARAGRAPH.RIGHT


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=Pt(8.5), color=BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = "Calibri"


def add_header_row(table, texts, col_aligns=None):
    row = table.rows[0]
    for i, txt in enumerate(texts):
        align = col_aligns[i] if col_aligns else L
        set_cell_text(row.cells[i], txt, bold=True, color=WHITE, alignment=align)
        set_cell_shading(row.cells[i], DARK_BLUE_HEX)


def add_data_row(table, texts, row_idx, col_aligns=None, is_total=False, is_highlight=False):
    row = table.rows[row_idx]
    bg = DARK_BLUE_HEX if is_total else (YELLOW_HEX if is_highlight else (ALT_ROW_HEX if row_idx % 2 == 0 else WHITE_HEX))
    txt_color = WHITE if is_total else BLACK
    for i, txt in enumerate(texts):
        align = col_aligns[i] if col_aligns else L
        set_cell_text(row.cells[i], txt, bold=is_total, color=txt_color, alignment=align)
        set_cell_shading(row.cells[i], bg)


def make_table(doc, headers, data, col_aligns=None, total_row_indices=None, highlight_row_indices=None):
    total_row_indices = total_row_indices or []
    highlight_row_indices = highlight_row_indices or []

    tbl = doc.add_table(rows=1 + len(data), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:top w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:left w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:bottom w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:right w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '</w:tcBorders>'
            )
            tcPr.append(borders)

    add_header_row(tbl, headers, col_aligns)
    for idx, row_data in enumerate(data):
        add_data_row(
            tbl,
            row_data,
            idx + 1,
            col_aligns,
            is_total=idx in total_row_indices,
            is_highlight=idx in highlight_row_indices,
        )
    return tbl


def h1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    r = p.add_run(text)
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(16)
    r.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    r = p.add_run(text)
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(13)
    r.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def sh(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(11)
    r.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Calibri"


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Calibri"


def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GRAY
    r.font.name = "Calibri"


def setup_doc(header_text, title_line1, title_line2, prepared_line):
    doc = Document()
    s = doc.sections[0]
    s.page_width = Inches(8.5)
    s.page_height = Inches(11)
    s.top_margin = Inches(1.0)
    s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(0.63)
    s.right_margin = Inches(0.63)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    hp = s.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(header_text)
    hr.font.size = Pt(8)
    hr.font.color.rgb = GRAY
    hr.font.name = "Calibri"

    fp = s.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Technijian Inc. - Confidential")
    fr.font.size = Pt(8)
    fr.font.color.rgb = GRAY
    fr.font.name = "Calibri"

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AAVA")
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title_line1)
    r.font.size = Pt(20)
    r.font.color.rgb = MED_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title_line2)
    r.font.size = Pt(20)
    r.font.color.rgb = MED_BLUE

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(prepared_line)
    r.font.size = Pt(12)
    r.font.color.rgb = GRAY

    return doc


def build_aava_2025():
    doc = setup_doc(
        "AAVA - 2025 IT Support Review and 2026 Planning",
        "2025 IT Support Review",
        "and 2026 Planning Readiness Assessment",
        "Prepared by Technijian Inc.  -  April 2026",
    )

    doc.add_paragraph()
    make_table(
        doc,
        ["2025 At a Glance", "Key Figures"],
        [
            ["Total 2025 Revenue", "$5,342.80"],
            ["Monthly Contract", "$5,342.80 (~$445/month)"],
            ["Recurring Licensing", "$0"],
            ["Projects / One-Time", "$0"],
            ["Labor Overage (2025)", "$0"],
            ["Unique Tickets", "162"],
            ["Time Entries", "163"],
            ["Total Support Hours", "78.8"],
            ["Desktops Managed", "8"],
            ["Servers / Cloud Managed", "0"],
            ["Proactive Ticket Ratio", "95.1%"],
        ],
        col_aligns=[L, R],
        total_row_indices=[0],
    )

    doc.add_page_break()

    h1(doc, "Section 1: Support Scope and Coverage")
    h2(doc, "1.1 Coverage Summary")
    body(doc, "AAVA operated under a single managed services agreement in 2025 with total billed revenue of $5,342.80.")
    bullet(doc, "Total hours: 78.8 across 162 unique tickets")
    bullet(doc, "Offshore NH (US overnight): 57.6 hours (73.1%)")
    bullet(doc, "Offshore AH (US daytime): 21.2 hours (26.9%)")
    bullet(doc, "Onsite: 0.0 hours (100% remote delivery)")
    bullet(doc, "Weekend entries: 77 of 163 (47.2%)")

    h2(doc, "1.2 Role Distribution")
    make_table(
        doc,
        ["Role", "NH", "AH", "Onsite", "Total Hours", "Entries"],
        [
            ["Off-Shore Tech Support", "57.6", "21.2", "0.0", "78.8", "163"],
            ["Tech Support (US)", "0.0", "0.0", "0.0", "0.0", "0"],
            ["TOTAL", "57.6", "21.2", "0.0", "78.8", "163"],
        ],
        col_aligns=[L, R, R, R, R, R],
        total_row_indices=[2],
    )

    doc.add_page_break()

    h1(doc, "Section 2: Effective Blended Rates")
    body(doc, "Rates are calculated from billed labor line items (Monthly + WeeklyOut).")
    make_table(
        doc,
        ["Role", "Labor Revenue", "Billed Hours", "Blended Rate"],
        [
            ["US Tech Support", "$1,875.00", "15.00", "$125.00/hr"],
            ["Offshore AH", "$674.10", "22.47", "$30.00/hr"],
            ["Offshore NH", "$533.70", "35.58", "$15.00/hr"],
            ["ALL LABOR", "$3,082.80", "73.05", "$42.20/hr"],
        ],
        col_aligns=[L, R, R, R],
        total_row_indices=[3],
    )
    note(doc, "Delivered labor hours in time entries (78.8) exceeded billed labor hours (73.05) by 5.75 hours.")

    h2(doc, "2.1 Monthly Billing Trend (2025)")
    make_table(
        doc,
        ["Month", "Monthly Total"],
        [
            ["Jan", "$411.25"],
            ["Feb", "$411.25"],
            ["Mar", "$411.25"],
            ["Apr", "$411.25"],
            ["May", "$411.25"],
            ["Jun", "$454.60"],
            ["Jul", "$460.60"],
            ["Aug", "$457.60"],
            ["Sep", "$457.60"],
            ["Oct", "$457.60"],
            ["Nov", "$508.50"],
            ["Dec", "$490.05"],
            ["2025 Total", "$5,342.80"],
        ],
        col_aligns=[L, R],
        total_row_indices=[12],
        highlight_row_indices=[5, 10],
    )

    doc.add_page_break()

    h1(doc, "Section 3: Service Mix and Ticket Profile")
    h2(doc, "3.1 Monthly Contract Categories")
    make_table(
        doc,
        ["Category", "Annual Total"],
        [
            ["Labor", "$3,082.80"],
            ["Security and Endpoint", "$1,296.00"],
            ["Monitoring and Management", "$480.00"],
            ["Secure Internet", "$384.00"],
            ["Assessment", "$100.00"],
            ["TOTAL", "$5,342.80"],
        ],
        col_aligns=[L, R],
        total_row_indices=[5],
    )

    h2(doc, "3.2 Ticket Categorization (2025)")
    make_table(
        doc,
        ["Category", "Type", "Tickets", "Hours", "Avg Hrs"],
        [
            ["Patch Management", "Proactive", "69", "29.8", "0.43"],
            ["Backup and DR", "Proactive", "30", "12.8", "0.43"],
            ["Security and Endpoint", "Proactive", "27", "14.8", "0.55"],
            ["RMM and Agent Management", "Proactive", "16", "11.2", "0.70"],
            ["Monitoring and Alerts", "Proactive", "12", "5.5", "0.46"],
            ["Reactive (all categories)", "Reactive", "8", "4.7", "0.59"],
            ["TOTAL", "", "162", "78.8", "0.49"],
        ],
        col_aligns=[L, L, R, R, R],
        total_row_indices=[6],
        highlight_row_indices=[0],
    )

    make_table(
        doc,
        ["Type", "Tickets", "% Tickets", "Hours", "% Hours"],
        [
            ["Proactive", "154", "95.1%", "74.1", "94.0%"],
            ["Reactive", "8", "4.9%", "4.7", "6.0%"],
        ],
        col_aligns=[L, R, R, R, R],
    )

    doc.add_page_break()

    h1(doc, "Section 4: 2026 Budget Projection")
    body(doc, "Projection uses observed Q1 2026 billing behavior and current March/April run rate.")
    make_table(
        doc,
        ["Month (Q1 2026)", "Monthly Total", "Driver"],
        [
            ["Jan 2026", "$461.25", "Baseline with Site Assessment"],
            ["Feb 2026", "$493.25", "AVMH/PMW/SI rate changes (+$32)"],
            ["Mar 2026", "$539.60", "Offshore labor quantity increase (+$46.35)"],
            ["Q1 Total", "$1,494.10", ""],
        ],
        col_aligns=[L, R, L],
        total_row_indices=[3],
        highlight_row_indices=[1, 2],
    )

    make_table(
        doc,
        ["Category", "2025 Actual", "2026 Projected", "Change"],
        [
            ["Monthly Contract", "$5,342.80", "$6,475.20", "+$1,132.40"],
            ["Recurring Licensing", "$0.00", "$0.00", "$0.00"],
            ["One-Time / Hardware Reserve", "$0.00", "$500.00", "+$500.00"],
            ["TOTAL", "$5,342.80", "$6,975.20", "+$1,632.40"],
        ],
        col_aligns=[L, R, R, R],
        total_row_indices=[3],
    )

    sh(doc, "Planning Notes")
    bullet(doc, "Current Mar/Apr run rate implies roughly 21.2% higher annual contract cost vs 2025.")
    bullet(doc, "If labor quantities return to January baseline, spend trends closer to ~$5,976 annual contract run rate.")
    bullet(doc, "Removing Site Assessment would reduce spend by about $600/year.")

    os.makedirs(os.path.dirname(OUT_2025), exist_ok=True)
    doc.save(OUT_2025)
    print(f"Saved: {OUT_2025}")


def build_aava_2026_q1():
    doc = setup_doc(
        "AAVA - 2026 Q1 IT Support Review",
        "2026 Q1 IT Support Review",
        "(January 1 - March 31, 2026)",
        "Prepared by Technijian Inc.  -  April 2026",
    )

    doc.add_paragraph()
    make_table(
        doc,
        ["Q1 2026 At a Glance", "Q1 2026", "vs 2025 Quarterly Avg"],
        [
            ["Unique Tickets", "64", "+58.0% (vs 40.5)"],
            ["Total Support Hours", "36.8", "+86.8% (vs 19.7)"],
            ["Avg Hours per Ticket", "0.58", "+18.4% (vs 0.49)"],
            ["Offshore NH", "21.1 (57.3%)", "-15.8pp share"],
            ["Offshore AH", "15.7 (42.7%)", "+15.8pp share"],
            ["Onsite", "0.0", "No change"],
            ["Unique Technicians", "12", "-"],
            ["Weekend Entries", "30 of 65 (46.2%)", "Near 2025 level"],
            ["Proactive Ticket Share", "93.8%", "Slightly below 95.1%"],
            ["Q1 Monthly Contract Billing", "$1,494.10", "+11.9% (vs $1,335.70)"],
        ],
        col_aligns=[L, R, R],
        highlight_row_indices=[1, 9],
    )

    doc.add_page_break()

    h1(doc, "Section 1: Support Volume and Trends")
    h2(doc, "1.1 Monthly Breakdown")
    make_table(
        doc,
        ["Month", "Tickets", "NH Hrs", "AH Hrs", "Total Hrs", "Hrs/Ticket"],
        [
            ["January", "16", "5.7", "3.6", "9.3", "0.58"],
            ["February", "21", "5.5", "9.7", "15.2", "0.73"],
            ["March", "27", "9.9", "2.4", "12.3", "0.46"],
            ["Q1 Total", "64", "21.1", "15.7", "36.8", "0.58"],
        ],
        col_aligns=[L, R, R, R, R, R],
        total_row_indices=[3],
        highlight_row_indices=[1],
    )
    bullet(doc, "Ticket volume increased month-over-month (16 -> 21 -> 27).")
    bullet(doc, "February had the highest effort because AH work spiked.")
    bullet(doc, "March delivered highest volume with best efficiency (0.46 hrs/ticket).")

    h2(doc, "1.2 Ticket Categorization")
    make_table(
        doc,
        ["Category", "Type", "Tickets", "Hours", "Avg Hrs"],
        [
            ["Patch Management", "Proactive", "25", "14.6", "0.59"],
            ["Backup and DR", "Proactive", "12", "3.4", "0.29"],
            ["Security and Endpoint", "Proactive", "11", "7.7", "0.70"],
            ["RMM and Agent Management", "Proactive", "8", "6.5", "0.81"],
            ["Monitoring and Alerts", "Proactive", "4", "1.8", "0.46"],
            ["Reactive (all categories)", "Reactive", "4", "2.8", "0.70"],
            ["TOTAL", "", "64", "36.8", "0.58"],
        ],
        col_aligns=[L, L, R, R, R],
        total_row_indices=[6],
    )

    make_table(
        doc,
        ["Type", "Tickets", "% Tickets", "Hours", "% Hours"],
        [
            ["Proactive", "60", "93.8%", "34.0", "92.4%"],
            ["Reactive", "4", "6.3%", "2.8", "7.6%"],
        ],
        col_aligns=[L, R, R, R, R],
    )

    doc.add_page_break()

    h1(doc, "Section 2: Staffing and Coverage")
    h2(doc, "2.1 Role Mix")
    make_table(
        doc,
        ["Role", "NH", "AH", "Onsite", "Total Hrs", "Entries"],
        [
            ["Off-Shore Tech Support", "20.6", "15.7", "0.0", "36.3", "64"],
            ["Tech Support (US)", "0.5", "0.0", "0.0", "0.5", "1"],
            ["TOTAL", "21.1", "15.7", "0.0", "36.8", "65"],
        ],
        col_aligns=[L, R, R, R, R, R],
        total_row_indices=[2],
    )

    h2(doc, "2.2 Top Technicians (Q1)")
    make_table(
        doc,
        ["Technician", "Hours"],
        [
            ["Rahul Uniyal", "10.7"],
            ["Sanjeev Kumar", "8.1"],
            ["Satish Sharma", "6.3"],
            ["Suresh Kumar Sharma", "4.1"],
            ["Aditya Saraf", "2.0"],
            ["Yogesh Kumar", "1.6"],
            ["Deepak Bhardwaj", "1.5"],
            ["Surinder Kumar", "1.0"],
            ["Gautam Chamoli", "0.5"],
            ["Rishad Mohamed", "0.5"],
        ],
        col_aligns=[L, R],
    )

    doc.add_page_break()

    h1(doc, "Section 3: Billing Snapshot and Cost Drivers")
    h2(doc, "3.1 Q1 Monthly Billing Trend")
    make_table(
        doc,
        ["Month", "Monthly Contract Total", "Primary Driver"],
        [
            ["Jan 2026", "$461.25", "Baseline with Site Assessment"],
            ["Feb 2026", "$493.25", "AVMH/PMW/SI rate changes (+$32)"],
            ["Mar 2026", "$539.60", "Offshore labor quantity increase (+$46.35)"],
            ["Q1 Total", "$1,494.10", ""],
        ],
        col_aligns=[L, R, L],
        total_row_indices=[3],
        highlight_row_indices=[1, 2],
    )

    h2(doc, "3.2 Q1 Billing by Category")
    make_table(
        doc,
        ["Category", "Q1 Total"],
        [
            ["Labor", "$740.10"],
            ["Security and Endpoint", "$340.00"],
            ["Assessment", "$150.00"],
            ["Monitoring and Management", "$136.00"],
            ["Secure Internet", "$128.00"],
            ["TOTAL", "$1,494.10"],
        ],
        col_aligns=[L, R],
        total_row_indices=[5],
    )

    h2(doc, "3.3 Labor Rate Context")
    make_table(
        doc,
        ["Role", "Q1 Revenue", "Billed Hours", "Effective Rate"],
        [
            ["US Tech Support", "$468.75", "3.75", "$125.00/hr"],
            ["Offshore AH", "$128.10", "4.27", "$30.00/hr"],
            ["Offshore NH", "$143.25", "9.55", "$15.00/hr"],
            ["ALL LABOR", "$740.10", "17.57", "$42.12/hr"],
        ],
        col_aligns=[L, R, R, R],
        total_row_indices=[3],
    )

    note(doc, "Blended labor rate is effectively flat vs 2025 ($42.20/hr), but monthly totals rose from January to March due to quantity and service-rate changes.")

    doc.add_page_break()

    h1(doc, "Section 4: Recommendations and Q2 Outlook")
    bullet(doc, "Validate March labor quantity increase against expected Q2 ticket cadence.")
    bullet(doc, "Monitor AH share (42.7% in Q1 vs 26.9% in 2025) to control blended support cost.")
    bullet(doc, "Sustain proactive posture above 90% while volume rises.")
    bullet(doc, "Track quarterly impact of AVMH/PMW/SI rate changes and labor quantity adjustments.")

    h2(doc, "4.1 2026 Run-Rate Outlook (if Q1 pace holds)")
    make_table(
        doc,
        ["Metric", "2025 Actual", "2026 Annualized (Q1 Pace)", "Change"],
        [
            ["Tickets", "162", "~256", "+58.0%"],
            ["Support Hours", "78.8", "~147.2", "+86.8%"],
            ["Monthly Contract", "$5,342.80", "~$5,976.40 to $6,475.20", "+11.9% to +21.2%"],
        ],
        col_aligns=[L, R, R, R],
    )

    note(doc, "Lower bound uses Q1 average monthly run rate; upper bound uses current March/April run rate.")

    os.makedirs(os.path.dirname(OUT_2026Q1), exist_ok=True)
    doc.save(OUT_2026Q1)
    print(f"Saved: {OUT_2026Q1}")


def main():
    build_aava_2025()
    build_aava_2026_q1()


if __name__ == "__main__":
    main()
