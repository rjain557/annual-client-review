"""
Generate BWH 2025 Annual IT Review — Professional Word Document
Technijian Inc. → Brandywine Homes
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ── Color palette ──────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
MED_BLUE = RGBColor(0x2E, 0x75, 0xB6)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

DARK_BLUE_HEX = "1F4E79"
ALT_ROW_HEX = "F2F7FB"
WHITE_HEX = "FFFFFF"
YELLOW_HEX = "FFF3CD"

OUT_DIR = r"c:\vscode\annual-client-review\annual-client-review\clients\bwh\2025"
OUT_FILE = os.path.join(OUT_DIR, "BWH - 2025 Annual Review.docx")


# ── Helpers ────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=Pt(8.5), color=BLACK,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, font_name="Calibri"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = font_name


def add_header_row(table, texts, col_aligns=None):
    row = table.rows[0]
    for i, txt in enumerate(texts):
        align = (col_aligns[i] if col_aligns else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[i], txt, bold=True, color=WHITE, alignment=align)
        set_cell_shading(row.cells[i], DARK_BLUE_HEX)


def add_data_row(table, texts, row_idx, col_aligns=None, bold=False,
                 is_total=False, is_highlight=False):
    row = table.rows[row_idx]
    bg = DARK_BLUE_HEX if is_total else (YELLOW_HEX if is_highlight else
         (ALT_ROW_HEX if row_idx % 2 == 0 else WHITE_HEX))
    txt_color = WHITE if is_total else BLACK
    for i, txt in enumerate(texts):
        align = (col_aligns[i] if col_aligns else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[i], txt, bold=(bold or is_total),
                      color=txt_color, alignment=align)
        set_cell_shading(row.cells[i], bg)


def make_table(doc, headers, data, col_aligns=None, total_row_indices=None,
               highlight_row_indices=None, col_widths=None):
    total_row_indices = total_row_indices or []
    highlight_row_indices = highlight_row_indices or []
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(data), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # Remove default borders, add thin gray
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:top w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:left w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:bottom w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:right w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '</w:tcBorders>')
            tcPr.append(borders)

    add_header_row(tbl, headers, col_aligns)
    for idx, row_data in enumerate(data):
        r = idx + 1
        is_total = idx in total_row_indices
        is_hl = idx in highlight_row_indices
        bold = is_total
        add_data_row(tbl, row_data, r, col_aligns, bold=bold,
                     is_total=is_total, is_highlight=is_hl)

    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                if w:
                    row.cells[i].width = w
    return tbl


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(16)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(13)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(11)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    return p


def add_page_break(doc):
    doc.add_page_break()


def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GRAY
    r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(4)
    return p


# ── Build document ─────────────────────────────────────────────
def build():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.63)
    section.right_margin = Inches(0.63)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # ── Header / Footer ──
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("BWH (Brandywine Homes) — 2025 IT Support Review & 2026 Planning")
    hr.font.size = Pt(8)
    hr.font.color.rgb = GRAY
    hr.font.name = "Calibri"

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Technijian Inc. — Confidential")
    fr.font.size = Pt(8)
    fr.font.color.rgb = GRAY
    fr.font.name = "Calibri"

    # ══════════════════════════════════════════════════════════════
    #  TITLE PAGE
    # ══════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BWH (Brandywine Homes)")
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2025 IT Support Review")
    r.font.size = Pt(20)
    r.font.color.rgb = MED_BLUE
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("& 2026 Planning Readiness Assessment")
    r.font.size = Pt(20)
    r.font.color.rgb = MED_BLUE
    r.font.name = "Calibri"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared by Technijian Inc.  —  March 2026")
    r.font.size = Pt(12)
    r.font.color.rgb = GRAY
    r.font.name = "Calibri"

    # ── At a Glance table ──
    doc.add_paragraph()
    R = WD_ALIGN_PARAGRAPH.RIGHT
    L = WD_ALIGN_PARAGRAPH.LEFT
    make_table(doc,
        ["2025 At a Glance", "Key Figures"],
        [
            ["Total 2025 Revenue", "$133,258"],
            ["Monthly Contract (Support + Services)", "$95,464"],
            ["Recurring Licensing", "$16,636"],
            ["Projects / One-Time", "$21,157"],
            ["Unique Tickets (2025)", "2,179"],
            ["Total Support Hours Delivered", "1,459"],
            ["Active Users (User List)", "101"],
            ["Cloud Servers Managed", "12"],
            ["Desktops Managed", "43 → 47 (Mar 2026)"],
        ],
        col_aligns=[L, R],
        total_row_indices=[0],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 1: Support Scope & Coverage
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 1: Support Scope & Coverage")

    add_heading2(doc, "1.1 What Is Included in Monthly Support")
    add_body(doc,
        "The monthly agreement covers a comprehensive managed IT package for BWH, "
        "totaling $95,464 across 2025 (~$7,955/month). This breaks down into two components:")

    add_body(doc,
        " Offshore technical support provides coverage across US business hours "
        "($30/hr, billed as offshore after-hours) and US overnight hours ($15/hr, billed "
        "as offshore normal hours). US-based technical support ($125/hr) handles escalations "
        "and onsite work. A Systems Architect advisory retainer ($200/hr, 5 hrs/month) "
        "provides strategic guidance.",
        bold_prefix="Labor (~$57,329/year, ~$4,777/month):")

    add_body(doc,
        " Security stack (CrowdStrike, Huntress, phishing training), monitoring "
        "(RMM, ManageEngine Ops Manager, patch management), backup (image backup, Veeam 365, "
        "cloud backup storage), email security (anti-spam, DKIM/DMARC), secure internet "
        "filtering, pen testing, and site assessment.",
        bold_prefix="Managed Services (~$38,135/year, ~$3,178/month):")

    add_heading2(doc, "1.2 Coverage Hours")
    add_bullet(doc,
        " 1,459 hours delivered across 2,179 unique tickets in 2025.",
        bold_prefix="Total Hours Delivered:")
    add_bullet(doc,
        " 974 hours (66.7% of total) at $15/hr — overnight monitoring and "
        "maintenance delivered by the offshore team during their normal business hours "
        "(India daytime = US overnight).",
        bold_prefix="US Overnight (Offshore NH):")
    add_bullet(doc,
        " 479 hours (32.9% of total) at $30/hr — US daytime support "
        "delivered by the offshore team during their after-hours (India night shift).",
        bold_prefix="US Business Hours (Offshore AH):")
    add_bullet(doc,
        " 5.9 hours (0.4%) across 3 visits — QNAP disk replacement, "
        "onsite network work, and a December visit.",
        bold_prefix="Onsite:")
    add_bullet(doc,
        " 659 time entries on Saturdays and Sundays (27.5% of all entries), "
        "confirming 7-day availability.",
        bold_prefix="Weekend Coverage:")
    add_bullet(doc,
        " 19 unique technicians, with 7 primary resources each contributing 100+ hours.",
        bold_prefix="Staffing Depth:")

    add_heading2(doc, "1.3 Remote vs. Onsite Mix")
    add_body(doc,
        "The model is 99.6% remote. Only 5.9 onsite hours were logged in 2025 (3 visits), "
        "with 6.0 hours of drive time. Onsite visits are reserved for hardware-related needs "
        "only (disk replacement, network equipment).")

    add_heading2(doc, "1.4 Services Billed Outside Monthly")
    add_bullet(doc,
        " Cybertraining, cloud backup, anti-spam (transitioned from monthly "
        "billing in April), M365 backup, Sophos firewall subscription, edge appliance, "
        "and Entra ID P1.",
        bold_prefix="Recurring Licensing ($16,636):")
    add_bullet(doc,
        " CTO IT Security Manual ($7,500), hardware purchases ($9,839), "
        "initial service setup fees ($2,161), and ad-hoc onsite labor ($1,650).",
        bold_prefix="One-Time Projects ($21,157):")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 2: Effective Blended Rates
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 2: Effective Blended Rates")
    add_body(doc,
        "The following table shows the effective blended rate per role, calculated as total "
        "billed revenue divided by total billed hours across all contract types (monthly "
        "contracted + one-time projects):")

    make_table(doc,
        ["Role", "Contracted", "One-Time", "Total Billed", "Hours", "Blended"],
        [
            ["IT Support", "$45,329", "$1,650", "$46,979", "1,287", "$36.50"],
            ["Systems Architect / CTO", "$12,000", "$7,500", "$19,500", "90", "$216.67"],
            ["ALL LABOR", "$57,329", "$9,150", "$66,479", "1,377", "$48.27"],
        ],
        col_aligns=[L, R, R, R, R, R],
        total_row_indices=[2],
    )

    add_note(doc,
        "Note: IT Support combines offshore NH ($15/hr, 584 hrs), offshore AH ($30/hr, "
        "526 hrs), US remote ($125/hr, 166 hrs), and one-time onsite/project labor "
        "($150/hr, 11 hrs). The overall blended rate of $36.50/hr reflects heavy offshore "
        "leverage. The Systems Architect role is a fixed 5 hrs/month retainer at $200/hr; "
        "the one-time CTO engagement (IT Security Manual) was billed at $250/hr.")

    add_subheading(doc, "Rate Card Breakdown")
    make_table(doc,
        ["Role / Tier", "Rate", "Coverage"],
        [
            ["Offshore Support (US Overnight)", "$15.00/hr", "US overnight monitoring/maintenance (India NH)"],
            ["Offshore Support (US Daytime)", "$30.00/hr", "US business hours support (India AH)"],
            ["US Tech Support (Remote)", "$125.00/hr", "Escalations, complex issues"],
            ["US Tech Support (Onsite)", "$150.00/hr", "On-premises hardware work"],
            ["Systems Architect (Retainer)", "$200.00/hr", "5 hrs/month advisory"],
            ["CTO Advisory (Project)", "$250.00/hr", "One-time project engagement"],
        ],
        col_aligns=[L, R, L],
    )

    add_body(doc,
        "IT Support achieves a low blended rate of $36.50/hr through heavy offshore leverage "
        "— 86% of support hours are offshore. US overnight support is delivered by the "
        "offshore team during their normal business hours at $15/hr, while US daytime "
        "support runs at $30/hr during India's after-hours shift. US tech support at "
        "$125–$150/hr is reserved for escalations and the rare onsite visit. The Systems "
        "Architect operates as a fixed 5 hrs/month retainer at $200/hr, providing strategic "
        "IT advisory and planning services.")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 3: Licensing & Recurring Services
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 3: Licensing & Recurring Services")

    add_heading2(doc, "3.1 Monthly Managed Services Breakdown")
    add_body(doc,
        "The monthly contract ($95,464/year, ~$7,955/month) covers both labor and managed "
        "services. The December 2025 invoice reflects the current service footprint:")

    add_subheading(doc, "Monthly Contract by Service Category (Annual)")
    make_table(doc,
        ["Service Category", "Annual", "Monthly"],
        [
            ["Labor (Support + Sys Architect)", "$57,329", "$4,777"],
            ["Security & Endpoint Protection", "$11,763", "$980"],
            ["Backup & Archiving", "$9,547", "$796"],
            ["Monitoring & Management (RMM/Ops)", "$7,702", "$642"],
            ["Email Security (Anti-Spam/DKIM)", "$4,771", "$397"],
            ["Secure Internet Filtering", "$2,596", "$216"],
            ["Pen Testing", "$1,757", "$146"],
            ["TOTAL", "$95,464", "$7,955"],
        ],
        col_aligns=[L, R, R],
        total_row_indices=[7],
    )

    add_subheading(doc, "December 2025 Monthly Invoice — Key Line Items")
    make_table(doc,
        ["Service", "Qty", "Unit Rate", "Monthly", "Category"],
        [
            ["US Tech Support (IRV-TS1)", "12.4", "$125.00", "$1,554", "Labor"],
            ["Offshore Support — US Daytime (CHD-TS1 AH)", "45.4", "$30.00", "$1,362", "Labor"],
            ["Systems Architect (IRV-AD1)", "5.0", "$200.00", "$1,000", "Labor"],
            ["Offshore Support — US Overnight (CHD-TS1 NH)", "59.3", "$15.00", "$889", "Labor"],
            ["Backup Storage (TB-BSTR)", "—", "—", "$600", "Backup"],
            ["Anti-Spam (ASA)", "89", "$6.25", "$556", "Email Security"],
            ["CrowdStrike — Desktop (AVD)", "43", "$8.50", "$366", "Security"],
            ["Veeam 365 Backup (V365)", "95", "$2.50", "$238", "Backup"],
            ["Huntress — Desktop (AVMH)", "43", "$5.00", "$215", "Security"],
            ["Phishing Training (PHT)", "38", "$5.00", "$190", "Security"],
            ["Ops Manager — Network (OPS-NET)", "55", "$3.25", "$179", "Monitoring"],
            ["Image Backup — Servers (IB)", "12", "$15.00", "$180", "Backup"],
            ["Secure Internet — Desktop (SI)", "43", "$4.00", "$172", "Secure Internet"],
            ["Patch Management (PMW)", "55", "$3.00", "$165", "Monitoring"],
            ["CrowdStrike — Server (AVS)", "12", "$10.50", "$126", "Security"],
            ["My Remote (MR)", "55", "$2.00", "$110", "Monitoring"],
            ["Network Assessment", "—", "—", "$63", "Monitoring"],
            ["Huntress — Server (AVHS)", "12", "$5.00", "$60", "Security"],
            ["Secure Internet — Server (SI)", "12", "$4.00", "$48", "Secure Internet"],
            ["Site Assessment (SA)", "1", "$50.00", "$50", "Monitoring"],
            ["Veeam ONE (VONE)", "—", "—", "$36", "Backup"],
            ["Ops Manager — Port (OPS-PRT)", "98", "$0.25", "$25", "Monitoring"],
            ["RTPT (Pen Testing)", "6", "$3.50", "$21", "Pen Testing"],
            ["DKIM/DMARC", "1", "$20.00", "$20", "Email Security"],
            ["Config Backup — Switches (OPS-BKP)", "3", "$6.00", "$18", "Monitoring"],
            ["Traffic Monitor — Firewall (OPS-TR)", "1", "$14.00", "$14", "Monitoring"],
            ["Storage Monitor (OPS-ST)", "—", "—", "$10", "Monitoring"],
            ["My Disk (MDU)", "1", "$5.00", "$5", "Monitoring"],
            ["WiFi Monitor (OPS-WF)", "4", "$1.00", "$4", "Monitoring"],
        ],
        col_aligns=[L, R, R, R, L],
    )

    add_page_break(doc)

    add_heading2(doc, "3.2 Recurring Licensing (Billed Separately)")
    add_body(doc,
        "Recurring items are billed on a separate monthly invoice. Several services "
        "transitioned from Monthly to Recurring billing in April 2025, and Sophos/Edge "
        "were added in September 2025.")

    make_table(doc,
        ["License / Service", "Seats", "$/Seat/Mo", "Monthly", "Annual (2025)", "Notes"],
        [
            ["Cybertraining", "83", "$5.00", "$415", "$3,735", "Apr–Dec (9 months)"],
            ["Server Cloud Backup", "8", "$50.00", "$400", "$3,600", "Apr–Dec (9 months)"],
            ["Anti-Spam Standard", "83", "$6.25", "$519", "$3,631", "Apr–Dec (split)"],
            ["M365 Backup Storage (est.)", "5", "$50.00", "$250", "$2,250", "Apr–Dec (9 months)"],
            ["Veeam 365 Backup", "83", "$2.50", "$208", "$1,868", "Apr–Dec (9 months)"],
            ["Sophos Firewall Sub (2C-4G)", "1", "$270.00", "$270", "$1,080", "Sep–Dec (4 months)"],
            ["Edge Appliance (MPC)", "1", "$100.00", "$100", "$400", "Sep–Dec (4 months)"],
            ["Microsoft Entra ID P1", "1", "$6.00", "$6", "$72", "Full year"],
            ["TOTAL", "", "", "$1,649", "$16,636", ""],
        ],
        col_aligns=[L, R, R, R, R, L],
        total_row_indices=[8],
    )

    add_note(doc,
        "Note: Anti-Spam and Phishing Training/Cybertraining were included in the Monthly "
        "invoice during Q1, then transitioned to the Recurring invoice starting April 2025. "
        "The Sophos firewall and Edge appliance subscriptions began in September after "
        "hardware installation in August. The 2025 annual totals reflect partial-year billing.")

    add_heading2(doc, "3.3 Service Count Tracking: December 2025 vs. March 2026")
    add_body(doc,
        "The following tables show the current 3/1/2026 service inventory alongside "
        "December 2025 billed quantities. The March 2026 counts establish the current "
        "baseline for 2026 billing projections and planning. Counts naturally fluctuate "
        "month-to-month as users and devices are added or removed.")

    add_subheading(doc, "M365 Licensing (Current Inventory)")
    make_table(doc,
        ["License", "In Use", "Dec Count", "Change"],
        [
            ["Office 365 E3", "60", "—", "New tracking"],
            ["Exchange Online (Plan 1)", "31", "—", "New tracking"],
            ["Audio Conferencing", "7", "—", "New tracking"],
            ["Planner & Project Plan 3", "10", "—", "New tracking"],
            ["Entra ID P1", "1", "1", "0 (flat)"],
        ],
        col_aligns=[L, R, R, L],
    )

    add_subheading(doc, "Security & Endpoint Protection (Dec → Mar)")
    make_table(doc,
        ["Service", "3/1/26 Count", "Dec Count", "Q1 Change", "Note"],
        [
            ["CrowdStrike — Desktop", "47", "43", "+4", "New workstations added"],
            ["CrowdStrike — Server", "13", "12", "+1", "FS01 marked missing"],
            ["Huntress — Desktop", "40", "43", "-3", "Current: 40 desktops covered"],
            ["Huntress — Server", "12", "12", "0", "Flat"],
            ["Cisco Umbrella — Desktop", "37", "43", "-6", "Current: 37 desktops covered"],
            ["Cisco Umbrella — Server", "12", "12", "0", "Other Services shows 11"],
            ["Phishing Training (PHT)", "85", "38", "+47", "Enrollment expanded"],
        ],
        col_aligns=[L, R, R, R, L],
        highlight_row_indices=[2, 4],
    )

    add_body(doc,
        "As of March 2026, CrowdStrike covers 47 desktops — the most complete deployment "
        "and the best indicator of actual device count. Huntress currently covers 40 desktops "
        "and Cisco Umbrella covers 37 desktops. Phishing training enrollment expanded "
        "significantly from 38 to 85 learners. These March 2026 counts form the billing "
        "baseline for 2026 projections.")

    add_subheading(doc, "Monitoring & Management (Dec → Mar)")
    make_table(doc,
        ["Service", "3/1/26 Count", "Dec Count", "Q1 Change", "Note"],
        [
            ["My Remote — Desktop", "46", "(part of 55)", "—", "Current: 46 desktops"],
            ["My Remote — Server", "12", "(part of 55)", "—", ""],
            ["My Remote — Total", "58", "55", "+3", ""],
            ["Patch Mgmt — Desktop", "42", "(part of 55)", "—", "Current: 42 desktops"],
            ["Patch Mgmt — Server", "12", "(part of 55)", "—", ""],
            ["Patch Mgmt — Total", "54", "55", "-1", ""],
            ["Ops Manager — Total", "44", "55", "-11", "Reconciled to actual"],
            ["WiFi APs (Ops Mgr)", "5", "4", "+1", ""],
        ],
        col_aligns=[L, R, R, R, L],
        highlight_row_indices=[3],
    )

    add_body(doc,
        "The Ops Manager device count dropped from 55 to 44, reflecting a reconciliation "
        "to actual monitored devices (22 desktops, 13 servers, 5 APs, 2 switches, 1 NAS, "
        "1 print server). Patch Management currently covers 42 desktops and 12 servers "
        "(54 total). These counts establish the monitoring baseline for 2026 projections.")

    add_subheading(doc, "Email & Backup Services (Dec → Mar)")
    make_table(doc,
        ["Service", "3/1/26 Count", "Dec Count", "Q1 Change", "Note"],
        [
            ["Anti-Spam", "83", "89", "-6", "Mailboxes cleaned up"],
            ["Image Backup (Servers)", "12", "12", "0", "Flat"],
            ["Cloud Backup Storage", "12 TB", "8 TB", "+4 TB", "Storage expanded"],
            ["Veeam ONE", "12", "—", "—", "Server monitoring"],
            ["My Disk", "1", "1", "0", "Flat"],
            ["DKIM/DMARC", "1", "1", "0", "Flat"],
        ],
        col_aligns=[L, R, R, R, L],
        highlight_row_indices=[2],
    )

    add_body(doc,
        "Anti-Spam decreased from 89 to 83 as inactive mailboxes were cleaned up. "
        "Cloud backup storage expanded significantly from 8 TB to 12 TB (+50%), "
        "reflecting growth in backup data volume.")

    add_subheading(doc, "Infrastructure Summary (3/1/2026)")
    make_table(doc,
        ["Category", "Count", "Detail"],
        [
            ["Servers", "12–13", "13 in CrowdStrike, 12 in other tools"],
            ["Desktops", "47", "Per CrowdStrike (best coverage indicator)"],
            ["Network Devices", "72", "46 DT, 12 svr, 5 AP, 4 ptr, 2 sw, 1 FW, 1 MPC, 1 NAS"],
            ["Users (User List)", "101", "Active user accounts"],
            ["M365 Licenses", "60 E3 + 31 Exch P1", "Plus 7 Audio, 10 Planner, 1 Entra"],
            ["Anti-Spam Protected", "83", "Protected mailboxes"],
            ["Phishing Training", "85", "Enrolled learners"],
            ["Cloud Backup Storage", "12 TB", "Server cloud backup"],
        ],
        col_aligns=[L, R, L],
    )

    add_subheading(doc, "Q1 Growth Summary")
    add_body(doc,
        "The March 2026 baseline reflects the current operational environment: 47 desktops "
        "(per CrowdStrike), 12\u201313 servers, 101 users, and 12 TB cloud backup storage. Key "
        "Q1 changes include 4 new workstations added, expanded cloud backup storage (+4 TB), "
        "and significantly expanded phishing training enrollment (38 \u2192 85 learners). Anti-spam "
        "was cleaned up (-6 mailboxes). The current per-tool coverage levels \u2014 CrowdStrike "
        "(47 desktops), Huntress (40), Cisco Umbrella (37), Patch Management (42) \u2014 represent "
        "the billing baseline for 2026 projections. The estimated net monthly cost change "
        "from count adjustments is approximately +$358/month, driven primarily by cloud "
        "backup expansion (+$200/mo) and phishing training enrollment (+$235/mo), partially "
        "offset by reduced anti-spam (-$38/mo) and monitoring consolidation.")

    add_body(doc,
        "Establish quarterly count reconciliation to keep invoices aligned with actual "
        "deployments as devices are added or removed. The March 2026 service inventory "
        "provides the baseline for all 2026 budget projections in Section 8.",
        bold_prefix="Recommendation: ")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 4: Projects & One-Time Spend
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 4: Projects & One-Time Spend (2025)")
    add_body(doc,
        "Total non-contract spend in 2025 was $21,157, broken into hardware purchases, "
        "a major CTO project, service setup costs, and ad-hoc labor.")

    add_heading2(doc, "4.1 Labor / Project Spend ($9,150)")
    make_table(doc,
        ["Item", "Rate", "Hours", "Amount", "Date", "Recurring?"],
        [
            ["CTO — IT Security Manual", "$250/hr", "30.0", "$7,500", "Mar 2025", "No"],
            ["Onsite — QNAP Disk Replacement", "$150/hr", "2.0", "$300", "May 2025", "No"],
            ["Onsite — Seagate EXOS Install", "$150/hr", "2.0", "$300", "May 2025", "No"],
            ["Sophos Firewall Installation", "$150/hr", "3.0", "$450", "Aug 2025", "No"],
            ["Meraki Sensor Installation", "$150/hr", "1.0", "$150", "Sep 2025", "No"],
            ["Other Installation Labor", "$150/hr", "3.0", "$450", "Various", "No"],
        ],
        col_aligns=[L, R, R, R, L, L],
    )

    add_note(doc,
        "Note: The IT Security Manual was the single largest one-time project in 2025 — "
        "a 30-hour CTO engagement at $250/hr to create comprehensive IT security "
        "documentation. This is a one-time deliverable that will not recur.")

    add_heading2(doc, "4.2 Hardware Purchases ($9,839)")
    make_table(doc,
        ["Item", "Amount", "Date", "One-Time?"],
        [
            ["Dell XPS 16 Laptop", "$3,269", "May 2025", "Yes"],
            ["CyberPower UPS System", "$3,163", "Jun 2025", "Yes"],
            ["Dell Tower Desktop (OptiPlex)", "$1,210", "Oct 2025", "Yes"],
            ["Seagate EXOS 20TB Drive", "$474", "May 2025", "Yes"],
            ["Docking Stations (x3)", "$786", "May–Jun 2025", "Yes"],
            ["QNAP NAS Disk", "$369", "May 2025", "Yes"],
            ["Meraki Environmental Sensor", "$370", "Sep 2025", "Yes"],
            ["Tripp Lite Rackmount Power Strip", "$200", "Jun 2025", "Yes"],
        ],
        col_aligns=[L, R, L, L],
    )

    add_heading2(doc, "4.3 Software / Service Setup ($2,161)")
    make_table(doc,
        ["Item", "Amount", "Date", "Recurs?"],
        [
            ["Initial Service Setup (Mar transition)", "$1,891", "Mar 2025", "No (one-time)"],
            ["Sophos Firewall Subscription (initial)", "$270", "Aug 2025", "→ Recurring"],
        ],
        col_aligns=[L, R, L, L],
    )

    add_heading2(doc, "4.4 Other One-Time Items ($7)")
    add_bullet(doc, " $7 — hardware disposal.", bold_prefix="E-Recycling Fees:")

    add_heading2(doc, "4.5 Predictability Recommendations")
    add_bullet(doc,
        " With 43 desktops on a 4–5 year cycle, plan for 8–10 rolling "
        "replacements per year rather than reactive purchases.",
        bold_prefix="Hardware Lifecycle Budget ($5K–$8K/yr):")
    add_bullet(doc,
        " Unlike some clients, BWH has no annual QuickBooks, 3CX, or "
        "SSL renewal obligations currently. If any are added, include them in a renewal calendar.",
        bold_prefix="No Recurring Software Renewals Detected:")
    add_bullet(doc,
        " The $7,500 IT Security Manual was a one-time engagement. If "
        "similar strategic projects are planned for 2026 (compliance audits, policy updates), "
        "budget $5K–$10K for ad-hoc CTO work.",
        bold_prefix="CTO/Project Budgeting:")
    add_bullet(doc,
        " The CyberPower UPS ($3,163) is a 5–7 year asset. Factor "
        "replacement into long-term budgets.",
        bold_prefix="UPS Replacement Cycle:")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 5: Ticket Categorization
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 5: Ticket Categorization (2025)")
    add_body(doc,
        "All 2,179 unique tickets were categorized by analyzing ticket titles and "
        "descriptions. The following table summarizes the work performed across 17 categories:")

    make_table(doc,
        ["Category", "Tickets", "% Tix", "Hours", "% Hrs", "Avg Hrs", "Proactive?"],
        [
            ["Patch Management",       "808", "37.1%", "405", "27.7%", "0.50", "Yes"],
            ["Security & Endpoint",    "209", "9.6%",  "131", "9.0%",  "0.63", "Yes"],
            ["Software Install/Update","201", "9.2%",  "176", "12.1%", "0.88", "Partial"],
            ["Monitoring & Alerts",    "179", "8.2%",  "87",  "6.0%",  "0.49", "Yes"],
            ["RMM & Agent Mgmt",       "157", "7.2%",  "108", "7.4%",  "0.69", "Yes"],
            ["Server Management",      "110", "5.0%",  "92",  "6.3%",  "0.84", "Partial"],
            ["Backup & DR",            "108", "5.0%",  "70",  "4.8%",  "0.65", "Yes"],
            ["Email & Microsoft 365",   "97", "4.5%",  "78",  "5.3%",  "0.80", "Reactive"],
            ["General IT Support",      "79", "3.6%",  "61",  "4.2%",  "0.78", "Reactive"],
            ["Workstation & Hardware",  "72", "3.3%",  "54",  "3.7%",  "0.75", "Reactive"],
            ["Firewall & Network",      "39", "1.8%",  "56",  "3.8%",  "1.43", "Reactive"],
            ["File & Permissions",      "38", "1.7%",  "43",  "3.0%",  "1.14", "Reactive"],
            ["Printing & Scanning",     "29", "1.3%",  "38",  "2.6%",  "1.30", "Reactive"],
            ["Phone / VoIP",            "18", "0.8%",  "26",  "1.8%",  "1.44", "Reactive"],
            ["User Onboard/Offboard",   "17", "0.8%",  "18",  "1.3%",  "1.09", "Reactive"],
            ["Password & Account Mgmt", "12", "0.6%",   "6",  "0.4%",  "0.50", "Reactive"],
            ["Domain & SSL",             "6", "0.3%",   "9",  "0.6%",  "1.42", "Reactive"],
            ["TOTAL",                "2,179","100%","1,459","100%",  "0.67", ""],
        ],
        col_aligns=[L, R, R, R, R, R, L],
        total_row_indices=[17],
    )

    add_heading2(doc, "5.1 Key Observations by Category")

    add_subheading(doc, "Patch Management — 808 tickets (37%), 405 hours (28%)")
    add_body(doc,
        "The single largest category by both ticket count and hours. This is entirely "
        "proactive work — identifying and remediating failed or missing patches on "
        "workstations and servers. Each patch ticket averages only 0.50 hours, reflecting "
        "a systematic, efficient process. This represents the core of the preventive "
        "maintenance program.")

    add_subheading(doc, "Security & Endpoint Protection — 209 tickets (10%), 131 hours (9%)")
    add_body(doc,
        "CrowdStrike and Huntress agent updates, version deployments, detection responses, "
        "and phishing-related alerts. Proactive security maintenance ensures the endpoint "
        "protection stack remains current across all 43 desktops and 12 servers.")

    add_subheading(doc, "Software Installation & Updates — 201 tickets (9%), 176 hours (12%)")
    add_body(doc,
        "Application deployments, updates, and upgrades across the environment. The higher "
        "average hours per ticket (0.88) reflects the complexity of coordinating software "
        "changes across multiple systems. Classified as \"Partial\" — some updates are "
        "proactive (scheduled) while others are reactive (user-requested).")

    add_subheading(doc, "Monitoring & Alerts — 179 tickets (8%), 87 hours (6%)")
    add_body(doc,
        "Automated monitoring alerts (memory utilization, disk utilization, device response "
        "time) from Ops Manager that are triaged and resolved. Very low average resolution "
        "time (0.49 hrs = ~29 minutes) indicates efficient triage.")

    add_subheading(doc, "RMM & Agent Management — 157 tickets (7%), 108 hours (7%)")
    add_body(doc,
        "Remote monitoring and management agent maintenance — ensuring RMM agents are "
        "running, updated, and reporting correctly across all managed endpoints. Core "
        "infrastructure management work.")

    add_subheading(doc, "Email & Microsoft 365 — 97 tickets (4%), 78 hours (5%)")
    add_body(doc,
        "Outlook configuration, Teams support, mailbox management, distribution list "
        "changes, and general M365 administration. The largest reactive category by "
        "ticket count.")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 6: Service Metrics & Trends
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 6: Service Metrics & Trends")

    add_heading2(doc, "6.1 Monthly Ticket Volume & Hours")
    make_table(doc,
        ["Month", "Tickets", "US Overnight (NH)", "US Daytime (AH)", "Total Hrs", "Hrs/Ticket"],
        [
            ["Jan",  "262", "82.9",  "55.8", "138.7", "0.53"],
            ["Feb",  "264", "71.6",  "43.2", "114.8", "0.43"],
            ["Mar",  "223", "97.2",  "37.1", "134.3", "0.60"],
            ["Apr",  "185", "79.1",  "44.7", "123.8", "0.67"],
            ["May",  "179", "73.8",  "37.8", "113.6", "0.63"],
            ["Jun",  "182", "86.8",  "49.2", "136.0", "0.75"],
            ["Jul",  "209", "91.9",  "47.2", "139.0", "0.67"],
            ["Aug",  "214", "113.1", "38.5", "151.6", "0.71"],
            ["Sep",  "208", "102.5", "67.0", "169.5", "0.81"],
            ["Oct*", "50",  "42.9",   "8.4",  "51.3", "1.03"],
            ["Nov",  "100", "59.6",  "18.4",  "79.3", "0.79"],
            ["Dec",  "116", "72.6",  "32.0", "107.3", "0.92"],
            ["Average","183","81.2", "40.0", "121.6", "0.67"],
        ],
        col_aligns=[L, R, R, R, R, R],
        highlight_row_indices=[9],
        total_row_indices=[12],
    )

    add_note(doc,
        "*October: The sharp drop to 50 tickets and 51.3 hours reflects the aftermath "
        "of a malware attack affecting other Technijian clients in late September. "
        "Support resources were temporarily redirected to incident response and recovery "
        "efforts across the affected client base, resulting in reduced routine support "
        "volume for BWH during October. This is not a reporting gap — it represents a "
        "genuine shift in resource allocation during the incident response period. Normal "
        "volumes resumed through November and December.")

    add_heading2(doc, "6.2 Notable Trends")
    add_bullet(doc,
        " Approximately 67% of all tickets (1,461) are proactive — patch "
        "management, monitoring alerts, security updates, backup management, and RMM agent "
        "maintenance. This indicates a strongly preventive maintenance program.",
        bold_prefix="Proactive vs. Reactive Split:")
    add_bullet(doc,
        " 974 offshore NH hours provide US overnight monitoring at $15/hr, "
        "while 479 offshore AH hours cover US daytime support at $30/hr. This is the "
        "intended 24/7 coverage model — the 'after-hours' label reflects India time zones, "
        "not a premium-rate issue.",
        bold_prefix="Coverage Model:")
    add_bullet(doc,
        " 808 tickets (37% of all work) is substantial. While essential, "
        "this high volume suggests opportunities for automation — reducing patch failure "
        "rates would directly lower remediation ticket volume.",
        bold_prefix="Patch Management Volume:")
    add_bullet(doc,
        " Oct–Dec saw 266 total tickets vs. 749 in Q1 and 571 in Q2. The "
        "October dip was caused by a late-September malware attack across other Technijian "
        "clients, which temporarily redirected support resources to incident response. "
        "November and December volumes recovered to 100 and 116 tickets respectively as "
        "normal operations resumed.",
        bold_prefix="Q4 Volume Drop:")
    add_bullet(doc,
        " 27.5% of entries fall on weekends, confirming robust 7-day "
        "monitoring coverage.",
        bold_prefix="Weekend Coverage:")

    add_heading2(doc, "6.3 Proactive Recommendations for 2026")
    add_bullet(doc,
        " With 808 patch tickets consuming 405 hours, investing in "
        "improved patch automation (auto-remediation, better scheduling) could significantly "
        "reduce manual remediation effort.",
        bold_prefix="Patch Automation (est. reduction: 100–200 tickets/yr):")
    add_bullet(doc,
        " The late-September malware attack across other Technijian clients "
        "temporarily reduced BWH support coverage in October. Ensure incident response "
        "resource planning accounts for multi-client events so that routine support for "
        "unaffected clients is maintained.",
        bold_prefix="Post-Incident Review:")
    add_bullet(doc,
        " The 479 offshore AH hours (US daytime at $30/hr) and 974 offshore "
        "NH hours (US overnight at $15/hr) reflect the intended 24/7 coverage model. "
        "Monitor the ratio to ensure US daytime demand is adequately staffed.",
        bold_prefix="Coverage Balance Review:")
    add_bullet(doc,
        " Budget $5K–$8K annually for rolling workstation replacements "
        "(43 desktops on a 4–5 year lifecycle).",
        bold_prefix="Hardware Refresh Planning:")
    add_bullet(doc,
        " Establish quarterly reviews to ensure billed service counts match "
        "actual deployments as devices are added or removed.",
        bold_prefix="Quarterly Count Reconciliation:")
    add_bullet(doc,
        " Implement structured QBRs to review ticket trends, service "
        "utilization, and project pipeline.",
        bold_prefix="Quarterly Business Reviews:")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 7: Summary
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 7: Summary")
    add_body(doc,
        "BWH's IT support structure is comprehensive and well-functioning. The model "
        "delivers 1,459 hours of support across 2,179 tickets annually, with 7-day "
        "coverage and deep expertise across infrastructure, security, patching, and "
        "end-user support.")

    add_body(doc,
        "At $133,258 total spend for 101 active users (47 desktops, 12 servers), "
        "the per-user IT cost is approximately $1,319/year. This covers fully managed "
        "infrastructure, security, backup, monitoring, email protection, and advisory services.")

    add_body(doc,
        "The ticket categorization reveals a healthy proactive/reactive balance: 67% of "
        "tickets are proactive (patching, monitoring, security, backups, RMM). The primary "
        "optimization opportunities are:")

    add_bullet(doc,
        " 808 tickets (37% of all work) in patching. Reducing patch failure "
        "rates through better automation would meaningfully reduce ticket volume.",
        bold_prefix="Patch Automation:")
    add_bullet(doc,
        " The blended labor rate of $36.50/hr for IT support reflects "
        "strong offshore leverage. 86% of support hours are offshore, keeping costs low "
        "while maintaining 24/7 coverage.",
        bold_prefix="Coverage Model Efficiency:")
    add_bullet(doc,
        " Establish a rolling replacement budget ($5K–$8K/yr) to avoid "
        "reactive hardware purchases.",
        bold_prefix="Hardware Lifecycle:")
    add_bullet(doc,
        " As of March 2026, CrowdStrike covers 47 desktops, Huntress covers "
        "40, Cisco Umbrella covers 37, and Patch Management covers 42. These counts "
        "form the billing baseline for 2026 projections.",
        bold_prefix="Current Service Baseline:")
    add_bullet(doc,
        " Q1 2026 shows +4 desktops, +4 TB backup storage, and +47 "
        "phishing training enrollments. Net monthly increase of ~$358/mo. Quarterly "
        "reconciliation recommended.",
        bold_prefix="Service Count Alignment:")
    add_bullet(doc,
        " The $7,500 CTO security manual and $9,839 hardware spend were "
        "the two largest one-time items. Budgeting for ad-hoc projects and hardware refresh "
        "reduces surprises.",
        bold_prefix="Predictability:")

    add_body(doc,
        "We are confident that with these adjustments, BWH will enter 2026 well-supported "
        "and operating efficiently.")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 8: 2026 Budget Projections
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 8: 2026 Budget Projections")
    add_body(doc,
        "The following projections use 2025 actuals as the baseline, adjusted for the "
        "contract adjustments that occurred mid-2025, the 3/1/2026 service inventory "
        "true-up, and known changes.")

    add_heading2(doc, "8.1 Contract Cycle & Adjustments")
    add_body(doc,
        "A single contract adjustment occurred in April/May 2025:",
        bold_prefix="Detected Pattern: ")
    add_bullet(doc, "US Tech Support hours decreased from ~18.2 hrs/month to ~12.4 hrs/month (-32%)")
    add_bullet(doc, "Offshore NH hours increased from ~37.5 hrs/month to ~51.6 hrs/month (+38%)")
    add_bullet(doc, "Offshore AH hours increased from ~38.9 hrs/month to ~45.4 hrs/month (+17%)")
    add_bullet(doc, "Rates remained unchanged across all roles")
    add_bullet(doc, "Several service items (Anti-Spam, Phishing Training) moved from Monthly to Recurring billing")
    add_bullet(doc, "Pen Testing was removed from the monthly invoice after July")

    add_body(doc,
        "This represents a cost optimization shift — moving labor from US-based ($125/hr) "
        "to offshore ($15–30/hr) while expanding the offshore team's coverage hours. The net "
        "effect increased total managed services while reducing per-hour labor costs.")

    add_note(doc,
        "No fixed 3/6/12-month cycle was detected. The adjustment appears to have been a "
        "one-time contract right-sizing. We recommend establishing a semi-annual review "
        "cadence going forward.")

    add_heading2(doc, "8.2 Support Cycle Analysis: Billed vs. Actual Hours")
    add_body(doc,
        "A critical input for 2026 projections is whether the current labor allotments "
        "match actual demand. The April/May 2025 contract adjustment shifted hours from "
        "US-based to offshore support. Analyzing the current cycle (May\u2013Dec 2025) "
        "reveals significant misalignment:")

    add_subheading(doc, "Billed vs. Actual Hours \u2014 Current Cycle (May\u2013Dec 2025 Average)")
    make_table(doc,
        ["Role", "Billed Hrs/Mo", "Actual Hrs/Mo", "Utilization", "Status"],
        [
            ["IT Support ($125/hr)", "12.4", "20.9", "166%", "Under-provisioned"],
            ["Offshore NH ($15/hr)", "51.6", "60.2", "115%", "Slightly under"],
            ["Offshore AH ($30/hr)", "45.4", "37.3", "82%", "Over-provisioned"],
            ["Sys Architect ($200/hr)", "5.0", "N/A (retainer)", "\u2014", "Fixed retainer"],
        ],
        col_aligns=[L, R, R, R, L],
        highlight_row_indices=[0],
    )

    add_body(doc,
        "The April contract adjustment reduced US Tech Support from ~18.2 to ~12.4 "
        "hrs/month, but actual demand averaged 20.9 hrs/month during the current cycle "
        "\u2014 a 166% utilization rate. Several months were extreme: August hit 31.8 "
        "actual hours (255%) and September hit 38.8 hours (312%). At $125/hr, the ~8.5 "
        "hrs/month shortfall represents approximately $1,063/month ($12,750/year) in labor "
        "delivered above the contracted amount.",
        bold_prefix="Key Finding \u2014 IT Support is significantly under-billed. ")

    add_body(doc,
        "Offshore NH hours are running slightly above contracted (115%), which is "
        "manageable at $15/hr. Offshore AH hours are over-provisioned at 82% utilization "
        "\u2014 the contract allots more US-daytime offshore hours than are being consumed.")

    add_note(doc,
        "Note: October actual hours were depressed across all roles due to the "
        "late-September malware attack redirecting resources. Excluding October, the "
        "IT Support utilization rate would be even higher.")

    add_subheading(doc, "Recommended Right-Sizing for Next Contract Cycle")
    make_table(doc,
        ["Role", "Current Allot.", "Recommended", "Rate", "Monthly Cost", "Change/Mo"],
        [
            ["IT Support", "12.4 hrs", "23.0 hrs", "$125/hr", "$2,875", "+$1,325"],
            ["Offshore NH", "51.6 hrs", "66.5 hrs", "$15/hr", "$998", "+$224"],
            ["Offshore AH", "45.4 hrs", "41.0 hrs", "$30/hr", "$1,230", "-$132"],
            ["Sys Architect", "5.0 hrs", "5.0 hrs", "$200/hr", "$1,000", "$0"],
            ["Total", "", "", "", "$6,103", "+$1,417"],
        ],
        col_aligns=[L, R, R, R, R, R],
        total_row_indices=[4],
    )

    add_body(doc,
        "The recommended allotments include a ~10% buffer above trailing averages. "
        "The net monthly labor increase would be approximately +$1,417/month "
        "(+$17,000/year), driven almost entirely by right-sizing the US Tech Support "
        "allotment. The offshore AH reduction partially offsets the increases.")

    add_heading2(doc, "8.3 Labor Projections")
    add_body(doc,
        "The labor projection presents two scenarios: (A) current contract unchanged, "
        "and (B) right-sized to match actual utilization.")

    make_table(doc,
        ["Component", "2025 Actual", "Scenario A", "Scenario B (Rec.)"],
        [
            ["Monthly Contracted Labor", "$57,329", "$57,660", "$73,230"],
            ["One-Time Labor (CTO/Projects)", "$9,150", "$2,000", "$2,000"],
            ["Total Labor", "$66,479", "$59,660", "$75,230"],
        ],
        col_aligns=[L, R, R, R],
        total_row_indices=[2],
    )

    add_body(doc,
        "Scenario A continues the December 2025 run rate (~$4,805/month). However, "
        "this perpetuates the IT Support under-provisioning \u2014 Technijian would "
        "deliver ~$12,750 in labor above the contracted amount annually.",
        bold_prefix="Scenario A (Current Contract): ")

    add_body(doc,
        "Adjusts allotments to match actual utilization with a 10% buffer. Monthly "
        "labor increases to ~$6,103/month ($73,230/year). The IT Support increase "
        "(+$1,325/mo) is the primary driver. We recommend Scenario B for the next "
        "contract adjustment.",
        bold_prefix="Scenario B (Right-Sized \u2014 Recommended): ")

    add_heading2(doc, "8.4 Managed Services Projections")
    make_table(doc,
        ["Component", "2025 Actual", "2026 Projected", "Change"],
        [
            ["Monthly Services (Dec rate)", "$38,135", "$41,628", "+$3,493"],
            ["Q1 2026 Count Adjustments", "—", "+$1,896", "+$1,896"],
            ["Total Managed Services", "$38,135", "$43,524", "+$5,389"],
        ],
        col_aligns=[L, R, R, R],
        total_row_indices=[2],
    )

    add_body(doc,
        "The December 2025 services run rate of ~$3,469/month is the baseline. The "
        "3/1/2026 service inventory true-up adds approximately +$158/month from net count "
        "changes: +4 CrowdStrike desktops, +1 server, expanded phishing training (+47 "
        "learners), partially offset by reduced anti-spam (-6), Ops Manager consolidation "
        "(-11 devices), and Cisco Umbrella gaps (-6 desktops). A full 12 months at the "
        "adjusted rate of ~$3,627/month yields $43,524.")

    add_heading2(doc, "8.5 Recurring Licensing Projections")
    make_table(doc,
        ["Component", "2025 Actual", "2026 Projected", "Change"],
        [
            ["Cybertraining", "$3,735", "$4,980", "+$1,245"],
            ["Server Cloud Backup (8→12 TB)", "$3,600", "$7,200", "+$3,600"],
            ["Anti-Spam Standard (83 users)", "$3,631", "$6,225", "+$2,594"],
            ["M365 Backup Storage", "$2,250", "$3,000", "+$750"],
            ["Veeam 365 Backup", "$1,868", "$2,490", "+$622"],
            ["Sophos Firewall", "$1,080", "$3,240", "+$2,160"],
            ["Edge Appliance", "$400", "$1,200", "+$800"],
            ["Entra ID P1", "$72", "$72", "$0"],
            ["Total Recurring", "$16,636", "$28,407", "+$11,771"],
        ],
        col_aligns=[L, R, R, R],
        total_row_indices=[8],
    )

    add_note(doc,
        "The large year-over-year increase reflects: (1) partial-year billing in 2025 — "
        "many items transitioned from Monthly to Recurring in April (9 months billed), "
        "Sophos/Edge started in September (4 months); (2) cloud backup storage expansion "
        "from 8 TB to 12 TB (+$200/month). The corresponding Monthly invoice decreased "
        "when items moved to Recurring. The combined total is the better comparison.")

    add_heading2(doc, "8.6 One-Time & Project Spend")
    make_table(doc,
        ["Component", "2025 Actual", "2026 Budget", "Change"],
        [
            ["Hardware Lifecycle", "$9,839", "$7,500", "-$2,339"],
            ["CTO / Project Labor", "$9,150", "$2,000", "-$7,150"],
            ["Service Setup", "$2,161", "$0", "-$2,161"],
            ["Other", "$7", "$0", "-$7"],
            ["Total One-Time", "$21,157", "$9,500", "-$11,657"],
        ],
        col_aligns=[L, R, R, R],
        total_row_indices=[4],
    )

    add_body(doc,
        "The CTO IT Security Manual ($7,500) and initial service setup charges ($2,161) "
        "were one-time items that will not recur. Hardware budget of $7,500 funds 2–3 "
        "workstation replacements or infrastructure needs.")

    add_page_break(doc)

    add_heading2(doc, "8.7 Total 2026 Projection Summary")
    make_table(doc,
        ["Category", "2025 Actual", "Scenario A", "Scenario B (Rec.)"],
        [
            ["Monthly Contract (Labor)", "$57,329", "$57,660", "$73,230"],
            ["Monthly Contract (Services)", "$38,135", "$43,524", "$43,524"],
            ["Recurring Licensing", "$16,636", "$28,407", "$28,407"],
            ["One-Time / Projects", "$21,157", "$9,500", "$9,500"],
            ["TOTAL 2026", "$133,258", "$139,091", "$154,661"],
        ],
        col_aligns=[L, R, R, R],
        total_row_indices=[4],
    )

    add_body(doc,
        "Maintains the current labor allotments. Total projected at $139,091 (+4.4%). "
        "This scenario continues the under-provisioning of IT Support at $125/hr, meaning "
        "Technijian delivers ~$12,750/year in labor above the contracted amount.",
        bold_prefix="Scenario A (Current Contract): ")

    add_body(doc,
        "Adjusts labor allotments to match actual utilization with a 10% buffer. "
        "Total projected at $154,661 (+16.1%). The increase is driven by right-sizing "
        "the IT Support allotment from 12.4 to 23.0 hrs/month (+$1,325/mo) to reflect "
        "actual demand. This eliminates the under-billing gap and ensures the contract "
        "accurately reflects the services being delivered.",
        bold_prefix="Scenario B (Right-Sized \u2014 Recommended): ")

    add_subheading(doc, "Key Assumptions")
    add_bullet(doc,
        " Scenario A continues current rates (~$4,805/month). "
        "Scenario B right-sizes to ~$6,103/month based on actual utilization. "
        "IT Support increases from 12.4 to 23.0 hrs/mo; Offshore AH decreases "
        "from 45.4 to 41.0 hrs/mo.",
        bold_prefix="Labor:")
    add_bullet(doc,
        " Adjusted monthly on the invoice to match actual counts. The "
        "March 2026 service inventory is used as the baseline. Each new desktop adds "
        "approximately $50–55/month across the full security, monitoring, and management stack.",
        bold_prefix="Device Services:")
    add_bullet(doc,
        " Projection assumes stable headcount at ~101 users. Each new user "
        "adds approximately $15–20/month in incremental per-user services (anti-spam, "
        "phishing training, Veeam backup).",
        bold_prefix="Headcount:")
    add_bullet(doc,
        " No major infrastructure expansions assumed. Server counts remain "
        "stable at 12–13. Cloud backup storage at 12 TB.",
        bold_prefix="Infrastructure:")
    add_bullet(doc,
        " $7,500 budget assumes normal replacement cycle (2–3 units). "
        "A large-scale refresh (e.g., Windows 11 migration for all 47 desktops) would "
        "significantly exceed this.",
        bold_prefix="Hardware:")
    add_bullet(doc,
        " Minimal ad-hoc project work assumed ($2,000 buffer). Significant "
        "strategic initiatives would require separate budgeting.",
        bold_prefix="CTO/Projects:")
    add_bullet(doc,
        " Projections use March 2026 service counts as the baseline \u2014 "
        "CrowdStrike (47 desktops), Huntress (40), Cisco Umbrella (37), Patch Management "
        "(42). Per-device costs are projected at current coverage levels.",
        bold_prefix="Service Counts:")

    # ── Save ───────────────────────────────────────────────────
    doc.save(OUT_FILE)
    print(f"Document saved to: {OUT_FILE}")


if __name__ == "__main__":
    build()
