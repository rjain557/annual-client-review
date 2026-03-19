"""
Generate HHOC 2026 Q1 IT Review — Word Document (with rate context from 2025 invoices)
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
MED_BLUE = RGBColor(0x2E, 0x75, 0xB6)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_BLUE_HEX = "1F4E79"
ALT_ROW_HEX = "F2F7FB"
WHITE_HEX = "FFFFFF"
YELLOW_HEX = "FFF3CD"

OUT = r"c:\vscode\annual-client-review\annual-client-review\clients\hhoc\2026\HHOC - 2026 Q1 Review.docx"

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, size=Pt(8.5), color=BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = alignment
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text)); run.font.size = size; run.font.color.rgb = color; run.font.bold = bold; run.font.name = "Calibri"

def add_header_row(table, texts, col_aligns=None):
    row = table.rows[0]
    for i, txt in enumerate(texts):
        align = (col_aligns[i] if col_aligns else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[i], txt, bold=True, color=WHITE, alignment=align)
        set_cell_shading(row.cells[i], DARK_BLUE_HEX)

def add_data_row(table, texts, row_idx, col_aligns=None, is_total=False, is_highlight=False):
    row = table.rows[row_idx]
    bg = DARK_BLUE_HEX if is_total else (YELLOW_HEX if is_highlight else (ALT_ROW_HEX if row_idx % 2 == 0 else WHITE_HEX))
    txt_color = WHITE if is_total else BLACK
    for i, txt in enumerate(texts):
        align = (col_aligns[i] if col_aligns else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[i], txt, bold=is_total, color=txt_color, alignment=align)
        set_cell_shading(row.cells[i], bg)

def make_table(doc, headers, data, col_aligns=None, total_row_indices=None, highlight_row_indices=None):
    total_row_indices = total_row_indices or []; highlight_row_indices = highlight_row_indices or []
    tbl = doc.add_table(rows=1 + len(data), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; tbl.style = "Table Grid"
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:color="D0D0D0"/><w:left w:val="single" w:sz="4" w:color="D0D0D0"/><w:bottom w:val="single" w:sz="4" w:color="D0D0D0"/><w:right w:val="single" w:sz="4" w:color="D0D0D0"/></w:tcBorders>')
            tcPr.append(borders)
    add_header_row(tbl, headers, col_aligns)
    for idx, row_data in enumerate(data):
        add_data_row(tbl, row_data, idx + 1, col_aligns, is_total=idx in total_row_indices, is_highlight=idx in highlight_row_indices)
    return tbl

def h1(doc, text):
    p = doc.add_paragraph(); p.style = doc.styles["Heading 1"]
    r = p.add_run(text); r.font.color.rgb = DARK_BLUE; r.font.size = Pt(16); r.font.bold = True

def h2(doc, text):
    p = doc.add_paragraph(); p.style = doc.styles["Heading 2"]
    r = p.add_run(text); r.font.color.rgb = DARK_BLUE; r.font.size = Pt(13); r.font.bold = True

def sh(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.color.rgb = DARK_BLUE; r.font.size = Pt(11); r.font.bold = True

def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.font.size = Pt(10); r.font.bold = True; r.font.name = "Calibri"
    r = p.add_run(text); r.font.size = Pt(10); r.font.name = "Calibri"

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.font.size = Pt(10); r.font.bold = True; r.font.name = "Calibri"
    r = p.add_run(text); r.font.size = Pt(10); r.font.name = "Calibri"

def note(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = GRAY; r.font.name = "Calibri"

L = WD_ALIGN_PARAGRAPH.LEFT
R = WD_ALIGN_PARAGRAPH.RIGHT

def build():
    doc = Document()
    s = doc.sections[0]
    s.page_width = Inches(8.5); s.page_height = Inches(11)
    s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(0.63); s.right_margin = Inches(0.63)
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)
    hp = s.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("HHOC (Housing for Health OC) — 2026 Q1 IT Support Review"); hr.font.size = Pt(8); hr.font.color.rgb = GRAY
    fp = s.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Technijian Inc. — Confidential"); fr.font.size = Pt(8); fr.font.color.rgb = GRAY

    # Title
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HHOC (Housing for Health OC)"); r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = DARK_BLUE
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026 Q1 IT Support Review"); r.font.size = Pt(20); r.font.color.rgb = MED_BLUE
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("(January – March 2026)"); r.font.size = Pt(20); r.font.color.rgb = MED_BLUE
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared by Technijian Inc.  —  March 2026"); r.font.size = Pt(12); r.font.color.rgb = GRAY

    doc.add_paragraph()
    make_table(doc,
        ["Q1 2026 At a Glance", "Q1 2026", "2025 Q Avg", "Change"],
        [
            ["Unique Tickets", "54", "42", "+28%"],
            ["Total Hours", "69.0", "37.3", "+85%"],
            ["Hours/Ticket", "1.28", "0.88", "+45%"],
            ["US Tech Support Hours", "26.4", "10.6", "+149%"],
            ["Proactive %", "13.0%", "35.5%", "-22.5pp"],
            ["Email & M365 % of Hours", "67.4%", "46.4%", "+21.0pp"],
            ["Estimated Labor Cost*", "~$5,286", "~$2,058/q", "+157%"],
        ],
        col_aligns=[L, R, R, R], highlight_row_indices=[1, 4, 5, 6])

    note(doc, "* Estimated using Dec 2025 rates: US Tech $150/hr, Offshore AH $45/hr, Offshore NH $15/hr.")

    doc.add_page_break()

    # Section 1
    h1(doc, "Section 1: Q1 Overview — Key Concerns")
    body(doc,
        "Q1 2026 shows significant changes from 2025 patterns. Total hours nearly doubled "
        "compared to the 2025 quarterly average (69.0 vs 37.3), driven primarily by complex "
        "Email & M365 work. The proactive/reactive balance shifted dramatically toward "
        "reactive (87% vs 65% in 2025).")

    h2(doc, "1.1 Estimated Financial Impact")
    body(doc,
        "Using December 2025 contracted rates:")
    make_table(doc,
        ["Role", "Q1 Hours", "Rate", "Est. Cost"],
        [
            ["US Tech Support (NH)", "26.4", "$150/hr", "$3,960"],
            ["Offshore AH (US Daytime)", "21.2", "$45/hr", "$954"],
            ["Offshore NH (US Overnight)", "21.4", "$15/hr", "$321"],
            ["Q1 TOTAL", "69.0", "", "$5,235"],
            ["Annualized", "", "", "$20,940"],
        ],
        col_aligns=[L, R, R, R], total_row_indices=[3], highlight_row_indices=[4])

    note(doc,
        "If Q1 consumption continues, annualized labor alone would be ~$20,940 vs $8,231 "
        "actual in 2025 — a 154% increase. The contracted allotment will need review at "
        "the next quarterly adjustment.")

    h2(doc, "1.2 Monthly Volume")
    make_table(doc,
        ["Month", "Tickets", "NH", "AH", "Total", "Hrs/Tkt"],
        [
            ["Jan 2026", "19", "19.9", "7.2", "27.1", "1.43"],
            ["Feb 2026", "15", "16.8", "6.3", "23.1", "1.54"],
            ["Mar 2026", "20", "11.1", "7.7", "18.9", "0.94"],
        ],
        col_aligns=[L, R, R, R, R, R], highlight_row_indices=[1])
    body(doc,
        "Hours/ticket started very high (1.43-1.54 in Jan-Feb) and normalized in March "
        "(0.94, close to 2025 avg of 0.88). This suggests a burst of complex work early "
        "in Q1 that is starting to stabilize.")

    doc.add_page_break()

    # Section 2
    h1(doc, "Section 2: Ticket Categorization — Email & M365 Deep Dive")
    make_table(doc,
        ["Category", "Type", "Tickets", "Hours", "Avg Hrs"],
        [
            ["Email & M365", "Reactive", "22", "46.5", "2.11"],
            ["General IT Support", "Reactive", "9", "3.8", "0.42"],
            ["Patch Management", "Proactive", "5", "3.0", "0.59"],
            ["Firewall & Network", "Reactive", "3", "2.2", "0.75"],
            ["File & Permissions", "Reactive", "3", "1.8", "0.58"],
            ["Password & Account Mgmt", "Reactive", "2", "1.0", "0.50"],
            ["Software Install & Updates", "Reactive", "2", "1.8", "0.88"],
            ["Security & Endpoint", "Proactive", "2", "1.4", "0.72"],
            ["Phone / VoIP", "Reactive", "2", "1.6", "0.78"],
            ["Other (App, Onboard, HW, Server)", "Mixed", "4", "6.0", "1.50"],
            ["TOTAL", "", "54", "69.0", "1.28"],
        ],
        col_aligns=[L, L, R, R, R], total_row_indices=[10], highlight_row_indices=[0])

    sh(doc, "Critical Finding: Email & M365")
    body(doc,
        "Email & M365 consumed 46.5 hours across 22 tickets — 67.4% of all Q1 hours. "
        "The average hours/ticket jumped from 0.99 in 2025 to 2.11 in Q1 2026, indicating "
        "significantly more complex issues, not just more volume.")
    body(doc,
        "At the US Tech Support rate of $150/hr, Email & M365 work alone may be consuming "
        "~$3,000-4,000 per quarter in labor, exceeding the total 2025 quarterly labor spend.")

    sh(doc, "Proactive Work Collapse")
    make_table(doc,
        ["Metric", "2025 Full Year", "Q1 2026", "Concern"],
        [
            ["Proactive Tickets", "60 (35.5%)", "7 (13.0%)", "Down 22.5pp"],
            ["Patch Mgmt Tickets", "17", "5", "On pace for 20 (ok)"],
            ["Monitoring Tickets", "10", "0", "Missing entirely"],
            ["RMM Tickets", "13", "0", "Missing entirely"],
            ["Backup Tickets", "7", "0", "Missing entirely"],
        ],
        col_aligns=[L, R, R, L], highlight_row_indices=[2, 3, 4])

    body(doc,
        "Monitoring, RMM, and Backup generated zero tickets in Q1. This could mean "
        "these systems are running smoothly, or it could indicate reduced oversight "
        "during the staffing transition. Verification is recommended.")

    doc.add_page_break()

    # Section 3
    h1(doc, "Section 3: Staffing Transition")
    make_table(doc,
        ["Technician", "Q1 2026 Hrs", "2025 Hrs", "Notes"],
        [
            ["Rishad Mohamed", "19.6", "0", "New primary US Tech"],
            ["Deepak Bhardwaj", "16.7", "0", "New offshore lead"],
            ["Surinder Kumar", "7.7", "10.1", "Continuing"],
            ["Sai Revanth", "6.8", "3.0", "Increased role"],
            ["Gurdeep Kumar", "4.2", "12.2", "Reduced"],
            ["Other (6 techs)", "14.0", "—", "Various"],
            ["Sunny Sark", "0", "38.8", "No longer assigned"],
        ],
        col_aligns=[L, R, R, L], highlight_row_indices=[0, 1, 6])

    body(doc,
        "Sunny Sark (2025 primary US Tech, 38.8 hrs) is no longer assigned to HHOC. "
        "Rishad Mohamed (19.6 hrs) has taken over as primary US resource, with Deepak "
        "Bhardwaj (16.7 hrs) as the new offshore lead.")

    body(doc,
        "US Tech Support hours increased as a percentage (38.3% in Q1 vs 28.4% in 2025). "
        "All Q1 entries are under the 'Monthly Service with India-Night' contract "
        "(vs 83.4% on standard 'Monthly Service' in 2025), indicating the contract "
        "structure has consolidated.")

    doc.add_page_break()

    # Section 4
    h1(doc, "Section 4: Recommendations & Q2 Action Items")

    h2(doc, "Immediate Actions")
    bullet(doc, " Review the 22 Email & M365 tickets to identify root causes. "
           "The 2.11 hrs/ticket avg suggests complex projects or systemic M365 issues. "
           "Consider a dedicated M365 optimization engagement.",
           bold_prefix="1. Audit Email & M365 drivers:")
    bullet(doc, " Verify that patching, monitoring, RMM, and backup routines "
           "are still running. Zero tickets in these categories may indicate reduced "
           "oversight rather than perfect health.",
           bold_prefix="2. Restore proactive maintenance:")
    bullet(doc, " Ensure Rishad Mohamed and Deepak Bhardwaj have full "
           "HHOC environment context. Complete knowledge transfer from prior team.",
           bold_prefix="3. Validate staffing transition:")
    bullet(doc, " At Q1 run rate, labor spend is tracking 154% above 2025. "
           "Review the contracted allotment at the next quarterly adjustment.",
           bold_prefix="4. Review labor allotment:")

    h2(doc, "Q2 Monitoring Targets")
    make_table(doc,
        ["Metric", "Q1 Actual", "Q2 Target", "Rationale"],
        [
            ["Email & M365 hrs/ticket", "2.11", "< 1.5", "Reduce complexity"],
            ["Proactive ratio", "13.0%", "> 25%", "Restore maintenance"],
            ["Patch Mgmt tickets", "5", "8+", "Ensure coverage"],
            ["Total quarterly hours", "69.0", "< 50", "Normalize to 2025 levels"],
            ["Total quarterly est. cost", "~$5,235", "< $3,500", "Align with budget"],
        ],
        col_aligns=[L, R, R, L])

    note(doc,
        "March 2026 showed encouraging normalization (0.94 hrs/ticket vs 1.43-1.54 in Jan-Feb). "
        "If this trend continues, Q2 should approach these targets. However, if Email & M365 "
        "complexity persists, a dedicated M365 health check engagement is recommended.")

    doc.save(OUT)
    print(f"Saved: {OUT}")

build()
