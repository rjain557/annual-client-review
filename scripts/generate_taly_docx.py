"""
Generate TALY 2025 Annual IT Review — Professional Word Document
Technijian Inc. → Talley & Associates
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Color palette ──────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
MED_BLUE = RGBColor(0x2E, 0x75, 0xB6)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

DARK_BLUE_HEX = "1F4E79"
ALT_ROW_HEX = "F2F7FB"
WHITE_HEX = "FFFFFF"
YELLOW_HEX = "FFF3CD"
GREEN_HEX = "D5F5E3"

OUT_DIR = r"c:\vscode\annual-client-review\annual-client-review\clients\taly\2025"
OUT_FILE = os.path.join(OUT_DIR, "TALY - 2025 Annual Review.docx")


# ── Helpers ────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=Pt(8.5), color=BLACK,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, font_name="Calibri"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = font_name


def add_header_row(table, texts, col_aligns=None):
    row = table.rows[0]
    for i, txt in enumerate(texts):
        align = (col_aligns[i] if col_aligns else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[i], txt, bold=True, color=WHITE, alignment=align)
        set_cell_shading(row.cells[i], DARK_BLUE_HEX)


def add_data_row(table, texts, row_idx, col_aligns=None, bold=False,
                 is_total=False, is_highlight=False, is_green=False):
    row = table.rows[row_idx]
    if is_total:
        bg = DARK_BLUE_HEX
    elif is_green:
        bg = GREEN_HEX
    elif is_highlight:
        bg = YELLOW_HEX
    else:
        bg = ALT_ROW_HEX if row_idx % 2 == 0 else WHITE_HEX
    txt_color = WHITE if is_total else BLACK
    for i, txt in enumerate(texts):
        align = (col_aligns[i] if col_aligns else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[i], txt, bold=(bold or is_total),
                      color=txt_color, alignment=align)
        set_cell_shading(row.cells[i], bg)


def make_table(doc, headers, data, col_aligns=None, total_row_indices=None,
               highlight_row_indices=None, green_row_indices=None, col_widths=None):
    total_row_indices = total_row_indices or []
    highlight_row_indices = highlight_row_indices or []
    green_row_indices = green_row_indices or []
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(data), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:top w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:left w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:bottom w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:right w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '</w:tcBorders>')
            tcPr.append(borders)

    add_header_row(tbl, headers, col_aligns)
    for idx, row_data in enumerate(data):
        r = idx + 1
        is_total = idx in total_row_indices
        is_hl = idx in highlight_row_indices
        is_gr = idx in green_row_indices
        add_data_row(tbl, row_data, r, col_aligns, bold=is_total,
                     is_total=is_total, is_highlight=is_hl, is_green=is_gr)

    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                if w:
                    row.cells[i].width = w
    return tbl


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(16)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(13)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(11)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    return p


def add_page_break(doc):
    doc.add_page_break()


def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GRAY
    r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(4)
    return p


# ── Build document ─────────────────────────────────────────────
def build():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.63)
    section.right_margin = Inches(0.63)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # ── Header / Footer ──
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("TALY (Talley & Associates) \u2014 2025 IT Support Review & 2026 Planning")
    hr.font.size = Pt(8)
    hr.font.color.rgb = GRAY
    hr.font.name = "Calibri"

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Technijian Inc. \u2014 Confidential")
    fr.font.size = Pt(8)
    fr.font.color.rgb = GRAY
    fr.font.name = "Calibri"

    R = WD_ALIGN_PARAGRAPH.RIGHT
    L = WD_ALIGN_PARAGRAPH.LEFT

    # ══════════════════════════════════════════════════════════════
    #  TITLE PAGE
    # ══════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TALY (Talley & Associates)")
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2025 IT Support Review")
    r.font.size = Pt(20)
    r.font.color.rgb = MED_BLUE
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("& 2026 Planning Readiness Assessment")
    r.font.size = Pt(20)
    r.font.color.rgb = MED_BLUE
    r.font.name = "Calibri"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared by Technijian Inc.  \u2014  March 2026")
    r.font.size = Pt(12)
    r.font.color.rgb = GRAY
    r.font.name = "Calibri"

    # ── At a Glance table ──
    doc.add_paragraph()
    make_table(doc,
        ["2025 At a Glance", "Key Figures"],
        [
            ["Total 2025 Revenue", "$9,336"],
            ["Monthly Contract (Support + Services)", "$8,261"],
            ["Hourly / Onsite Labor", "$1,075"],
            ["Unique Tickets (2025)", "54"],
            ["Total Support Hours Delivered", "63.9"],
            ["Active Users", "~5"],
            ["Desktops Managed", "9 \u2192 7 (Mar 2026)"],
            ["Network Devices Monitored", "16"],
            ["Backup Storage", "1 TB"],
        ],
        col_aligns=[L, R],
        total_row_indices=[0],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 1: Support Scope & Coverage
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 1: Support Scope & Coverage")

    add_heading2(doc, "1.1 What Is Included in Monthly Support")
    add_body(doc,
        "The monthly agreement covers a managed IT package for TALY, totaling "
        "$8,261 across 2025 (~$688/month). This breaks down into two components:")

    add_body(doc,
        " US-based technical support ($125/hr remote, $150\u2013200/hr onsite) handles "
        "escalations, advisory, and hands-on work. Offshore technical support provides "
        "US daytime coverage ($30/hr) and US overnight coverage ($15/hr).",
        bold_prefix="Labor (~$3,907/year, ~$326/month):")

    add_body(doc,
        " Security stack (CrowdStrike, Huntress), monitoring (RMM, syslog/health "
        "monitoring, patch management), backup (cloud backup storage), secure internet "
        "filtering, pen testing, and site assessment.",
        bold_prefix="Managed Services (~$4,354/year, ~$363/month):")

    add_heading2(doc, "1.2 Coverage Hours")
    add_bullet(doc,
        " 63.9 hours delivered across 54 unique tickets in 2025.",
        bold_prefix="Total Hours Delivered:")
    add_bullet(doc,
        " 28.8 hours (45.1%) at $125/hr remote, plus 6.2 hours onsite "
        "at $150\u2013200/hr billed separately.",
        bold_prefix="US Tech Support:")
    add_bullet(doc,
        " 15.0 hours (23.5%) at $15/hr \u2014 overnight monitoring and "
        "maintenance (India daytime = US overnight).",
        bold_prefix="Offshore US Overnight (NH):")
    add_bullet(doc,
        " 20.0 hours (31.3%) at $30/hr \u2014 US daytime support "
        "(India night shift).",
        bold_prefix="Offshore US Daytime (AH):")
    add_bullet(doc,
        " 6.2 hours across 4 visits \u2014 device onboarding (Jan), "
        "network troubleshooting (Nov\u2013Dec).",
        bold_prefix="Onsite:")
    add_bullet(doc,
        " 15 unique technicians, with Gurdeep Kumar (25.1 hrs) and "
        "Hamid Yaghoubi (17.6 hrs) as primary resources.",
        bold_prefix="Staffing Depth:")

    add_heading2(doc, "1.3 Remote vs. Onsite Mix")
    add_body(doc,
        "The model is 90.3% remote. 6.2 onsite hours were logged in 2025 across "
        "4 visits, with 2.5 hours of drive time. Onsite visits covered device "
        "onboarding and network troubleshooting.")

    add_heading2(doc, "1.4 Services Billed Outside Monthly")
    add_bullet(doc,
        " 4 onsite visits billed at $150\u2013200/hr. January visit "
        "at $150/hr (2.5 hrs = $375); November and December visits at "
        "$150\u2013200/hr ($700 combined).",
        bold_prefix="Hourly / Onsite Labor ($1,075):")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 2: Effective Blended Rates
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 2: Effective Blended Rates")

    make_table(doc,
        ["Role", "Contracted", "Hourly/Onsite", "Total Billed", "Hours", "Blended"],
        [
            ["US Tech Support", "$2,915", "$1,075", "$3,990", "28.8", "$138.54"],
            ["Offshore Support", "$992", "$0", "$992", "35.0", "$28.34"],
            ["ALL LABOR", "$3,907", "$1,075", "$4,982", "63.9", "$77.97"],
        ],
        col_aligns=[L, R, R, R, R, R],
        total_row_indices=[2],
    )

    add_note(doc,
        "Note: US Tech Support includes contracted remote hours ($125/hr) and "
        "separately billed onsite labor ($150\u2013200/hr). Offshore support combines "
        "NH ($15/hr, 15.0 hrs) and AH ($30/hr, 20.0 hrs). The overall blended rate "
        "of $77.97/hr reflects a roughly 55:45 split between US and offshore labor "
        "\u2014 a significantly higher US Tech ratio than optimal for cost efficiency.")

    add_subheading(doc, "Rate Card")
    make_table(doc,
        ["Role / Tier", "Rate", "Coverage"],
        [
            ["US Tech Support (Remote)", "$125.00/hr", "Escalations, advisory, complex issues"],
            ["US Tech Support (Onsite)", "$150\u2013200/hr", "On-premises work (hourly)"],
            ["Offshore Support (US Overnight)", "$15.00/hr", "US overnight (India NH)"],
            ["Offshore Support (US Daytime)", "$30.00/hr", "US business hours (India AH)"],
        ],
        col_aligns=[L, R, L],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 3: Licensing & Recurring Services
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 3: Licensing & Recurring Services")

    add_heading2(doc, "3.1 Monthly Managed Services Breakdown")
    add_body(doc,
        "The monthly contract ($8,261/year, ~$688/month) covers both labor and "
        "managed services:")

    make_table(doc,
        ["Service Category", "Annual", "Monthly"],
        [
            ["Labor (US Tech + Offshore)", "$3,907", "$326"],
            ["Security & Endpoint Protection", "$1,458", "$122"],
            ["Network Monitoring (Syslog/Health)", "$608", "$51"],
            ["Backup & Storage", "$600", "$50"],
            ["Site Assessment", "$450", "$38"],
            ["Secure Internet Filtering", "$432", "$36"],
            ["Patch Management", "$324", "$27"],
            ["Pen Testing", "$266", "$22"],
            ["Remote Monitoring (MyRemote)", "$216", "$18"],
            ["TOTAL", "$8,261", "$688"],
        ],
        col_aligns=[L, R, R],
        total_row_indices=[9],
    )

    add_subheading(doc, "December 2025 Monthly Invoice \u2014 Key Line Items")
    make_table(doc,
        ["Service", "Qty", "Unit Rate", "Monthly", "Category"],
        [
            ["US Tech Support (IRV-TS1)", "1.4", "$125.00", "$173", "Labor"],
            ["CrowdStrike \u2014 Desktop (AVD)", "9", "$8.50", "$77", "Security"],
            ["Site Assessment (SA)", "1", "$50.00", "$50", "Assessment"],
            ["Backup Storage (TB-BSTR)", "1", "$50.00", "$50", "Backup"],
            ["Huntress \u2014 Desktop (AVMH)", "9", "$5.00", "$45", "Security"],
            ["Offshore \u2014 US Daytime (AH)", "1.4", "$30.00", "$41", "Labor"],
            ["Offshore \u2014 US Overnight (NH)", "2.8", "$15.00", "$41", "Labor"],
            ["Syslog Monitoring (SSM)", "18", "$2.00", "$36", "Monitoring"],
            ["Health Monitoring (SHM)", "18", "$2.00", "$36", "Monitoring"],
            ["Secure Internet (SI)", "9", "$4.00", "$36", "Secure Internet"],
            ["Patch Management (PMW)", "9", "$3.00", "$27", "Monitoring"],
            ["Pen Testing (RTPT)", "6", "$3.50", "$21", "Pen Testing"],
            ["My Remote (MR)", "9", "$2.00", "$18", "Monitoring"],
        ],
        col_aligns=[L, R, R, R, L],
    )

    add_heading2(doc, "3.2 Security Stack Transition")
    add_body(doc,
        "In May 2025, TALY\u2019s secondary endpoint protection was transitioned from "
        "Malwarebytes (AVMD) to Huntress (AVMH) across all desktops. This upgrade "
        "provides stronger managed detection and response (MDR) capabilities alongside "
        "CrowdStrike\u2019s primary EDR protection.")

    make_table(doc,
        ["Period", "CrowdStrike", "Secondary EDR", "Desktops"],
        [
            ["Jan\u2013Apr 2025", "CrowdStrike (AVD)", "Malwarebytes (AVMD)", "9"],
            ["May\u2013Dec 2025", "CrowdStrike (AVD)", "Huntress (AVMH)", "9"],
            ["Mar 2026", "CrowdStrike (AVD)", "Huntress (AVMH)", "7"],
        ],
        col_aligns=[L, L, L, R],
    )

    add_page_break(doc)

    add_heading2(doc, "3.3 Service Count Tracking: December 2025 vs. March 2026")
    add_body(doc,
        "The following tables show the current 3/1/2026 service inventory alongside "
        "December 2025 counts. The March 2026 counts establish the current baseline "
        "for 2026 billing projections. Device counts decreased from 9 to 7 desktops "
        "as two systems were decommissioned.")

    add_subheading(doc, "Security & Endpoint Protection")
    make_table(doc,
        ["Service", "3/1/26 Count", "Dec Count", "Change", "Note"],
        [
            ["CrowdStrike \u2014 Desktop", "7", "9", "-2", "2 desktops decommissioned"],
            ["Huntress \u2014 Desktop", "7", "9", "-2", "Matches CrowdStrike"],
        ],
        col_aligns=[L, R, R, R, L],
    )

    add_subheading(doc, "Monitoring & Management")
    make_table(doc,
        ["Service", "3/1/26 Count", "Dec Count", "Change", "Note"],
        [
            ["Secure Internet", "7", "9", "-2", "Matches desktop count"],
            ["Patch Management", "7", "9", "-2", "Matches desktop count"],
            ["My Remote", "7", "9", "-2", "Matches desktop count"],
            ["Syslog Monitoring", "16", "18", "-2", "Network device adjustment"],
            ["Health Monitoring", "16", "18", "-2", "Network device adjustment"],
        ],
        col_aligns=[L, R, R, R, L],
    )

    add_subheading(doc, "Backup & Other")
    make_table(doc,
        ["Service", "3/1/26 Count", "Dec Count", "Change", "Note"],
        [
            ["Backup Storage", "1 TB", "1 TB", "0", "Flat"],
            ["Site Assessment", "1", "1", "0", "Flat"],
            ["Pen Testing IPs", "6", "6", "0", "Flat"],
        ],
        col_aligns=[L, R, R, R, L],
    )

    add_subheading(doc, "Pricing Changes (2026)")
    add_body(doc, "Several per-unit rates increased in the 2026 contract cycle:")
    make_table(doc,
        ["Service", "2025 Rate", "2026 Rate", "Change"],
        [
            ["Huntress (AVMH)", "$5.00/device", "$6.00/device", "+$1.00"],
            ["Secure Internet (SI)", "$4.00/device", "$6.00/device", "+$2.00"],
            ["Patch Management (PMW)", "$3.00/device", "$4.00/device", "+$1.00"],
            ["Pen Testing (RTPT)", "$3.50/IP", "$7.00/IP", "+$3.50"],
        ],
        col_aligns=[L, R, R, R],
        highlight_row_indices=[3],
    )

    add_subheading(doc, "Infrastructure Summary (3/1/2026)")
    make_table(doc,
        ["Category", "Count", "Detail"],
        [
            ["Desktops", "7", "Down from 9 (2 decommissioned)"],
            ["Network Devices", "16", "Monitored via Syslog/Health"],
            ["Backup Storage", "1 TB", "Cloud backup"],
            ["Users", "~5", "Active users"],
            ["Pen Test IPs", "6", "Real-time pen testing"],
        ],
        col_aligns=[L, R, L],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 4: Projects & One-Time Spend
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 4: Projects & One-Time Spend (2025)")
    add_body(doc,
        "Total non-contract spend in 2025 was $1,075, consisting entirely of "
        "hourly onsite labor.")

    add_heading2(doc, "4.1 Onsite Labor ($1,075)")
    make_table(doc,
        ["Visit", "Rate", "Hours", "Amount", "Date"],
        [
            ["Laptop & Mac Mini Onboarding", "$150/hr", "2.5", "$375", "Jan 2025"],
            ["Network Troubleshooting (Nov)", "$200/hr", "1.5", "$300", "Nov 2025"],
            ["Network Troubleshooting (Nov)", "$200/hr", "1.2", "$250", "Nov 2025"],
            ["Network / QNAP Work (Dec)", "$150/hr", "1.0", "$150", "Dec 2025"],
        ],
        col_aligns=[L, R, R, R, L],
    )

    add_note(doc,
        "Note: All onsite work was billed hourly outside the monthly contract. "
        "The January visit was for device onboarding at $150/hr. November visits "
        "were billed at $200/hr (after-hours rate). No hardware purchases or "
        "project-based work occurred in 2025.")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 5: Ticket Categorization
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 5: Ticket Categorization (2025)")
    add_body(doc,
        "All 54 unique tickets were categorized by analyzing ticket titles "
        "and descriptions:")

    make_table(doc,
        ["Category", "Tickets", "% Tix", "Hours", "% Hrs", "Avg Hrs", "Proactive?"],
        [
            ["Security Assessment & Compliance", "8", "14.8%", "13.0", "20.3%", "1.63", "Yes"],
            ["Backup & Storage (QNAP)", "16", "29.6%", "11.4", "17.8%", "0.71", "Yes"],
            ["General IT Support", "3", "5.6%", "11.3", "17.7%", "3.77", "Reactive"],
            ["Device Onboarding & Setup", "3", "5.6%", "8.8", "13.8%", "2.93", "Reactive"],
            ["File Access & Remote Issues", "6", "11.1%", "6.2", "9.7%", "1.03", "Reactive"],
            ["Network & Connectivity", "3", "5.6%", "5.5", "8.6%", "1.83", "Reactive"],
            ["Advisory / Touch Meetings", "4", "7.4%", "3.0", "4.7%", "0.75", "Proactive"],
            ["Email & Cloud Apps", "3", "5.6%", "2.3", "3.6%", "0.77", "Reactive"],
            ["Security Agent Updates", "3", "5.6%", "1.0", "1.6%", "0.33", "Yes"],
            ["RMM & Agent Maintenance", "3", "5.6%", "0.8", "1.3%", "0.27", "Yes"],
            ["Patch Management", "2", "3.7%", "0.6", "0.9%", "0.30", "Yes"],
            ["TOTAL", "54", "100%", "63.9", "100%", "1.18", ""],
        ],
        col_aligns=[L, R, R, R, R, R, L],
        total_row_indices=[11],
    )

    add_heading2(doc, "5.1 Key Observations")

    add_subheading(doc, "Backup & Storage (QNAP) \u2014 16 tickets (30%), 11.4 hours (18%)")
    add_body(doc,
        "The largest category by ticket count. Most are proactive QNAP firmware "
        "upgrades and HBS3 sync verification tasks between the TALY NAS and "
        "Technijian DC. These are routine maintenance tasks \u2014 8 QNAP firmware "
        "upgrades and 8 sync verification checks \u2014 each averaging under 1 hour.")

    add_subheading(doc, "Security Assessment & Compliance \u2014 8 tickets (15%), 13.0 hours (20%)")
    add_body(doc,
        "Network Detective scans, risk report reviews, site assessments, and attack "
        "surface analysis. The most time-intensive category per ticket (1.63 hrs avg), "
        "reflecting the thoroughness of security assessments.")

    add_subheading(doc, "General IT Support \u2014 3 tickets (6%), 11.3 hours (18%)")
    add_body(doc,
        "Only 3 tickets but 11.3 hours, driven by the large \"Re: Talley and "
        "Associates\" ticket (12.1 hrs) in December \u2014 a comprehensive engagement "
        "covering tool installations, Time Machine backup setup, QNAP configuration, "
        "and EasyDMARC setup across multiple users.")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 6: Service Metrics & Trends
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 6: Service Metrics & Trends")

    add_heading2(doc, "6.1 Monthly Ticket Volume & Hours")
    make_table(doc,
        ["Month", "Tickets", "US Tech", "Offshore", "Total Hrs", "Hrs/Ticket"],
        [
            ["Jan", "17", "14.1", "3.4", "17.5", "1.03"],
            ["Feb", "8", "0.5", "0.0", "0.5", "0.06"],
            ["Mar", "3", "0.0", "3.4", "3.4", "1.13"],
            ["Apr", "4", "0.0", "3.0", "3.0", "0.75"],
            ["May", "8", "0.5", "4.6", "5.1", "0.64"],
            ["Jun", "4", "1.5", "2.2", "3.7", "0.93"],
            ["Jul", "5", "1.0", "3.3", "4.3", "0.86"],
            ["Aug", "4", "0.0", "1.8", "1.8", "0.45"],
            ["Sep", "2", "0.0", "0.7", "0.7", "0.35"],
            ["Oct", "0", "0.0", "0.0", "0.0", "\u2014"],
            ["Nov", "6", "5.5", "1.5", "7.0", "1.17"],
            ["Dec", "7", "5.8", "11.1", "16.9", "2.41"],
            ["Average", "5.7", "2.4", "2.9", "5.3", "0.93"],
        ],
        col_aligns=[L, R, R, R, R, R],
        highlight_row_indices=[0, 11],
        total_row_indices=[12],
    )

    add_note(doc,
        "January (17.5 hrs) and December (16.9 hrs) accounted for 54% of the "
        "year\u2019s total hours. January was device onboarding; December was a "
        "comprehensive tool deployment and backup configuration project.")

    add_heading2(doc, "6.2 Notable Trends")
    add_bullet(doc,
        " 54 tickets and 63.9 hours across the full year \u2014 an average of "
        "~5.3 hours/month. TALY is a small, stable environment that does not "
        "generate heavy support demand.",
        bold_prefix="Low Overall Volume:")
    add_bullet(doc,
        " Approximately 50% of tickets are proactive (QNAP maintenance, "
        "security assessments, patching, agent updates). This is appropriate for "
        "a small environment.",
        bold_prefix="Proactive Work:")
    add_bullet(doc,
        " 45% of all hours are US Tech Support at $125/hr. For a 7-desktop "
        "environment, this is disproportionately expensive. Most of this work "
        "(QNAP checks, file access troubleshooting, advisory meetings) can be "
        "handled by offshore support.",
        bold_prefix="US Tech Support Concentration:")
    add_bullet(doc,
        " January (17.5 hrs) and December (16.9 hrs) accounted for "
        "54% of the year\u2019s total hours, driven by device onboarding and a "
        "major tool deployment project respectively.",
        bold_prefix="Seasonal Spikes:")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 7: Summary & Cost Optimization Strategy
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 7: Summary & Cost Optimization Strategy")
    add_body(doc,
        "TALY is a small, stable managed IT environment \u2014 7 desktops, ~5 users, "
        "16 network devices \u2014 generating modest support volume of 54 tickets and "
        "63.9 hours in 2025.")

    add_body(doc,
        "At $9,336 total spend (~$778/month), the per-user IT cost is approximately "
        "$1,867/year. This covers fully managed security (CrowdStrike + Huntress), "
        "backup, monitoring, patch management, pen testing, and support labor.")

    add_body(doc, "The primary optimization opportunities are:")

    add_bullet(doc,
        " US Tech Support accounts for 45% of hours at $125/hr. The majority "
        "of TALY\u2019s work \u2014 QNAP maintenance, security scans, patching, agent "
        "updates \u2014 is proactive and can be performed during US overnight hours by "
        "the offshore team at $15/hr. This represents a potential 8x cost reduction "
        "on shifted hours.",
        bold_prefix="Shift Labor to Overnight:")
    add_bullet(doc,
        " After the current contract cycle, US Tech Support transitions to "
        "hourly billing only \u2014 reserved for the rare escalation or onsite visit. "
        "This removes the contracted US Tech allotment (~$2,915/yr) from the monthly "
        "invoice, replacing it with pay-per-use at $125/hr only when needed.",
        bold_prefix="Minimize US Tech Support:")
    add_bullet(doc,
        " The offshore AH hours ($30/hr, US daytime) should also be shifted "
        "to NH ($15/hr, US overnight) wherever possible. Most offshore work for TALY "
        "is already proactive maintenance that does not require real-time user "
        "interaction.",
        bold_prefix="Move Offshore AH to NH:")
    add_bullet(doc,
        " The reduction from 9 to 7 desktops decreases per-device service "
        "costs. Ensure billing reflects the current 7-desktop count.",
        bold_prefix="Device Count Reduction:")
    add_bullet(doc,
        " At ~5 hours/month average, TALY does not need significant "
        "contracted labor. A small offshore overnight allotment plus hourly US Tech "
        "billing is the most cost-effective model.",
        bold_prefix="Leverage Low Volume:")

    add_heading2(doc, "7.1 Billed vs. Actual \u2014 Support Cycle Analysis")
    make_table(doc,
        ["Role", "Allotted/Mo", "Actual/Mo (Avg)", "Utilization", "Status"],
        [
            ["US Tech ($125/hr)", "1.7 hrs", "2.4 hrs", "141%", "Over-utilized"],
            ["Offshore NH ($15/hr)", "2.8 hrs", "1.3 hrs", "45%", "Under-utilized"],
            ["Offshore AH ($30/hr)", "1.4 hrs", "1.7 hrs", "119%", "Slightly over"],
        ],
        col_aligns=[L, R, R, R, L],
        highlight_row_indices=[0],
        green_row_indices=[1],
    )

    add_body(doc,
        "US Tech Support is over-utilized at 141% while Offshore NH is only at "
        "45% utilization. This confirms the opportunity to shift work from US Tech "
        "($125/hr) to Offshore NH ($15/hr). The Offshore NH allotment is being "
        "underused \u2014 exactly the capacity that should absorb more of TALY\u2019s "
        "routine work.",
        bold_prefix="Key Finding: ")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    #  SECTION 8: 2026 Budget Projections
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "Section 8: 2026 Budget Projections")

    add_heading2(doc, "8.1 Cost Optimization Strategy")
    add_body(doc,
        "The 2026 strategy for TALY focuses on three levers to reduce costs:")

    add_subheading(doc, "1. Move Work to Overnight ($15/hr)")
    add_body(doc,
        "The bulk of TALY\u2019s support work is proactive \u2014 QNAP maintenance, "
        "security assessments, firmware upgrades, patching, and agent updates. None "
        "of this requires real-time user interaction. By routing these tasks to the "
        "offshore overnight team (India daytime = US overnight at $15/hr), the cost "
        "per hour drops from $30\u2013$125 down to $15.")

    add_subheading(doc, "2. Minimize US Tech Support")
    add_body(doc,
        "After the current contract cycle, US Tech Support transitions from contracted "
        "to hourly billing only. This eliminates the ~$173/month contracted US Tech "
        "allotment. US Tech is reserved for:")
    add_bullet(doc, "Onsite visits (hardware, network)")
    add_bullet(doc, "Complex escalations requiring US-based expertise")
    add_bullet(doc, "Billed at $125/hr remote or $150\u2013200/hr onsite, only when used")

    add_subheading(doc, "3. Keep Per-Device Costs Low")
    add_body(doc,
        "The reduction from 9 to 7 desktops saves approximately $50\u201360/month on "
        "per-device services. Pricing increases on some services (Huntress, SI, PMW, "
        "Pen Testing) partially offset this reduction.")

    add_heading2(doc, "8.2 Service Cost Projections")

    add_subheading(doc, "Per-Device Services (7 desktops)")
    make_table(doc,
        ["Service", "2025 Rate", "2026 Rate", "Qty", "2026 Monthly", "2026 Annual"],
        [
            ["CrowdStrike (AVD)", "$8.50", "$8.50", "7", "$59.50", "$714"],
            ["Huntress (AVMH)", "$5.00", "$6.00", "7", "$42.00", "$504"],
            ["Secure Internet (SI)", "$4.00", "$6.00", "7", "$42.00", "$504"],
            ["Patch Management (PMW)", "$3.00", "$4.00", "7", "$28.00", "$336"],
            ["My Remote (MR)", "$2.00", "$2.00", "7", "$14.00", "$168"],
            ["Desktop subtotal", "", "", "", "$185.50", "$2,226"],
        ],
        col_aligns=[L, R, R, R, R, R],
        total_row_indices=[5],
    )

    add_subheading(doc, "Fixed / Network Services")
    make_table(doc,
        ["Service", "2026 Monthly", "2026 Annual"],
        [
            ["Backup Storage (1 TB)", "$50.00", "$600"],
            ["Site Assessment", "$50.00", "$600"],
            ["Pen Testing (6 IPs)", "$42.00", "$504"],
            ["Syslog Monitoring (16 devices)", "$32.00", "$384"],
            ["Health Monitoring (16 devices)", "$32.00", "$384"],
            ["Network/fixed subtotal", "$206.00", "$2,472"],
        ],
        col_aligns=[L, R, R],
        total_row_indices=[5],
    )

    add_body(doc,
        "$391.50/month ($4,698/year)",
        bold_prefix="Total Managed Services: ")

    add_heading2(doc, "8.3 Labor Projections \u2014 Optimized Model")
    make_table(doc,
        ["Component", "2025 Actual", "2026 Optimized", "Change"],
        [
            ["US Tech Support (contracted)", "$2,915", "$0", "-$2,915"],
            ["US Tech Support (hourly, est.)", "$1,075", "$750", "-$325"],
            ["Offshore NH ($15/hr)", "$495", "$1,080", "+$585"],
            ["Offshore AH ($30/hr)", "$497", "$360", "-$137"],
            ["Total Labor", "$4,982", "$2,190", "-$2,792"],
        ],
        col_aligns=[L, R, R, R],
        total_row_indices=[4],
        green_row_indices=[0],
    )

    add_bullet(doc,
        " US Tech contracts out after current cycle; estimated 6 hours/year "
        "onsite at $125/hr = $750.",
        bold_prefix="US Tech:")
    add_bullet(doc,
        " Increases to 6 hrs/month ($90/mo) to absorb shifted work.",
        bold_prefix="Offshore NH:")
    add_bullet(doc,
        " Decreases to 1 hr/month ($30/mo) for minimal daytime coordination.",
        bold_prefix="Offshore AH:")
    add_bullet(doc,
        " Most proactive work (QNAP, patching, security scans, agent updates) "
        "moves to overnight.",
        bold_prefix="Work Shift:")

    add_heading2(doc, "8.4 Total 2026 Projection Summary")
    make_table(doc,
        ["Category", "2025 Actual", "2026 Projected", "Change"],
        [
            ["Managed Services", "$4,354", "$4,698", "+$344"],
            ["Contracted Labor", "$3,907", "$1,440", "-$2,467"],
            ["Hourly Labor", "$1,075", "$750", "-$325"],
            ["TOTAL", "$9,336", "$6,888", "-$2,448 (-26.2%)"],
        ],
        col_aligns=[L, R, R, R],
        total_row_indices=[3],
        green_row_indices=[1],
    )

    add_body(doc,
        "The optimized model projects a 26% cost reduction ($9,336 \u2192 $6,888) "
        "driven almost entirely by the labor restructuring. Managed services increase "
        "slightly (+$344) due to per-unit price increases on Huntress, Secure Internet, "
        "Patch Management, and Pen Testing, partially offset by the device count "
        "reduction (9 \u2192 7 desktops). The labor savings of $2,792/year come from "
        "eliminating the contracted US Tech allotment and shifting work to $15/hr "
        "overnight support.")

    add_subheading(doc, "Key Assumptions")
    add_bullet(doc,
        " US Tech moves to hourly after current cycle. Offshore NH absorbs "
        "routine work at 6 hrs/mo. Offshore AH reduced to 1 hr/mo.",
        bold_prefix="Labor:")
    add_bullet(doc,
        " Based on March 2026 counts (7 desktops, 16 network devices) at "
        "updated 2026 rates.",
        bold_prefix="Device Services:")
    add_bullet(doc,
        " No changes to backup storage (1 TB), pen testing IPs (6), or "
        "site assessment.",
        bold_prefix="Infrastructure:")
    add_bullet(doc,
        " Estimated 6 hours/year for rare hardware or network visits at $125/hr.",
        bold_prefix="Onsite:")
    add_bullet(doc,
        " Assumes similar volume to 2025 (~54 tickets, ~64 hours). If volume "
        "increases, the overnight model keeps incremental cost at $15/hr rather "
        "than $125/hr.",
        bold_prefix="Volume:")

    # ── Save ───────────────────────────────────────────────────
    doc.save(OUT_FILE)
    print(f"Document saved to: {OUT_FILE}")


if __name__ == "__main__":
    build()
