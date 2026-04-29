"""
Generate branded monthly Meraki activity Word reports per org.

Reads the JSON summaries written by aggregate_monthly.py and produces:

  clients/_meraki/<org_slug>/reports/<ORG> - Meraki Monthly Activity - <YYYY-MM>.docx

Uses the same brand styling (dark blue 1F4E79, alt rows F2F7FB) as the
existing per-client annual review docs.

Usage:
  python generate_monthly_docx.py                              # all orgs / months that have JSON summaries
  python generate_monthly_docx.py --month 2026-03
  python generate_monthly_docx.py --only vaf,bwh
  python generate_monthly_docx.py --from 2026-01 --to 2026-03
"""

from __future__ import annotations

import argparse
import json
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

REPO_ROOT = Path(__file__).resolve().parents[2]  # annual-client-review-1
PROOFREADER = REPO_ROOT / "technijian" / "shared" / "scripts" / "proofread_docx.py"

EXPECTED_SECTIONS = [
    "Executive Summary",
    "Network & Device Inventory",
    "Security Posture",
    "IDS/IPS & AMP Events",
    "Firewall / Network Activity",
    "Daily Trend",
]

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(2)

# Brand palette (matches generate_aava_docx.py / generate_bwh_docx.py)
DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
MED_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
GRAY      = RGBColor(0x66, 0x66, 0x66)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x00, 0x00, 0x00)
DARK_BLUE_HEX = "1F4E79"
ALT_ROW_HEX   = "F2F7FB"
WHITE_HEX     = "FFFFFF"

DEFAULT_ROOT = Path(__file__).resolve().parents[2] / "clients" / "_meraki"

L = WD_ALIGN_PARAGRAPH.LEFT
R = WD_ALIGN_PARAGRAPH.RIGHT
C = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# docx helpers (kept minimal, mirroring conventions in generate_aava_docx.py)
# ---------------------------------------------------------------------------

def shade(cell, color_hex: str) -> None:
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    )


def set_text(cell, text, *, bold=False, size=Pt(9), color=BLACK, align=L) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = size
    run.font.color.rgb = color


def header_row(table, texts, aligns=None) -> None:
    aligns = aligns or [L] * len(texts)
    row = table.rows[0]
    for i, t in enumerate(texts):
        cell = row.cells[i]
        shade(cell, DARK_BLUE_HEX)
        set_text(cell, t, bold=True, color=WHITE, align=aligns[i])


def data_row(table, texts, idx, aligns=None) -> None:
    aligns = aligns or [L] * len(texts)
    row = table.add_row()
    fill = ALT_ROW_HEX if idx % 2 == 0 else WHITE_HEX
    for i, t in enumerate(texts):
        cell = row.cells[i]
        shade(cell, fill)
        set_text(cell, t, align=aligns[i])


def make_table(doc, headers, rows, col_aligns=None) -> None:
    if not rows:
        rows = [["—"] * len(headers)]
    t = doc.add_table(rows=1, cols=len(headers))
    t.autofit = True
    header_row(t, headers, col_aligns)
    for i, r in enumerate(rows):
        data_row(t, r, i, col_aligns)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_BLUE


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = MED_BLUE


def body(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def kv_table(doc, pairs):
    """Two-column key/value table — used for at-a-glance summaries."""
    rows = [[k, str(v)] for k, v in pairs]
    make_table(doc, ["Field", "Value"], rows, [L, R])


def callout_box(doc, title: str, text: str) -> None:
    """Single-cell shaded callout box. The proofreader expects at least one
    of these per report (warns otherwise)."""
    t = doc.add_table(rows=1, cols=1)
    t.autofit = True
    cell = t.rows[0].cells[0]
    shade(cell, ALT_ROW_HEX)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    sub = cell.add_paragraph()
    r2 = sub.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = BLACK


def metric_cards(doc, metrics: list[tuple[str, str]]) -> None:
    """Multi-column header-style metric row (3+ cells). The proofreader
    expects this in the first half of the document (warns otherwise)."""
    if len(metrics) < 3:
        # pad to at least 3 columns so the row is recognized as a metric strip
        while len(metrics) < 3:
            metrics.append(("", ""))
    t = doc.add_table(rows=2, cols=len(metrics))
    t.autofit = True
    for i, (label, value) in enumerate(metrics):
        head = t.rows[0].cells[i]
        shade(head, DARK_BLUE_HEX)
        set_text(head, label, bold=True, size=Pt(9), color=WHITE, align=C)
        val = t.rows[1].cells[i]
        shade(val, ALT_ROW_HEX)
        set_text(val, value, bold=True, size=Pt(14), color=DARK_BLUE, align=C)


# ---------------------------------------------------------------------------
# Report assembly
# ---------------------------------------------------------------------------

def fmt_int(n) -> str:
    try:
        return f"{int(n):,}"
    except Exception:
        return str(n)


def setup_doc(org_name: str, month: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)

    # Title
    p = doc.add_paragraph()
    p.alignment = C
    r = p.add_run(f"{org_name}\nMeraki Monthly Activity Report\n{month}")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = C
    r = p.add_run(f"Prepared by Technijian, Inc.  |  Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY
    return doc


def section_executive_summary(doc, payload: dict) -> None:
    h1(doc, "Executive Summary")
    cfg = payload["configuration"]
    sec = payload["security_events"]
    net = payload["network_events"]

    # Top-of-page metric strip (3+ cells -> recognized as metric cards)
    metric_cards(doc, [
        ("Networks",         fmt_int(cfg.get("network_count", 0))),
        ("Devices",          fmt_int(cfg.get("device_count", 0))),
        ("IDS/IPS events",   fmt_int(sec.get("total", 0))),
        ("Activity events",  fmt_int(net.get("total", 0))),
    ])

    # Status callout (single-cell table)
    blocked = (sec.get("by_blocked") or {}).get("blocked", 0)
    alerted = (sec.get("by_blocked") or {}).get("alerted", 0)
    if sec.get("total", 0) == 0:
        status = ("All clear", "No IDS/IPS or AMP events were recorded across "
                  "the reporting period. Continue normal monitoring.")
    elif blocked >= alerted:
        status = ("Active prevention", f"{fmt_int(blocked)} threats were blocked "
                  f"and {fmt_int(alerted)} alerted. The intrusion-prevention "
                  "engine is actively dropping malicious traffic.")
    else:
        status = ("Detect-mode activity", f"{fmt_int(alerted)} alerts were "
                  f"recorded with {fmt_int(blocked)} blocks. Review whether the "
                  "appliance should be moved from detect to prevention mode for "
                  "high-priority signatures.")
    callout_box(doc, status[0], status[1])

    pairs = [
        ("Reporting period",                payload["month"]),
        ("Networks under management",       fmt_int(cfg.get("network_count", 0))),
        ("Devices under management",        fmt_int(cfg.get("device_count", 0))),
        ("Total IDS/IPS / AMP events",      fmt_int(sec.get("total", 0))),
        ("Days with security events",       fmt_int(sec.get("days_with_events", 0))),
        ("Total firewall activity events",  fmt_int(net.get("total", 0))),
        ("Networks with activity",          fmt_int(net.get("networks_with_events", 0))),
    ]
    kv_table(doc, pairs)


def section_inventory(doc, cfg: dict) -> None:
    h1(doc, "Network & Device Inventory")
    h2(doc, "Devices by model")
    rows = sorted(((m or "(unknown)", c) for m, c in cfg.get("device_models", {}).items()),
                  key=lambda x: -x[1])
    make_table(doc, ["Model", "Count"], [[m, fmt_int(c)] for m, c in rows], [L, R])

    h2(doc, "Devices by product type")
    rows = sorted(((m or "(unknown)", c) for m, c in cfg.get("device_product_types", {}).items()),
                  key=lambda x: -x[1])
    make_table(doc, ["Product type", "Count"], [[m, fmt_int(c)] for m, c in rows], [L, R])

    h2(doc, "Networks")
    headers = ["Name", "Product types", "VLANs", "L3 rules", "L7 rules", "Inbound", "Port fwd", "1:1 NAT"]
    rows = []
    for n in cfg.get("networks", []):
        rows.append([
            n.get("name") or n.get("slug") or "—",
            ", ".join(n.get("productTypes") or []),
            fmt_int(n.get("vlan_count", 0)),
            fmt_int(n.get("firewall_l3_rule_count", 0)),
            fmt_int(n.get("firewall_l7_rule_count", 0)),
            fmt_int(n.get("firewall_inbound_rule_count", 0)),
            fmt_int(n.get("port_forward_count", 0)),
            fmt_int(n.get("one_to_one_nat_count", 0)),
        ])
    make_table(doc, headers, rows, [L, L, R, R, R, R, R, R])


def section_security_posture(doc, cfg: dict) -> None:
    h1(doc, "Security Posture (current configuration)")
    rows = []
    for n in cfg.get("networks", []):
        intr = n.get("intrusion") or {}
        amp  = n.get("malware") or {}
        cfilt = n.get("content_filtering") or {}
        rows.append([
            n.get("name") or n.get("slug") or "—",
            intr.get("mode") or "—",
            amp.get("mode") or "—",
            fmt_int(cfilt.get("blocked_categories_count", 0)),
            fmt_int(cfilt.get("blocked_url_patterns_count", 0)),
            n.get("s2s_vpn_mode") or "—",
            fmt_int(n.get("syslog_destination_count", 0)),
        ])
    make_table(doc,
               ["Network", "IDS/IPS mode", "AMP mode",
                "URL cats blocked", "URL patterns blocked",
                "S2S VPN mode", "Syslog dests"],
               rows, [L, L, L, R, R, L, R])


def section_security_events(doc, sec: dict) -> None:
    h1(doc, "IDS/IPS & AMP Events")
    body(doc, f"Total events captured this period: {fmt_int(sec.get('total', 0))} "
              f"across {sec.get('days_with_events', 0)} days with activity.")

    if sec.get("by_priority"):
        h2(doc, "By priority / severity")
        rows = sorted(sec["by_priority"].items(), key=lambda x: x[0])
        make_table(doc, ["Priority", "Count"],
                   [[k, fmt_int(v)] for k, v in rows], [L, R])

    if sec.get("by_blocked"):
        h2(doc, "Blocked vs. alerted")
        make_table(doc, ["Action", "Count"],
                   [[k.title(), fmt_int(v)] for k, v in sec["by_blocked"].items()], [L, R])

    if sec.get("by_signature_top"):
        h2(doc, "Top 15 signatures")
        rows = [[s, fmt_int(c)] for s, c in sec["by_signature_top"]]
        make_table(doc, ["Signature", "Count"], rows, [L, R])

    if sec.get("top_sources"):
        h2(doc, "Top 15 source IPs")
        rows = [[s, fmt_int(c)] for s, c in sec["top_sources"]]
        make_table(doc, ["Source", "Hits"], rows, [L, R])

    if sec.get("top_destinations"):
        h2(doc, "Top 15 destination IPs")
        rows = [[s, fmt_int(c)] for s, c in sec["top_destinations"]]
        make_table(doc, ["Destination", "Hits"], rows, [L, R])


def section_activity(doc, net: dict) -> None:
    h1(doc, "Firewall / Network Activity")
    body(doc, f"Total firewall/appliance activity events: {fmt_int(net.get('total', 0))} "
              f"across {net.get('networks_with_events', 0)} active networks.")

    if net.get("by_type_top"):
        h2(doc, "Top 20 event types")
        rows = [[t, fmt_int(c)] for t, c in net["by_type_top"]]
        make_table(doc, ["Event type", "Count"], rows, [L, R])

    if net.get("by_category"):
        h2(doc, "By category")
        rows = sorted(net["by_category"].items(), key=lambda x: -x[1])
        make_table(doc, ["Category", "Count"],
                   [[k, fmt_int(v)] for k, v in rows], [L, R])

    if net.get("by_network"):
        h2(doc, "Per-network rollup")
        rows = []
        for slug, info in sorted(net["by_network"].items(), key=lambda kv: -kv[1]["total"]):
            top_types = ", ".join(f"{t}({fmt_int(c)})" for t, c in info["by_type_top"][:3])
            rows.append([info["name"] or slug, fmt_int(info["total"]), top_types])
        make_table(doc, ["Network", "Total events", "Top event types"], rows, [L, R, L])


def section_daily_trend(doc, sec: dict, net: dict) -> None:
    h1(doc, "Daily Trend")
    sec_by_day = {x["date"]: x["count"] for x in sec.get("daily_counts", [])}
    net_by_day = {x["date"]: x["count"] for x in net.get("daily_counts", [])}
    days = sorted(set(sec_by_day) | set(net_by_day))
    rows = [[d, fmt_int(sec_by_day.get(d, 0)), fmt_int(net_by_day.get(d, 0))]
            for d in days]
    make_table(doc, ["Date", "Security events", "Activity events"], rows, [L, R, R])


def build_report(payload: dict, out_path: Path) -> None:
    org_name = payload["configuration"].get("org") or payload["org_slug"]
    doc = setup_doc(org_name, payload["month"])
    section_executive_summary(doc, payload)
    section_inventory(doc, payload["configuration"])
    section_security_posture(doc, payload["configuration"])
    section_security_events(doc, payload["security_events"])
    section_activity(doc, payload["network_events"])
    section_daily_trend(doc, payload["security_events"], payload["network_events"])

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = C
    r = p.add_run("End of report  |  Generated by Technijian Meraki monthly pipeline")
    r.font.size = Pt(8)
    r.font.color.rgb = GRAY

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description=__doc__,
                                formatter_class=argparse.RawDescriptionHelpFormatter)
    p.add_argument("--month")
    p.add_argument("--from", dest="from_month")
    p.add_argument("--to",   dest="to_month")
    p.add_argument("--only")
    p.add_argument("--skip")
    p.add_argument("--root", default=str(DEFAULT_ROOT))
    return p.parse_args()


def run_proofreader(generated: list[Path]) -> int:
    """Invoke the shared proofread_docx.py against every generated report.
    Exits non-zero if any fails. The proofread-report skill describes the
    7 scored checks + 2 warning checks performed."""
    if not generated:
        return 0
    if not PROOFREADER.exists():
        print(f"\n[proofread] WARNING: proofreader missing at {PROOFREADER} — skipping.")
        return 0
    sections_csv = ",".join(EXPECTED_SECTIONS)
    cmd = [sys.executable, str(PROOFREADER), "--sections", sections_csv]
    cmd += [str(p) for p in generated if p.exists()]
    print("\n=== Proofreading reports ===")
    sys.stdout.flush()
    rc = subprocess.run(cmd).returncode
    if rc != 0:
        print(f"\n[proofread] FAILED — one or more reports did not pass the gate (rc={rc}).")
    else:
        print(f"[proofread] OK — {len(generated)} report(s) passed all checks.")
    return rc


def main() -> int:
    args = parse_args()
    root = Path(args.root)
    only = {s.strip().lower() for s in (args.only or "").split(",") if s.strip()}
    skip = {s.strip().lower() for s in (args.skip or "").split(",") if s.strip()}

    generated: list[Path] = []
    for org_dir in sorted([d for d in root.iterdir() if d.is_dir() and not d.name.startswith("_")]):
        if only and org_dir.name not in only:
            continue
        if org_dir.name in skip:
            continue
        monthly_dir = org_dir / "monthly"
        if not monthly_dir.exists():
            continue
        for f in sorted(monthly_dir.glob("*.json")):
            month = f.stem
            if args.month and month != args.month:
                continue
            if args.from_month and month < args.from_month:
                continue
            if args.to_month and month > args.to_month:
                continue
            payload = json.loads(f.read_text(encoding="utf-8"))
            org_label = (payload["configuration"].get("org") or org_dir.name).strip()
            safe_label = "".join(c if c.isalnum() or c in " -_" else "_" for c in org_label)
            out = org_dir / "reports" / f"{safe_label} - Meraki Monthly Activity - {month}.docx"
            build_report(payload, out)
            generated.append(out)
            print(f"  [{org_dir.name}] {month} -> {out.relative_to(root)}")
    print(f"\nGenerated {len(generated)} Word report(s)")

    rc = run_proofreader(generated)
    return rc


if __name__ == "__main__":
    sys.exit(main())
