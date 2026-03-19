"""
Generate Word documents for VAF and HHOC — 2025 Annual + 2026 Q1 Reviews
Technijian Inc.
"""
import os
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Color palette ──────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
MED_BLUE = RGBColor(0x2E, 0x75, 0xB6)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

DARK_BLUE_HEX = "1F4E79"
ALT_ROW_HEX = "F2F7FB"
WHITE_HEX = "FFFFFF"
YELLOW_HEX = "FFF3CD"

BASE = r"c:\vscode\annual-client-review\annual-client-review"

# ── Helpers ────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, size=Pt(8.5), color=BLACK,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, font_name="Calibri"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = font_name

def add_header_row(table, texts, col_aligns=None):
    row = table.rows[0]
    for i, txt in enumerate(texts):
        align = (col_aligns[i] if col_aligns else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[i], txt, bold=True, color=WHITE, alignment=align)
        set_cell_shading(row.cells[i], DARK_BLUE_HEX)

def add_data_row(table, texts, row_idx, col_aligns=None, bold=False,
                 is_total=False, is_highlight=False):
    row = table.rows[row_idx]
    bg = DARK_BLUE_HEX if is_total else (YELLOW_HEX if is_highlight else
         (ALT_ROW_HEX if row_idx % 2 == 0 else WHITE_HEX))
    txt_color = WHITE if is_total else BLACK
    for i, txt in enumerate(texts):
        align = (col_aligns[i] if col_aligns else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[i], txt, bold=(bold or is_total),
                      color=txt_color, alignment=align)
        set_cell_shading(row.cells[i], bg)

def make_table(doc, headers, data, col_aligns=None, total_row_indices=None,
               highlight_row_indices=None):
    total_row_indices = total_row_indices or []
    highlight_row_indices = highlight_row_indices or []
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(data), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:top w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:left w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:bottom w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:right w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '</w:tcBorders>')
            tcPr.append(borders)
    add_header_row(tbl, headers, col_aligns)
    for idx, row_data in enumerate(data):
        r = idx + 1
        is_total = idx in total_row_indices
        is_hl = idx in highlight_row_indices
        add_data_row(tbl, row_data, r, col_aligns, bold=is_total,
                     is_total=is_total, is_highlight=is_hl)
    return tbl

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(16)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(13)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def add_subheading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(11)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Calibri"

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Calibri"

def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GRAY
    r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(4)

def add_page_break(doc):
    doc.add_page_break()

def setup_doc(header_text, footer_text="Technijian Inc. — Confidential"):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.63)
    section.right_margin = Inches(0.63)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(header_text)
    hr.font.size = Pt(8)
    hr.font.color.rgb = GRAY
    hr.font.name = "Calibri"
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(footer_text)
    fr.font.size = Pt(8)
    fr.font.color.rgb = GRAY
    fr.font.name = "Calibri"
    return doc

def add_title_page(doc, client_name, title, subtitle, date_text):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(client_name)
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.name = "Calibri"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(20)
    r.font.color.rgb = MED_BLUE
    r.font.name = "Calibri"
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.font.size = Pt(20)
        r.font.color.rgb = MED_BLUE
        r.font.name = "Calibri"
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(date_text)
    r.font.size = Pt(12)
    r.font.color.rgb = GRAY
    r.font.name = "Calibri"

L = WD_ALIGN_PARAGRAPH.LEFT
R = WD_ALIGN_PARAGRAPH.RIGHT

def fmt(v):
    """Format number as currency string."""
    if isinstance(v, (int, float)):
        return f"${v:,.2f}"
    return str(v)

def fmtk(v):
    """Format number with commas."""
    if isinstance(v, (int, float)):
        return f"{v:,.0f}" if v == int(v) else f"{v:,.1f}"
    return str(v)

# ══════════════════════════════════════════════════════════════
#  VAF 2025 Annual Review
# ══════════════════════════════════════════════════════════════
def build_vaf_2025(data):
    te = data['vaf_2025_te']
    inv = data['vaf_2025_inv']
    doc = setup_doc("VAF (Via Auto Finance) — 2025 IT Support Review & 2026 Planning")

    add_title_page(doc, "VAF (Via Auto Finance)",
                   "2025 IT Support Review",
                   "& 2026 Planning Readiness Assessment",
                   "Prepared by Technijian Inc.  —  March 2026")

    doc.add_paragraph()
    make_table(doc,
        ["2025 At a Glance", "Key Figures"],
        [
            ["Total 2025 Revenue", "$117,229"],
            ["Monthly Contract (Support + Services)", "$75,802"],
            ["Recurring Licensing (M365 + SFTP)", "$9,793"],
            ["Projects / One-Time", "$31,485"],
            ["Labor Overage (WeeklyOut)", "$150"],
            ["Unique Tickets (2025)", "462"],
            ["Total Support Hours Delivered", "334"],
            ["Cloud Servers Managed", "8"],
            ["Desktops Managed", "48"],
            ["Users (MyDisk)", "50"],
        ],
        col_aligns=[L, R],
        total_row_indices=[0],
    )

    add_page_break(doc)

    # Section 1
    add_heading1(doc, "Section 1: Support Scope & Coverage")
    add_heading2(doc, "1.1 What Is Included in Monthly Support")
    add_body(doc,
        "The monthly agreement covers a comprehensive managed IT package for VAF, "
        "totaling $75,802 across 2025 (~$6,317/month). This breaks down into two components:")
    add_body(doc,
        " Offshore technical support provides coverage across US business hours "
        "($30/hr, billed as offshore after-hours) and US overnight hours ($15/hr, billed "
        "as offshore normal hours). US-based technical support ($125/hr) handles escalations "
        "and onsite work.",
        bold_prefix="Labor (~$10,843/year, ~$904/month):")
    add_body(doc,
        " Security stack (CrowdStrike, Huntress, Malwarebytes), monitoring "
        "(RMM, ManageEngine Ops Manager, patch management), backup (image backup, Veeam ONE), "
        "cloud hosting (VMs, storage), VoIP/telephony (SIP, DIDs), network equipment "
        "(Sophos, Velocloud, Edge), secure internet filtering, pen testing, and site assessment.",
        bold_prefix="Managed Services (~$64,959/year, ~$5,413/month):")

    add_heading2(doc, "1.2 Coverage Hours")
    nh_pct = te['hours']['nh'] / te['hours']['total'] * 100
    ah_pct = te['hours']['ah'] / te['hours']['total'] * 100
    on_pct = te['hours']['onsite'] / te['hours']['total'] * 100
    add_bullet(doc, f" {fmtk(te['hours']['total'])} hours delivered across {fmtk(te['unique_tickets'])} unique tickets in 2025.",
               bold_prefix="Total Hours Delivered:")
    add_bullet(doc,
        f" {fmtk(te['hours']['nh'])} hours ({nh_pct:.1f}% of total) at $15/hr — overnight monitoring and "
        "maintenance delivered by the offshore team during their normal business hours "
        "(India daytime = US overnight).",
        bold_prefix="US Overnight (Offshore NH):")
    add_bullet(doc,
        f" {fmtk(te['hours']['ah'])} hours ({ah_pct:.1f}% of total) at $30/hr — US daytime support "
        "delivered by the offshore team during their after-hours (India night shift).",
        bold_prefix="US Business Hours (Offshore AH):")
    add_bullet(doc,
        f" {fmtk(te['hours']['onsite'])} hours ({on_pct:.1f}%) — 1 onsite visit for labor/hardware work.",
        bold_prefix="Onsite:")
    add_bullet(doc,
        f" {te['weekend_entries']} time entries on Saturdays and Sundays ({te['weekend_pct']}% of all entries), "
        "confirming 7-day availability.",
        bold_prefix="Weekend Coverage:")
    add_bullet(doc,
        f" {te['tech_count']} unique technicians. Top resources: Ajay Bhardwaj (59.1 hrs), "
        "Gurdeep Kumar (49.2 hrs), Sanjeev Kumar (39.3 hrs).",
        bold_prefix="Staffing Depth:")

    add_heading2(doc, "1.3 Remote vs. Onsite Mix")
    add_body(doc,
        "The model is 99.4% remote. Only 2.0 onsite hours were logged in 2025 (1 visit). "
        "Onsite visits are reserved for hardware-related needs only.")

    add_heading2(doc, "1.4 Services Billed Outside Monthly")
    add_bullet(doc, " M365 Business Standard ($8,833) and SFTP Server ($960).",
               bold_prefix="Recurring Licensing ($9,793):")
    add_bullet(doc,
        " Office recabling ($19,931), 3CX v18-to-v20 upgrade ($4,650), "
        "Meraki AP + UPS hardware ($4,207), Cisco Meraki 3YR license ($1,355), "
        "and other items.",
        bold_prefix="One-Time Projects ($31,485):")

    add_page_break(doc)

    # Section 2: Blended Rates
    add_heading1(doc, "Section 2: Effective Blended Rates")
    add_body(doc,
        "The following table shows the effective blended rate per role, calculated from "
        "billed hours on Monthly + WeeklyOut invoices:")

    make_table(doc,
        ["Role", "Contracted", "Overage", "Total Billed", "Billed Hours", "Blended Rate"],
        [
            ["IT Support (All Tiers)", "$10,843", "$150", "$10,993", "~394", "$27.90"],
            ["ALL LABOR", "$10,843", "$150", "$10,993", "~394", "$27.90"],
        ],
        col_aligns=[L, R, R, R, R, R],
        total_row_indices=[1],
    )

    add_note(doc,
        "Note: IT Support combines offshore NH ($15/hr, ~234 hrs), offshore AH ($30/hr, ~132 hrs), "
        "US remote ($125/hr, ~27 hrs), and one onsite visit ($150/hr, 1 hr). The overall "
        "blended rate of $27.90/hr reflects heavy offshore leverage — 93% of hours are offshore.")

    add_subheading(doc, "Rate Card Breakdown")
    make_table(doc,
        ["Role / Tier", "Rate", "Coverage"],
        [
            ["Offshore Support (US Overnight)", "$15.00/hr", "US overnight monitoring/maintenance (India NH)"],
            ["Offshore Support (US Daytime)", "$30.00/hr", "US business hours support (India AH)"],
            ["US Tech Support (Remote)", "$125.00/hr", "Escalations, complex issues"],
            ["US Tech Support (Onsite)", "$150.00/hr", "On-premises hardware work"],
        ],
        col_aligns=[L, R, L],
    )

    add_page_break(doc)

    # Section 3: Licensing & Recurring
    add_heading1(doc, "Section 3: Licensing & Recurring Services")
    add_heading2(doc, "3.1 Monthly Managed Services Breakdown")
    add_body(doc,
        "The monthly contract ($75,802/year, ~$6,317/month) covers both labor and managed "
        "services:")

    cats = inv['monthly_categories']
    cat_rows = []
    for cat, val in sorted(cats.items(), key=lambda x: -x[1]):
        cat_rows.append([cat, fmt(val), fmt(val / 12)])
    cat_rows.append(["TOTAL", fmt(inv['monthly_total']), fmt(inv['monthly_total'] / 12)])

    make_table(doc,
        ["Service Category", "Annual", "Monthly"],
        cat_rows,
        col_aligns=[L, R, R],
        total_row_indices=[len(cat_rows) - 1],
    )

    add_heading2(doc, "3.2 Recurring Licensing (Billed Separately)")
    make_table(doc,
        ["License / Service", "Annual (2025)", "Notes"],
        [
            ["M365 Business Standard (46 seats)", "$8,833", "$15/seat/mo — full year"],
            ["SFTP Server", "$960", "$80/mo — full year"],
            ["TOTAL", "$9,793", ""],
        ],
        col_aligns=[L, R, L],
        total_row_indices=[2],
    )

    add_heading2(doc, "3.3 Monthly Run Rate Trend")
    rate_rows = []
    for m, v in sorted(inv['monthly_run_rate'].items()):
        rate_rows.append([m, fmt(v)])
    make_table(doc,
        ["Month", "Monthly Total"],
        rate_rows,
        col_aligns=[L, R],
        highlight_row_indices=[len(rate_rows) - 1] if rate_rows else [],
    )
    add_note(doc,
        "Note: December 2025 appears elevated due to duplicate line items in the invoice data. "
        "The normalized December run rate is approximately $5,310. The monthly run rate trended "
        "downward during 2025 as labor allotments were right-sized quarterly.")

    add_page_break(doc)

    # Section 4: Projects & One-Time
    add_heading1(doc, "Section 4: Projects & One-Time Spend")
    add_body(doc,
        "Total non-contract spend in 2025 was $31,485, broken into infrastructure upgrades, "
        "phone system modernization, hardware, and licensing.")

    add_heading2(doc, "4.1 Infrastructure Projects")
    make_table(doc,
        ["Project", "Date", "Amount"],
        [
            ["Office Recabling (contractor)", "Aug 2025", "$19,931"],
            ["Meraki MR44 AP", "Sep 2025", "$1,298"],
            ["Cyber Power UPS", "Sep 2025", "$2,229"],
            ["Rackmount PDU", "Sep 2025", "$220"],
            ["UPS Management Card", "Sep 2025", "$426"],
            ["NEMA Power Cord", "Sep 2025", "$34"],
            ["AP and UPS SNMP Configuration", "Sep 2025", "$450"],
        ],
        col_aligns=[L, L, R],
    )

    add_heading2(doc, "4.2 Phone System Upgrade")
    make_table(doc,
        ["Item", "Date", "Amount"],
        [
            ["3CX v18 to v20 Upgrade — Phase 1", "Sep 2025", "$1,800"],
            ["3CX v18 to v20 Upgrade — Phase 2", "Sep 2025", "$750"],
            ["3CX v18 to v20 Upgrade — Phase 3", "Oct 2025", "$2,100"],
            ["TOTAL", "", "$4,650"],
        ],
        col_aligns=[L, L, R],
        total_row_indices=[3],
    )

    add_heading2(doc, "4.3 Licensing & Other")
    make_table(doc,
        ["Item", "Date", "Amount"],
        [
            ["Cisco Meraki 3YR Enterprise License", "Jun 2025", "$1,355"],
            ["3CX Renewal", "Jul 2025", "$350"],
            ["SSL Certificate (remote.viaautofinance.com)", "Jul 2025", "$19"],
            ["M365 License Adjustments (various)", "2025", "$72"],
            ["Late Fee", "Nov 2025", "$0"],
        ],
        col_aligns=[L, L, R],
    )

    add_heading2(doc, "4.4 Labor Overage")
    add_body(doc, "Only one WeeklyOut overage event in 2025: 1 hour of onsite support at "
             "$150/hr in September 2025. This confirms the contracted labor allotment was "
             "well-sized for VAF's needs.")

    add_page_break(doc)

    # Section 5: Ticket Categories
    add_heading1(doc, "Section 5: Ticket Categorization")
    add_body(doc,
        "All 462 unique tickets were categorized by analyzing title and notes fields. "
        "Each category is flagged as Proactive (scheduled maintenance) or Reactive (user-initiated).")

    cat_data = []
    for cat, s in te['categories'].items():
        cat_data.append([cat, s['type'], str(s['count']), fmtk(s['hours']), str(s['avg'])])
    cat_data.append(["TOTAL", "", str(te['unique_tickets']), fmtk(te['hours']['total']), ""])

    make_table(doc,
        ["Category", "Type", "Tickets", "Hours", "Avg Hrs"],
        cat_data,
        col_aligns=[L, L, R, R, R],
        total_row_indices=[len(cat_data) - 1],
    )

    pro = te['proactive_count']
    rea = te['reactive_count']
    tot = pro + rea
    add_body(doc,
        f"Proactive work accounts for {pro} tickets ({pro/tot*100:.1f}%) and reactive for "
        f"{rea} tickets ({rea/tot*100:.1f}%). The high proactive ratio reflects a mature "
        "managed services environment with robust patch management, monitoring, and backup routines.")

    add_page_break(doc)

    # Section 6: Service Metrics & Trends
    add_heading1(doc, "Section 6: Service Metrics & Trends")
    add_subheading(doc, "Monthly Volume Table")

    monthly_data = []
    for m, d in sorted(te['monthly'].items()):
        monthly_data.append([m, str(d['tickets']), fmtk(d['nh']), fmtk(d['ah']),
                            fmtk(d['onsite']), fmtk(d['total']), str(d['hrs_per_ticket'])])

    make_table(doc,
        ["Month", "Tickets", "US Overnight", "US Daytime", "Onsite", "Total Hrs", "Hrs/Ticket"],
        monthly_data,
        col_aligns=[L, R, R, R, R, R, R],
    )

    add_subheading(doc, "Notable Trends")
    add_bullet(doc, " 66.7% of tickets are proactive (patch management, monitoring, backup, "
               "security, RMM). Patch Management alone accounts for 210 tickets (45.5%).",
               bold_prefix="Proactive-heavy workload:")
    add_bullet(doc, " Volume peaked in January (69 tickets, 65.6 hrs) and normalized "
               "to ~30-40 tickets/month by Q2.",
               bold_prefix="Volume normalization:")
    add_bullet(doc, " Only 27 hours of US Tech Support in 2025. 92% of all support "
               "hours are offshore-delivered.",
               bold_prefix="High offshore leverage:")
    add_bullet(doc, " The labor allotment adjusted every quarter, trending downward "
               "from $2,231/mo (Jan) to $470/mo (Dec) as utilization data informed right-sizing.",
               bold_prefix="3-month contract cycle:")

    add_subheading(doc, "Recommendations")
    add_bullet(doc, " Patch Management (210 tickets, 110 hrs) is the largest category. "
               "Evaluate automation options to reduce manual patch handling.",
               bold_prefix="Automation opportunity:")
    add_bullet(doc, " Server Management tickets (48) consumed disproportionate hours "
               "(39.4 hrs). Monitor for recurring server issues.",
               bold_prefix="Server health monitoring:")
    add_bullet(doc, " Phone/VoIP (18 tickets) and Firewall/Network (20 tickets) both "
               "show above-average hours/ticket, suggesting complex configurations.",
               bold_prefix="Network & VoIP complexity:")
    add_bullet(doc, " The 3CX v18-to-v20 upgrade was the major phone system project. "
               "Monitor post-upgrade stability in Q1 2026.",
               bold_prefix="3CX upgrade follow-up:")

    add_page_break(doc)

    # Section 7: Summary
    add_heading1(doc, "Section 7: Executive Summary")
    add_body(doc,
        "VAF spent $117,229 on IT services in 2025, covering a comprehensive managed IT "
        "environment of 8 cloud servers, 48 desktops, and 50 users.")
    add_bullet(doc, " $75,802 monthly contract covers labor, security, monitoring, backup, "
               "cloud hosting, VoIP, and network services.",
               bold_prefix="Monthly Contract:")
    add_bullet(doc, " $9,793 for M365 Business Standard (46 seats) and SFTP server.",
               bold_prefix="Recurring Licensing:")
    add_bullet(doc, " $31,485 for office recabling ($19,931), 3CX phone system upgrade "
               "($4,650), hardware ($4,207), and licensing ($1,796).",
               bold_prefix="One-Time Projects:")
    add_bullet(doc, " $27.90/hr blended rate through 93% offshore leverage.",
               bold_prefix="Blended Rate:")
    add_bullet(doc, " 462 tickets resolved, 66.7% proactive, 333.9 total hours.",
               bold_prefix="Service Delivery:")

    add_subheading(doc, "Top Optimization Opportunities")
    add_bullet(doc, " Automate high-volume, low-touch patch management workflows.",
               bold_prefix="1.")
    add_bullet(doc, " Consolidate network monitoring and alerting to reduce noise.",
               bold_prefix="2.")
    add_bullet(doc, " Establish quarterly count reconciliation for device-based billing.",
               bold_prefix="3.")

    add_page_break(doc)

    # Section 8: Budget Projections
    add_heading1(doc, "Section 8: 2026 Budget Projections")
    add_body(doc,
        "Projections use the latest monthly run rate as baseline. VAF operates on a "
        "3-month labor contract cycle — allotments adjust quarterly based on utilization.")

    # Use the normalized latest run rate (Dec was doubled, so use ~$5,310)
    make_table(doc,
        ["Category", "2025 Actual", "2026 Projected", "Change"],
        [
            ["Monthly Contract", "$75,802", "$63,720", "-$12,082"],
            ["  Labor (at current rate)", "$10,843", "$5,640", "-$5,203"],
            ["  Managed Services", "$64,959", "$58,080", "-$6,879"],
            ["Recurring Licensing", "$9,793", "$10,080", "+$287"],
            ["Estimated Overage", "$150", "$150", "$0"],
            ["One-Time (Renewals)", "$31,485", "$3,500", "-$27,985"],
            ["PROJECTED TOTAL", "$117,229", "$77,450", "-$39,779"],
        ],
        col_aligns=[L, R, R, R],
        total_row_indices=[6],
        highlight_row_indices=[5],
    )

    add_subheading(doc, "Key Assumptions")
    add_bullet(doc, " Labor allotment stays at the current Q4 2025 level (~$470/mo) "
               "and adjusts quarterly. This is conservative — actual may be higher if "
               "server management work increases.",
               bold_prefix="Labor:")
    add_bullet(doc, " Device-based services projected at current counts (48 desktops, "
               "8 servers, 50 users). Counts adjust monthly on invoices.",
               bold_prefix="Managed Services:")
    add_bullet(doc, " M365 at 46 seats x $15/mo ($8,280) + SFTP ($960) + expected "
               "growth buffer.",
               bold_prefix="Recurring:")
    add_bullet(doc, " 2025 was an exceptional year for one-time spend (office recabling, "
               "3CX upgrade). 2026 projects estimated at $3,500 for SSL renewals, Meraki "
               "license, and ad-hoc needs.",
               bold_prefix="One-Time:")
    add_note(doc,
        "Note: The significant year-over-year decrease is driven primarily by the non-recurrence "
        "of the $19,931 office recabling project and $4,650 3CX upgrade. Core managed services "
        "costs remain stable.")

    return doc

# ══════════════════════════════════════════════════════════════
#  HHOC 2025 Annual Review
# ══════════════════════════════════════════════════════════════
def build_hhoc_2025(data):
    te = data['hhoc_2025_te']
    doc = setup_doc("HHOC (Housing for Health OC) — 2025 IT Support Review")

    add_title_page(doc, "HHOC (Housing for Health OC)",
                   "2025 IT Support Review",
                   "& 2026 Planning Readiness Assessment",
                   "Prepared by Technijian Inc.  —  March 2026")

    doc.add_paragraph()
    make_table(doc,
        ["2025 At a Glance", "Key Figures"],
        [
            ["Unique Tickets (2025)", "169"],
            ["Total Support Hours Delivered", "149"],
            ["US Overnight Hours (Offshore NH)", "98"],
            ["US Daytime Hours (Offshore AH + US Tech)", "51"],
            ["US Tech Support Hours", "42"],
            ["Unique Technicians", "18"],
            ["Weekend Entries", "44 (20.3%)"],
            ["Primary US Resource", "Sunny Sark (38.8 hrs)"],
        ],
        col_aligns=[L, R],
    )

    add_note(doc,
        "Note: Invoice data for HHOC was unavailable for this review (the export contained "
        "another client's data). Financial sections (blended rates, contract breakdown, budget "
        "projections) will be updated once the correct HHOC invoice export is provided.")

    add_page_break(doc)

    # Section 1
    add_heading1(doc, "Section 1: Support Scope & Coverage")
    add_heading2(doc, "1.1 Coverage Model")
    add_body(doc,
        "HHOC operates under two service contracts: a standard Monthly Service agreement and "
        "a Monthly Service with India-Night contract that adds US daytime coverage through "
        "the offshore team's after-hours shift. Together, these provide comprehensive 24/7 "
        "IT support coverage.")
    add_body(doc,
        "The support team delivered 149.1 hours across 169 unique tickets in 2025 — an average "
        "of 14.1 tickets and 12.4 hours per month. The model is 100% remote with zero onsite "
        "visits during the review period.")

    add_heading2(doc, "1.2 Coverage Hours")
    nh_pct = te['hours']['nh'] / te['hours']['total'] * 100
    ah_pct = te['hours']['ah'] / te['hours']['total'] * 100
    add_bullet(doc, f" {fmtk(te['hours']['total'])} hours across {te['unique_tickets']} tickets.",
               bold_prefix="Total Hours Delivered:")
    add_bullet(doc,
        f" {fmtk(te['hours']['nh'])} hours ({nh_pct:.1f}%) — overnight monitoring and "
        "maintenance (India daytime = US overnight), plus US-based tech support remote hours.",
        bold_prefix="Normal Hours (NH):")
    add_bullet(doc,
        f" {fmtk(te['hours']['ah'])} hours ({ah_pct:.1f}%) — US daytime support delivered "
        "by the offshore team during their after-hours (India night shift).",
        bold_prefix="After Hours (AH / US Daytime):")
    add_bullet(doc, " Zero onsite hours in 2025 — entirely remote delivery.",
               bold_prefix="Onsite:")
    add_bullet(doc,
        f" {te['weekend_entries']} entries on weekends ({te['weekend_pct']}%), confirming "
        "7-day availability.",
        bold_prefix="Weekend Coverage:")

    add_heading2(doc, "1.3 Role Distribution")
    role_data = []
    for role, h in sorted(te['by_role'].items(), key=lambda x: -(x[1]['nh']+x[1]['ah']+x[1]['onsite'])):
        t = h['nh'] + h['ah'] + h['onsite']
        pct = t / te['hours']['total'] * 100
        role_data.append([role, fmtk(h['nh']), fmtk(h['ah']), fmtk(t), f"{pct:.1f}%", str(h['count'])])
    make_table(doc,
        ["Role", "NH", "AH", "Total", "% of Total", "Entries"],
        role_data,
        col_aligns=[L, R, R, R, R, R],
    )
    add_body(doc,
        "Off-Shore Tech Support handles 71.5% of total hours, with Tech Support (US-based) "
        "contributing 28.4%. Sunny Sark (38.8 hrs) was the primary US resource, providing "
        "continuity and escalation handling. The split reflects a balanced coverage model.")

    add_heading2(doc, "1.4 Staffing")
    tech_data = []
    for name, hrs in list(te['technicians'].items())[:10]:
        tech_data.append([name, fmtk(hrs)])
    make_table(doc, ["Technician", "Hours"], tech_data, col_aligns=[L, R])

    add_page_break(doc)

    # Section 2-4: Financial (placeholder)
    add_heading1(doc, "Sections 2-4: Financial Analysis")
    add_body(doc,
        "Invoice data for HHOC was unavailable for this review period. The invoice export "
        "provided contained another client's billing data. The following sections will be "
        "completed once the correct HHOC invoice export is provided:")
    add_bullet(doc, " Section 2: Effective Blended Rates")
    add_bullet(doc, " Section 3: Licensing & Recurring Services")
    add_bullet(doc, " Section 4: Projects & One-Time Spend")

    add_page_break(doc)

    # Section 5: Ticket Categories
    add_heading1(doc, "Section 5: Ticket Categorization")
    add_body(doc,
        f"All {te['unique_tickets']} unique tickets were categorized by analyzing title and notes fields.")

    cat_data = []
    for cat, s in te['categories'].items():
        cat_data.append([cat, s['type'], str(s['count']), fmtk(s['hours']), str(s['avg'])])
    cat_data.append(["TOTAL", "", str(te['unique_tickets']), fmtk(te['hours']['total']), ""])

    make_table(doc,
        ["Category", "Type", "Tickets", "Hours", "Avg Hrs"],
        cat_data,
        col_aligns=[L, L, R, R, R],
        total_row_indices=[len(cat_data) - 1],
        highlight_row_indices=[0],  # Email & M365 dominant
    )

    pro = te['proactive_count']
    rea = te['reactive_count']
    tot = pro + rea
    add_body(doc,
        f"Proactive work accounts for {pro} tickets ({pro/tot*100:.1f}%) and reactive for "
        f"{rea} tickets ({rea/tot*100:.1f}%). The reactive-leaning ratio is driven by the "
        "dominance of Email & M365 tickets (70 tickets, 41.4% of all tickets).")

    add_subheading(doc, "Key Category Insights")
    add_bullet(doc,
        " 70 tickets consuming 69.2 hours (46.4% of total hours). This is by far "
        "the largest category, indicating heavy M365 administration work including "
        "user management, SharePoint, Teams, and email configuration.",
        bold_prefix="Email & M365 dominance:")
    add_bullet(doc,
        " File & Permissions tickets, while only 3 in count, averaged 3.14 hours each — "
        "the highest per-ticket effort, suggesting complex permission structures.",
        bold_prefix="File & Permissions complexity:")
    add_bullet(doc,
        " Patch Management (17 tickets, 7.2 hrs), RMM (13 tickets, 8.5 hrs), "
        "Security (13 tickets, 10.9 hrs), Monitoring (10 tickets, 7.4 hrs), "
        "and Backup (7 tickets, 5.0 hrs) collectively form the proactive maintenance base.",
        bold_prefix="Proactive maintenance:")

    add_page_break(doc)

    # Section 6: Trends
    add_heading1(doc, "Section 6: Service Metrics & Trends")
    add_subheading(doc, "Monthly Volume Table")

    monthly_data = []
    for m, d in sorted(te['monthly'].items()):
        monthly_data.append([m, str(d['tickets']), fmtk(d['nh']), fmtk(d['ah']),
                            fmtk(d['total']), str(d['hrs_per_ticket'])])
    make_table(doc,
        ["Month", "Tickets", "NH", "AH", "Total Hrs", "Hrs/Ticket"],
        monthly_data,
        col_aligns=[L, R, R, R, R, R],
    )

    add_subheading(doc, "Notable Trends")
    add_bullet(doc, " Ticket volume was lower in Q2 (May: 5 tickets) and higher "
               "in Q3-Q4 (Oct-Nov peak). Hours/ticket also increased in the second half, "
               "suggesting more complex work.",
               bold_prefix="Seasonal variation:")
    add_bullet(doc, " Email & M365 work is the primary driver of support volume. "
               "Consider whether dedicated M365 admin sessions could be more efficient "
               "than ticket-based handling.",
               bold_prefix="M365 concentration:")
    add_bullet(doc, " The two-contract structure (Monthly Service + India-Night) provides "
               "24/7 coverage. 83.4% of entries were under the standard contract.",
               bold_prefix="Contract mix:")
    add_bullet(doc, " 100% remote delivery with 20.3% weekend entries confirms "
               "always-available support posture.",
               bold_prefix="Fully remote model:")

    add_subheading(doc, "Recommendations")
    add_bullet(doc, " Evaluate recurring M365 tasks for standardization or automation. "
               "70 tickets is a significant volume that may benefit from runbooks.",
               bold_prefix="M365 workflow optimization:")
    add_bullet(doc, " At 35.5% proactive, there's room to increase scheduled maintenance. "
               "More automated patching and monitoring could shift the ratio.",
               bold_prefix="Increase proactive coverage:")
    add_bullet(doc, " Establish documented processes for common M365 tasks to reduce "
               "per-ticket resolution time.",
               bold_prefix="Knowledge base development:")

    add_page_break(doc)

    # Section 7: Summary
    add_heading1(doc, "Section 7: Executive Summary")
    add_body(doc,
        "HHOC received 149.1 hours of IT support across 169 tickets in 2025, delivered "
        "100% remotely by a team of 18 technicians with 7-day coverage.")
    add_bullet(doc, " 71.5% offshore, 28.4% US Tech Support.",
               bold_prefix="Coverage Model:")
    add_bullet(doc, " Email & M365 (41.4%), Patch Management (10.1%), RMM (7.7%), "
               "Security (7.7%), General IT (5.9%).",
               bold_prefix="Top Categories:")
    add_bullet(doc, " 60 proactive (35.5%), 109 reactive (64.5%).",
               bold_prefix="Proactive/Reactive:")
    add_bullet(doc, " Sunny Sark (38.8 hrs), Parveen Biswal (25.6 hrs), "
               "Sanjeev Kumar (18.3 hrs).",
               bold_prefix="Primary Resources:")

    add_subheading(doc, "Top Priorities for 2026")
    add_bullet(doc, " Streamline Email & M365 support workflows.")
    add_bullet(doc, " Increase proactive maintenance ratio (target: 45%+).")
    add_bullet(doc, " Obtain correct invoice export for complete financial analysis.")

    return doc

# ══════════════════════════════════════════════════════════════
#  VAF 2026 Q1 Review
# ══════════════════════════════════════════════════════════════
def build_vaf_2026_q1(data):
    te = data['vaf_2026_te']
    te_2025 = data['vaf_2025_te']
    doc = setup_doc("VAF (Via Auto Finance) — 2026 Q1 IT Support Review")

    add_title_page(doc, "VAF (Via Auto Finance)",
                   "2026 Q1 IT Support Review",
                   "(January – March 2026)",
                   "Prepared by Technijian Inc.  —  March 2026")

    doc.add_paragraph()
    q_avg_tickets = te_2025['unique_tickets'] / 4
    q_avg_hours = te_2025['hours']['total'] / 4
    make_table(doc,
        ["Q1 2026 At a Glance", "Q1 2026", "2025 Quarterly Avg"],
        [
            ["Unique Tickets", "124", f"{q_avg_tickets:.0f}"],
            ["Total Hours", "89.4", f"{q_avg_hours:.1f}"],
            ["US Overnight (NH)", "59.7", f"{te_2025['hours']['nh']/4:.1f}"],
            ["US Daytime (AH)", "29.7", f"{te_2025['hours']['ah']/4:.1f}"],
            ["Technicians", "14", "18"],
            ["Proactive %", "60.5%", "66.7%"],
            ["Weekend Entries", "19.1%", "15.3%"],
        ],
        col_aligns=[L, R, R],
    )

    add_page_break(doc)

    add_heading1(doc, "Section 1: Q1 Overview")
    add_body(doc,
        "Q1 2026 tracked slightly above 2025 quarterly averages: 124 tickets (vs 116 avg) "
        "and 89.4 hours (vs 83.5 avg). The support model remains almost entirely offshore — "
        "only 0.2 hours of US Tech Support were logged, compared to 27 hours for all of 2025.")

    add_heading2(doc, "1.1 Monthly Consistency")
    monthly_data = []
    for m, d in sorted(te['monthly'].items()):
        monthly_data.append([m, str(d['tickets']), fmtk(d['nh']), fmtk(d['ah']),
                            fmtk(d['total']), str(d['hrs_per_ticket'])])
    make_table(doc,
        ["Month", "Tickets", "NH", "AH", "Total", "Hrs/Tkt"],
        monthly_data,
        col_aligns=[L, R, R, R, R, R],
    )
    add_body(doc, "Volume was remarkably consistent across all three months (~41 tickets each). "
             "AH (US Daytime) hours trended slightly upward month-over-month.")

    add_page_break(doc)

    add_heading1(doc, "Section 2: Ticket Categorization")
    cat_data = []
    for cat, s in te['categories'].items():
        cat_data.append([cat, s['type'], str(s['count']), fmtk(s['hours']), str(s['avg'])])
    cat_data.append(["TOTAL", "", str(te['unique_tickets']), fmtk(te['hours']['total']), ""])

    make_table(doc,
        ["Category", "Type", "Tickets", "Hours", "Avg Hrs"],
        cat_data,
        col_aligns=[L, L, R, R, R],
        total_row_indices=[len(cat_data) - 1],
    )

    add_subheading(doc, "Key Observations")
    add_bullet(doc,
        " 25 tickets in Q1 vs 48 for all of 2025, putting it on pace to double. "
        "Average hours/ticket also increased (1.07 vs 0.82). This warrants monitoring "
        "for recurring server issues or infrastructure changes.",
        bold_prefix="Server Management surge:")
    add_bullet(doc,
        " Still the #1 category (33 tickets) with higher avg hours (0.75 vs 0.52 in 2025). "
        "Patches may be requiring more remediation effort.",
        bold_prefix="Patch Management:")
    add_bullet(doc,
        " 22 tickets but only 6.4 hours (0.29 hrs/ticket) — highly efficient automated "
        "or scripted backup checks.",
        bold_prefix="Backup & DR efficiency:")
    add_bullet(doc,
        " Proactive work is 60.5% vs 66.7% in 2025. The shift toward server management "
        "(reactive) is pulling the ratio down.",
        bold_prefix="Proactive ratio dipped:")

    add_page_break(doc)

    add_heading1(doc, "Section 3: Staffing & Coverage")
    tech_data = []
    for name, hrs in te['technicians'].items():
        tech_data.append([name, fmtk(hrs)])
    make_table(doc, ["Technician", "Q1 Hours"], tech_data, col_aligns=[L, R])

    add_body(doc,
        "The top two resources — Ajay Bhardwaj (23.0 hrs) and Gurdeep Kumar (21.0 hrs) — "
        "account for 49.3% of Q1 hours, consistent with 2025 patterns. Aditya Saraf (13.2 hrs) "
        "is a new addition to the team not present in 2025 top contributors.")
    add_body(doc,
        "Notable: US Tech Support usage dropped to near-zero (0.2 hrs in Q1 vs 27 hrs in 2025). "
        "This represents a shift to nearly 100% offshore delivery, potentially reducing costs "
        "but worth monitoring for escalation gaps.")

    add_page_break(doc)

    add_heading1(doc, "Section 4: Recommendations & Q2 Outlook")
    add_bullet(doc,
        " The server management surge should be investigated. Identify recurring "
        "issues (disk alerts, performance, VM maintenance) and address root causes.",
        bold_prefix="Investigate server management spike:")
    add_bullet(doc,
        " With near-zero US Tech Support, ensure escalation paths remain clear "
        "for issues requiring US-based expertise.",
        bold_prefix="Monitor escalation paths:")
    add_bullet(doc,
        " Patch remediation hours are increasing. Evaluate whether newer OS/software "
        "versions or automation can reduce manual intervention.",
        bold_prefix="Optimize patch workflows:")
    add_bullet(doc,
        " Based on Q1 run rate, VAF is tracking toward ~496 tickets and ~358 hours "
        "for 2026, both above 2025 levels. If server management continues at Q1 pace, "
        "consider adjusting the labor allotment at the Q2 cycle.",
        bold_prefix="Q2 labor review:")

    return doc

# ══════════════════════════════════════════════════════════════
#  HHOC 2026 Q1 Review
# ══════════════════════════════════════════════════════════════
def build_hhoc_2026_q1(data):
    te = data['hhoc_2026_te']
    te_2025 = data['hhoc_2025_te']
    doc = setup_doc("HHOC (Housing for Health OC) — 2026 Q1 IT Support Review")

    add_title_page(doc, "HHOC (Housing for Health OC)",
                   "2026 Q1 IT Support Review",
                   "(January – March 2026)",
                   "Prepared by Technijian Inc.  —  March 2026")

    doc.add_paragraph()
    q_avg_tickets = te_2025['unique_tickets'] / 4
    q_avg_hours = te_2025['hours']['total'] / 4
    make_table(doc,
        ["Q1 2026 At a Glance", "Q1 2026", "2025 Quarterly Avg", "Change"],
        [
            ["Unique Tickets", "54", f"{q_avg_tickets:.0f}", "+28%"],
            ["Total Hours", "69.0", f"{q_avg_hours:.1f}", "+85%"],
            ["Hours/Ticket", "1.28", "0.88", "+45%"],
            ["US Tech Support Hours", "26.4", f"{42.4/4:.1f}", "+149%"],
            ["Proactive %", "13.0%", "35.5%", "-22.5pp"],
            ["Email & M365 % of Hours", "67.4%", "46.4%", "+21.0pp"],
            ["Weekend Entries", "12.2%", "20.3%", "-8.1pp"],
        ],
        col_aligns=[L, R, R, R],
        highlight_row_indices=[1, 4, 5],
    )

    add_page_break(doc)

    add_heading1(doc, "Section 1: Q1 Overview — Key Concerns")
    add_body(doc,
        "Q1 2026 shows significant changes from 2025 patterns that warrant attention. "
        "Total hours nearly doubled compared to the 2025 quarterly average (69.0 vs 37.3), "
        "driven primarily by complex Email & M365 work. The proactive/reactive balance "
        "shifted dramatically toward reactive (87% reactive vs 65% in 2025).")

    add_heading2(doc, "1.1 Volume Spike Analysis")
    monthly_data = []
    for m, d in sorted(te['monthly'].items()):
        monthly_data.append([m, str(d['tickets']), fmtk(d['nh']), fmtk(d['ah']),
                            fmtk(d['total']), str(d['hrs_per_ticket'])])
    make_table(doc,
        ["Month", "Tickets", "NH", "AH", "Total", "Hrs/Tkt"],
        monthly_data,
        col_aligns=[L, R, R, R, R, R],
    )

    add_body(doc,
        "January and February showed elevated hours/ticket (1.43 and 1.54 respectively), "
        "normalizing to 0.94 in March. This suggests a burst of complex work early in Q1 "
        "that has started to stabilize.")

    add_page_break(doc)

    add_heading1(doc, "Section 2: Ticket Categorization — Email & M365 Deep Dive")
    cat_data = []
    for cat, s in te['categories'].items():
        cat_data.append([cat, s['type'], str(s['count']), fmtk(s['hours']), str(s['avg'])])
    cat_data.append(["TOTAL", "", str(te['unique_tickets']), fmtk(te['hours']['total']), ""])

    make_table(doc,
        ["Category", "Type", "Tickets", "Hours", "Avg Hrs"],
        cat_data,
        col_aligns=[L, L, R, R, R],
        total_row_indices=[len(cat_data) - 1],
        highlight_row_indices=[0],
    )

    add_subheading(doc, "Critical Finding: Email & M365")
    add_body(doc,
        "Email & M365 consumed 46.5 hours across 22 tickets in Q1 alone — that is 67.4% of "
        "all support hours. The average hours/ticket jumped from 0.99 in 2025 to 2.11 in Q1 2026, "
        "indicating significantly more complex issues (not just more tickets).")
    add_body(doc,
        "This category alone accounts for nearly the entire quarterly average of all support "
        "hours from 2025 (46.5 vs 37.3). If this trend continues, it will drive substantial "
        "overage against any contracted labor allotment.")

    add_subheading(doc, "Proactive Work Collapse")
    add_body(doc,
        "Only 7 proactive tickets in Q1 (13.0%) compared to 60 for all of 2025 (35.5%). "
        "Patch Management dropped from 17 tickets in 2025 to just 5 in Q1 2026. "
        "Monitoring, RMM, and Backup tickets are near zero. This suggests proactive "
        "maintenance routines may need attention.")

    add_page_break(doc)

    add_heading1(doc, "Section 3: Staffing Transition")
    tech_data = []
    for name, hrs in te['technicians'].items():
        tech_data.append([name, fmtk(hrs)])
    make_table(doc, ["Technician", "Q1 Hours"], tech_data, col_aligns=[L, R])

    add_body(doc,
        "Major staffing change: Sunny Sark, who was the primary US Tech Support resource "
        "in 2025 (38.8 hours, #1 contributor), does not appear in Q1 2026 data. "
        "Rishad Mohamed (19.6 hrs) has taken over as the primary US resource, "
        "with Deepak Bhardwaj (16.7 hrs) as the new offshore lead.")

    add_body(doc,
        "US Tech Support hours actually increased as a percentage (38.3% in Q1 2026 vs "
        "28.4% in 2025), suggesting Rishad Mohamed is handling more complex escalations. "
        "All Q1 entries are under the 'Monthly Service with India-Night' contract "
        "(vs a split between two contracts in 2025).")

    add_page_break(doc)

    add_heading1(doc, "Section 4: Recommendations & Q2 Action Items")

    add_heading2(doc, "Immediate Actions")
    add_bullet(doc,
        " The 2.11 hrs/ticket average suggests either complex projects "
        "being handled through support tickets or systemic M365 configuration issues. "
        "Review the specific tickets to identify root causes and consider a dedicated "
        "M365 optimization engagement.",
        bold_prefix="1. Audit Email & M365 ticket drivers:")
    add_bullet(doc,
        " Patch Management, Monitoring, RMM, and Backup work has "
        "dropped precipitously. Verify that automated routines are still running and "
        "that the monitoring stack is properly configured.",
        bold_prefix="2. Restore proactive maintenance:")
    add_bullet(doc,
        " Ensure Rishad Mohamed and Deepak Bhardwaj have full "
        "context on HHOC's environment and that knowledge transfer from previous "
        "team members is complete.",
        bold_prefix="3. Validate staffing transition:")

    add_heading2(doc, "Q2 Monitoring Targets")
    add_bullet(doc, " Email & M365 hours/ticket below 1.5 (currently 2.11)")
    add_bullet(doc, " Proactive ratio above 25% (currently 13%)")
    add_bullet(doc, " Patch Management tickets at 8+ per quarter (currently 5)")
    add_bullet(doc, " Total quarterly hours below 50 (currently 69)")

    add_note(doc,
        "If Q1 trends continue unabated, HHOC is tracking toward 216 tickets and "
        "276 hours for 2026 — nearly double the 2025 levels. The labor allotment should "
        "be reviewed at the next contract cycle.")

    return doc


# ══════════════════════════════════════════════════════════════
#  Main
# ══════════════════════════════════════════════════════════════
def main():
    with open(os.path.join(BASE, 'scripts/analysis_results.json')) as f:
        data = json.load(f)

    docs = [
        ("VAF 2025 Annual", build_vaf_2025, os.path.join(BASE, "clients/vaf/2025/VAF - 2025 Annual Review.docx")),
        ("HHOC 2025 Annual", build_hhoc_2025, os.path.join(BASE, "clients/hhoc/2025/HHOC - 2025 Annual Review.docx")),
        ("VAF 2026 Q1", build_vaf_2026_q1, os.path.join(BASE, "clients/vaf/2026/VAF - 2026 Q1 Review.docx")),
        ("HHOC 2026 Q1", build_hhoc_2026_q1, os.path.join(BASE, "clients/hhoc/2026/HHOC - 2026 Q1 Review.docx")),
    ]

    for name, builder, path in docs:
        print(f"Generating {name}...")
        doc = builder(data)
        doc.save(path)
        print(f"  Saved: {path}")

    print("\nAll documents generated successfully!")

if __name__ == "__main__":
    main()
