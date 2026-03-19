"""
Generate HHOC 2025 Annual IT Review — Full Professional Word Document
Technijian Inc. -> Housing for Health OC
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
MED_BLUE = RGBColor(0x2E, 0x75, 0xB6)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_BLUE_HEX = "1F4E79"
ALT_ROW_HEX = "F2F7FB"
WHITE_HEX = "FFFFFF"
YELLOW_HEX = "FFF3CD"

OUT = r"c:\vscode\annual-client-review\annual-client-review\clients\hhoc\2025\HHOC - 2025 Annual Review.docx"

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, size=Pt(8.5), color=BLACK,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = "Calibri"

def add_header_row(table, texts, col_aligns=None):
    row = table.rows[0]
    for i, txt in enumerate(texts):
        align = (col_aligns[i] if col_aligns else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[i], txt, bold=True, color=WHITE, alignment=align)
        set_cell_shading(row.cells[i], DARK_BLUE_HEX)

def add_data_row(table, texts, row_idx, col_aligns=None, is_total=False, is_highlight=False):
    row = table.rows[row_idx]
    bg = DARK_BLUE_HEX if is_total else (YELLOW_HEX if is_highlight else
         (ALT_ROW_HEX if row_idx % 2 == 0 else WHITE_HEX))
    txt_color = WHITE if is_total else BLACK
    for i, txt in enumerate(texts):
        align = (col_aligns[i] if col_aligns else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[i], txt, bold=is_total, color=txt_color, alignment=align)
        set_cell_shading(row.cells[i], bg)

def make_table(doc, headers, data, col_aligns=None, total_row_indices=None, highlight_row_indices=None):
    total_row_indices = total_row_indices or []
    highlight_row_indices = highlight_row_indices or []
    tbl = doc.add_table(rows=1 + len(data), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:top w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:left w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:bottom w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:right w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '</w:tcBorders>')
            tcPr.append(borders)
    add_header_row(tbl, headers, col_aligns)
    for idx, row_data in enumerate(data):
        add_data_row(tbl, row_data, idx + 1, col_aligns,
                     is_total=idx in total_row_indices,
                     is_highlight=idx in highlight_row_indices)
    return tbl

def h1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    r = p.add_run(text); r.font.color.rgb = DARK_BLUE; r.font.size = Pt(16); r.font.bold = True
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)

def h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    r = p.add_run(text); r.font.color.rgb = DARK_BLUE; r.font.size = Pt(13); r.font.bold = True
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)

def sh(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = DARK_BLUE; r.font.size = Pt(11); r.font.bold = True
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)

def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.font.size = Pt(10); r.font.bold = True; r.font.name = "Calibri"
    r = p.add_run(text); r.font.size = Pt(10); r.font.name = "Calibri"

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.font.size = Pt(10); r.font.bold = True; r.font.name = "Calibri"
    r = p.add_run(text); r.font.size = Pt(10); r.font.name = "Calibri"

def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = GRAY; r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(4)

L = WD_ALIGN_PARAGRAPH.LEFT
R = WD_ALIGN_PARAGRAPH.RIGHT

def build():
    doc = Document()
    s = doc.sections[0]
    s.page_width = Inches(8.5); s.page_height = Inches(11)
    s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(0.63); s.right_margin = Inches(0.63)
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)

    hp = s.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("HHOC (Housing for Health OC) — 2025 IT Support Review & 2026 Planning")
    hr.font.size = Pt(8); hr.font.color.rgb = GRAY; hr.font.name = "Calibri"
    fp = s.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Technijian Inc. — Confidential")
    fr.font.size = Pt(8); fr.font.color.rgb = GRAY; fr.font.name = "Calibri"

    # Title page
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HHOC (Housing for Health OC)"); r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = DARK_BLUE
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2025 IT Support Review"); r.font.size = Pt(20); r.font.color.rgb = MED_BLUE
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("& 2026 Planning Readiness Assessment"); r.font.size = Pt(20); r.font.color.rgb = MED_BLUE
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared by Technijian Inc.  —  March 2026"); r.font.size = Pt(12); r.font.color.rgb = GRAY

    doc.add_paragraph()
    make_table(doc,
        ["2025 At a Glance", "Key Figures"],
        [
            ["Total 2025 Revenue", "$12,192"],
            ["Monthly Contract (Labor + Services)", "$11,598"],
            ["Recurring Licensing (M365 Non-Profit)", "$594"],
            ["Projects / One-Time", "$0"],
            ["Labor Overage", "$0"],
            ["Unique Tickets (2025)", "169"],
            ["Total Support Hours Delivered", "149"],
            ["Desktops Managed", "8"],
            ["Active Mailboxes (Anti-Spam)", "9"],
        ],
        col_aligns=[L, R], total_row_indices=[0])

    doc.add_page_break()

    # ── Section 1 ──
    h1(doc, "Section 1: Support Scope & Coverage")
    h2(doc, "1.1 What Is Included in Monthly Support")
    body(doc,
        "The monthly agreement covers a managed IT package for HHOC, "
        "totaling $11,598 across 2025 (~$966/month). This breaks down into two components:")
    body(doc,
        " US Tech Support ($150/hr, 2.75 hrs/mo contracted), offshore support providing "
        "US business hours coverage ($45/hr AH) and US overnight monitoring ($30/hr NH, "
        "reduced to $15/hr in November). Labor adjusts quarterly based on utilization.",
        bold_prefix="Labor (~$8,231/year, ~$686/month):")
    body(doc,
        " Security (CrowdStrike, Huntress, Malwarebytes), email security "
        "(anti-spam, DKIM/DMARC, phishing training), monitoring (RMM, patch management, "
        "My Remote), secure internet filtering, and site assessment.",
        bold_prefix="Managed Services (~$3,366/year, ~$281/month):")

    h2(doc, "1.2 Coverage Hours")
    bullet(doc, " 149.1 hours across 169 unique tickets.", bold_prefix="Total:")
    bullet(doc, " 98.3 hours (65.9%) — includes offshore NH (55.8 hrs for "
           "overnight monitoring) and US Tech Support NH (42.4 hrs).",
           bold_prefix="Normal Hours (NH):")
    bullet(doc, " 50.8 hours (34.1%) — US daytime support delivered by offshore "
           "team during India after-hours shift at $45/hr.",
           bold_prefix="After Hours (AH / US Daytime):")
    bullet(doc, " Zero onsite hours — entirely remote delivery model.",
           bold_prefix="Onsite:")
    bullet(doc, " 44 entries on weekends (20.3%), confirming 7-day availability.",
           bold_prefix="Weekend Coverage:")
    bullet(doc, " 18 unique technicians. Primary: Sunny Sark (38.8 hrs, US Tech), "
           "Parveen Biswal (25.6 hrs), Sanjeev Kumar (18.3 hrs).",
           bold_prefix="Staffing Depth:")

    h2(doc, "1.3 Role Distribution")
    make_table(doc,
        ["Role", "NH Hrs", "AH Hrs", "Total Hrs", "% of Total", "Entries"],
        [
            ["Off-Shore Tech Support", "55.8", "50.8", "106.6", "71.5%", "157"],
            ["Tech Support (US)", "42.4", "0.0", "42.4", "28.4%", "60"],
            ["TOTAL", "98.3", "50.8", "149.1", "100%", "217"],
        ],
        col_aligns=[L, R, R, R, R, R], total_row_indices=[2])

    h2(doc, "1.4 Contract Structure")
    body(doc,
        "HHOC operates under two service contracts: 'Monthly Service' (181 entries, 83.4%) "
        "covers standard support, while 'Monthly Service with India-Night' (36 entries, 16.6%) "
        "adds US daytime coverage through the offshore team's after-hours shift. Together they "
        "provide comprehensive 24/7 coverage.")

    doc.add_page_break()

    # ── Section 2 ──
    h1(doc, "Section 2: Effective Blended Rates")
    body(doc,
        "Blended rates are calculated from billed hours on Monthly invoices:")
    make_table(doc,
        ["Role", "Contracted Revenue", "Billed Hours", "Blended Rate"],
        [
            ["US Tech Support", "$4,050", "~27 hrs", "$150.00/hr"],
            ["Offshore AH (US Daytime)", "$2,430", "~54 hrs", "$45.00/hr"],
            ["Offshore NH (US Overnight)", "$1,751", "~89 hrs", "$19.67/hr*"],
            ["ALL LABOR", "$8,231", "~170 hrs", "$48.42/hr"],
        ],
        col_aligns=[L, R, R, R], total_row_indices=[3])

    note(doc,
        "* NH rate was $30/hr Jan-Oct and $15/hr Nov-Dec, producing a blended NH rate of $19.67. "
        "US Tech Support at $150/hr and Offshore AH at $45/hr are higher than the standard "
        "Technijian schedule ($125 and $30 respectively).")

    sh(doc, "Rate Card")
    make_table(doc,
        ["Role / Tier", "Rate", "Coverage"],
        [
            ["Offshore Support (US Overnight)", "$15.00/hr*", "Overnight monitoring/maintenance (India NH)"],
            ["Offshore Support (US Daytime)", "$45.00/hr", "US business hours support (India AH)"],
            ["US Tech Support (Remote)", "$150.00/hr", "Escalations, complex issues"],
        ],
        col_aligns=[L, R, L])
    note(doc, "* NH rate was $30/hr through October 2025; adjusted to $15/hr from November 2025.")

    doc.add_page_break()

    # ── Section 3 ──
    h1(doc, "Section 3: Licensing & Recurring Services")
    h2(doc, "3.1 Monthly Managed Services Breakdown")
    body(doc,
        "The monthly contract ($11,598/year, ~$966/month) covers labor and managed services:")
    make_table(doc,
        ["Service Category", "Annual", "Monthly"],
        [
            ["Labor (US Tech + Offshore)", "$8,231", "$686"],
            ["Security & Endpoint (CrowdStrike/Huntress/Malwarebytes)", "$1,107", "$92"],
            ["Email Security (Anti-Spam/DKIM/Phishing)", "$1,071", "$89"],
            ["Assessment (Site Assessment)", "$450", "$38"],
            ["Monitoring & Management (RMM/Remote/Patch)", "$410", "$34"],
            ["Secure Internet", "$328", "$27"],
            ["TOTAL", "$11,598", "$966"],
        ],
        col_aligns=[L, R, R], total_row_indices=[6])

    sh(doc, "December 2025 Monthly Invoice Detail")
    make_table(doc,
        ["Service", "Qty", "Rate", "Monthly", "Category"],
        [
            ["US Tech Support (IRV-TS1)", "2.75", "$150.00", "$413", "Labor"],
            ["Offshore AH — US Daytime (CHD-TS1)", "4.51", "$45.00", "$203", "Labor"],
            ["Offshore NH — US Overnight (CHD-TS1)", "5.50", "$15.00", "$83", "Labor"],
            ["CrowdStrike — Desktop (AVD)", "8", "$8.50", "$68", "Security"],
            ["Anti-Spam Standard (ASA)", "9", "$6.25", "$56", "Email Security"],
            ["Site Assessment (SA)", "1", "$50.00", "$50", "Assessment"],
            ["Huntress — Desktop (AVMH)", "8", "$5.00", "$40", "Security"],
            ["Secure Internet (SI)", "8", "$4.00", "$32", "Secure Internet"],
            ["Patch Management (PMW)", "8", "$3.00", "$24", "Monitoring"],
            ["DKIM/DMARC", "1", "$20.00", "$20", "Email Security"],
            ["My Remote (MR)", "8", "$2.00", "$16", "Monitoring"],
        ],
        col_aligns=[L, R, R, R, L])

    h2(doc, "3.2 Recurring Licensing")
    make_table(doc,
        ["License / Service", "Annual (2025)", "Notes"],
        [
            ["M365 Business Standard (Non-Profit)", "$594", "Nonprofit pricing"],
            ["TOTAL", "$594", ""],
        ],
        col_aligns=[L, R, L], total_row_indices=[1])

    h2(doc, "3.3 Monthly Run Rate Trend")
    make_table(doc,
        ["Month", "Monthly Total", "Notes"],
        [
            ["Jan 2025", "$658", "Initial run rate"],
            ["Feb 2025", "$639", ""],
            ["Mar 2025", "$783", ""],
            ["Apr 2025", "$1,330", "Labor allotment increased"],
            ["May 2025", "$1,092", ""],
            ["Jun 2025", "$1,057", ""],
            ["Jul 2025", "$958", "Credit adjustments"],
            ["Aug 2025", "$1,039", ""],
            ["Sep 2025", "$1,017", ""],
            ["Oct 2025", "$1,017", ""],
            ["Nov 2025", "$1,004", "NH rate reduced to $15/hr"],
            ["Dec 2025", "$1,004", "Stabilized"],
        ],
        col_aligns=[L, R, L], highlight_row_indices=[3, 10])

    note(doc,
        "The run rate stepped up significantly in April (from ~$660 to ~$1,330) when labor "
        "allotments were increased. It then stabilized around $1,000-1,050/month. The NH "
        "rate reduction in November ($30 to $15/hr) was offset by increased contracted hours.")

    h2(doc, "3.4 Infrastructure Summary")
    make_table(doc,
        ["Category", "Count", "Detail"],
        [
            ["Desktops", "8", "CrowdStrike, Huntress, Malwarebytes, RMM, Patch Mgmt"],
            ["Mailboxes (Anti-Spam)", "9", "Anti-Spam Standard protection"],
            ["Domains", "1", "DKIM/DMARC email authentication"],
            ["Servers", "0", "No servers under management"],
            ["Cloud Hosting", "None", "No cloud VMs or storage"],
            ["VoIP", "None", "No telephony services"],
        ],
        col_aligns=[L, R, L])

    doc.add_page_break()

    # ── Section 4 ──
    h1(doc, "Section 4: Projects & One-Time Spend")
    body(doc,
        "HHOC had zero one-time or project spend in 2025 and zero labor overage. "
        "All work was handled within the contracted monthly allotments.")
    body(doc,
        "This clean billing profile reflects a stable, well-sized engagement. However, "
        "as the desktop fleet ages, budget $2,000-4,000 for potential hardware refresh "
        "or replacement needs in 2026.")

    doc.add_page_break()

    # ── Section 5 ──
    h1(doc, "Section 5: Ticket Categorization")
    body(doc,
        "All 169 unique tickets were categorized by analyzing title and notes fields:")
    make_table(doc,
        ["Category", "Type", "Tickets", "Hours", "Avg Hrs"],
        [
            ["Email & M365", "Reactive", "70", "69.2", "0.99"],
            ["Patch Management", "Proactive", "17", "7.2", "0.42"],
            ["RMM & Agent Management", "Proactive", "13", "8.5", "0.65"],
            ["Security & Endpoint", "Proactive", "13", "10.9", "0.84"],
            ["General IT Support", "Reactive", "10", "11.2", "1.12"],
            ["Monitoring & Alerts", "Proactive", "10", "7.4", "0.74"],
            ["Backup & DR", "Proactive", "7", "5.0", "0.72"],
            ["Software Installation & Updates", "Reactive", "7", "4.1", "0.59"],
            ["Workstation & Hardware", "Reactive", "4", "4.4", "1.09"],
            ["File & Permissions", "Reactive", "3", "9.4", "3.14"],
            ["Firewall & Network", "Reactive", "3", "2.0", "0.68"],
            ["Domain & SSL", "Reactive", "2", "3.2", "1.62"],
            ["Password & Account Mgmt", "Reactive", "2", "1.0", "0.52"],
            ["Phone / VoIP", "Reactive", "2", "0.6", "0.30"],
            ["User Onboarding/Offboarding", "Reactive", "2", "1.6", "0.79"],
            ["Printing & Scanning", "Reactive", "1", "1.3", "1.33"],
            ["Other (Server, App, Dev)", "Reactive", "3", "1.8", "0.60"],
            ["TOTAL", "", "169", "149.1", "0.88"],
        ],
        col_aligns=[L, L, R, R, R],
        total_row_indices=[17], highlight_row_indices=[0])

    body(doc,
        "Proactive work accounts for 60 tickets (35.5%) and reactive for 109 tickets "
        "(64.5%). The reactive tilt is driven by Email & M365 dominance — 70 tickets "
        "consuming 69.2 hours (46.4% of all support hours).")

    sh(doc, "Key Insights")
    bullet(doc, " 70 tickets (41.4%) and 69.2 hours (46.4%). This includes M365 admin, "
           "SharePoint, Teams, Outlook configuration, and domain work. Average 0.99 hrs/ticket "
           "indicates moderate complexity per issue.",
           bold_prefix="Email & M365 dominance:")
    bullet(doc, " Only 3 tickets but 9.4 hours (3.14 hrs/ticket avg) — the highest "
           "per-ticket effort of any category. Suggests complex permission structures.",
           bold_prefix="File & Permissions complexity:")
    bullet(doc, " Patch Management (17), RMM (13), Security (13), Monitoring (10), "
           "Backup (7) collectively form a solid proactive base of 60 tickets.",
           bold_prefix="Proactive maintenance base:")

    doc.add_page_break()

    # ── Section 6 ──
    h1(doc, "Section 6: Service Metrics & Trends")
    sh(doc, "Monthly Volume Table")
    make_table(doc,
        ["Month", "Tickets", "NH", "AH", "Total", "Hrs/Tkt"],
        [
            ["Jan", "16", "9.4", "3.2", "12.6", "0.79"],
            ["Feb", "16", "5.2", "2.2", "7.4", "0.46"],
            ["Mar", "15", "3.5", "3.1", "6.6", "0.44"],
            ["Apr", "8", "3.1", "5.3", "8.4", "1.06"],
            ["May", "5", "6.3", "2.8", "9.1", "1.82"],
            ["Jun", "20", "10.3", "5.3", "15.6", "0.78"],
            ["Jul", "11", "6.9", "4.1", "11.0", "1.00"],
            ["Aug", "12", "9.2", "2.3", "11.5", "0.96"],
            ["Sep", "16", "10.3", "3.2", "13.5", "0.84"],
            ["Oct", "20", "10.7", "8.3", "19.0", "0.95"],
            ["Nov", "14", "13.2", "6.5", "19.7", "1.41"],
            ["Dec", "16", "10.2", "4.5", "14.7", "0.92"],
        ],
        col_aligns=[L, R, R, R, R, R],
        highlight_row_indices=[4, 9, 10])

    sh(doc, "Notable Trends")
    bullet(doc, " Q1 averaged 15.7 tickets/month, Q2 dipped (11/month), "
           "then Q3-Q4 rebounded to 14.6/month. May was the lightest month (5 tickets).",
           bold_prefix="Seasonal pattern:")
    bullet(doc, " H2 averaged 15.6 hours/month vs H1's 9.9 hours — a 58% increase. "
           "This reflects both more tickets and more complex work in the second half.",
           bold_prefix="H2 volume increase:")
    bullet(doc, " The contract stepped up in April (from ~$660 to ~$1,330/mo), likely "
           "in response to Q1 utilization data. It then stabilized around $1,000/mo. "
           "This confirms a 3-month adjustment cycle.",
           bold_prefix="Contract cycle:")
    bullet(doc, " NH rate changed from $30 to $15/hr in November. This was offset by "
           "increased NH hours (5.5 vs ~2.5), keeping the NH cost stable while providing "
           "more monitoring hours.",
           bold_prefix="Rate optimization:")

    sh(doc, "Technician Utilization")
    make_table(doc,
        ["Technician", "Hours", "Role"],
        [
            ["Sunny Sark", "38.8", "US Tech Support (primary)"],
            ["Parveen Biswal", "25.6", "Offshore"],
            ["Sanjeev Kumar", "18.3", "Offshore"],
            ["Gurdeep Kumar", "12.2", "Offshore"],
            ["Surinder Kumar", "10.1", "Offshore"],
            ["Rahul Uniyal", "7.9", "Offshore"],
            ["Abhishek Kumar Poddar", "7.6", "Offshore"],
            ["Other (11 techs)", "28.6", "Various"],
        ],
        col_aligns=[L, R, L])

    sh(doc, "Recommendations")
    bullet(doc, " Email & M365 consumes 46% of all hours. Evaluate recurring tasks "
           "for runbooks, automation, or scheduled admin sessions.",
           bold_prefix="1. Optimize M365 workflows:")
    bullet(doc, " At 35.5% proactive, there is room to grow. Target 45%+ through "
           "increased automated patching and monitoring coverage.",
           bold_prefix="2. Increase proactive ratio:")
    bullet(doc, " File & Permissions averaged 3.14 hrs/ticket. Review "
           "permission structures to reduce resolution complexity.",
           bold_prefix="3. Simplify permission structures:")
    bullet(doc, " Reconcile service inventory quarterly to keep billing aligned "
           "with actual deployments.",
           bold_prefix="4. Quarterly count reconciliation:")

    doc.add_page_break()

    # ── Section 7 ──
    h1(doc, "Section 7: Executive Summary")
    body(doc,
        "HHOC spent $12,192 on IT services in 2025, covering a lean managed environment "
        "of 8 desktops, 9 mailboxes, and 1 domain. The engagement is 100% remote.")

    make_table(doc,
        ["Metric", "Value"],
        [
            ["Total Annual Spend", "$12,192"],
            ["Monthly Run Rate (Dec)", "$1,004"],
            ["Per-User Annual Cost", "~$1,355 (9 users)"],
            ["Blended Labor Rate", "$48.42/hr"],
            ["Tickets Resolved", "169"],
            ["Proactive/Reactive Split", "35.5% / 64.5%"],
            ["Primary US Resource", "Sunny Sark (38.8 hrs)"],
            ["Top Category", "Email & M365 (41.4% of tickets)"],
        ],
        col_aligns=[L, R])

    sh(doc, "Key Findings")
    bullet(doc, " Clean billing — zero overage and zero one-time spend.")
    bullet(doc, " Email & M365 dominates the workload (70 tickets, 69.2 hours).")
    bullet(doc, " Labor contract right-sized quarterly with a major step-up in April.")
    bullet(doc, " NH rate optimized from $30 to $15/hr in November.")
    bullet(doc, " 100% remote, 7-day coverage with 18 technicians available.")

    doc.add_page_break()

    # ── Section 8 ──
    h1(doc, "Section 8: 2026 Budget Projections")
    body(doc,
        "Projections use the December 2025 run rate as baseline. HHOC operates on a "
        "3-month labor contract cycle with quarterly adjustments.")

    make_table(doc,
        ["Category", "2025 Actual", "2026 Projected", "Change"],
        [
            ["Monthly Contract", "$11,598", "$12,050", "+$452"],
            ["  Labor", "$8,231", "$8,376", "+$145"],
            ["  Managed Services", "$3,367", "$3,674", "+$307"],
            ["Recurring (M365)", "$594", "$594", "$0"],
            ["One-Time / Hardware Budget", "$0", "$1,500", "+$1,500"],
            ["PROJECTED TOTAL", "$12,192", "$14,144", "+$1,952"],
        ],
        col_aligns=[L, R, R, R],
        total_row_indices=[5], highlight_row_indices=[4])

    sh(doc, "Key Assumptions")
    bullet(doc, " Labor stays at December 2025 rate ($698/mo = $8,376/yr). "
           "Quarterly adjustments may move this up or down.",
           bold_prefix="Labor:")
    bullet(doc, " Projected at current December rate ($306/mo). Device-based "
           "services adjust monthly with actual counts.",
           bold_prefix="Managed Services:")
    bullet(doc, " M365 Non-Profit carried forward at current rate.",
           bold_prefix="Recurring:")
    bullet(doc, " $1,500 budgeted for potential desktop replacement or "
           "hardware refresh as the 8-device fleet ages.",
           bold_prefix="Hardware Budget:")

    note(doc,
        "The projected increase of $1,952 (+16%) is driven entirely by the recommended "
        "hardware reserve budget. Core monthly costs are expected to remain stable. "
        "If no hardware purchases are needed, 2026 total would be approximately $12,644 (+3.7%).")

    doc.save(OUT)
    print(f"Saved: {OUT}")

build()
