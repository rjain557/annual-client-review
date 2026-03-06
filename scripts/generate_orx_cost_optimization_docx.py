"""
Generate ORX Cost Optimization Analysis — Professional Word Document
Technijian Inc. → OrthoXpress (ORX)
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Color palette ──────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
MED_BLUE = RGBColor(0x2E, 0x75, 0xB6)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x27, 0xAE, 0x60)

DARK_BLUE_HEX = "1F4E79"
ALT_ROW_HEX = "F2F7FB"
WHITE_HEX = "FFFFFF"
YELLOW_HEX = "FFF3CD"
RED_LIGHT_HEX = "FADBD8"
GREEN_LIGHT_HEX = "D5F5E3"

OUT_DIR = r"c:\vscode\annual-client-review\annual-client-review\clients\orx\2025"
OUT_FILE = os.path.join(OUT_DIR, "ORX - Cost Optimization Analysis.docx")


# ── Helpers ────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=Pt(8.5), color=BLACK,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, font_name="Calibri"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = font_name


def add_header_row(table, texts, col_aligns=None):
    row = table.rows[0]
    for i, txt in enumerate(texts):
        align = (col_aligns[i] if col_aligns else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[i], txt, bold=True, color=WHITE, alignment=align)
        set_cell_shading(row.cells[i], DARK_BLUE_HEX)


def add_data_row(table, texts, row_idx, col_aligns=None, bold=False,
                 is_total=False, is_highlight=False, is_red=False, is_green=False):
    row = table.rows[row_idx]
    if is_total:
        bg = DARK_BLUE_HEX
    elif is_red:
        bg = RED_LIGHT_HEX
    elif is_green:
        bg = GREEN_LIGHT_HEX
    elif is_highlight:
        bg = YELLOW_HEX
    else:
        bg = ALT_ROW_HEX if row_idx % 2 == 0 else WHITE_HEX
    txt_color = WHITE if is_total else BLACK
    for i, txt in enumerate(texts):
        align = (col_aligns[i] if col_aligns else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[i], txt, bold=(bold or is_total),
                      color=txt_color, alignment=align)
        set_cell_shading(row.cells[i], bg)


def make_table(doc, headers, data, col_aligns=None, total_row_indices=None,
               highlight_row_indices=None, col_widths=None, red_row_indices=None,
               green_row_indices=None):
    total_row_indices = total_row_indices or []
    highlight_row_indices = highlight_row_indices or []
    red_row_indices = red_row_indices or []
    green_row_indices = green_row_indices or []
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(data), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:top w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:left w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:bottom w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '  <w:right w:val="single" w:sz="4" w:color="D0D0D0"/>'
                '</w:tcBorders>')
            tcPr.append(borders)

    add_header_row(tbl, headers, col_aligns)
    for idx, row_data in enumerate(data):
        r = idx + 1
        is_total = idx in total_row_indices
        is_hl = idx in highlight_row_indices
        is_red = idx in red_row_indices
        is_green = idx in green_row_indices
        add_data_row(tbl, row_data, r, col_aligns, bold=is_total,
                     is_total=is_total, is_highlight=is_hl,
                     is_red=is_red, is_green=is_green)

    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                if w:
                    row.cells[i].width = w
    return tbl


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(16)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(13)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(11)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GRAY
    r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(4)
    return p


def add_page_break(doc):
    doc.add_page_break()


# ── Build document ─────────────────────────────────────────────
def build():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = BLACK

    # ═══════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TECHNIJIAN INC.")
    r.font.size = Pt(12)
    r.font.color.rgb = MED_BLUE
    r.font.bold = True
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Cost Optimization Analysis")
    r.font.size = Pt(28)
    r.font.color.rgb = DARK_BLUE
    r.font.bold = True
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("IT Spend Reduction Recommendations")
    r.font.size = Pt(16)
    r.font.color.rgb = MED_BLUE
    r.font.name = "Calibri"

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared for: OrthoXpress (ORX)")
    r.font.size = Pt(14)
    r.font.color.rgb = DARK_BLUE
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("March 5, 2026")
    r.font.size = Pt(12)
    r.font.color.rgb = GRAY
    r.font.name = "Calibri"

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONFIDENTIAL")
    r.font.size = Pt(10)
    r.font.color.rgb = RED
    r.font.bold = True

    # Key parameters box
    for _ in range(2):
        doc.add_paragraph()

    make_table(doc,
        ["Parameter", "Value"],
        [
            ["Current Annual Spend", "$292,988"],
            ["Reduction Target", "35\u201350%  ($146K\u2013$191K)"],
            ["Compliance Framework", "HIPAA Security Rule"],
            ["Hosting", "Technijian Data Center"],
            ["Analysis Period", "January\u2013December 2025"],
            ["Service Inventory", "March 1, 2026"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
        col_widths=[Inches(2.5), Inches(3.5)]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════════
    add_heading1(doc, "1. Executive Summary")

    add_body(doc, "After reviewing the full 2025 annual data and applying HIPAA compliance constraints, "
             "we\u2019ve identified $26K\u2013$52K in actionable savings (9\u201318%) through advisory restructuring, "
             "development billing optimization, licensing right-sizing, and modest infrastructure consolidation.")

    add_body(doc, "Key factors limiting deeper cuts: HIPAA mandates dual-layer endpoint protection, "
             "dev/prod server segmentation, email archiving retention, and layered email security. "
             "The overnight offshore model ($15/hr) is already ORX\u2019s primary cost optimization lever \u2014 "
             "it routes non-urgent work (patching, monitoring, maintenance) to the lowest rate tier. "
             "Cutting overnight hours would increase daytime costs.")

    make_table(doc,
        ["Scenario", "Annual Savings", "New Spend", "% Reduction"],
        [
            ["Phase 1 only (quick wins)", "$10,192", "$282,796", "3%"],
            ["Phase 1 + 2 (realistic target)", "$26K\u2013$52K", "$241K\u2013$267K", "9\u201318%"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER],
        highlight_row_indices=[1]
    )

    # ═══════════════════════════════════════════════════════════
    # CURRENT SPEND BREAKDOWN
    # ═══════════════════════════════════════════════════════════
    add_heading1(doc, "2. Current Spend Breakdown")

    make_table(doc,
        ["Category", "Annual", "% of Total"],
        [
            ["IT Support Labor", "$53,485", "18.3%"],
            ["CTO Advisory", "$36,688", "12.5%"],
            ["Software Development", "$27,659", "9.4%"],
            ["Cloud Hosting (Compute/Storage)", "$54,962", "18.8%"],
            ["Security & Endpoint (CrowdStrike + Huntress)", "$14,080", "4.8%"],
            ["Backup & Archiving", "$13,801", "4.7%"],
            ["Monitoring & RMM", "$9,237", "3.2%"],
            ["M365 & SPLA Licensing", "$30,461", "10.4%"],
            ["Email Security (Anti-Spam + DMARC)", "$6,377", "2.2%"],
            ["Network & Firewall", "$6,120", "2.1%"],
            ["VoIP / Telephony", "$5,919", "2.0%"],
            ["DNS Filtering (Umbrella / Secure Internet)", "$7,580", "2.6%"],
            ["Pen Testing", "$1,981", "0.7%"],
            ["SSL/Domains/Other Licensing", "$4,313", "1.5%"],
            ["Hardware & One-Time", "$24,165", "8.2%"],
            ["TOTAL", "$292,988", "100%"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.CENTER],
        total_row_indices=[15]
    )

    add_note(doc, "Note: Cisco Umbrella and Secure Internet are the same DNS filtering service "
             "billed as two components (per-device agent + platform fee). Combined: $7,580/yr.")

    # ═══════════════════════════════════════════════════════════
    # HIPAA COMPLIANCE CONSTRAINTS
    # ═══════════════════════════════════════════════════════════
    add_heading1(doc, "3. HIPAA Compliance Constraints")

    add_body(doc, "The following HIPAA Security Rule requirements constrain what can be consolidated or eliminated:")

    make_table(doc,
        ["HIPAA Requirement", "Impact on Cost Optimization"],
        [
            ["\u00a7164.312(c) \u2014 Integrity Controls",
             "Dev/prod environment segmentation must be maintained; dev SQL and app servers cannot be merged with production"],
            ["\u00a7164.310(a)(2)(ii) \u2014 Facility Security",
             "Redundant AD domain controllers required for availability of access controls"],
            ["\u00a7164.312(a)(1) \u2014 Access Control",
             "RDP CALs must cover all users accessing ePHI systems"],
            ["\u00a7164.312(e)(1) \u2014 Transmission Security",
             "Layered email security (anti-spam) must remain on all mailboxes handling PHI"],
            ["\u00a7164.312(c)(1) \u2014 Audit Controls",
             "Email archiving: minimum 6-year retention; aggressive reduction not advisable"],
            ["\u00a7164.308(a)(1)(ii)(A) \u2014 Risk Analysis",
             "Dual EDR (CrowdStrike + Huntress) provides defense-in-depth; removing a layer increases residual risk"],
            ["\u00a7164.308(a)(6) \u2014 Contingency Plan",
             "Backup coverage must include all systems processing/storing ePHI"],
            ["\u00a7164.312(d) \u2014 Authentication",
             "Pen testing validates access controls \u2014 must be maintained"],
        ],
        col_widths=[Inches(2.5), Inches(4.0)]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════
    # AREA 1: INFRASTRUCTURE
    # ═══════════════════════════════════════════════════════════
    add_heading1(doc, "4. Area 1: Infrastructure Consolidation")

    add_body(doc, "25 servers (down from 27 in Dec \u2014 VLAN14 and OPSPRB decommissioned), "
             "hosted in Technijian DC.",
             bold_prefix="Current State (March 2026): ")
    add_body(doc, "$54,962/yr ($4,580/mo)", bold_prefix="Current Cost: ")

    add_heading2(doc, "4.1 Compute & Storage Breakdown")

    make_table(doc,
        ["Resource", "Qty", "Rate", "Monthly"],
        [
            ["Production Storage", "7 TB", "$200/TB", "$1,400"],
            ["Cloud VM vCores", "176", "$6.25/core", "$1,100"],
            ["Backup Storage", "18 TB", "$50/TB", "$900"],
            ["Replicated Storage", "7 TB", "$100/TB", "$700"],
            ["Cloud VM Memory", "360 GB", "$0.63/GB", "$227"],
            ["Shared Bandwidth", "1", "$15.00", "$15"],
            ["Total", "", "", "$4,342"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                    WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT],
        total_row_indices=[6]
    )

    add_heading2(doc, "4.2 HIPAA Segmentation \u2014 Required Servers")

    make_table(doc,
        ["Server Group", "Servers", "Reason"],
        [
            ["Production SQL", "SQL-01, SQL-02", "Production ePHI database workloads"],
            ["Dev/Test SQL", "SQL-03", "HIPAA-required dev/prod segmentation"],
            ["Production App", "APP-01 or APP-04", "Production application serving"],
            ["Dev/Test App", "APP-01 or APP-04", "Separate dev environment for compliance"],
            ["Domain Controllers", "AD-03, AD-04", "HIPAA availability \u2014 redundant access controls"],
            ["Production IIS", "IIS1, IIS2", "Web-facing production (HA pair)"],
            ["File Servers", "FS-01, FS-02", "ePHI document storage + redundancy"],
            ["Terminal Servers", "TS-03, TS-04", "User access to applications"],
            ["3CX / VoIP", "3CX-01", "Phone system"],
            ["Production CRM", "OXPLIVE", "OXPLive production"],
            ["Test CRM", "OXPTEST", "Dev/test environment"],
        ],
        col_widths=[Inches(1.6), Inches(1.8), Inches(3.1)]
    )

    add_heading2(doc, "4.3 Consolidation Opportunities")

    add_bullet(doc, " Only 2\u20134 active VDI licenses. Consolidate to 2 active hosts + 1 golden image. "
               "Savings: ~$150/mo ($1,800/yr).",
               bold_prefix="VDI Consolidation (6 \u2192 3):")
    add_bullet(doc, " Verify workload. If legacy or idle, decommission. "
               "Potential savings: $100\u2013$150/mo ($1,200\u2013$1,800/yr).",
               bold_prefix="JRMEDSVR01 Review:")
    add_bullet(doc, " Move older backups to cold tier (HIPAA requires retention, not instant access). "
               "Verify 7 TB replicated needs real-time vs. daily sync. "
               "Potential savings: $200\u2013$400/mo ($2,400\u2013$4,800/yr).",
               bold_prefix="Storage Tier Optimization:")

    add_subheading(doc, "Infrastructure Savings Summary")

    make_table(doc,
        ["Action", "Timeline", "Annual Savings"],
        [
            ["VDI consolidation (6 \u2192 3)", "1\u20133 months", "$1,800"],
            ["JRMEDSVR01 review", "1 month", "$0\u2013$1,800"],
            ["Storage tier optimization", "1\u20133 months", "$2,400\u2013$4,800"],
            ["Total Infrastructure", "", "$4,200\u2013$8,400"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                    WD_ALIGN_PARAGRAPH.RIGHT],
        total_row_indices=[3]
    )

    add_note(doc, "2 servers (VLAN14, OPSPRB) already decommissioned between Dec 2025 and Mar 2026. "
             "March baseline is 25 servers.")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════
    # AREA 2: SECURITY
    # ═══════════════════════════════════════════════════════════
    add_heading1(doc, "5. Area 2: Security Stack")

    add_body(doc, "Dual EDR + DNS filtering + pen testing", bold_prefix="Current State: ")
    add_body(doc, "$14,080 (endpoint) + $7,580 (DNS filtering) + $1,981 (pen testing) = $23,641/yr",
             bold_prefix="Current Cost: ")

    add_heading2(doc, "5.1 Current Security Layers")

    make_table(doc,
        ["Tool", "Endpoints", "Monthly", "Annual", "Function"],
        [
            ["CrowdStrike (Desktop + Server)", "84", "$760", "$9,120", "Automated EDR, threat detection, AV"],
            ["Huntress (Desktop + Server)", "83", "$415", "$4,980", "Human-managed threat hunting, EDR"],
            ["Cisco Umbrella / Secure Internet", "80 + platform", "$632", "$7,580", "DNS security, web filtering"],
            ["Pen Testing", "18 IPs", "$63", "$756", "Continuous pen testing"],
            ["Total", "", "$1,870", "$22,436", ""],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                    WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.LEFT],
        total_row_indices=[4]
    )

    add_note(doc, "Cisco Umbrella and Secure Internet are the same service \u2014 per-device agent ($316/mo) "
             "+ platform subscription ($312/mo). Previously counted as two separate services.")

    add_heading2(doc, "5.2 HIPAA-Compliant Recommendations")

    add_body(doc, " Defense-in-depth is a HIPAA best practice and compliance necessity. "
             "CrowdStrike provides automated EDR and threat intelligence; Huntress adds human-managed "
             "threat hunting that catches what automation misses. For a healthcare org handling ePHI, "
             "dual-layer endpoint protection is non-negotiable. A single healthcare breach averages "
             "$10.9M (IBM 2025). The $14,100/yr is compliance insurance.",
             bold_prefix="CrowdStrike + Huntress: KEEP BOTH.")

    add_body(doc, " The current Umbrella/SI service ($7,580/yr) is being migrated to CloudBrink. "
             "This is a compliance upgrade, not a cost play \u2014 CloudBrink enables taking OXPLive "
             "(ORX\u2019s CRM) off direct internet access for HIPAA compliance. Cost should be comparable.",
             bold_prefix="Cisco Umbrella / Secure Internet \u2192 CloudBrink Migration:")

    add_body(doc, " Required for HIPAA risk analysis (\u00a7164.308(a)(1)(ii)(A)).",
             bold_prefix="Pen Testing ($756/yr): KEEP.")

    add_subheading(doc, "Security Savings Summary")

    make_table(doc,
        ["Action", "Annual Savings"],
        [
            ["CloudBrink migration (Umbrella/SI replacement)", "Cost-neutral (compliance upgrade)"],
            ["Total Security", "$0"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT],
        total_row_indices=[1]
    )

    add_note(doc, "Security is not the right place to find cost savings under HIPAA. "
             "The current $23.6K/yr protects against breach costs that would dwarf the annual IT spend.")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════
    # AREA 3: DEV & ADVISORY
    # ═══════════════════════════════════════════════════════════
    add_heading1(doc, "6. Area 3: Dev & Advisory Restructuring")

    add_body(doc, "CTO and Dev work mixed into the support contract", bold_prefix="Current State: ")
    add_body(doc, "$64,347/yr (22% of total spend)", bold_prefix="Current Cost: ")

    add_heading2(doc, "6.1 2025 Breakdown")

    make_table(doc,
        ["Role", "Contracted", "Overage", "Total", "Hours", "Eff. Rate"],
        [
            ["CTO Advisory", "$30,375", "$6,313", "$36,688", "160", "$229/hr"],
            ["Software Dev", "$14,850", "$12,809", "$27,659", "285", "$97/hr"],
            ["Total", "$45,225", "$19,122", "$64,347", "445", "$145/hr"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT],
        total_row_indices=[2]
    )

    add_heading2(doc, "6.2 CTO Retainer Restructuring: 15 hrs/mo \u2192 8 hrs/mo")

    make_table(doc,
        ["", "Current", "Proposed"],
        [
            ["Monthly Hours", "15", "8"],
            ["Hourly Rate", "$225/hr (volume)", "$250/hr (standard)"],
            ["Monthly Cost", "$3,375", "$2,000"],
            ["Overage Rate", "$250/hr", "$250/hr"],
            ["Annual Contracted", "$40,500", "$24,000"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.RIGHT]
    )

    add_body(doc, "Reducing from 15 to 8 hours eliminates the volume discount \u2014 "
             "the rate goes from $225/hr back to the standard $250/hr. "
             "Savings come from fewer contracted hours, not from rate reduction.",
             bold_prefix="Important: ")

    add_body(doc, "Post-Q1 2025, actual CTO usage averaged ~10 hrs/mo. "
             "At 8 contracted + ~2 overage/mo: projected annual $30,000 vs. current $36,688.")
    add_body(doc, "$6,700\u2013$16,500/yr", bold_prefix="Realistic savings range: ")

    add_heading2(doc, "6.3 Development Billing Restructuring")

    add_body(doc, "Currently, 374 hours of dev work is buried in support tickets \u2014 making IT support "
             "costs impossible to evaluate independently. The restructuring separates dev into two tracks:")

    add_subheading(doc, "Track 1: Booked Development (Offshore, Tiered Volume Pricing)")

    make_table(doc,
        ["Monthly Hours", "Rate", "Monthly Cost"],
        [
            ["Up to 40 hrs", "$45/hr", "Up to $1,800"],
            ["40\u201380 hrs", "$40/hr", "$1,600\u2013$3,200"],
            ["80\u2013120 hrs", "$35/hr", "$2,800\u2013$4,200"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.RIGHT],
        green_row_indices=[2]
    )

    add_bullet(doc, " set based on the average actual hours from the previous 6-month billing cycle",
               bold_prefix="Rate tier is")
    add_bullet(doc, " \u2014 incentivizing committed volume",
               bold_prefix="As development scales, the per-hour rate decreases")
    add_bullet(doc, " \u2192 starts at $45/hr tier",
               bold_prefix="2025 offshore dev averaged ~23.75 hrs/mo")
    add_bullet(doc, " in the next cycle, rate drops to $40/hr",
               bold_prefix="If volume increases to 40+ hrs/mo")

    add_subheading(doc, "Track 2: Proposal-Based Development (Fixed Price)")

    add_bullet(doc, "Fixed-price engagements scoped per project")
    add_bullet(doc, " to cover unforeseen hours and project risk",
               bold_prefix="Billed at $150/hr effective rate")
    add_bullet(doc, "Appropriate for defined deliverables with clear acceptance criteria")

    add_subheading(doc, "Development Project Management")

    add_bullet(doc, " for each phase (GSD format)",
               bold_prefix="Each project gets a project ticket with child tickets")
    add_bullet(doc, " with week-to-date (WTD) and month-to-date (MTD) hours worked",
               bold_prefix="End-of-week (EOW) status emails")
    add_bullet(doc, "Full visibility into dev spend separate from IT support")

    add_subheading(doc, "Development Cost Projection")

    make_table(doc,
        ["Scenario", "Monthly", "Annual", "vs. 2025 ($27,659)"],
        [
            ["Current volume (24 hrs/mo @ $45)", "$1,080", "$12,960", "\u2212$14,699 (\u221253%)"],
            ["Growth (50 hrs/mo @ $40 blended)", "$2,000", "$24,000", "\u2212$3,659 (\u221213%)"],
            ["High volume (90 hrs/mo @ $35 blended)", "$3,150", "$37,800", "+$10,141 (+37%)"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT],
        green_row_indices=[0]
    )

    add_note(doc, "2025 dev included $12,809 in overage. The tiered model eliminates surprise overage "
             "by setting rates based on actual usage patterns from the prior cycle.")

    add_subheading(doc, "Dev/Advisory Savings Summary")

    make_table(doc,
        ["Action", "Annual Savings"],
        [
            ["CTO retainer reduction (15 \u2192 8 hrs, rate to $250)", "$6,700\u2013$16,500"],
            ["Dev restructuring (tiered + visibility)", "$3,600\u2013$14,700"],
            ["Total Dev/Advisory", "$10,300\u2013$31,200"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT],
        total_row_indices=[2]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════
    # AREA 4: LICENSING & BACKUP
    # ═══════════════════════════════════════════════════════════
    add_heading1(doc, "7. Area 4: Licensing & Backup Optimization")

    add_body(doc, "$34,774 recurring + $13,801 backup = $48,575/yr", bold_prefix="Current State: ")
    add_body(doc, "M365 right-sizing, SPLA reduction, backup optimization within HIPAA retention rules",
             bold_prefix="Focus: ")

    add_heading2(doc, "7.1 M365 License Audit")

    make_table(doc,
        ["License", "Seats", "In Use", "Unused", "Monthly Waste"],
        [
            ["M365 Business Standard", "69", "64", "5", "$75/mo"],
            ["M365 Business Basic", "78", "72", "6", "$43/mo"],
            ["Power BI Pro", "18", "16", "2", "$34/mo"],
            ["Total Unused", "", "", "13", "$152/mo ($1,824/yr)"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                    WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                    WD_ALIGN_PARAGRAPH.RIGHT],
        highlight_row_indices=[3]
    )

    add_body(doc, "Reclaim 13 unused licenses \u2192 $1,824/yr savings", bold_prefix="Quick win: ")

    add_heading2(doc, "7.2 M365 Tier Optimization")

    add_body(doc, "Not all 64 Standard users may need desktop Office apps. "
             "Each downgrade from Standard ($15/mo) to Basic ($7.20/mo) saves $7.80/mo.")
    add_body(doc, "$1,404/yr savings", bold_prefix="If 15 users can downgrade: ")

    add_heading2(doc, "7.3 SPLA & Server Licensing")

    make_table(doc,
        ["License", "Annual", "Tied To"],
        [
            ["RDP/User CAL", "$5,520", "50 CALs for terminal server access"],
            ["Server Std 2-Core", "$5,418", "86 cores (server licensing)"],
            ["SQL Server Standard", "$222", "1 license"],
            ["Total", "$11,160", ""],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.LEFT],
        total_row_indices=[3]
    )

    add_bullet(doc, " Under HIPAA, all users accessing ePHI systems must have proper licensed access. "
               "VDI 6\u21923 reduces CALs 50\u219245. Savings: $552/yr.",
               bold_prefix="RDP CALs:")
    add_bullet(doc, " With 25 servers and VDI consolidation to 23, cores 86\u219276. Savings: $525/yr.",
               bold_prefix="Server cores:")

    add_heading2(doc, "7.4 Copilot Evaluation")

    add_bullet(doc, "1 seat at $360/mo ($4,320/yr) \u2014 started Sep 2025")
    add_bullet(doc, " Confirm Microsoft Copilot is covered under your existing BAA.",
               bold_prefix="HIPAA concern:")
    add_bullet(doc, " Is this single seat delivering $360/mo in productivity?",
               bold_prefix="ROI question:")
    add_body(doc, "If BAA-covered and ROI-justified, keep. Otherwise cancel. "
             "Potential savings: $4,320/yr.", bold_prefix="Recommendation: ")

    add_heading2(doc, "7.5 Backup & Archiving (HIPAA-Constrained)")

    make_table(doc,
        ["Service", "Count", "Rate", "Monthly", "Annual"],
        [
            ["Veeam 365 Backup", "199", "$2.50", "$498", "$5,970"],
            ["Image Backup (Servers)", "27", "$15.00", "$405", "$4,860"],
            ["Email Archiving", "142", "$2.50", "$355", "$4,260"],
            ["Total", "", "", "$1,258", "$15,090"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                    WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.RIGHT],
        total_row_indices=[3]
    )

    add_body(doc, "HIPAA retention rules limit aggressive cuts:",
             bold_prefix="")
    add_bullet(doc, " HIPAA requires 6-year retention. Conservative 142\u2192120 "
               "(non-PHI mailboxes only). Savings: $660/yr.",
               bold_prefix="Email Archiving:")
    add_bullet(doc, " Audit shared/inactive mailboxes without PHI. "
               "Conservative 199\u2192175. Savings: $720/yr.",
               bold_prefix="Veeam 365:")
    add_bullet(doc, " Tracks server count. 27\u219223 with decommissioning. Savings: $720/yr.",
               bold_prefix="Image Backup:")

    add_heading2(doc, "7.6 Other Licensing")

    make_table(doc,
        ["Item", "Annual", "Recommendation"],
        [
            ["Sophos 1C-4G (Irvine)", "$1,320", "Keep \u2014 perimeter security required"],
            ["Edge Appliance (Irvine)", "$1,200", "Review if SD-WAN covers this function"],
            ["SSL Certificates", "$843", "Consolidate: wildcard covers most subdomains"],
            ["Domains", "$177", "Keep \u2014 minimal cost"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.LEFT]
    )

    add_body(doc, "Wildcard cert (*.orthoxpress.com) covers most subdomains, "
             "but 7+ individual certs are also active. Savings: ~$400/yr.",
             bold_prefix="SSL Consolidation: ")

    add_subheading(doc, "Licensing & Backup Savings Summary")

    make_table(doc,
        ["Action", "Annual Savings"],
        [
            ["Reclaim unused M365 licenses", "$1,824"],
            ["Downgrade 15 Standard \u2192 Basic", "$1,404"],
            ["Cancel Copilot (if no ROI)", "$4,320"],
            ["Reduce RDP CALs (50 \u2192 45)", "$552"],
            ["Reduce Server cores (86 \u2192 76)", "$525"],
            ["Email archiving audit (142 \u2192 120)", "$660"],
            ["Veeam 365 audit (199 \u2192 175)", "$720"],
            ["Image backup (27 \u2192 23)", "$720"],
            ["SSL consolidation", "$400"],
            ["Total Licensing & Backup", "$11,125"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT],
        total_row_indices=[9]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════
    # AREA 5: ADDITIONAL LEVERS
    # ═══════════════════════════════════════════════════════════
    add_heading1(doc, "8. Area 5: Additional Optimization Levers")

    add_heading2(doc, "8.1 VoIP Audit")

    add_body(doc, "80 DID phone numbers at $3.60/mo = $3,456/yr. "
             "Audit for unused DIDs \u2014 if 20 are inactive: savings $864/yr.")

    add_heading2(doc, "8.2 Overnight Offshore Model \u2014 Already Optimized")

    add_body(doc, "The current overnight model is NOT a cost to cut \u2014 "
             "it IS ORX\u2019s primary cost optimization.",
             bold_prefix="Key insight: ")

    make_table(doc,
        ["Shift", "Rate", "2025 Hours", "Annual Cost", "Work Type"],
        [
            ["US Daytime (Offshore AH)", "$30/hr", "840", "$25,200",
             "Same-day support, user-facing"],
            ["US Overnight (Offshore NH)", "$15/hr", "1,263", "$18,945",
             "Patch mgmt, monitoring, maintenance"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.LEFT],
        green_row_indices=[1]
    )

    add_body(doc, "The overnight shift deliberately routes non-urgent operational work \u2014 "
             "patch management failures, monitoring alerts, scheduled maintenance \u2014 "
             "to the $15/hr rate. This is 60% of all offshore hours at half the daytime rate. "
             "Reducing overnight hours would push this work to the $30/hr daytime shift, "
             "increasing costs by up to $18,945/yr.")

    add_heading2(doc, "8.3 Anti-Spam / Email Security \u2014 HIPAA-Required")

    add_body(doc, "136 mailboxes at $4.25/mo = $6,936/yr. Email is the #1 PHI exposure vector. "
             "Third-party anti-spam provides a layer beyond M365\u2019s built-in EOP. "
             "Under HIPAA\u2019s defense-in-depth posture, this must remain on all mailboxes. No savings.")

    add_subheading(doc, "Additional Savings Summary")

    make_table(doc,
        ["Action", "Annual Savings"],
        [
            ["VoIP DID audit", "$864"],
            ["Total Additional", "$864"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT],
        total_row_indices=[1]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════
    # CONSOLIDATED SAVINGS
    # ═══════════════════════════════════════════════════════════
    add_heading1(doc, "9. Consolidated Savings Summary")

    add_heading2(doc, "9.1 Phase 1: Quick Wins (0\u20133 months)")

    make_table(doc,
        ["Action", "Annual Savings"],
        [
            ["Reclaim 13 unused M365 licenses", "$1,824"],
            ["Downgrade 15 Standard \u2192 Basic", "$1,404"],
            ["Cancel Copilot (if no ROI)", "$4,320"],
            ["Email archiving audit (142 \u2192 120)", "$660"],
            ["Veeam 365 audit (199 \u2192 175)", "$720"],
            ["SSL consolidation", "$400"],
            ["VoIP DID audit", "$864"],
            ["Phase 1 Total", "$10,192"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT],
        total_row_indices=[7]
    )

    add_heading2(doc, "9.2 Phase 2: Contract & Infrastructure (3\u20136 months)")

    make_table(doc,
        ["Action", "Annual Savings"],
        [
            ["CTO retainer restructuring (15 \u2192 8 hrs @ $250/hr)", "$6,700\u2013$16,500"],
            ["Dev billing restructuring (tiered + project-based)", "$3,600\u2013$14,700"],
            ["VDI consolidation (6 \u2192 3)", "$1,800"],
            ["JRMEDSVR01 review", "$0\u2013$1,800"],
            ["Storage tier optimization", "$2,400\u2013$4,800"],
            ["Reduce RDP CALs & Server cores", "$1,077"],
            ["Image backup (27 \u2192 23)", "$720"],
            ["Phase 2 Total", "$16,297\u2013$41,397"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT],
        total_row_indices=[7]
    )

    add_heading2(doc, "9.3 Additional Ongoing Optimization (6+ months)")

    make_table(doc,
        ["Action", "Annual Savings"],
        [
            ["Edge Appliance review (if SD-WAN covers)", "$1,200"],
            ["Dev volume scaling into lower rate tiers", "Variable"],
            ["Quarterly license audits (prevent waste)", "Variable"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════
    # OVERALL PROJECTION
    # ═══════════════════════════════════════════════════════════
    add_heading1(doc, "10. Overall Projection")

    make_table(doc,
        ["Scenario", "Annual Savings", "New Spend", "% Reduction"],
        [
            ["Phase 1 only (quick wins)", "$10,192", "$282,796", "3%"],
            ["Phase 1 + 2 (realistic target)", "$26K\u2013$52K", "$241K\u2013$267K", "9\u201318%"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER],
        highlight_row_indices=[1]
    )

    add_heading2(doc, "10.1 Why 35\u201350% Is Extremely Difficult Under HIPAA")

    add_body(doc, "The gap between the 35\u201350% target and the realistic 9\u201318% comes from "
             "HIPAA-mandated costs that cannot be reduced:")

    make_table(doc,
        ["HIPAA-Protected Cost", "Annual", "Why It Can\u2019t Be Cut"],
        [
            ["Dual EDR (CrowdStrike + Huntress)", "$14,100", "Defense-in-depth for ePHI endpoints"],
            ["DNS Filtering (\u2192 CloudBrink)", "~$7,580", "HIPAA transmission security + CRM isolation"],
            ["Dev/test server infrastructure", "~$8,000", "Dev/prod segmentation required"],
            ["AD redundancy (2 controllers)", "~$3,600", "Access control availability"],
            ["Email archiving (core 120 users)", "$3,600", "6-year retention requirement"],
            ["Anti-spam (all mailboxes)", "$6,936", "PHI transmission security"],
            ["Pen testing", "$756", "Risk analysis requirement"],
            ["Overnight offshore model", "$18,945", "Already the lowest-cost lever (cutting it increases costs)"],
            ["Total HIPAA floor", "~$63,500", ""],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.LEFT],
        total_row_indices=[8],
        red_row_indices=[0, 1, 2, 3, 4, 5, 6, 7]
    )

    add_body(doc, "Additionally, core managed services (monitoring, RMM, backup, patching) "
             "at ~$32K/yr are operational necessities for a 136-user environment.")

    add_heading2(doc, "10.2 Path to Further Savings Beyond Phase 2")

    add_bullet(doc, " \u2014 As hours increase past 40/mo, "
               "rate drops from $45 to $40/hr; past 80/mo it drops to $35/hr.",
               bold_prefix="Dev volume scaling")
    add_bullet(doc, " \u2014 As servers age out and workloads evolve, "
               "further consolidation opportunities may emerge.",
               bold_prefix="Continued infrastructure optimization")
    add_bullet(doc, " \u2014 Reviewing seats every quarter rather than "
               "annually prevents waste accumulation.",
               bold_prefix="Quarterly license audits")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════
    # RECOMMENDED APPROACH
    # ═══════════════════════════════════════════════════════════
    add_heading1(doc, "11. Recommended Approach")

    add_body(doc, "15\u201318% reduction (~$245K) within 6 months.",
             bold_prefix="Target: ")

    make_table(doc,
        ["Month", "Action", "Impact"],
        [
            ["Month 1", "Execute Phase 1 quick wins", "\u2212$10K/yr"],
            ["Months 2\u20133", "Restructure CTO retainer (8 hrs @ $250)", "\u2212$7K\u2013$17K/yr"],
            ["Months 2\u20133", "Implement dev billing tiers + GSD project tracking", "\u2212$4K\u2013$15K/yr"],
            ["Months 3\u20136", "VDI consolidation + storage optimization", "\u2212$4K\u2013$8K/yr"],
            ["Ongoing", "Quarterly license audits", "Prevent waste"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                    WD_ALIGN_PARAGRAPH.RIGHT]
    )

    add_heading2(doc, "11.1 Cost Optimization vs. Risk Trade-offs")

    make_table(doc,
        ["Savings Tier", "Annual Spend", "Risk Level", "HIPAA Impact"],
        [
            ["Conservative (Phase 1)", "$283K", "None", "Fully compliant"],
            ["Moderate (Phase 1+2)", "$241K\u2013$266K", "Low", "Compliant, advisory coverage reduced"],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        green_row_indices=[0, 1]
    )

    # ── Footer note ──
    for _ in range(2):
        doc.add_paragraph()

    add_note(doc, "Prepared for discussion \u2014 all figures based on 2025 actual spend data, "
             "March 2026 service inventory (25 servers), and HIPAA Security Rule requirements "
             "(45 CFR \u00a7164.308\u2013312). Technijian DC hosting assumed throughout.")

    add_note(doc, "CONFIDENTIAL \u2014 Technijian Inc. \u00a9 2026")

    # ── Save ──
    os.makedirs(OUT_DIR, exist_ok=True)
    doc.save(OUT_FILE)
    print(f"Saved: {OUT_FILE}")


if __name__ == "__main__":
    build()
