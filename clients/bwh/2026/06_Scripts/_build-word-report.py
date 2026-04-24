"""Build Technijian-branded Word report for BWH hours/savings review.

Follows technijian-report SKILL.md brand rules:
- Core Blue #006DB6, Core Orange #F67D4B, Teal #1EAAC8
- Dark Charcoal #1A1A2E, Brand Grey #59595B, Off White #F8F9FA
- Font: Open Sans
- Logo: technijian-logo-full-color-600x125.png
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = Path(__file__).parent.parent                       # clients/BWH/
LOGO = Path(r"C:\VSCode\tech-branding\tech-branding\assets\logos\png\technijian-logo-full-color-600x125.png")
OUT = HERE / "05_Reports" / "BWH-Hours-Savings-Analysis.docx"

# ---- Brand palette ----
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
TEAL = RGBColor(0x1E, 0xAA, 0xC8)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
OFF_WHITE = "F8F9FA"
LIGHT_GREY = "E9ECEF"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x28, 0xA7, 0x45)
RED = RGBColor(0xCC, 0x00, 0x00)

FONT = "Open Sans"


def shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def remove_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for b in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{b}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tc_pr.append(borders)


def set_cell_border_color(cell, hex_color=LIGHT_GREY, sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for b in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{b}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:color"), hex_color)
        borders.append(e)
    tc_pr.append(borders)


def add_run(paragraph, text, *, bold=False, size=11, color=BRAND_GREY, font=FONT, italic=False):
    run = paragraph.add_run(text)
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    if level == 1:
        size = 22
        clr = color or CORE_BLUE
        bold = True
    elif level == 2:
        size = 16
        clr = color or CORE_BLUE
        bold = True
    else:
        size = 13
        clr = color or DARK_CHARCOAL
        bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, text, bold=bold, size=size, color=clr)
    return p


def add_body(doc, text, *, bold=False, size=11, color=BRAND_GREY, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bold=bold, size=size, color=color)
    return p


def add_bullet(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=11, color=DARK_CHARCOAL)
        add_run(p, text, size=11, color=BRAND_GREY)
    else:
        add_run(p, text, size=11, color=BRAND_GREY)
    return p


def add_color_bar(doc, hex_color, height_pt=4):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade(cell, hex_color)
    remove_borders(cell)
    # empty paragraph but small
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("")
    r.font.size = Pt(height_pt)
    return table


def add_section_header(doc, title, accent_color=CORE_BLUE):
    """Left-bar + title treatment from skill."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    col_bar, col_title = table.columns
    col_bar.width = Emu(60000)
    col_title.width = Emu(5700000)
    bar_cell, title_cell = table.rows[0].cells
    bar_cell.width = Emu(60000)
    title_cell.width = Emu(5700000)
    bar_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shade(bar_cell, "006DB6" if accent_color == CORE_BLUE else "F67D4B")
    remove_borders(bar_cell)
    remove_borders(title_cell)
    p = title_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "  " + title, bold=True, size=14, color=accent_color)
    return table


def set_col_widths(table, widths_in_inches):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_in_inches):
                cell.width = Inches(widths_in_inches[i])


def styled_table(doc, headers, rows, *, col_widths=None, header_fill="006DB6",
                 zebra_fill=OFF_WHITE, money_cols=(), bold_last_row=False):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if col_widths:
        set_col_widths(table, col_widths)
    # header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        shade(cell, header_fill)
        set_cell_border_color(cell, LIGHT_GREY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, h, bold=True, size=10, color=WHITE)
    # data
    for i, row in enumerate(rows):
        is_total = bold_last_row and i == len(rows) - 1
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            if is_total:
                shade(cell, "E9ECEF")
            elif i % 2 == 1:
                shade(cell, zebra_fill)
            set_cell_border_color(cell, LIGHT_GREY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            text = str(val)
            color = DARK_CHARCOAL if is_total else BRAND_GREY
            # Right-align money and numeric columns
            if j in money_cols or (isinstance(val, (int, float))):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, text, size=10, color=color, bold=is_total)
    return table


def add_metric_card_row(doc, cards):
    """cards = [(big_value, label, color), ...]"""
    table = doc.add_table(rows=1, cols=len(cards))
    table.autofit = False
    for i, (value, label, color) in enumerate(cards):
        cell = table.rows[0].cells[i]
        shade(cell, "F8F9FA")
        remove_borders(cell)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(12)
        p1.paragraph_format.space_after = Pt(0)
        add_run(p1, value, bold=True, size=26, color=color)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(12)
        add_run(p2, label, size=10, color=BRAND_GREY)
    return table


def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.font.color.rgb = BRAND_GREY
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Technijian  |  18 Technology Dr., Ste 141, Irvine, CA 92618  |  949.379.8500  |  technijian.com",
            size=9, color=BRAND_GREY)


def add_header_logo(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    if LOGO.exists():
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.7))
    else:
        add_run(p, "TECHNIJIAN", bold=True, size=14, color=CORE_BLUE)


# ==================== BUILD DOC ====================

doc = Document()
set_default_style(doc)

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

add_header_logo(doc)
add_footer(doc)

# ==================== COVER PAGE ====================

add_color_bar(doc, "006DB6", height_pt=6)

# lots of spacing
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
if LOGO.exists():
    run = p.add_run()
    run.add_picture(str(LOGO), width=Inches(3.0))
else:
    add_run(p, "TECHNIJIAN", bold=True, size=32, color=CORE_BLUE)

for _ in range(2):
    doc.add_paragraph()

# Orange divider (short centered bar)
divider_tbl = doc.add_table(rows=1, cols=3)
divider_tbl.autofit = False
for i, cell in enumerate(divider_tbl.rows[0].cells):
    remove_borders(cell)
    if i == 1:
        shade(cell, "F67D4B")
        cell.width = Inches(1.5)
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        r = cp.add_run("")
        r.font.size = Pt(3)
    else:
        cell.width = Inches(2.6)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "BWH Hours Analysis", bold=True, size=32, color=DARK_CHARCOAL)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "& Savings Review", bold=True, size=32, color=DARK_CHARCOAL)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Life-of-Contract Review (May 2023 – April 2026)", size=14, color=BRAND_GREY)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Prepared for: Dave — Brandywine Homes", bold=True, size=13, color=DARK_CHARCOAL)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Prepared by: Technijian  |  April 24, 2026", size=11, color=BRAND_GREY)

for _ in range(5):
    doc.add_paragraph()

add_color_bar(doc, "F67D4B", height_pt=6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "CONFIDENTIAL — FOR CLIENT REVIEW", bold=True, size=10, color=BRAND_GREY)

add_page_break(doc)

# ==================== EXECUTIVE SUMMARY ====================

add_section_header(doc, "Executive Summary", accent_color=CORE_BLUE)
doc.add_paragraph()

add_body(doc,
    "Over the 36-month life of Brandywine Homes' IT Services contract (May 2, 2023 → present), "
    "Technijian delivered 4,050.10 billable hours of support, pulled directly from the client portal. "
    "This review explains what those hours were spent on, and — most importantly — identifies 588 hours "
    "of project-style work (ERP upgrades, VM/server rebuilds, OneDrive migration, Windows 11 refresh, firewall install, "
    "and more) that were absorbed into your monthly support contract rather than quoted as separate proposals.",
)

add_body(doc,
    "Because that project work was delivered under your monthly support model, it was never billed at the contract's "
    "over-contract / project-proposal rate of $150 per hour. If those same deliverables had been quoted as separate "
    "Statements of Work — as is standard practice for server/VM upgrades, ERP migrations, and network buildouts — you "
    "would have received project invoices totaling $88,251.",
)

add_body(doc, "In short: keeping this work inside your monthly support model saved BWH approximately "
              "$88,251 in project-proposal billing over the life of the contract.", bold=True, color=DARK_CHARCOAL)

doc.add_paragraph()
add_metric_card_row(doc, [
    ("4,050", "Total Hours Delivered", CORE_BLUE),
    ("588", "Project-Type Hours", CORE_ORANGE),
    ("$88,251", "Proposal Billing Avoided", GREEN),
])

doc.add_paragraph()

# ==================== CONTEXT: CONTRACT RATE ====================

add_page_break(doc)
add_section_header(doc, "The Savings Mechanism", accent_color=CORE_ORANGE)
doc.add_paragraph()

add_body(doc,
    "Your active contract (Contract ID 4924, signed 2023-05-02) specifies two billing modes:")
add_bullet(doc, "A monthly managed-services allocation covering routine IT support (patching, AV, monitoring, user help-desk, agent management, etc.)",
           bold_prefix="Monthly Service: ")
add_bullet(doc, "$150.00 per hour — the rate that applies to any work delivered outside the monthly allocation, and the rate that would have applied to any separately-scoped project SOW.",
           bold_prefix="Over-Contract / Proposal Rate: ")

doc.add_paragraph()
add_body(doc,
    "In a traditional MSP engagement, certain categories of work — operating system upgrades, server rebuilds, "
    "ERP migrations, OneDrive / SharePoint data migrations, firewall replacements, new-site buildouts — would be "
    "quoted as Statements of Work at the $150/hr project rate, billed separately, and delivered outside the monthly "
    "support pool. Under BWH's contract, those hours were instead delivered by the same India + USA support pods as "
    "routine work, and logged against the monthly support stream. No separate proposal invoice was ever generated.",
)
add_body(doc,
    "The net effect: BWH received $88,251 of project-rate deliverables without receiving $88,251 of project-rate invoices. "
    "The full life-of-contract breakdown, per project, is below.",
    bold=True, color=DARK_CHARCOAL,
)

# ==================== PER-PROJECT SAVINGS TABLE ====================

add_page_break(doc)
add_section_header(doc, "Savings by Project — at $150/hr Proposal Rate", accent_color=CORE_BLUE)
doc.add_paragraph()

project_rows = [
    ("NewStar ERP upgrade & support",            125.21,  63.05, 188.26, 28239.00),
    ("RMM / tooling install on new machines",    103.37,   0.00, 103.37, 15505.50),
    ("Server / VM / ESXi / VMware upgrades",      79.16,   8.39,  87.55, 13132.50),
    ("Windows 11 / PC refresh / laptop deploy",   13.77,  52.18,  65.95,  9892.50),
    ("OneDrive / SharePoint data migration",      60.49,   5.00,  65.49,  9823.50),
    ("Backup / Veeam / Replication projects",     31.39,   2.00,  33.39,  5008.50),
    ("Firewall / VPN / Network buildout",          2.00,  16.83,  18.83,  2824.50),
    ("File server / data migration",               8.50,   5.00,  13.50,  2025.00),
    ("Security / EDR / SSL / MFA rollouts",       11.50,   0.50,  12.00,  1800.00),
]
total_india = sum(r[1] for r in project_rows)
total_usa   = sum(r[2] for r in project_rows)
total_hrs   = sum(r[3] for r in project_rows)
total_cost  = sum(r[4] for r in project_rows)

table_rows = []
for name, ind, usa, tot, cost in project_rows:
    table_rows.append([name,
                       f"{ind:,.2f}",
                       f"{usa:,.2f}",
                       f"{tot:,.2f}",
                       f"${cost:,.2f}"])
table_rows.append(["TOTAL",
                   f"{total_india:,.2f}",
                   f"{total_usa:,.2f}",
                   f"{total_hrs:,.2f}",
                   f"${total_cost:,.2f}"])

styled_table(doc,
    headers=["Project", "India hrs", "USA hrs", "Total hrs", "Proposal cost @ $150/hr"],
    rows=table_rows,
    col_widths=[2.7, 0.85, 0.85, 0.85, 1.35],
    money_cols=(1, 2, 3, 4),
    bold_last_row=True,
)

doc.add_paragraph()
add_body(doc,
    "Every row above represents real, named deliverables found in the portal ticket record. "
    "Full ticket-level detail is available in the attached project-candidate-tickets.csv "
    "(588 rows, every one date-stamped, with requestor, role, and hours).",
    size=10,
)

# ==================== INDIA vs USA HOURS ====================

add_page_break(doc)
add_section_header(doc, "Where the Labor Came From", accent_color=CORE_BLUE)
doc.add_paragraph()

add_body(doc,
    "The contract's $150/hr proposal rate is the same whether the work is delivered from the India or USA pod. "
    "For BWH specifically, the India pod carried the majority of the project workload:")

doc.add_paragraph()
india_pct = total_india / total_hrs * 100
usa_pct = total_usa / total_hrs * 100
add_metric_card_row(doc, [
    (f"{total_india:.1f}", f"India Hours ({india_pct:.0f}%)", CORE_BLUE),
    (f"{total_usa:.1f}", f"USA Hours ({usa_pct:.0f}%)", TEAL),
    (f"${total_india*150:,.0f}", "India proposal value", CORE_ORANGE),
])

doc.add_paragraph()
add_body(doc,
    "A standard proposal SOW would not have discounted the India-delivered portion; the contract rate is flat. "
    "That is why 74% of the avoided-proposal value ($65,309) corresponds to India-delivered project hours that, "
    "in a traditional project-billing model, BWH would have been invoiced for at the same $150 rate as U.S. hours.",
)

# ==================== PROJECT NARRATIVES ====================

add_page_break(doc)
add_section_header(doc, "What Each Project Actually Was", accent_color=CORE_ORANGE)
doc.add_paragraph()

projects = [
    ("NewStar ERP Upgrades — 188.26 hrs — $28,239",
     "Multiple rounds of NewStar ERP upgrade work: the original 'NewStar upgrade project' thread (Nov 2023 → Feb 2024, "
     "18+ hours), the 'Updating NewStar batches to 2023 version' work (Mar–Apr 2024), the 'NewStar upgrade needed' "
     "push in Aug 2024, and continuous NewStar configuration/login/integration support. Under a proposal model, "
     "each upgrade wave would have been a separate SOW."),
    ("RMM / Tooling Install on New Machines — 103.37 hrs — $15,505",
     "Technijian Tools installation, MyRMM/ManageEngine agent provisioning, Passportal agent configuration, SNMP setup, "
     "and Network Detective scans across the BWH fleet. Typically scoped as 'new device onboarding' at ~1.5 hrs/device, "
     "this would normally be quoted per batch."),
    ("Server / VM / ESXi / VMware Upgrades — 87.55 hrs — $13,133",
     "HP server refresh work, multiple ESXi host reboots, virtual-disk and snapshot consolidation on BWH-HQ-ESXI-01, "
     "and various VM rebuild/reconfigure efforts. Traditional MSPs scope VMware host maintenance as a project."),
    ("Windows 11 / PC Refresh / Laptop Deploy — 65.95 hrs — $9,893",
     "Dedicated 'Brandywine Homes Windows 11 Upgrade' initiative, onsite new-PC configuration, preconfigured Dell "
     "computer deployment, plus named-user laptop builds (Nancy Hayden, Chris, Traci, and others)."),
    ("OneDrive / SharePoint Data Migration — 65.49 hrs — $9,824",
     "'Projects folder migration to OneDrive' (16.3 hrs, Sept 2023) and 'Share folder migration to OneDrive' "
     "(14.8 hrs, Sept 2023) were the core of this work — a classic data-migration SOW. Additional ongoing OneDrive "
     "sync and configuration troubleshooting followed."),
    ("Backup / Veeam / Replication Projects — 33.39 hrs — $5,009",
     "QNAP firmware upgrade, Loaner QNAP configuration, Veeam backup and replication rebuild/configuration work, "
     "and Datto server setup tasks."),
    ("Firewall / VPN / Network Buildout — 18.83 hrs — $2,825",
     "New firewall install at BWH-HQ in September 2025, BWH VPN infrastructure updates, and related network reconfiguration."),
    ("File Server / Data Migration — 13.50 hrs — $2,025",
     "Dedicated file-server migration work and Datto workplace configuration."),
    ("Security / EDR / SSL / MFA Rollouts — 12.00 hrs — $1,800",
     "SSL certificate updates and renewals on BWH infrastructure, plus EDR/security tool configuration work."),
]

for title, desc in projects:
    add_body(doc, title, bold=True, size=12, color=DARK_CHARCOAL)
    add_body(doc, desc, size=10.5)

# ==================== THE 4,050 HOURS IN FULL CONTEXT ====================

add_page_break(doc)
add_section_header(doc, "How the Full 4,050 Hours Breaks Down", accent_color=CORE_BLUE)
doc.add_paragraph()

add_body(doc,
    "For complete transparency, below is the full categorization of all 4,050.10 hours delivered. "
    "The 588 project hours represent only 14.5% of total hours — the remaining 85% was routine monthly "
    "support, correctly covered under the existing allocation.")

doc.add_paragraph()

breakdown = [
    ["Routine monthly-support work",           "2,466", "60.9%", "Patching, AV, agent updates, monitoring alerts, user support"],
    ["Project-style work (SOW candidates)",    "588",   "14.5%", "Absorbed — see per-project table"],
    ["Short ad-hoc / uncategorized tickets",   "996",   "24.6%", "Long-tail routine support"],
    ["TOTAL",                                   "4,050", "100%",  ""],
]

styled_table(doc,
    headers=["Work Type", "Hours", "% of Total", "Description"],
    rows=breakdown,
    col_widths=[2.3, 0.8, 0.9, 2.7],
    money_cols=(1, 2),
    bold_last_row=True,
)

# ==================== CONCLUSION ====================

add_page_break(doc)
add_section_header(doc, "Bottom Line for BWH", accent_color=CORE_ORANGE)
doc.add_paragraph()

add_body(doc,
    "The 4,050 hours delivered over 36 months may sound like a large number in isolation, but it breaks down cleanly:")

add_bullet(doc, "An average of ~112 hours of IT work per month — consistent with a ~40-seat environment running full "
                "monitoring, patching, EDR, backup, and end-user support stacks.",
           bold_prefix="Operationally routine: ")
add_bullet(doc, "Delivered and billed entirely through monthly support at no extra project charge — instead of the "
                "$88,251 in separate proposal invoices this work would have generated at the contract's $150/hr rate.",
           bold_prefix="Project-type deliverables: ")
add_bullet(doc, "The structure of BWH's contract kept these upgrades, migrations, and deployments "
                "inside the monthly fee instead of adding them on top.",
           bold_prefix="Direct savings: $88,251 ")

doc.add_paragraph()
add_body(doc, "Going forward, Technijian is glad to formalize which kinds of work fall under monthly support vs. "
              "warrant separate SOWs, so there is no confusion about where each category of work is priced and billed. "
              "We appreciate BWH's partnership and welcome the opportunity to walk through this report together.",
         color=DARK_CHARCOAL)

doc.add_paragraph()
doc.add_paragraph()

# Signature block
add_body(doc, "Prepared by Technijian — Account Management", bold=True, size=11, color=DARK_CHARCOAL)
add_body(doc, "18 Technology Dr., Ste 141, Irvine, CA 92618", size=10)
add_body(doc, "949.379.8500  |  technijian.com", size=10)

doc.save(str(OUT))
print(f"Wrote: {OUT}")
