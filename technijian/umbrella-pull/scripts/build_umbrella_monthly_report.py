"""Build a client-facing Technijian-branded monthly DNS Security report
from the Cisco Umbrella per-client daily snapshots.

Inputs (per client, per target month YYYY-MM):
    clients/<code>/umbrella/YYYY-MM-DD/pull_summary.json
    clients/<code>/umbrella/YYYY-MM-DD/activity_summary.json
    clients/<code>/umbrella/YYYY-MM-DD/top_destinations.json
    clients/<code>/umbrella/YYYY-MM-DD/top_identities.json
    clients/<code>/umbrella/YYYY-MM-DD/blocked_threats.json
    clients/<code>/umbrella/YYYY-MM-DD/roaming_computers.json

Output:
    clients/<code>/umbrella/monthly/<YYYY-MM>/<CODE>-DNS-Security-<YYYY-MM>.docx

Usage:
    python build_umbrella_monthly_report.py --month 2026-02
    python build_umbrella_monthly_report.py --month 2026-02 --only VAF
    python build_umbrella_monthly_report.py --month 2026-03 --skip ORX
    python build_umbrella_monthly_report.py --all
"""
from __future__ import annotations

import argparse
import json
import sys
from collections import defaultdict
from datetime import datetime, timezone
from pathlib import Path

HERE = Path(__file__).resolve().parent
PIPELINE_ROOT = HERE.parent
REPO = PIPELINE_ROOT.parent.parent
CLIENTS_ROOT = REPO / "clients"

sys.path.insert(0, str(HERE))
sys.path.insert(0, str(REPO / "technijian" / "huntress-pull" / "scripts"))
import _brand as brand  # noqa: E402

from docx.shared import Pt  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

MONTH_NAMES = ["", "January", "February", "March", "April", "May", "June",
               "July", "August", "September", "October", "November", "December"]


def month_label(month: str) -> str:
    y, m = month.split("-")
    return f"{MONTH_NAMES[int(m)]} {y}"


def load_json(path: Path, default):
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError:
        return default
    except Exception as e:
        print(f"  WARN: could not parse {path}: {e}")
        return default


def fmt_int(n: int | float) -> str:
    return f"{int(n):,}"


def pct(num: float, den: float) -> str:
    if den == 0:
        return "0.0%"
    return f"{num / den * 100:.1f}%"


# ---------------------------------------------------------------------------
# Monthly data aggregator
# ---------------------------------------------------------------------------

class ClientUmbrellaMonth:
    def __init__(self, code: str, month: str, umbrella_dir: Path):
        self.code = code.upper()
        self.month = month
        self.umbrella_dir = umbrella_dir

        # Collect all daily dirs for the month
        prefix = month  # "2026-02"
        self.daily_dirs: list[Path] = sorted(
            p for p in umbrella_dir.iterdir()
            if p.is_dir() and p.name.startswith(prefix) and len(p.name) == 10
        )

        self._aggregate()

    def _aggregate(self) -> None:
        total_requests = 0
        total_blocked = 0
        total_allowed = 0
        dest_counts: dict[str, int] = defaultdict(int)
        category_counts: dict[str, int] = defaultdict(int)
        identity_counts: dict[str, int] = defaultdict(int)
        errors: list[str] = []
        location_name = self.code

        latest_agents: list[dict] = []
        latest_agents_date = ""
        latest_agents_status: dict[str, int] = {}

        for day_dir in self.daily_dirs:
            summary = load_json(day_dir / "pull_summary.json", {})
            if not summary:
                continue

            if summary.get("Location_Name") and summary["Location_Name"] != self.code:
                location_name = summary["Location_Name"]

            req = summary.get("client_requests_total", 0) or 0
            blk = summary.get("client_blocked_threats", 0) or 0
            total_requests += req
            total_blocked += blk
            total_allowed += max(0, req - blk)

            for e in (summary.get("errors") or []):
                errors.append(f"{day_dir.name}: {e}")

            # Track most recent agent inventory and status (point-in-time, not cumulative)
            if day_dir.name > latest_agents_date:
                agents_raw = load_json(day_dir / "roaming_computers.json", [])
                if agents_raw:
                    latest_agents = agents_raw
                    latest_agents_date = day_dir.name
                status_raw = summary.get("agents_status") or {}
                if status_raw:
                    latest_agents_status = {k.lower(): v for k, v in status_raw.items()}

            # Aggregate top destinations
            for dest in load_json(day_dir / "top_destinations.json", []):
                domain = dest.get("domain", "")
                if domain:
                    dest_counts[domain] += dest.get("count", 0)

            # Aggregate blocked threat categories
            for rec in load_json(day_dir / "blocked_threats.json", []):
                for cat in (rec.get("categories") or []):
                    label = cat.get("label", "")
                    if label and not cat.get("deprecated", False):
                        category_counts[label] += 1

            # Aggregate identity request counts
            for ident in load_json(day_dir / "top_identities.json", []):
                name = (ident.get("identity") or {}).get("label", "")
                if name:
                    identity_counts[name] += ident.get("requests", 0)

        self.location_name = location_name
        self.days_with_data = len(self.daily_dirs)
        self.total_requests = total_requests
        self.total_blocked = total_blocked
        self.total_allowed = total_allowed
        self.errors = errors
        self.latest_agents = latest_agents
        self.latest_agents_date = latest_agents_date
        self.agents_status_latest = latest_agents_status

        # Top 15 blocked destinations
        self.top_destinations = sorted(
            dest_counts.items(), key=lambda kv: kv[1], reverse=True
        )[:15]

        # Top 10 blocked categories
        self.top_categories = sorted(
            category_counts.items(), key=lambda kv: kv[1], reverse=True
        )[:10]

        # Top 10 identities by total requests for the month
        self.top_identities = sorted(
            identity_counts.items(), key=lambda kv: kv[1], reverse=True
        )[:10]

    @property
    def has_data(self) -> bool:
        return self.days_with_data > 0 and self.total_requests > 0

    @property
    def block_rate(self) -> float:
        if self.total_requests == 0:
            return 0.0
        return self.total_blocked / self.total_requests * 100

    def agent_summary(self) -> dict:
        """Aggregate status counts from latest day."""
        status_map: dict[str, int] = defaultdict(int)
        for a in self.latest_agents:
            status = (a.get("status") or "Unknown").title()
            status_map[status] += 1
        return dict(sorted(status_map.items(), key=lambda kv: kv[1], reverse=True))

    def stale_agents(self) -> list[dict]:
        """Agents whose lastSync is stale (no sync in last 30 days relative to latest_agents_date)."""
        ref_str = self.latest_agents_date or datetime.now(timezone.utc).strftime("%Y-%m-%d")
        try:
            ref = datetime.fromisoformat(ref_str)
        except ValueError:
            ref = datetime.now(timezone.utc)

        out = []
        for a in self.latest_agents:
            last_s = a.get("lastSync")
            if not last_s:
                out.append(a)
                continue
            try:
                last = datetime.fromisoformat(last_s.replace("Z", "+00:00"))
                if ref.tzinfo is None:
                    from datetime import timezone as tz
                    ref = ref.replace(tzinfo=tz.utc)
                delta = ref - last
                if delta.days > 30:
                    out.append(a)
            except Exception:
                out.append(a)
        return out


# ---------------------------------------------------------------------------
# Render
# ---------------------------------------------------------------------------

def render_report(data: ClientUmbrellaMonth, out_path: Path) -> dict:
    doc = brand.new_branded_document()

    # ---- COVER ----
    brand.render_cover(
        doc,
        title="DNS Security Report",
        subtitle=data.location_name,
        date_text=month_label(data.month),
        footer_note="Confidential — prepared by Technijian for the named client only.",
    )
    brand.add_page_break(doc)

    # ---- EXECUTIVE SUMMARY ----
    brand.add_section_header(doc, "Executive Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    brand.add_run(
        p,
        f"This report summarizes the Cisco Umbrella DNS security activity "
        f"Technijian managed for {data.location_name} during "
        f"{month_label(data.month)}. Umbrella acts as a protective DNS layer "
        f"for every covered device — every domain lookup is inspected before "
        f"a connection is established. Below is a plain-English summary of "
        f"DNS activity, what was blocked, and where the fleet stands today.",
        size=11,
    )

    # KPI cards
    agents_count = len(data.latest_agents)
    encrypted_count = data.agents_status_latest.get("encrypted", 0)
    block_pct_str = f"{data.block_rate:.1f}%"
    cards = [
        (fmt_int(data.total_requests), "DNS Queries This Month", brand.CORE_BLUE),
        (fmt_int(data.total_blocked), "Queries Blocked", brand.RED if data.total_blocked > 0 else brand.GREEN),
        (block_pct_str, "Block Rate", brand.CORE_ORANGE),
        (str(agents_count), "Protected Endpoints", brand.TEAL),
    ]
    brand.add_metric_card_row(doc, cards)

    # Health callout
    stale = data.stale_agents()
    off_count = data.agents_status_latest.get("off", 0)
    if stale or off_count > 0:
        bits = []
        if off_count > 0:
            bits.append(f"{off_count} agent(s) are in 'off' status")
        if stale:
            bits.append(f"{len(stale)} agent(s) have not synced in 30+ days")
        brand.add_callout_box(
            doc,
            "Attention: " + "; ".join(bits) +
            ". See the Recommendations section for the remediation list.",
            accent_hex=brand.CORE_ORANGE_HEX, bg_hex="FEF3EE",
        )
    else:
        brand.add_callout_box(
            doc,
            f"All {agents_count} protected endpoints are active and syncing "
            "with Cisco Umbrella. No action items this month.",
            accent_hex=brand.GREEN_HEX, bg_hex="EAF7EE",
        )

    brand.add_page_break(doc)

    # ---- DNS ACTIVITY OVERVIEW ----
    brand.add_section_header(doc, "DNS Activity Overview")
    brand.add_body(
        doc,
        f"Across {data.days_with_data} day(s) of data in "
        f"{month_label(data.month)}, Cisco Umbrella processed every DNS "
        f"lookup made by covered devices. The table below shows the overall "
        f"query volume, allowed traffic, and blocked traffic.",
    )

    allowed_total = data.total_requests - data.total_blocked
    doc.add_paragraph()
    brand.styled_table(
        doc,
        ["Metric", "Value"],
        [
            ("Total DNS queries", fmt_int(data.total_requests)),
            ("Queries allowed", fmt_int(max(0, allowed_total))),
            ("Queries blocked", fmt_int(data.total_blocked)),
            ("Block rate", block_pct_str),
            ("Days covered", str(data.days_with_data)),
            ("Average queries per day", fmt_int(data.total_requests // max(1, data.days_with_data))),
            ("Average blocks per day", fmt_int(data.total_blocked // max(1, data.days_with_data))),
        ],
        col_widths=[4.5, 1.5],
    )

    # Top blocked destinations
    if data.top_destinations:
        doc.add_paragraph()
        brand.add_body(doc, "Top blocked destinations", bold=True, size=12, color=brand.DARK_CHARCOAL)
        brand.add_body(
            doc,
            "These are the domains most frequently blocked for this client "
            "during the month. High counts on a single domain may indicate "
            "persistent software calling home to a blocked service.",
            size=10,
        )
        dest_rows = [
            (i + 1, row[0], fmt_int(row[1]))
            for i, row in enumerate(data.top_destinations)
        ]
        brand.styled_table(
            doc,
            ["#", "Blocked Domain", "Queries Blocked"],
            dest_rows,
            col_widths=[0.35, 4.35, 1.3],
        )

    # Top blocked categories
    if data.top_categories:
        doc.add_paragraph()
        brand.add_body(doc, "Top blocked categories", bold=True, size=12, color=brand.DARK_CHARCOAL)
        cat_rows = [
            (i + 1, cat, fmt_int(cnt))
            for i, (cat, cnt) in enumerate(data.top_categories)
        ]
        brand.styled_table(
            doc,
            ["#", "Category", "Blocked Records"],
            cat_rows,
            col_widths=[0.35, 4.35, 1.3],
        )

    brand.add_page_break(doc)

    # ---- PROTECTED ENDPOINTS ----
    brand.add_section_header(doc, "Protected Endpoints")
    agent_status_break = data.agent_summary()
    brand.add_body(
        doc,
        f"The table below shows the Cisco Umbrella roaming-computer agent "
        f"status for {data.location_name} as of {data.latest_agents_date or 'end of month'}. "
        f"Encrypted status means the agent is active and enforcing DNS "
        f"policy. Off status means the agent is installed but not currently "
        f"enforcing.",
    )

    # Agent status summary
    doc.add_paragraph()
    brand.add_body(doc, "Agent status summary", bold=True, size=12, color=brand.DARK_CHARCOAL)
    status_rows = list(agent_status_break.items())
    status_rows.append(("Total", str(agents_count)))
    brand.styled_table(
        doc,
        ["Status", "Count"],
        status_rows,
        col_widths=[4.5, 1.5],
        status_col=0,
        bold_last_row=True,
    )

    # Top active endpoints
    if data.top_identities:
        doc.add_paragraph()
        brand.add_body(doc, "Most active endpoints (by DNS queries this month)",
                       bold=True, size=12, color=brand.DARK_CHARCOAL)
        id_rows = [
            (i + 1, name, fmt_int(cnt))
            for i, (name, cnt) in enumerate(data.top_identities)
        ]
        brand.styled_table(
            doc,
            ["#", "Endpoint / Identity", "Total Queries"],
            id_rows,
            col_widths=[0.35, 4.35, 1.3],
        )

    # Full agent list
    if data.latest_agents:
        doc.add_paragraph()
        brand.add_body(doc, "Full agent inventory (point-in-time snapshot)",
                       bold=True, size=12, color=brand.DARK_CHARCOAL)
        agent_rows = []
        for a in sorted(data.latest_agents, key=lambda x: x.get("name", "")):
            last_sync = a.get("lastSync", "")
            if last_sync:
                try:
                    dt = datetime.fromisoformat(last_sync.replace("Z", "+00:00"))
                    last_sync = dt.strftime("%Y-%m-%d")
                except Exception:
                    pass
            agent_rows.append((
                a.get("name", "—"),
                (a.get("status") or "—").title(),
                a.get("version", "—"),
                a.get("osVersionName") or "—",
                last_sync or "—",
            ))
        brand.styled_table(
            doc,
            ["Hostname", "Status", "Version", "OS", "Last Sync"],
            agent_rows,
            col_widths=[1.7, 0.9, 0.8, 1.5, 1.1],
            status_col=1,
        )

    brand.add_page_break(doc)

    # ---- WHAT WE DID ----
    brand.add_section_header(doc, "What Technijian Did For You")
    bullets = [
        ("DNS protection — always on. ",
         f"Cisco Umbrella inspected every DNS lookup from every covered "
         f"device. {fmt_int(data.total_requests)} queries were evaluated; "
         f"{fmt_int(data.total_blocked)} ({block_pct_str}) were blocked "
         f"before a connection could be established."),
        ("Threat blocking. ",
         f"Domains associated with malware, phishing, ransomware, and "
         f"policy-restricted categories were blocked at the DNS layer — "
         f"before the browser or application could open a connection."),
        ("Endpoint coverage. ",
         f"The Cisco Umbrella roaming-computer agent was installed on "
         f"{agents_count} device(s), ensuring protection both on the "
         f"corporate network and on public or home Wi-Fi."),
        ("Posture tracking. ",
         "Agent version, sync status, and activity data were captured daily "
         "to support fleet health reviews and coverage gap identification."),
    ]
    for prefix, body in bullets:
        brand.add_bullet(doc, body, bold_prefix=prefix)

    # ---- RECOMMENDATIONS ----
    doc.add_paragraph()
    brand.add_section_header(doc, "Recommendations")
    if not stale and off_count == 0:
        brand.add_body(
            doc,
            "No action is required from your team this month. All "
            "Umbrella agents are syncing on schedule.",
        )
    else:
        if off_count > 0:
            brand.add_body(doc, f"Agents in 'Off' state ({off_count})",
                           bold=True, size=12, color=brand.DARK_CHARCOAL)
            off_agents = [
                a for a in data.latest_agents
                if (a.get("status") or "").lower() == "off"
            ]
            off_rows = [
                (a.get("name", "—"), (a.get("status") or "—").title(),
                 a.get("version", "—"), a.get("osVersionName") or "—")
                for a in off_agents[:20]
            ]
            brand.styled_table(
                doc,
                ["Hostname", "Status", "Version", "OS"],
                off_rows,
                col_widths=[2.0, 1.1, 0.9, 2.0],
                status_col=1,
            )
            brand.add_body(
                doc,
                "Recommended action: investigate why DNS enforcement is "
                "disabled on these devices. Common causes include the user "
                "temporarily disabling the agent or an OS change. Contact "
                "your Technijian representative to schedule remediation.",
            )

        if stale:
            doc.add_paragraph()
            brand.add_body(doc, f"Agents not syncing in 30+ days ({len(stale)})",
                           bold=True, size=12, color=brand.DARK_CHARCOAL)
            stale_rows = [
                (a.get("name", "—"), (a.get("status") or "—").title(),
                 a.get("version", "—"),
                 (a.get("lastSync") or "—")[:10])
                for a in stale[:20]
            ]
            brand.styled_table(
                doc,
                ["Hostname", "Status", "Version", "Last Sync"],
                stale_rows,
                col_widths=[2.0, 1.1, 0.9, 2.0],
                status_col=1,
            )
            if len(stale) > 20:
                brand.add_body(doc, f"... plus {len(stale) - 20} more. Full list available on request.",
                               size=10)
            brand.add_body(
                doc,
                "Recommended action: confirm whether each device is still in "
                "service. If decommissioned, notify Technijian to remove the "
                "agent record. If still active, the agent likely needs "
                "reinstallation or the device may have been offline.",
            )

    # ---- ABOUT ----
    doc.add_paragraph()
    brand.add_section_header(doc, "About This Report", accent_color=brand.CORE_ORANGE)
    brand.add_body(
        doc,
        "Technijian manages Cisco Umbrella DNS security as part of its "
        "layered cybersecurity stack for covered clients. Activity captured "
        "here is pulled directly from the Cisco Umbrella API for the "
        "Technijian-managed organization. Endpoint inventory reflects a "
        "point-in-time snapshot taken at the most recent daily pull within "
        "the month. DNS query and block counts are aggregated from the "
        "available daily snapshots for the reported month.",
        size=10,
    )
    brand.add_body(doc, brand.CONTACT_LINE, size=10, color=brand.BRAND_GREY)

    # Save
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    return {
        "code": data.code,
        "month": data.month,
        "out_path": str(out_path),
        "days": data.days_with_data,
        "total_requests": data.total_requests,
        "total_blocked": data.total_blocked,
        "agents": agents_count,
        "stale": len(stale),
        "off": off_count,
    }


# ---------------------------------------------------------------------------
# Discovery + driver
# ---------------------------------------------------------------------------

def discover_clients_with_month(month: str) -> list[str]:
    out: list[str] = []
    for child in CLIENTS_ROOT.iterdir():
        if not child.is_dir():
            continue
        umbrella_dir = child / "umbrella"
        if not umbrella_dir.exists():
            continue
        prefix = month
        if any(p.is_dir() and p.name.startswith(prefix) for p in umbrella_dir.iterdir()):
            out.append(child.name.lower())
    return sorted(out)


def discover_all_months_for_client(code: str) -> list[str]:
    umbrella_dir = CLIENTS_ROOT / code.lower() / "umbrella"
    if not umbrella_dir.exists():
        return []
    months: set[str] = set()
    for p in umbrella_dir.iterdir():
        if p.is_dir() and len(p.name) == 10 and p.name[4] == "-" and p.name[7] == "-":
            months.add(p.name[:7])
    return sorted(months)


def parse_args() -> argparse.Namespace:
    ap = argparse.ArgumentParser(
        description="Build branded Cisco Umbrella monthly DNS security report per client."
    )
    grp = ap.add_mutually_exclusive_group(required=True)
    grp.add_argument("--month", help="YYYY-MM target month")
    grp.add_argument("--all", action="store_true",
                     help="Generate reports for every client+month that has data")
    ap.add_argument("--only", help="comma-separated LocationCodes")
    ap.add_argument("--skip", action="append", default=[],
                    help="LocationCode to skip (repeatable)")
    return ap.parse_args()


def main() -> int:
    args = parse_args()
    skip_codes = {s.upper() for chunk in args.skip for s in chunk.split(",") if s.strip()}
    only_codes = None
    if args.only:
        only_codes = {s.strip().upper() for s in args.only.split(",") if s.strip()}

    # Build (client, month) work list
    work: list[tuple[str, str]] = []
    if args.all:
        for child in sorted(CLIENTS_ROOT.iterdir()):
            if not child.is_dir():
                continue
            code = child.name.lower()
            for m in discover_all_months_for_client(code):
                work.append((code, m))
    else:
        month = args.month
        codes = discover_clients_with_month(month)
        if not codes:
            print(f"No clients have Umbrella data for month {month}. "
                  f"Run pull_umbrella_daily.py or backfill_umbrella.py first.")
            return 1
        for code in codes:
            work.append((code, month))

    if not work:
        print("Nothing to generate.")
        return 0

    print(f"[{datetime.now():%H:%M:%S}] Building Umbrella monthly reports")
    print(f"  {len(work)} (client, month) pair(s) to generate")

    results: list[dict] = []
    for code, month in work:
        upper = code.upper()
        if upper in skip_codes:
            continue
        if only_codes is not None and upper not in only_codes:
            continue

        umbrella_dir = CLIENTS_ROOT / code / "umbrella"
        data = ClientUmbrellaMonth(upper, month, umbrella_dir)
        if not data.has_data:
            print(f"  [skip] {upper:<6s} {month}  no request data")
            continue

        out_path = umbrella_dir / "monthly" / month / f"{upper}-DNS-Security-{month}.docx"
        try:
            res = render_report(data, out_path)
        except Exception as e:
            import traceback
            print(f"  [ERR ] {upper:<6s} {month}  {e}")
            traceback.print_exc()
            continue

        print(
            f"  [ok  ] {upper:<6s} {month}  "
            f"days={res['days']:>2d}  "
            f"queries={res['total_requests']:>10,}  "
            f"blocked={res['total_blocked']:>7,}  "
            f"agents={res['agents']:>3d}  "
            f"off={res['off']:>2d}  "
            f"stale={res['stale']:>2d}  "
            f"-> {Path(res['out_path']).relative_to(REPO)}"
        )
        results.append(res)

    print()
    print(f"[{datetime.now():%H:%M:%S}] DONE — generated {len(results)} report(s)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
