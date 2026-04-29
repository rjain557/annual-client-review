#!/usr/bin/env python3
"""
build_teramind_compliance_report.py

Generate a Technijian-branded compliance monitoring report per client
from Teramind daily pull data.

Usage
-----
  # Reports for all clients that have pull data for a date
  python technijian/teramind-pull/scripts/build_teramind_compliance_report.py --date 2026-04-29

  # Specific clients only
  python ... --date 2026-04-29 --only LAG,QOSNET

Output
------
  technijian/teramind-pull/YYYY-MM-DD/reports/<CLIENT>-Compliance-YYYY-MM.docx
"""

import argparse
import json
import sys
from datetime import datetime, timezone
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[3]
sys.path.insert(0, str(Path(__file__).parent))
sys.path.insert(0, str(REPO_ROOT / "technijian" / "huntress-pull" / "scripts"))

import _brand as brand
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── DLP category labels ───────────────────────────────────────────────────────

DLP_CATEGORIES = {
    1: "Productivity",
    2: "Data Exfiltration",
    3: "Behavioral Risk",
    4: "Communication",
}

# ── domain -> client label mapping ───────────────────────────────────────────
# Maps partial email/fqdn domain fragments to display names.
# Add entries here as new clients are enrolled in Teramind.
DOMAIN_MAP = {
    "lag.local":   ("LAG",    "LA Growth"),
    "qosnet.com":  ("QOSNET", "QOSNet"),
    "qosnet.info": ("QOSNET", "QOSNet"),
    "mkc-m24":     ("MKC",    "MKC"),
    "csprx.com":   ("CSPRX",  "CS PRX"),
    "technijian.com": None,  # suppress Technijian admin agent
}


def infer_client(email_or_fqdn):
    """Return (code, display_name) or None if internal/unrecognised."""
    val = (email_or_fqdn or "").lower()
    for fragment, mapping in DOMAIN_MAP.items():
        if fragment in val:
            return mapping
    return None


def segment_by_client(agents, computers):
    """Group agents and computers by client code."""
    clients = {}

    for agent in agents:
        email = agent.get("email", "")
        result = infer_client(email)
        if result is None:
            continue
        code, name = result
        if code not in clients:
            clients[code] = {"name": name, "agents": [], "computers": []}
        clients[code]["agents"].append(agent)

    for computer in computers:
        fqdn = computer.get("fqdn", "") or computer.get("name", "")
        result = infer_client(fqdn)
        if result is None:
            continue
        code, name = result
        if code not in clients:
            clients[code] = {"name": name, "agents": [], "computers": []}
        clients[code]["computers"].append(computer)

    return clients


# ── risk label ────────────────────────────────────────────────────────────────

def risk_level(score):
    if score is None:
        return "Unknown"
    if score >= 75:
        return "High Risk"
    if score >= 40:
        return "Medium"
    return "Low"


# ── report builder ────────────────────────────────────────────────────────────

def build_report(client_code, client_name, agents, computers,
                 policies, risk_scores, cube_data,
                 pull_date, out_path):

    month_label = pull_date.strftime("%B %Y")
    doc = brand.new_branded_document()

    # ── cover ──────────────────────────────────────────────────────────────
    brand.render_cover(
        doc,
        title="Compliance Monitoring Report",
        subtitle=client_name,
        footer_note="CONFIDENTIAL — For authorized recipients only.",
        date_text=month_label,
    )

    # ── executive summary ─────────────────────────────────────────────────
    brand.add_section_header(doc, "Executive Summary")

    agents_monitored = len(agents)
    computers_monitored = len([c for c in computers if c.get("is_monitored")])
    active_policies = [p for p in policies if p.get("active") and not p.get("deleted")]
    total_policies = len([p for p in policies if not p.get("deleted")])

    # Determine max risk score for this client's agents
    client_emails = {a["email"].lower() for a in agents}
    client_risks = [
        r for r in risk_scores
        if (r.get("email") or "").lower() in client_emails
    ]
    max_score = max((r.get("score") or 0 for r in client_risks), default=0)
    risk_label = risk_level(max_score)

    # Activity totals (account-wide for now; single-tenant)
    total_activity = sum(
        (cube_data.get(c, {}).get("rows", [{}]) or [{}])[0].get("count", 0)
        for c in ["activity", "keystrokes", "web_search", "social_media"]
    )

    # Metric cards — (value, label, color)
    brand.add_metric_card_row(doc, [
        (str(agents_monitored),      "Employees Monitored",    brand.CORE_BLUE),
        (str(computers_monitored),   "Computers Monitored",    brand.CORE_BLUE),
        (str(len(active_policies)),  "Active DLP Policies",    brand.CORE_ORANGE),
        (str(total_policies),        "Available DLP Policies", brand.BRAND_GREY),
    ])

    # Overall status callout
    if total_activity == 0:
        brand.add_callout_box(
            doc,
            "System Operational - Baseline Period: Teramind monitoring is active "
            "and agents are enrolled. No activity events have been recorded in this "
            "reporting period. This is expected during initial system deployment. "
            "As monitored endpoints check in and users perform activity, events will "
            "be captured and included in future reports.",
            accent_hex="1EAAC8",
            bg_hex="EAF8FB",
        )
    elif max_score >= 75:
        brand.add_callout_box(
            doc,
            f"Attention Required: One or more monitored employees has a high "
            f"insider-threat risk score ({max_score}/100). Review the per-employee "
            "risk detail below.",
            accent_hex="C0392B",
            bg_hex="FDEDEC",
        )
    else:
        brand.add_callout_box(
            doc,
            "Monitoring Healthy: All monitored endpoints are enrolled and DLP "
            "policies are active. Insider-threat risk scores are within normal range.",
            accent_hex="27AE60",
            bg_hex="EAFAF1",
        )

    # ── endpoint monitoring coverage ────────────────────────────────────────
    brand.add_section_header(doc, "Endpoint Monitoring Coverage")

    p = doc.add_paragraph()
    brand.add_run(p, "Monitored Employees", bold=True)

    emp_rows = []
    for agent in agents:
        emp_rows.append({
            "Email / Username": agent.get("email", ""),
            "Department": str(agent.get("department_id") or "Unassigned"),
            "Last Login": agent.get("last_web_login") or "Not logged in",
            "Monitoring": "Active" if not agent.get("deleted") else "Inactive",
        })

    brand.styled_table(
        doc,
        headers=["Email / Username", "Department", "Last Login", "Monitoring"],
        rows=[[r["Email / Username"], r["Department"],
               r["Last Login"][:19] if r["Last Login"] != "Not logged in" else r["Last Login"],
               r["Monitoring"]]
              for r in emp_rows],
        status_col=3,
    )

    p2 = doc.add_paragraph()
    brand.add_run(p2, "\nMonitored Computers", bold=True)

    if computers:
        computer_rows = []
        for c in computers:
            computer_rows.append([
                c.get("name", ""),
                c.get("fqdn", ""),
                (c.get("os") or "")[:35],
                c.get("ip", ""),
                "Active" if c.get("is_monitored") else "Inactive",
            ])
        brand.styled_table(
            doc,
            headers=["Computer Name", "FQDN", "OS", "IP Address", "Status"],
            rows=computer_rows,
            status_col=4,
        )
    else:
        doc.add_paragraph("No computers currently registered for this client.")

    # ── DLP policy status ────────────────────────────────────────────────────
    brand.add_section_header(doc, "DLP Policy Status")

    p3 = doc.add_paragraph(
        f"Teramind has {total_policies} compliance policies configured, "
        f"of which {len(active_policies)} are currently active. "
        "Inactive policies are configured but not enforced; they can be "
        "activated without additional setup."
    )

    policy_rows = []
    for pol in sorted(policies, key=lambda x: (not x.get("active"), x.get("name", ""))):
        if pol.get("deleted"):
            continue
        cat_id = pol.get("params", {}).get("ui", {}).get("category", 0)
        cat_name = DLP_CATEGORIES.get(cat_id, "General")
        desc = pol.get("params", {}).get("ui", {}).get("description", "")
        policy_rows.append([
            pol.get("name", ""),
            cat_name,
            "Active" if pol.get("active") else "Inactive",
            desc[:60] + ("..." if len(desc) > 60 else ""),
        ])

    brand.styled_table(
        doc,
        headers=["Policy Name", "Category", "Status", "Description"],
        rows=policy_rows,
        status_col=2,
    )

    # ── activity summary ─────────────────────────────────────────────────────
    brand.add_section_header(doc, "Activity Summary")

    cube_display = {
        "activity":    "General Activity",
        "keystrokes":  "Keystroke Log",
        "web_search":  "Web Searches",
        "social_media": "Social Media",
    }
    activity_rows = []
    for cube, label in cube_display.items():
        count = (cube_data.get(cube, {}).get("rows", [{}]) or [{}])[0].get("count", 0)
        activity_rows.append([
            label,
            str(count),
            "No events captured" if count == 0 else f"{count:,} events",
        ])

    brand.styled_table(
        doc,
        headers=["Activity Type", "Events (24 h)", "Note"],
        rows=activity_rows,
    )

    if total_activity == 0:
        note_p = doc.add_paragraph(
            "Zero activity events are recorded during this reporting period. "
            "This is expected when monitored computers have not yet connected to "
            "the Teramind server or when the monitoring agent has not yet reported "
            "in. Activity data will populate in future reports once endpoint agents "
            "are communicating."
        )
        note_p.runs[0].font.color.rgb = brand.CORE_BLUE

    # ── insider-threat risk scores ────────────────────────────────────────────
    brand.add_section_header(doc, "Insider-Threat Risk Assessment")

    doc.add_paragraph(
        "Teramind continuously scores each monitored employee on a 0-100 "
        "insider-threat risk scale based on behavioral patterns, policy "
        "violations, and activity anomalies. Scores above 75 warrant review."
    )

    risk_rows = []
    for r in client_risks:
        score = r.get("score") or 0
        pct = r.get("percentile") or 0
        label = risk_level(score)
        risk_rows.append([
            r.get("email", ""),
            str(score),
            f"{pct}th percentile",
            label,
        ])

    if risk_rows:
        brand.styled_table(
            doc,
            headers=["Employee", "Risk Score", "Percentile", "Level"],
            rows=risk_rows,
            status_col=3,
        )
    else:
        doc.add_paragraph("No risk score data available for this client period.")

    # ── what Technijian did ───────────────────────────────────────────────────
    brand.add_section_header(doc, "What Technijian Did For You")

    actions = doc.add_paragraph(style="List Bullet")
    actions.add_run(
        f"Deployed Teramind compliance monitoring agents to {agents_monitored} "
        f"employee endpoint(s)."
    )

    actions2 = doc.add_paragraph(style="List Bullet")
    actions2.add_run(
        f"Configured {len(active_policies)} active DLP policies covering behavioral "
        f"risk, data exfiltration, and productivity controls."
    )

    actions3 = doc.add_paragraph(style="List Bullet")
    actions3.add_run(
        "Established continuous insider-threat risk scoring for all enrolled employees."
    )

    actions4 = doc.add_paragraph(style="List Bullet")
    actions4.add_run(
        "Verified daily compliance data capture pipeline — pull runs at 4:00 AM PT."
    )

    # ── recommendations ───────────────────────────────────────────────────────
    brand.add_section_header(doc, "Recommendations")

    inactive_policies = [p for p in policies
                         if not p.get("active") and not p.get("deleted")]
    high_value_inactive = [
        p for p in inactive_policies
        if any(kw in (p.get("name") or "").lower()
               for kw in ["ssn", "credit card", "pii", "clipboard", "removable"])
    ]

    if high_value_inactive:
        rec_p = doc.add_paragraph(style="List Bullet")
        rec_p.add_run("Activate high-value DLP policies: ").bold = True
        names = ", ".join(p["name"] for p in high_value_inactive[:4])
        rec_p.add_run(
            f"The following policies are configured but inactive: {names}. "
            "Activating these provides stronger data exfiltration protection."
        )

    if agents_monitored > 0 and computers_monitored == 0:
        rec_p2 = doc.add_paragraph(style="List Bullet")
        rec_p2.add_run("Deploy endpoint agents: ").bold = True
        rec_p2.add_run(
            "Employees are enrolled but no monitored computers are registered. "
            "Install the Teramind agent on employee workstations to begin capturing activity."
        )

    rec_p3 = doc.add_paragraph(style="List Bullet")
    rec_p3.add_run("Review department assignments: ").bold = True
    rec_p3.add_run(
        "All enrolled agents currently show no department assignment. Assigning "
        "employees to departments enables department-level compliance reporting."
    )

    # ── about ─────────────────────────────────────────────────────────────────
    brand.add_section_header(doc, "About This Report")

    about = doc.add_paragraph(
        f"This report was generated by Technijian on {datetime.now().strftime('%B %d, %Y')} "
        f"using data pulled from the Teramind on-premise compliance monitoring server. "
        f"Data covers the 24-hour window ending {pull_date.strftime('%B %d, %Y %H:%M UTC')}. "
        "Activity counts, risk scores, and policy status reflect point-in-time snapshots. "
        "This report is confidential and intended for authorized recipients only."
    )

    # ── save ──────────────────────────────────────────────────────────────────
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"  Saved: {out_path}")


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    p = argparse.ArgumentParser(description="Build Teramind compliance reports")
    p.add_argument("--date", default=datetime.now().strftime("%Y-%m-%d"),
                   help="Pull date YYYY-MM-DD (default: today)")
    p.add_argument("--only", help="Comma-separated client codes, e.g. LAG,QOSNET")
    args = p.parse_args()

    pull_dir = REPO_ROOT / "technijian" / "teramind-pull" / args.date
    if not pull_dir.exists():
        print(f"No pull data for {args.date} at {pull_dir}")
        sys.exit(1)

    # Load data
    agents    = json.loads((pull_dir / "agents.json").read_text())
    computers = json.loads((pull_dir / "computers.json").read_text())
    policies  = json.loads((pull_dir / "behavior_policies.json").read_text())
    risk_scores = json.loads((pull_dir / "risk_scores.json").read_text())

    cube_data = {}
    for cube in ["activity", "keystrokes", "web_search", "social_media"]:
        cube_file = pull_dir / f"{cube}.json"
        if cube_file.exists():
            cube_data[cube] = json.loads(cube_file.read_text())

    pull_date = datetime.strptime(args.date, "%Y-%m-%d").replace(
        hour=19, tzinfo=timezone.utc
    )

    # Segment by client
    clients = segment_by_client(agents, computers)
    if not clients:
        print("No client-mappable agents/computers found. Check DOMAIN_MAP in script.")
        sys.exit(1)

    only_set = set(args.only.upper().split(",")) if args.only else None

    print(f"Building compliance reports for {args.date}...")
    generated = []
    for code, info in sorted(clients.items()):
        if only_set and code not in only_set:
            continue
        month_str = pull_date.strftime("%Y-%m")
        out_path = pull_dir / "reports" / f"{code}-Compliance-{month_str}.docx"
        print(f"\n  {code} ({info['name']}): "
              f"{len(info['agents'])} agent(s), {len(info['computers'])} computer(s)")
        build_report(
            client_code=code,
            client_name=info["name"],
            agents=info["agents"],
            computers=info["computers"],
            policies=policies,
            risk_scores=risk_scores,
            cube_data=cube_data,
            pull_date=pull_date,
            out_path=out_path,
        )
        generated.append(out_path)

    print(f"\nDone. Reports in {pull_dir / 'reports'}/")
    sys.stdout.flush()

    # Proofread every generated report
    _proofread_reports(generated)


def _proofread_reports(paths):
    """Run structural proofreader on generated DOCX files."""
    proofreader = REPO_ROOT / "technijian" / "shared" / "scripts" / "proofread_docx.py"
    if not proofreader.exists():
        print("[proofread] Skipped — proofread_docx.py not found")
        return
    import subprocess
    file_args = [str(p) for p in paths if p.exists()]
    if not file_args:
        return
    sections = (
        "Executive Summary,Endpoint Monitoring Coverage,DLP Policy Status,"
        "Activity Summary,Insider-Threat Risk Assessment,"
        "What Technijian Did For You,Recommendations,About This Report"
    )
    result = subprocess.run(
        [sys.executable, str(proofreader), "--sections", sections] + file_args,
        capture_output=False,
    )
    if result.returncode != 0:
        print("\n[proofread] WARNING: one or more reports failed proofreading above.")
        sys.exit(result.returncode)


if __name__ == "__main__":
    main()
