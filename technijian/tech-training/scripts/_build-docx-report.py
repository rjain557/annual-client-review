"""Build Technijian-branded Word report for the tech-time-entry audit.

Reads outlier CSVs and emits docx at:
  technijian/tech-training/<client>/<year>/<CLIENT>-Tech-Time-Entry-Audit.docx

Usage: python _build-docx-report.py <client_code> <year>
"""
import csv
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SCRIPTS = Path(__file__).resolve().parent
REPO = SCRIPTS.parent.parent.parent
CLIENT = (sys.argv[1] if len(sys.argv) > 1 else "bwh").lower()
YEAR = sys.argv[2] if len(sys.argv) > 2 else "2026"
BASE = REPO / "technijian" / "tech-training" / CLIENT / YEAR
LOGO = Path(r"C:\VSCode\tech-branding\tech-branding\assets\logos\png\technijian-logo-full-color-600x125.png")
OUT = BASE / f"{CLIENT.upper()}-Tech-Time-Entry-Audit.docx"

# Brand
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
TEAL = RGBColor(0x1E, 0xAA, 0xC8)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
OFF_WHITE = "F8F9FA"
LIGHT_GREY = "E9ECEF"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x28, 0xA7, 0x45)
RED = RGBColor(0xCC, 0x00, 0x00)
FONT = "Open Sans"


def shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def remove_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for b in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{b}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tc_pr.append(borders)


def set_cell_border_color(cell, hex_color=LIGHT_GREY, sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for b in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{b}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:color"), hex_color)
        borders.append(e)
    tc_pr.append(borders)


def add_run(paragraph, text, *, bold=False, size=11, color=BRAND_GREY, font=FONT, italic=False):
    run = paragraph.add_run(text)
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def add_body(doc, text, *, bold=False, size=11, color=BRAND_GREY, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bold=bold, size=size, color=color)
    return p


def add_bullet(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=11, color=DARK_CHARCOAL)
        add_run(p, text, size=11, color=BRAND_GREY)
    else:
        add_run(p, text, size=11, color=BRAND_GREY)
    return p


def add_color_bar(doc, hex_color, height_pt=4):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade(cell, hex_color)
    remove_borders(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("")
    r.font.size = Pt(height_pt)
    return t


def add_section_header(doc, title, accent_color=CORE_BLUE):
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Emu(60000)
    t.columns[1].width = Emu(5700000)
    bar_cell, title_cell = t.rows[0].cells
    bar_cell.width = Emu(60000)
    title_cell.width = Emu(5700000)
    shade(bar_cell, "006DB6" if accent_color == CORE_BLUE else "F67D4B")
    remove_borders(bar_cell)
    remove_borders(title_cell)
    p = title_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "  " + title, bold=True, size=14, color=accent_color)
    return t


def set_col_widths(table, widths_in_inches):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_in_inches):
                cell.width = Inches(widths_in_inches[i])


def styled_table(doc, headers, rows, *, col_widths=None, money_cols=(),
                 bold_last_row=False, flag_col=None, hrs_col=None, cap_col=None,
                 hrs_red_if_over_cap=True):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = False
    if col_widths:
        set_col_widths(t, col_widths)
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        shade(cell, "006DB6")
        set_cell_border_color(cell, LIGHT_GREY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, h, bold=True, size=10, color=WHITE)
    for i, row in enumerate(rows):
        is_total = bold_last_row and i == len(rows) - 1
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            if is_total:
                shade(cell, "E9ECEF")
            elif i % 2 == 1:
                shade(cell, OFF_WHITE)
            set_cell_border_color(cell, LIGHT_GREY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            text = str(val)
            color = DARK_CHARCOAL if is_total else BRAND_GREY
            if j in money_cols or (isinstance(val, (int, float))):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # colour-code hours red if over cap
            if hrs_red_if_over_cap and j == hrs_col and cap_col is not None:
                try:
                    h = float(val)
                    c = float(row[cap_col])
                    if h > c * 1.5:
                        color = RED
                    elif h > c:
                        color = CORE_ORANGE
                except (ValueError, TypeError):
                    pass
            add_run(p, text, size=9.5, color=color, bold=is_total)
    return t


def add_metric_card_row(doc, cards):
    t = doc.add_table(rows=1, cols=len(cards))
    t.autofit = False
    for i, (value, label, color) in enumerate(cards):
        cell = t.rows[0].cells[i]
        shade(cell, "F8F9FA")
        remove_borders(cell)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(12)
        p1.paragraph_format.space_after = Pt(0)
        add_run(p1, value, bold=True, size=24, color=color)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(12)
        add_run(p2, label, size=10, color=BRAND_GREY)
    return t


def set_default_style(doc):
    s = doc.styles["Normal"]
    s.font.name = FONT
    s.font.size = Pt(11)
    s.font.color.rgb = BRAND_GREY


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Technijian  |  Internal Use — Tech Training  |  18 Technology Dr., Ste 141, Irvine, CA 92618",
            size=9, color=BRAND_GREY)


def add_header_logo(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    if LOGO.exists():
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.7))
    else:
        add_run(p, "TECHNIJIAN", bold=True, size=14, color=CORE_BLUE)


# --- Load CSVs ---
def read_csv(path: Path):
    with path.open("r", encoding="utf-8", newline="") as f:
        rows = list(csv.reader(f))
    return rows[0], rows[1:]


bytech_hdr, bytech_rows = read_csv(BASE / "tech-outliers-by-tech.csv")
detail_hdr, detail_rows = read_csv(BASE / "tech-outliers-detail.csv")

# Re-derive the scanned totals
te_csv = REPO / "clients" / CLIENT / YEAR / "03_Accounting" / "time-entries.csv"
scanned_entries = 0
scanned_hours = 0.0
with te_csv.open("r", encoding="utf-8", newline="") as f:
    reader = csv.DictReader(f)
    for row in reader:
        try:
            h = abs(float(row.get("Hours") or 0))
        except ValueError:
            h = 0.0
        if h > 0:
            scanned_entries += 1
            scanned_hours += h

flagged_entries = len(detail_rows)
flagged_hours = sum(float(r[detail_hdr.index("Hours")]) for r in detail_rows)

# ============ BUILD DOC ============
doc = Document()
set_default_style(doc)
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

add_header_logo(doc)
add_footer(doc)

# ---- Cover ----
add_color_bar(doc, "006DB6", height_pt=6)
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
if LOGO.exists():
    run = p.add_run()
    run.add_picture(str(LOGO), width=Inches(3.0))

for _ in range(2):
    doc.add_paragraph()

div_t = doc.add_table(rows=1, cols=3)
div_t.autofit = False
for i, cell in enumerate(div_t.rows[0].cells):
    remove_borders(cell)
    if i == 1:
        shade(cell, "F67D4B")
        cell.width = Inches(1.5)
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        r = cp.add_run("")
        r.font.size = Pt(3)
    else:
        cell.width = Inches(2.6)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Tech Time-Entry Audit", bold=True, size=28, color=DARK_CHARCOAL)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, f"{CLIENT.upper()} — {YEAR}", bold=True, size=20, color=CORE_BLUE)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Life-of-contract review of time entries against reasonable-hours caps per work category",
        size=12, color=BRAND_GREY)

for _ in range(6):
    doc.add_paragraph()

add_color_bar(doc, "F67D4B", height_pt=6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "INTERNAL USE — TECH TRAINING", bold=True, size=10, color=BRAND_GREY)

add_page_break(doc)

# ---- Executive Summary ----
add_section_header(doc, "Executive Summary", CORE_BLUE)
doc.add_paragraph()

add_body(doc,
    f"This audit scanned all {scanned_entries:,} time entries ({scanned_hours:,.2f} hours) "
    f"logged by Technijian pods against {CLIENT.upper()} over the life of the {YEAR} annual-review "
    f"period. Each entry was classified into one of the 32 work categories used in the master "
    f"accounting, then compared against a reasonable-hours cap derived from the nature of the "
    f"work. {flagged_entries} entries totalling {flagged_hours:,.2f} hours "
    f"({flagged_hours/scanned_hours*100:.1f}% of total) fell outside the expected range.")

doc.add_paragraph()
add_metric_card_row(doc, [
    (f"{scanned_entries:,}", "Entries Scanned", CORE_BLUE),
    (f"{flagged_entries}", "Flagged Entries", CORE_ORANGE),
    (f"{flagged_hours:,.0f}", "Flagged Hours", RED),
    (f"{flagged_hours/scanned_hours*100:.1f}%", "% of Hours Flagged", GREEN),
])

doc.add_paragraph()
add_body(doc, "The overall flagged rate is low, which supports the accuracy of the hour totals. "
              "However, the audit identified specific repeating patterns by tech that warrant "
              "coaching — particularly around vague titles, same-ticket same-day duplicate entries, "
              "and over-claiming on recurring maintenance windows.",
         color=DARK_CHARCOAL)

# ---- Flag Codes ----
add_page_break(doc)
add_section_header(doc, "Flag Codes Used in This Audit", CORE_ORANGE)
doc.add_paragraph()

add_bullet(doc, "Routine/low-complexity work with hours exceeding the reasonable category cap (e.g. a password reset logged at 2 hours, a single CPU-alert ack logged at 1.5 hours).",
           bold_prefix="H1 — Trivial-work-high-hours: ")
add_bullet(doc, "Generic/vague titles (\"Help\", \"Test\", \"Fix\", \"Issue\", \"Support\") with more than 0.5 hour claimed.",
           bold_prefix="H2 — Vague-title-high-hours: ")
add_bullet(doc, "Any single time-block greater than 8 hours (suggests a whole-day dump rather than an itemized entry).",
           bold_prefix="H3 — Single-entry-too-long: ")
add_bullet(doc, "One tech's total on one date greater than 12 hours (cross-ticket over-claim).",
           bold_prefix="H4 — Tech-day-too-long: ")
add_bullet(doc, "Same tech + same ticket + same day with multiple entries summing to more than 2× the category cap.",
           bold_prefix="H5 — Duplicate-day-entries: ")

doc.add_paragraph()
add_body(doc, "Category caps are configurable in _flag-outliers.py. The Appendix lists every cap used.")

# ---- Techs ranked ----
add_page_break(doc)
add_section_header(doc, "Techs Ranked by Flagged Hours", CORE_BLUE)
doc.add_paragraph()

# Top 15 techs
table_rows = []
for r in bytech_rows[:15]:
    table_rows.append([
        r[0],                       # Tech
        r[2],                       # Total hours
        r[4],                       # Flagged hours
        r[5],                       # Flagged %
        r[3],                       # # flagged
        r[6], r[7], r[8], r[9], r[10],  # H1-H5
    ])
styled_table(doc,
    headers=["Tech", "Total hrs", "Flagged hrs", "% flagged", "# entries", "H1", "H2", "H3", "H4", "H5"],
    rows=table_rows,
    col_widths=[1.4, 0.85, 0.95, 0.85, 0.85, 0.4, 0.4, 0.4, 0.4, 0.4],
)

# ---- Top 25 worst entries ----
add_page_break(doc)
add_section_header(doc, "Top 25 Worst Individual Entries", CORE_ORANGE)
doc.add_paragraph()

hours_col = detail_hdr.index("Hours")
cap_col = detail_hdr.index("CategoryCap")
flags_col = detail_hdr.index("Flags")
top25 = sorted(detail_rows, key=lambda r: -float(r[hours_col]))[:25]
t25 = []
for r in top25:
    title = r[4][:70] + ("…" if len(r[4]) > 70 else "")
    t25.append([
        r[0],                   # Date
        r[1],                   # Tech
        r[hours_col],           # Hours
        r[cap_col],             # Cap
        title,                  # Title
        r[flags_col],           # Flags
    ])
styled_table(doc,
    headers=["Date", "Tech", "Hours", "Cap", "Title", "Flags"],
    rows=t25,
    col_widths=[0.9, 1.3, 0.65, 0.55, 3.7, 0.7],
    hrs_col=2, cap_col=3,
)

# ---- Per-flag deep dives ----
flag_labels = {
    "H1": "Trivial/routine work with excessive hours",
    "H2": "Vague title with too many hours",
    "H3": "Single entries over 8 hours",
    "H4": "Tech days over 12 hours total",
    "H5": "Duplicate same-ticket/same-day stacks",
}
for code in ["H1", "H2", "H3", "H4", "H5"]:
    subset = [r for r in detail_rows if code in r[flags_col].split(";")]
    if not subset:
        continue
    add_page_break(doc)
    add_section_header(doc, f"Flag {code} — {flag_labels[code]} ({len(subset)} entries)", CORE_BLUE)
    doc.add_paragraph()
    top = sorted(subset, key=lambda r: -float(r[hours_col]))[:12]
    rs = []
    for r in top:
        title = r[4][:70] + ("…" if len(r[4]) > 70 else "")
        reason = r[detail_hdr.index("Reasons")][:100] + ("…" if len(r[detail_hdr.index("Reasons")]) > 100 else "")
        rs.append([r[0], r[1], r[hours_col], title, reason])
    styled_table(doc,
        headers=["Date", "Tech", "Hours", "Title", "Reason"],
        rows=rs,
        col_widths=[0.9, 1.3, 0.65, 3.1, 2.5],
        hrs_col=2,
    )

# ---- Training recommendations ----
add_page_break(doc)
add_section_header(doc, "Training Recommendations", CORE_ORANGE)
doc.add_paragraph()

recs = [
    ('Require descriptive ticket titles. ',
     'No standalone "Help", "Test", "Fix", or "Issue" titles. Every entry title must describe what '
     'the tech actually did — the customer-specific subject, the system or user affected, and the '
     'action taken (e.g. "Help — Sarah\'s OneDrive sync stuck" rather than just "Help").'),
    ('Consolidate same-ticket same-day work into one entry. ',
     'If multiple blocks of time are spent on the same ticket in one day, merge them into one entry '
     'with a concise note covering the full work. Leaving 3–4 separate entries under the same title '
     'creates the visual appearance of duplicate charging.'),
    ('Cap routine-alert tickets at 1 hour unless genuinely non-routine. ',
     'Monitoring alerts (CPU/memory/disk thresholds, agent offline, device-not-responding) should '
     'usually resolve in under 1 hour. If an alert triggers an investigation that takes longer, '
     'the title must indicate why (e.g. "Critical CPU — investigated SQL runaway process").'),
    ('Split Weekly Maintenance Window hours across covered clients. ',
     'The Weekly Maintenance Window is a shared 6-hour block covering multiple clients. Logging '
     'the whole block to one client is over-claiming.'),
    ('Document agent-update tickets with scope, not effort. ',
     'A ScreenConnect/MyRMM/CrowdStrike version bump across N machines should take ~0.25 hr/machine. '
     'If an update ticket goes >1 hr, the title should explain why (e.g. "MyRMM update — 4 machines '
     'required manual reinstall after service-stop failure").'),
    ('Review your own weekly entries before submission. ',
     'Each tech should spot-check their own week for titles that would not make sense to a client '
     'reviewing the invoice. If "Help" or "Server Issue" is the most specific title on the row, '
     'rewrite it.'),
]
for prefix, body in recs:
    add_bullet(doc, body, bold_prefix=prefix)

# ---- Appendix: Category caps ----
add_page_break(doc)
add_section_header(doc, "Appendix — Category Caps Used", CORE_BLUE)
doc.add_paragraph()

add_body(doc, "These are the single-entry hour caps above which an H1 flag is raised. "
              "Adjust `CATEGORY_CAP` in `_flag-outliers.py` to tighten or loosen these rules.")
doc.add_paragraph()

fo_src = (SCRIPTS / "_flag-outliers.py").read_text(encoding="utf-8")
import re as _re
cap_rows = _re.findall(r'"(Routine: .+?|Project: .+?|Uncategorized)":\s*([\d.]+)', fo_src)
cap_table = [[c, v, "Project" if c.startswith("Project:") else "Routine"]
             for c, v in cap_rows]
styled_table(doc,
    headers=["Category", "Cap (hours)", "Type"],
    rows=cap_table,
    col_widths=[4.5, 1.0, 1.3],
)

doc.save(str(OUT))
print(f"Wrote: {OUT}")
