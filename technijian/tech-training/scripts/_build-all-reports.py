"""Batch-build branded Excel + Word reports for every client and every tech
under the 2026 audit output. Assumes _audit-all-clients.py has already run.

Produces:
  technijian/tech-training/2026/by-client/<client>/<CLIENT>-Tech-Time-Entry-Audit.xlsx
  technijian/tech-training/2026/by-client/<client>/<CLIENT>-Tech-Time-Entry-Audit.docx
  technijian/tech-training/2026/by-tech/<slug>/<slug>-Training.docx

Usage: python _build-all-reports.py [YEAR]
"""
import csv
import re
import sys
from pathlib import Path
from collections import defaultdict

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.formatting.rule import ColorScaleRule, CellIsRule

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

csv.field_size_limit(10_000_000)

SCRIPTS = Path(__file__).resolve().parent
REPO = SCRIPTS.parent.parent.parent
YEAR = sys.argv[1] if len(sys.argv) > 1 else "2026"
ROOT = REPO / "technijian" / "tech-training" / YEAR
LOGO = Path(r"C:\VSCode\tech-branding\tech-branding\assets\logos\png\technijian-logo-full-color-600x125.png")

# --- brand ---
CORE_BLUE = "006DB6"
CORE_ORANGE = "F67D4B"
TEAL = "1EAAC8"
DARK_CHARCOAL = "1A1A2E"
BRAND_GREY = "59595B"
OFF_WHITE = "F8F9FA"
LIGHT_GREY = "E9ECEF"
WHITE = "FFFFFF"
GREEN = "28A745"
RED = "CC0000"
FONT = "Open Sans"

# docx-shared
CORE_BLUE_R = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE_R = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL_R = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY_R = RGBColor(0x59, 0x59, 0x5B)
WHITE_R = RGBColor(0xFF, 0xFF, 0xFF)
GREEN_R = RGBColor(0x28, 0xA7, 0x45)
RED_R = RGBColor(0xCC, 0x00, 0x00)


# ===== docx helpers =====

def shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def remove_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for b in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{b}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tc_pr.append(borders)


def cell_border(cell, hex_color=LIGHT_GREY):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for b in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{b}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), hex_color)
        borders.append(e)
    tc_pr.append(borders)


def add_run(p, text, *, bold=False, size=11, color=BRAND_GREY_R, italic=False):
    run = p.add_run(text)
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def add_body(doc, text, *, bold=False, size=11, color=BRAND_GREY_R, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bold=bold, size=size, color=color)
    return p


def add_bullet(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=11, color=DARK_CHARCOAL_R)
        add_run(p, text, size=11, color=BRAND_GREY_R)
    else:
        add_run(p, text, size=11, color=BRAND_GREY_R)
    return p


def add_color_bar(doc, hex_color, height_pt=4):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade(cell, hex_color)
    remove_borders(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("")
    r.font.size = Pt(height_pt)
    return t


def add_section_header(doc, title, accent="blue"):
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Emu(60000)
    t.columns[1].width = Emu(5700000)
    bar, title_cell = t.rows[0].cells
    bar.width = Emu(60000)
    title_cell.width = Emu(5700000)
    shade(bar, CORE_BLUE if accent == "blue" else CORE_ORANGE)
    remove_borders(bar)
    remove_borders(title_cell)
    p = title_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "  " + title, bold=True, size=14,
            color=CORE_BLUE_R if accent == "blue" else CORE_ORANGE_R)


def set_col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Inches(widths[i])


def styled_table(doc, headers, rows, *, col_widths=None, bold_last=False,
                 hrs_col=None, cap_col=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = False
    if col_widths:
        set_col_widths(t, col_widths)
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        shade(c, CORE_BLUE)
        cell_border(c)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, h, bold=True, size=10, color=WHITE_R)
    for i, row in enumerate(rows):
        is_total = bold_last and i == len(rows) - 1
        for j, val in enumerate(row):
            c = t.rows[i + 1].cells[j]
            if is_total:
                shade(c, LIGHT_GREY)
            elif i % 2 == 1:
                shade(c, OFF_WHITE)
            cell_border(c)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            text = str(val)
            color = DARK_CHARCOAL_R if is_total else BRAND_GREY_R
            try:
                if j == hrs_col and cap_col is not None:
                    hh = float(val)
                    cc = float(row[cap_col])
                    if hh > cc * 1.5:
                        color = RED_R
                    elif hh > cc:
                        color = CORE_ORANGE_R
            except (ValueError, TypeError):
                pass
            add_run(p, text, size=9.5, color=color, bold=is_total)


def add_metric_cards(doc, cards):
    t = doc.add_table(rows=1, cols=len(cards))
    t.autofit = False
    for i, (value, label, color) in enumerate(cards):
        c = t.rows[0].cells[i]
        shade(c, OFF_WHITE)
        remove_borders(c)
        p1 = c.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(12)
        p1.paragraph_format.space_after = Pt(0)
        add_run(p1, value, bold=True, size=24, color=color)
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(12)
        add_run(p2, label, size=10, color=BRAND_GREY_R)


def set_default_style(doc):
    s = doc.styles["Normal"]
    s.font.name = FONT
    s.font.size = Pt(11)
    s.font.color.rgb = BRAND_GREY_R


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def add_header_footer(doc, footer_text):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    if LOGO.exists():
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.6))
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(fp, footer_text, size=9, color=BRAND_GREY_R)


def read_csv(path: Path):
    with path.open("r", encoding="utf-8", newline="") as f:
        rows = list(csv.reader(f))
    return rows[0], rows[1:]


# ===== XLSX helpers =====
HEADER_FILL = PatternFill("solid", fgColor=CORE_BLUE)
HEADER_FONT = Font(name=FONT, color=WHITE, bold=True, size=11)
BODY_FONT = Font(name=FONT, color=BRAND_GREY, size=10)
TITLE_FONT = Font(name=FONT, color=CORE_BLUE, bold=True, size=18)
VALUE_FONT = Font(name=FONT, color=CORE_BLUE, bold=True, size=20)
LABEL_FONT = Font(name=FONT, color=BRAND_GREY, size=10)
THIN = Side(style="thin", color=LIGHT_GREY)
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def xlsx_style_header(ws, row_idx, n_cols):
    for c in range(1, n_cols + 1):
        cell = ws.cell(row=row_idx, column=c)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(vertical="center", horizontal="left", wrap_text=True)
        cell.border = BORDER


def xlsx_style_body(ws, row_idx, n_cols, zebra=False):
    fill = PatternFill("solid", fgColor=OFF_WHITE) if zebra else None
    for c in range(1, n_cols + 1):
        cell = ws.cell(row=row_idx, column=c)
        cell.font = BODY_FONT
        cell.border = BORDER
        if fill:
            cell.fill = fill


def xlsx_col_widths(ws, widths):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w


# ===== PER-CLIENT BUILDERS =====
def build_client_xlsx(client: str, base: Path):
    bytech_hdr, bytech_rows = read_csv(base / "tech-outliers-by-tech.csv")
    detail_hdr, detail_rows = read_csv(base / "tech-outliers-detail.csv")

    wb = openpyxl.Workbook()

    # Summary
    ws = wb.active
    ws.title = "Summary"
    ws.sheet_view.showGridLines = False
    ws.merge_cells("A1:F1")
    ws["A1"] = f"Tech Time-Entry Audit — {client.upper()} {YEAR}"
    ws["A1"].font = TITLE_FONT
    ws.row_dimensions[1].height = 32

    total_entries = sum(int(r[1]) for r in bytech_rows) if bytech_rows else 0
    total_hours = sum(float(r[2]) for r in bytech_rows) if bytech_rows else 0.0
    flagged_entries = sum(int(r[3]) for r in bytech_rows) if bytech_rows else 0
    flagged_hours = sum(float(r[4]) for r in bytech_rows) if bytech_rows else 0.0
    pct = flagged_hours / total_hours * 100 if total_hours else 0

    kpi = 4
    kpis = [
        ("A", f"{total_entries:,}", "Entries scanned"),
        ("B", f"{total_hours:,.1f}", "Hours scanned"),
        ("C", f"{flagged_entries}", "Flagged entries"),
        ("D", f"{flagged_hours:,.1f}", "Flagged hours"),
        ("E", f"{pct:.1f}%", "% hours flagged"),
    ]
    for col, v, l in kpis:
        vc = ws[f"{col}{kpi}"]
        vc.value = v
        vc.font = VALUE_FONT
        vc.alignment = Alignment(horizontal="center", vertical="center")
        vc.border = BORDER
        vc.fill = PatternFill("solid", fgColor=OFF_WHITE)
        lc = ws[f"{col}{kpi+1}"]
        lc.value = l
        lc.font = LABEL_FONT
        lc.alignment = Alignment(horizontal="center")
        lc.border = BORDER
        lc.fill = PatternFill("solid", fgColor=OFF_WHITE)
    xlsx_col_widths(ws, [18] * 6)

    # Techs Ranked
    if bytech_rows:
        ws2 = wb.create_sheet("Techs Ranked")
        ws2.sheet_view.showGridLines = False
        for i, h in enumerate(bytech_hdr, 1):
            ws2.cell(row=1, column=i, value=h)
        xlsx_style_header(ws2, 1, len(bytech_hdr))
        for i, row in enumerate(bytech_rows, 2):
            for j, v in enumerate(row, 1):
                cell = ws2.cell(row=i, column=j)
                if j in (2, 3, 4, 5):
                    try:
                        cell.value = float(v) if "." in v else int(v)
                        cell.number_format = "#,##0.0" if "." in v else "#,##0"
                    except ValueError:
                        cell.value = v
                else:
                    cell.value = v
            xlsx_style_body(ws2, i, len(bytech_hdr), zebra=(i % 2 == 0))
        xlsx_col_widths(ws2, [18, 12, 12, 12, 12, 12, 8, 8, 8, 8, 8])
        ws2.freeze_panes = "A2"
        ws2.auto_filter.ref = f"A1:K{len(bytech_rows)+1}"

    # Detail
    if detail_rows:
        ws3 = wb.create_sheet("Flagged Detail")
        ws3.sheet_view.showGridLines = False
        for i, h in enumerate(detail_hdr, 1):
            ws3.cell(row=1, column=i, value=h)
        xlsx_style_header(ws3, 1, len(detail_hdr))
        for i, row in enumerate(detail_rows, 2):
            for j, v in enumerate(row, 1):
                cell = ws3.cell(row=i, column=j)
                cell.value = v
                if detail_hdr[j - 1] in ("Hours", "CategoryCap", "DailyTotal"):
                    try:
                        cell.value = float(v)
                        cell.number_format = "0.00"
                    except ValueError:
                        pass
            xlsx_style_body(ws3, i, len(detail_hdr), zebra=(i % 2 == 0))
        xlsx_col_widths(ws3, [11, 18, 10, 8, 48, 28, 9, 9, 11, 9, 42, 20])
        ws3.freeze_panes = "A2"
        ws3.auto_filter.ref = f"A1:{get_column_letter(len(detail_hdr))}{len(detail_rows)+1}"
        # Color-scale Hours
        hrs_col_letter = get_column_letter(detail_hdr.index("Hours") + 1)
        ws3.conditional_formatting.add(
            f"{hrs_col_letter}2:{hrs_col_letter}{len(detail_rows)+1}",
            ColorScaleRule(start_type="min", start_color="FFF1DF",
                           mid_type="percentile", mid_value=50, mid_color="FBBF77",
                           end_type="max", end_color="E85A3A"),
        )

    out = base / f"{client.upper()}-Tech-Time-Entry-Audit.xlsx"
    wb.save(str(out))
    return out


def build_client_docx(client: str, base: Path):
    bytech_hdr, bytech_rows = read_csv(base / "tech-outliers-by-tech.csv")
    detail_hdr, detail_rows = read_csv(base / "tech-outliers-detail.csv")

    doc = Document()
    set_default_style(doc)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    add_header_footer(doc, "Technijian  |  Internal — Tech Training  |  18 Technology Dr., Ste 141, Irvine, CA 92618")

    add_color_bar(doc, CORE_BLUE, 6)
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(2.6))
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Tech Time-Entry Audit", bold=True, size=28, color=DARK_CHARCOAL_R)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"{client.upper()} — {YEAR}", bold=True, size=18, color=CORE_BLUE_R)
    for _ in range(6):
        doc.add_paragraph()
    add_color_bar(doc, CORE_ORANGE, 6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "INTERNAL USE — TECH TRAINING", bold=True, size=10, color=BRAND_GREY_R)
    add_page_break(doc)

    # Exec summary
    total_entries = sum(int(r[1]) for r in bytech_rows) if bytech_rows else 0
    total_hours = sum(float(r[2]) for r in bytech_rows) if bytech_rows else 0.0
    flagged_entries = sum(int(r[3]) for r in bytech_rows) if bytech_rows else 0
    flagged_hours = sum(float(r[4]) for r in bytech_rows) if bytech_rows else 0.0
    pct = flagged_hours / total_hours * 100 if total_hours else 0

    add_section_header(doc, "Executive Summary", "blue")
    doc.add_paragraph()
    add_body(doc,
             f"This audit scanned all {YEAR} time entries for {client.upper()} and flagged entries whose "
             f"hours claimed fell outside the expected range for the work described. {flagged_entries} "
             f"entries totalling {flagged_hours:,.2f} hours ({pct:.1f}% of scanned hours) were flagged.")
    doc.add_paragraph()
    add_metric_cards(doc, [
        (f"{total_entries:,}", "Entries Scanned", CORE_BLUE_R),
        (f"{flagged_entries}", "Flagged Entries", CORE_ORANGE_R),
        (f"{flagged_hours:,.1f}", "Flagged Hours", RED_R),
        (f"{pct:.1f}%", "% Hours Flagged", GREEN_R),
    ])

    add_page_break(doc)
    add_section_header(doc, "Flag Codes", "orange")
    doc.add_paragraph()
    add_bullet(doc, "Routine/low-complexity work with hours exceeding the reasonable category cap.",
               bold_prefix="H1 — Trivial-work-high-hours: ")
    add_bullet(doc, "Generic/vague titles (\"Help\", \"Test\", \"Fix\") with > 0.5 hour.",
               bold_prefix="H2 — Vague-title-high-hours: ")
    add_bullet(doc, "Single time-block > 8 hours.", bold_prefix="H3 — Single-entry-too-long: ")
    add_bullet(doc, "Tech daily total > 12 hours across tickets.", bold_prefix="H4 — Tech-day-too-long: ")
    add_bullet(doc, "Same tech + ticket + day with multiple entries summing > 2× cap.",
               bold_prefix="H5 — Duplicate-day-entries: ")

    if bytech_rows:
        add_page_break(doc)
        add_section_header(doc, "Techs ranked by flagged hours", "blue")
        doc.add_paragraph()
        tr = []
        for r in bytech_rows[:20]:
            tr.append([r[0], r[2], r[4], r[5], r[3], r[6], r[7], r[8], r[9], r[10]])
        styled_table(doc,
                     ["Tech", "Total hrs", "Flagged hrs", "% flagged", "#", "H1", "H2", "H3", "H4", "H5"],
                     tr, col_widths=[1.4, 0.85, 0.9, 0.8, 0.5, 0.4, 0.4, 0.4, 0.4, 0.4])

    if detail_rows:
        add_page_break(doc)
        add_section_header(doc, "Top 25 worst individual entries", "orange")
        doc.add_paragraph()
        hrs_idx = detail_hdr.index("Hours")
        cap_idx = detail_hdr.index("CategoryCap")
        flags_idx = detail_hdr.index("Flags")
        top25 = sorted(detail_rows, key=lambda r: -float(r[hrs_idx]))[:25]
        tr = []
        for r in top25:
            t = r[4][:70] + ("…" if len(r[4]) > 70 else "")
            tr.append([r[0], r[1], r[hrs_idx], r[cap_idx], t, r[flags_idx]])
        styled_table(doc,
                     ["Date", "Tech", "Hours", "Cap", "Title", "Flags"],
                     tr, col_widths=[0.9, 1.3, 0.65, 0.55, 3.7, 0.7],
                     hrs_col=2, cap_col=3)

    out = base / f"{client.upper()}-Tech-Time-Entry-Audit.docx"
    doc.save(str(out))
    return out


# ===== PER-TECH BUILDER =====
def build_tech_docx(slug: str, base: Path):
    """Build personalized training.docx for a tech."""
    hdr, rows = read_csv(base / "flagged-entries.csv")
    # Parse training.md header line for totals
    md_text = (base / "training.md").read_text(encoding="utf-8")
    m_tot_e = re.search(r"Total entries logged:\*\*\s*([\d,]+)", md_text)
    m_tot_h = re.search(r"Total hours logged:\*\*\s*([\d,\.]+)", md_text)
    m_fn = re.search(r"Flagged entries:\*\*\s*(\d+)", md_text)
    m_fh = re.search(r"Flagged hours:\*\*\s*([\d,\.]+)", md_text)
    tot_e = int((m_tot_e.group(1) if m_tot_e else "0").replace(",", ""))
    tot_h = float((m_tot_h.group(1) if m_tot_h else "0").replace(",", ""))
    fn = int(m_fn.group(1) if m_fn else "0")
    fh = float((m_fh.group(1) if m_fh else "0").replace(",", ""))

    # tech display name = slug with dashes replaced and first-letter fix
    disp = slug.replace("-", ". ").replace("..", ".")
    # crude: if single dash between initial and name
    m = re.match(r"([A-Z])-(.+)", slug)
    if m:
        disp = f"{m.group(1)}. {m.group(2)}"

    doc = Document()
    set_default_style(doc)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    add_header_footer(doc, f"Technijian  |  Tech Training: {disp}  |  CONFIDENTIAL")

    add_color_bar(doc, CORE_BLUE, 6)
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(2.6))
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Personalized Tech Training", bold=True, size=26, color=DARK_CHARCOAL_R)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, disp, bold=True, size=22, color=CORE_BLUE_R)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"{YEAR} — Time Entry Quality Review", size=12, color=BRAND_GREY_R)
    for _ in range(6):
        doc.add_paragraph()
    add_color_bar(doc, CORE_ORANGE, 6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "CONFIDENTIAL — INDIVIDUAL TRAINING DOCUMENT", bold=True, size=10, color=BRAND_GREY_R)
    add_page_break(doc)

    # Summary
    add_section_header(doc, "Your 2026 Summary", "blue")
    doc.add_paragraph()
    add_body(doc, f"Across all clients you logged {tot_e:,} time entries totalling {tot_h:,.2f} hours "
                  f"in {YEAR}. Of these, {fn} entries ({fh:,.2f} hours) were flagged as outside the "
                  f"reasonable range for the work described. The purpose of this document is to help "
                  f"you understand those patterns so future entries are cleaner.")
    doc.add_paragraph()
    add_metric_cards(doc, [
        (f"{tot_e:,}", "Your Entries", CORE_BLUE_R),
        (f"{tot_h:,.0f}", "Your Hours", CORE_BLUE_R),
        (f"{fn}", "Flagged Entries", CORE_ORANGE_R),
        (f"{fh:,.1f}", "Flagged Hours", RED_R),
    ])

    # Flag breakdown by client
    flag_by_code = defaultdict(int)
    flag_by_client = defaultdict(lambda: {"n": 0, "h": 0.0})
    flag_by_cat = defaultdict(lambda: {"n": 0, "h": 0.0})
    hrs_idx = hdr.index("Hours")
    flags_idx = hdr.index("Flags")
    cap_idx = hdr.index("CategoryCap")
    cat_idx = hdr.index("Category")
    reasons_idx = hdr.index("Reasons")
    client_idx = hdr.index("Client")

    for r in rows:
        for c in r[flags_idx].split(";"):
            flag_by_code[c] += 1
        h = float(r[hrs_idx])
        flag_by_client[r[client_idx]]["n"] += 1
        flag_by_client[r[client_idx]]["h"] += h
        flag_by_cat[r[cat_idx]]["n"] += 1
        flag_by_cat[r[cat_idx]]["h"] += h

    add_page_break(doc)
    add_section_header(doc, "Flags breakdown", "orange")
    doc.add_paragraph()
    legend = {"H1": "Routine work > category cap",
              "H2": "Vague title with too many hours",
              "H3": "Single entry > 8 hours",
              "H4": "Daily total > 12 hours",
              "H5": "Duplicate same-ticket/same-day entries"}
    rows_code = [[c, flag_by_code[c], legend[c]] for c in ["H1", "H2", "H3", "H4", "H5"] if flag_by_code[c]]
    if rows_code:
        styled_table(doc, ["Code", "Count", "Meaning"], rows_code, col_widths=[0.8, 0.8, 5.0])

    doc.add_paragraph()
    add_section_header(doc, "Your flagged entries by client", "blue")
    doc.add_paragraph()
    rows_cli = [[k, v["n"], f"{v['h']:,.2f}"]
                for k, v in sorted(flag_by_client.items(), key=lambda kv: -kv[1]["h"])]
    if rows_cli:
        styled_table(doc, ["Client", "# Flagged", "Flagged Hours"], rows_cli, col_widths=[1.6, 1.2, 1.6])

    doc.add_paragraph()
    add_section_header(doc, "Your flagged work categories", "orange")
    doc.add_paragraph()
    rows_cat = [[k, v["n"], f"{v['h']:,.2f}"]
                for k, v in sorted(flag_by_cat.items(), key=lambda kv: -kv[1]["h"])[:10]]
    if rows_cat:
        styled_table(doc, ["Category", "# Flagged", "Flagged Hours"], rows_cat, col_widths=[3.8, 1.0, 1.4])

    # Top individual entries
    add_page_break(doc)
    add_section_header(doc, "Your 12 most-flagged entries", "blue")
    doc.add_paragraph()
    top = sorted(rows, key=lambda r: -float(r[hrs_idx]))[:12]
    rs = []
    for r in top:
        title = r[hdr.index("Title")][:60] + ("…" if len(r[hdr.index("Title")]) > 60 else "")
        reason = r[reasons_idx][:70] + ("…" if len(r[reasons_idx]) > 70 else "")
        rs.append([r[client_idx], r[0] if r[0].count("-") == 2 else r[hdr.index("Date")],
                   r[hrs_idx], r[cap_idx], title, reason])
    if rs:
        styled_table(doc,
                     ["Client", "Date", "Hours", "Cap", "Title", "Reason"],
                     rs, col_widths=[0.7, 0.85, 0.55, 0.45, 2.6, 2.2],
                     hrs_col=2, cap_col=3)

    # Personalized advice
    add_page_break(doc)
    add_section_header(doc, "Personalized Training Focus", "orange")
    doc.add_paragraph()
    advice = {
        "H1": ("Your most common issue is over-claiming time on routine work. ",
               "For patch-management alerts, agent version updates (CrowdStrike, ScreenConnect, "
               "MyRMM, ManageEngine), CPU/memory/disk threshold alerts, and similar auto-generated "
               "monitoring tickets, the expected resolution time is 0.25–1.0 hours. If an alert "
               "genuinely takes longer, your title must explain why (e.g. \"Critical CPU — "
               "investigated runaway SQL process\" instead of just \"Critical - CPU Utilization\")."),
        "H2": ("Your most common issue is vague ticket titles. ",
               "Titles like \"Help\", \"Fix\", \"Issue\", \"Test\" are not acceptable on any entry "
               "greater than 0.5 hour. Every title must describe what you actually did: the "
               "system/user affected and the action taken. Example — instead of \"Help\", use "
               "\"Help — Sarah's OneDrive sync stuck, rebuilt local cache\"."),
        "H3": ("Your most common issue is logging whole-day dumps into a single entry. ",
               "If you worked 8 or more hours across a day, break it into separate entries per "
               "activity with individual time blocks and descriptive titles."),
        "H4": ("Your daily totals are exceeding 12 hours across tickets. ",
               "This typically happens when overnight work is double-counted or when hours are "
               "claimed against the wrong date. Review your day-end entries before submitting "
               "the timesheet."),
        "H5": ("Your most common issue is creating multiple entries on the same ticket on the same day. ",
               "If you spent 4 hours on a single ticket in three separate blocks, consolidate them "
               "into one entry with a clear note covering the full scope of work. Multiple entries "
               "on the same ticket look like duplicate billing to a client reviewing the invoice."),
    }
    if flag_by_code:
        top_flag = max(flag_by_code.items(), key=lambda kv: kv[1])[0]
        prefix, body = advice[top_flag]
        add_body(doc, prefix, bold=True, color=DARK_CHARCOAL_R)
        add_body(doc, body)
        doc.add_paragraph()

    add_body(doc, "General time-entry rules for all techs:", bold=True, size=12, color=DARK_CHARCOAL_R)
    rules = [
        ("Descriptive titles required. ", "No standalone \"Help\", \"Test\", \"Fix\", \"Issue\" titles."),
        ("One entry per ticket per day. ", "Consolidate all work on the same ticket into one entry."),
        ("Cap routine alerts at 1 hour. ", "If it takes longer, the title must explain why."),
        ("Weekly Maintenance Window time must be split. ", "Across the covered clients, not wholesale logged to one."),
        ("Agent updates are ~0.25 hour per machine. ", "If an update ticket goes over, describe the problem."),
        ("Spot-check your own week before submitting. ",
         "If a title would confuse a client, rewrite it."),
    ]
    for prefix, body in rules:
        add_bullet(doc, body, bold_prefix=prefix)

    out = base / f"{slug}-Training.docx"
    doc.save(str(out))
    return out


# ===== MAIN =====
def main():
    by_client = ROOT / "by-client"
    by_tech = ROOT / "by-tech"

    n_xlsx = 0
    n_docx_client = 0
    n_docx_tech = 0

    # per-client
    for client_dir in sorted(by_client.iterdir()):
        if not client_dir.is_dir():
            continue
        try:
            build_client_xlsx(client_dir.name, client_dir)
            n_xlsx += 1
            build_client_docx(client_dir.name, client_dir)
            n_docx_client += 1
        except Exception as ex:
            print(f"  {client_dir.name}: ERR {ex}")

    # per-tech
    for tech_dir in sorted(by_tech.iterdir()):
        if not tech_dir.is_dir():
            continue
        try:
            build_tech_docx(tech_dir.name, tech_dir)
            n_docx_tech += 1
        except Exception as ex:
            print(f"  tech {tech_dir.name}: ERR {ex}")

    print(f"Done.")
    print(f"  Client XLSX: {n_xlsx}")
    print(f"  Client DOCX: {n_docx_client}")
    print(f"  Tech DOCX:   {n_docx_tech}")


if __name__ == "__main__":
    main()
