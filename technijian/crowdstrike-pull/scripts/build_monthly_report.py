"""Build a client-facing Technijian-branded monthly cybersecurity activity
report from the CrowdStrike per-client snapshots.

Inputs (per client, per target month YYYY-MM):
    clients/<code>/crowdstrike/monthly/<YYYY-MM>/alerts.json
    clients/<code>/crowdstrike/monthly/<YYYY-MM>/incidents.json
    clients/<code>/crowdstrike/monthly/<YYYY-MM>/pull_summary.json
    clients/<code>/crowdstrike/<latest-YYYY-MM-DD>/hosts.json  (optional)
    clients/<code>/crowdstrike/<latest-YYYY-MM-DD>/hosts.csv   (optional)

Output:
    clients/<code>/crowdstrike/monthly/<YYYY-MM>/<CODE>-CrowdStrike-Activity-<YYYY-MM>.docx

Usage:
    python build_monthly_report.py --month 2026-03
    python build_monthly_report.py --month 2026-03 --only BWH,ORX
    python build_monthly_report.py --month 2026-03 --skip ANI
    python build_monthly_report.py --all-months   # build every month with data

Each report is written for a non-technical business reader: here is what
CrowdStrike Falcon detected, here is what Technijian's team did, here is
what you should act on.
"""
from __future__ import annotations

import argparse
import json
import sys
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

HERE = Path(__file__).resolve().parent
PIPELINE_ROOT = HERE.parent
REPO = PIPELINE_ROOT.parent.parent
CLIENTS_ROOT = REPO / "clients"

sys.path.insert(0, str(HERE.parent.parent / "huntress-pull" / "scripts"))
import _brand as brand  # noqa: E402

# Shared vendor-news helper (lives under technijian/shared/scripts/)
SHARED_SCRIPTS = REPO / "technijian" / "shared" / "scripts"
sys.path.insert(0, str(SHARED_SCRIPTS))
import vendor_news  # noqa: E402

from docx.shared import Pt, RGBColor  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

MONTH_NAMES = ["", "January", "February", "March", "April", "May", "June",
               "July", "August", "September", "October", "November", "December"]


def month_label(month: str) -> str:
    y, m = month.split("-")
    return f"{MONTH_NAMES[int(m)]} {y}"


def load_json(path: Path, default):
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError:
        return default
    except Exception as e:
        print(f"  WARN: could not parse {path}: {e}")
        return default


def parse_iso(s: str | None) -> datetime | None:
    if not s:
        return None
    try:
        if s.endswith("Z"):
            s = s[:-1] + "+00:00"
        return datetime.fromisoformat(s)
    except Exception:
        return None


def fmt_dt(s: str | None) -> str:
    dt = parse_iso(s)
    return dt.strftime("%Y-%m-%d %H:%M UTC") if dt else "—"


def fmt_date(s: str | None) -> str:
    dt = parse_iso(s)
    return dt.strftime("%Y-%m-%d") if dt else "—"


def _severity_bucket(alert: dict) -> str:
    """Return canonical severity string: Critical / High / Medium / Low / Informational."""
    sn = alert.get("severity_name") or ""
    if sn:
        return sn.title()
    score = alert.get("severity") or 0
    if score >= 90:
        return "Critical"
    if score >= 70:
        return "High"
    if score >= 40:
        return "Medium"
    if score > 0:
        return "Low"
    return "Informational"


def _hostname(alert: dict) -> str:
    """Best-effort host extraction from an alert record."""
    dev = alert.get("device") or {}
    if isinstance(dev, dict):
        h = dev.get("hostname") or ""
        if h:
            return h
    dn = alert.get("display_name") or ""
    if " at " in dn:
        return dn.split(" at ")[0].strip()
    uname = alert.get("user_name") or ""
    if uname.endswith("$"):
        return uname[:-1]
    return dn or "Unknown"


def _sev_color(sev: str) -> RGBColor:
    s = sev.lower()
    if s == "critical":
        return brand.RED
    if s == "high":
        return RGBColor(0xCC, 0x55, 0x00)
    if s == "medium":
        return brand.CORE_ORANGE
    if s == "low":
        return brand.TEAL
    return brand.BRAND_GREY


def find_latest_daily_dir(cs_dir: Path) -> Path | None:
    candidates = [p for p in cs_dir.iterdir()
                  if p.is_dir() and p.name != "monthly"
                  and len(p.name) == 10 and p.name[4] == "-" and p.name[7] == "-"]
    if not candidates:
        return None
    candidates.sort(reverse=True)
    return candidates[0]


# ---------------------------------------------------------------------------
# Data container
# ---------------------------------------------------------------------------

class CsMonthData:
    def __init__(self, code: str, month: str, cs_dir: Path):
        self.code = code.upper()
        self.month = month
        self.cs_dir = cs_dir
        monthly_dir = cs_dir / "monthly" / month

        self.alerts: list[dict] = load_json(monthly_dir / "alerts.json", []) or []
        self.incidents: list[dict] = load_json(monthly_dir / "incidents.json", []) or []
        self.summary: dict = load_json(monthly_dir / "pull_summary.json", {})

        daily_dir = find_latest_daily_dir(cs_dir) if cs_dir.exists() else None
        self.hosts: list[dict] = load_json(daily_dir / "hosts.json", []) if daily_dir else []
        self.daily_date: str = daily_dir.name if daily_dir else ""

    @property
    def location_name(self) -> str:
        return (self.summary.get("Location_Name")
                or self.summary.get("child_name")
                or self.code)

    @property
    def has_data(self) -> bool:
        return bool(self.alerts or self.incidents or self.hosts)

    # --- Alert analysis ---

    def severity_counts(self) -> dict[str, int]:
        c: dict[str, int] = {}
        for a in self.alerts:
            s = _severity_bucket(a)
            c[s] = c.get(s, 0) + 1
        order = ["Critical", "High", "Medium", "Low", "Informational"]
        return {k: c[k] for k in order if k in c}

    def type_counts(self) -> dict[str, int]:
        c: dict[str, int] = {}
        labels = {
            "signal": "Detection Signal",
            "ldt": "Large Detection",
            "automated-lead": "Correlated Alert",
        }
        for a in self.alerts:
            raw = a.get("type") or "unknown"
            label = labels.get(raw, raw.replace("-", " ").title())
            c[label] = c.get(label, 0) + 1
        return dict(sorted(c.items(), key=lambda kv: kv[1], reverse=True))

    def top_hosts(self, n: int = 10) -> list[tuple[str, int]]:
        c: Counter[str] = Counter()
        for a in self.alerts:
            h = _hostname(a)
            if h and h != "Unknown":
                c[h] += 1
        return c.most_common(n)

    def mitre_tactics(self) -> dict[str, int]:
        c: dict[str, int] = {}
        for a in self.alerts:
            tactic = a.get("tactic") or ""
            if tactic:
                c[tactic] = c.get(tactic, 0) + 1
            else:
                for m in a.get("mitre_attack") or []:
                    t = m.get("tactic") or ""
                    if t:
                        c[t] = c.get(t, 0) + 1
                        break
        return dict(sorted(c.items(), key=lambda kv: kv[1], reverse=True))

    def alert_detail_rows(self, limit: int = 40) -> list[tuple]:
        rows = []
        for a in self.alerts[:limit]:
            rows.append((
                fmt_date(a.get("created_timestamp") or a.get("timestamp")),
                _severity_bucket(a),
                (a.get("name") or a.get("display_name") or "—")[:50],
                _hostname(a)[:30],
                (a.get("tactic") or "—")[:30],
                ("Closed" if a.get("is_closed") else (a.get("status") or "New")).title(),
            ))
        return rows

    # --- Host analysis (if daily snapshot available) ---

    def os_breakdown(self) -> dict[str, int]:
        c: dict[str, int] = {}
        for h in self.hosts:
            os = h.get("os_version") or h.get("platform_name") or "Unknown"
            c[os] = c.get(os, 0) + 1
        return dict(sorted(c.items(), key=lambda kv: kv[1], reverse=True))

    def agent_version_breakdown(self) -> dict[str, int]:
        c: dict[str, int] = {}
        for h in self.hosts:
            v = h.get("agent_version") or "Unknown"
            c[v] = c.get(v, 0) + 1
        return dict(sorted(c.items(), key=lambda kv: kv[1], reverse=True))

    def host_status_breakdown(self) -> dict[str, int]:
        c: dict[str, int] = {}
        for h in self.hosts:
            s = (h.get("status") or "unknown").title()
            c[s] = c.get(s, 0) + 1
        return dict(sorted(c.items(), key=lambda kv: kv[1], reverse=True))


# ---------------------------------------------------------------------------
# Report renderer
# ---------------------------------------------------------------------------

def render_report(data: CsMonthData, out_path: Path) -> dict:
    sev_counts = data.severity_counts()
    type_counts = data.type_counts()
    top_hosts = data.top_hosts()
    mitre = data.mitre_tactics()
    alert_rows = data.alert_detail_rows()

    critical_high = sev_counts.get("Critical", 0) + sev_counts.get("High", 0)
    unique_hosts = len({_hostname(a) for a in data.alerts
                        if _hostname(a) not in ("Unknown", "")})
    inc_total = len(data.incidents)
    alert_total = len(data.alerts)

    has_hosts = bool(data.hosts)
    host_count = len(data.hosts)

    doc = brand.new_branded_document()

    # ── COVER ──────────────────────────────────────────────────────────────
    brand.render_cover(
        doc,
        title="CrowdStrike Security Activity Report",
        subtitle=data.location_name,
        date_text=month_label(data.month),
        footer_note="Confidential — prepared by Technijian for the named client only.",
    )
    brand.add_page_break(doc)

    # ── EXECUTIVE SUMMARY ──────────────────────────────────────────────────
    brand.add_section_header(doc, "Executive Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    brand.add_run(
        p,
        f"This report summarizes the cybersecurity activity Technijian delivered "
        f"for {data.location_name} during {month_label(data.month)}. The "
        f"CrowdStrike Falcon endpoint detection and response (EDR) platform "
        f"protected your environment and Technijian's security operations team "
        f"reviewed every alert. Below is a plain-English summary of what was "
        f"detected, what we did about it, and where you stand today.",
        size=11,
    )

    # KPI cards
    cards = [
        (str(alert_total), "Total Alerts", brand.CORE_BLUE),
        (str(critical_high), "Critical / High", brand.RED if critical_high else brand.GREEN),
        (str(unique_hosts), "Hosts Affected", brand.CORE_ORANGE if unique_hosts else brand.GREEN),
        (str(len(type_counts)), "Alert Categories", brand.TEAL),
    ]
    brand.add_metric_card_row(doc, cards)

    if critical_high > 0:
        brand.add_callout_box(
            doc,
            f"{critical_high} Critical or High severity alert(s) were raised this month. "
            f"See the Threat Activity section for details and recommended actions.",
            accent_hex=brand.RED_HEX, bg_hex="FFEEEE",
        )
    else:
        brand.add_callout_box(
            doc,
            f"No Critical or High severity alerts fired this month. "
            f"Falcon continued to monitor your environment 24x7 and Technijian's "
            f"team reviewed all {alert_total} detection(s) raised.",
            accent_hex=brand.GREEN_HEX, bg_hex="EAF7EE",
        )

    brand.add_page_break(doc)

    # ── HOST INVENTORY (if daily snapshot available) ────────────────────────
    if has_hosts:
        brand.add_section_header(doc, "Endpoint Inventory")
        brand.add_body(
            doc,
            f"Endpoint inventory is a point-in-time snapshot taken {data.daily_date}. "
            f"CrowdStrike Falcon is installed and reporting on {host_count} endpoint(s) "
            f"in your environment.",
        )

        os_break = data.os_breakdown()
        status_break = data.host_status_breakdown()
        ver_break = data.agent_version_breakdown()

        if status_break:
            brand.styled_table(
                doc,
                ["Agent Status", "Endpoints"],
                list(status_break.items()),
                col_widths=[4.6, 1.4],
                status_col=0,
            )

        if os_break:
            doc.add_paragraph()
            brand.add_body(doc, "Operating system distribution",
                           bold=True, size=12, color=brand.DARK_CHARCOAL)
            brand.styled_table(
                doc,
                ["Operating System", "Endpoints"],
                list(os_break.items()),
                col_widths=[4.6, 1.4],
            )

        if ver_break and len(ver_break) > 1:
            doc.add_paragraph()
            brand.add_body(doc, "Falcon sensor version distribution",
                           bold=True, size=12, color=brand.DARK_CHARCOAL)
            brand.styled_table(
                doc,
                ["Sensor Version", "Endpoints"],
                list(ver_break.items()),
                col_widths=[4.6, 1.4],
            )

        brand.add_page_break(doc)

    # ── THREAT ACTIVITY ────────────────────────────────────────────────────
    brand.add_section_header(doc, "Threat Activity This Month")

    if alert_total == 0:
        brand.add_body(
            doc,
            f"No detections were raised in your environment during "
            f"{month_label(data.month)}. Falcon's continuous behavioral "
            f"monitoring remained active throughout the period.",
        )
    else:
        brand.add_body(
            doc,
            f"During {month_label(data.month)}, CrowdStrike Falcon raised "
            f"{alert_total} alert(s) across {unique_hosts} endpoint(s). "
            f"Each alert was reviewed by Technijian's security operations team. "
            f"The breakdown below shows what was detected and at what severity.",
        )

        # Severity breakdown
        if sev_counts:
            doc.add_paragraph()
            brand.add_body(doc, "Alerts by severity",
                           bold=True, size=12, color=brand.DARK_CHARCOAL)
            brand.styled_table(
                doc,
                ["Severity", "Count"],
                [(k, v) for k, v in sev_counts.items()],
                col_widths=[3.5, 2.5],
                status_col=0,
            )

        # Alert type breakdown
        if type_counts:
            doc.add_paragraph()
            brand.add_body(doc, "Detection categories",
                           bold=True, size=12, color=brand.DARK_CHARCOAL)
            brand.styled_table(
                doc,
                ["Category", "Count"],
                list(type_counts.items()),
                col_widths=[4.6, 1.4],
            )

        # MITRE tactics
        if mitre:
            doc.add_paragraph()
            brand.add_body(doc, "MITRE ATT&CK tactics observed",
                           bold=True, size=12, color=brand.DARK_CHARCOAL)
            brand.styled_table(
                doc,
                ["Tactic", "Alerts"],
                list(mitre.items()),
                col_widths=[4.6, 1.4],
            )

        # Top hosts
        if top_hosts:
            doc.add_paragraph()
            brand.add_body(doc, "Most active endpoints",
                           bold=True, size=12, color=brand.DARK_CHARCOAL)
            brand.styled_table(
                doc,
                ["Hostname", "Alert Count"],
                top_hosts,
                col_widths=[4.6, 1.4],
            )

        # Alert detail table
        if alert_rows:
            doc.add_paragraph()
            brand.add_body(doc, "Alert detail",
                           bold=True, size=12, color=brand.DARK_CHARCOAL)
            brand.styled_table(
                doc,
                ["Date", "Severity", "Alert Name", "Host", "Tactic", "Status"],
                alert_rows,
                col_widths=[0.9, 0.8, 1.7, 1.2, 1.2, 0.7],
                status_col=1,
            )
            if alert_total > 40:
                brand.add_body(
                    doc,
                    f"Showing 40 of {alert_total} alerts. Full list available on request — email support@technijian.com.",
                    size=10,
                )

    brand.add_page_break(doc)

    # ── WHAT WE DID ────────────────────────────────────────────────────────
    brand.add_section_header(doc, "What Technijian Did For You")
    bullets = [
        ("Continuous monitoring. ",
         f"CrowdStrike Falcon monitored your endpoints 24x7 using behavioral "
         f"AI and threat intelligence. Every process, network connection, and "
         f"file operation was inspected in real time."),
        ("Alert triage. ",
         f"Our security operations team reviewed all {alert_total} alert(s) "
         f"raised this month and classified each by severity and risk."),
        ("Threat intelligence. ",
         f"Detections were cross-referenced with CrowdStrike's global threat "
         f"intelligence database and MITRE ATT&CK framework to identify "
         f"adversary tactics and techniques."),
        ("Incident coordination. ",
         f"{'Technijian coordinated response actions for the ' + str(inc_total) + ' incident(s) flagged this month.' if inc_total else 'No escalated incidents required response this month. All detections were resolved within normal triage workflow.'}"),
        ("Posture data captured. ",
         f"Endpoint inventory, sensor version, and detection telemetry were "
         f"captured to support your security posture reviews and compliance "
         f"reporting."),
    ]
    for prefix, body in bullets:
        brand.add_bullet(doc, body, bold_prefix=prefix)

    # ── INDUSTRY NEWS & VENDOR INNOVATIONS ────────────────────────────────
    doc.add_paragraph()
    try:
        year_int, month_int = (int(x) for x in data.month.split("-"))
        vendor_news.render_section(doc, "crowdstrike", year_int, month_int, brand)
    except Exception:
        pass

    # ── RECOMMENDATIONS ────────────────────────────────────────────────────
    doc.add_paragraph()
    brand.add_section_header(doc, "Recommendations")

    if critical_high == 0 and not data.hosts:
        brand.add_body(
            doc,
            "No Critical or High severity alerts fired this month. "
            "Continue operating under your current Falcon policy and "
            "email support@technijian.com if you plan to add or remove "
            "endpoints.",
        )
    else:
        if critical_high > 0:
            brand.add_body(
                doc,
                f"Review and close open Critical/High alerts",
                bold=True, size=12, color=brand.DARK_CHARCOAL,
            )
            ch_rows = []
            for a in data.alerts:
                if _severity_bucket(a) in ("Critical", "High"):
                    ch_rows.append((
                        fmt_date(a.get("created_timestamp") or a.get("timestamp")),
                        _severity_bucket(a),
                        (a.get("name") or a.get("display_name") or "—")[:50],
                        _hostname(a)[:25],
                        ("Closed" if a.get("is_closed") else (a.get("status") or "New")).title(),
                    ))
            brand.styled_table(
                doc,
                ["Date", "Severity", "Alert Name", "Host", "Status"],
                ch_rows[:20],
                col_widths=[0.9, 0.85, 2.3, 1.5, 0.85],
                status_col=1,
            )
            brand.add_body(
                doc,
                "Recommended action: email support@technijian.com to review "
                "each open Critical/High alert with our team, confirm no "
                "further attacker activity is present, and close or suppress "
                "confirmed false positives.",
            )

        if not has_hosts:
            doc.add_paragraph()
            brand.add_body(
                doc,
                "Enable daily host inventory snapshots",
                bold=True, size=12, color=brand.DARK_CHARCOAL,
            )
            brand.add_body(
                doc,
                "The CrowdStrike daily pull scheduled task has not yet "
                "been registered on the Technijian monitoring workstation. "
                "Once registered, future reports will include full endpoint "
                "inventory with sensor version, OS distribution, and "
                "check-in age analysis. Email support@technijian.com to "
                "schedule this one-time setup step.",
            )

    # ── HOW THIS PROTECTION WORKS ──────────────────────────────────────────
    doc.add_paragraph()
    brand.add_section_header(doc, "How This Protection Works")
    brand.add_body(
        doc,
        "Your endpoints are protected by a layered security stack — "
        "CrowdStrike Falcon and Huntress working alongside each other — "
        "chosen and managed by Technijian as an MSP best-practice "
        "configuration. Here's what each piece does and why we run both.",
    )

    brand.add_body(doc, "CrowdStrike Falcon: AI-driven EDR with the Falcon Overwatch SOC", bold=True, size=12, color=brand.DARK_CHARCOAL)
    brand.add_body(
        doc,
        "CrowdStrike Falcon is the EDR (Endpoint Detection and Response) layer. "
        "Falcon's behavioral AI inspects every process, network connection, "
        "and file operation on your endpoints in real time and stops "
        "identified threats before they can execute. On top of the AI, "
        "CrowdStrike's Falcon Overwatch managed-threat-hunting team — the "
        "same elite NOC that protects Fortune 500s and government agencies — "
        "watches your environment 24×7 looking for the kind of slow, "
        "low-and-slow adversary tradecraft that automated detection alone can "
        "miss. Overwatch hunters reach out directly when they see something "
        "that needs immediate action.",
    )

    brand.add_body(doc, "Huntress: managed detection & response with a 24×7 SOC", bold=True, size=12, color=brand.DARK_CHARCOAL)
    brand.add_body(
        doc,
        "Huntress runs its own Security Operations Center staffed with threat "
        "hunters who review every signal that fires on your endpoints around "
        "the clock — not just nights and weekends, but 24×7×365. When the "
        "Huntress agent flags suspicious behavior (persistence mechanisms, "
        "process injection, credential abuse, known-bad files), the activity "
        "is examined by Huntress threat hunters BEFORE it ever reaches "
        "Technijian. False positives are filtered out at the source so that "
        "what does reach our team is signal, not noise. When Huntress sees "
        "something that needs attention, they coordinate response with "
        "Technijian and provide guided remediation playbooks.",
    )

    brand.add_body(doc, "Why both? Defense in depth.", bold=True, size=12, color=brand.DARK_CHARCOAL)
    brand.add_body(
        doc,
        "CrowdStrike and Huntress are deliberately complementary. CrowdStrike's "
        "Falcon is industry-leading nation-state-grade EDR with the broadest "
        "telemetry collection and the most mature threat-intel feed in the "
        "market. Huntress is tuned for the SMB and mid-market threat "
        "landscape — phishing-driven footholds, persistence tradecraft, "
        "credential theft, and the kinds of attacks that target managed-"
        "service customers. They use different detection styles, different "
        "telemetry sources, and different threat-intelligence pipelines.",
    )
    brand.add_body(
        doc,
        "Running them together gives your endpoints two independent sets of "
        "expert eyes — CrowdStrike's Overwatch team and Huntress's SOC — "
        "both watching 24×7. If one platform misses something, the other is "
        "extremely likely to catch it. This redundant, layered approach is "
        "the configuration Technijian recommends as best-in-class endpoint "
        "protection for clients who can't afford a security incident.",
    )

    # ── ABOUT ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    brand.add_section_header(doc, "About This Report", accent_color=brand.CORE_ORANGE)
    brand.add_body(
        doc,
        "Technijian delivers 24x7 endpoint detection and response (EDR) "
        "powered by the CrowdStrike Falcon platform. Activity captured here "
        "is pulled directly from the Falcon API for your tenant. Alert and "
        "incident data is filtered to the calendar month above. Endpoint "
        "inventory, when available, is a point-in-time snapshot taken at "
        "report generation.",
        size=10,
    )
    brand.add_body(doc, brand.CONTACT_LINE, size=10, color=brand.BRAND_GREY)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    return {
        "code": data.code,
        "month": data.month,
        "out_path": str(out_path),
        "alerts": alert_total,
        "critical_high": critical_high,
        "unique_hosts": unique_hosts,
        "incidents": inc_total,
        "has_host_snapshot": has_hosts,
    }


# ---------------------------------------------------------------------------
# Driver
# ---------------------------------------------------------------------------

def discover_months() -> list[str]:
    months: set[str] = set()
    for client in CLIENTS_ROOT.iterdir():
        mdir = client / "crowdstrike" / "monthly"
        if mdir.exists():
            for m in mdir.iterdir():
                if m.is_dir() and len(m.name) == 7 and m.name[4] == "-":
                    months.add(m.name)
    return sorted(months)


def discover_clients_with_month(month: str) -> list[str]:
    out: list[str] = []
    for child in CLIENTS_ROOT.iterdir():
        if not child.is_dir():
            continue
        if (child / "crowdstrike" / "monthly" / month).exists():
            out.append(child.name.lower())
    return sorted(out)


def parse_args() -> argparse.Namespace:
    ap = argparse.ArgumentParser(
        description="Build branded CrowdStrike monthly activity report per client."
    )
    grp = ap.add_mutually_exclusive_group(required=True)
    grp.add_argument("--month", help="YYYY-MM target month")
    grp.add_argument("--all-months", action="store_true",
                     help="Build every month that has data")
    ap.add_argument("--only", help="comma-separated LocationCodes")
    ap.add_argument("--skip", action="append", default=[],
                    help="LocationCode to skip (repeatable)")
    return ap.parse_args()


def run_month(month: str, only_codes: set[str] | None,
              skip_codes: set[str]) -> list[dict]:
    codes = discover_clients_with_month(month)
    if not codes:
        print(f"  No clients have CrowdStrike data for {month}")
        return []

    print(f"\n[{datetime.now():%H:%M:%S}] {month} — {len(codes)} client(s)")
    results: list[dict] = []
    for code in codes:
        upper = code.upper()
        if upper in skip_codes:
            continue
        if only_codes is not None and upper not in only_codes:
            continue
        cs_dir = CLIENTS_ROOT / code / "crowdstrike"
        data = CsMonthData(upper, month, cs_dir)
        if not data.has_data:
            print(f"  [skip] {upper:<6s} no data")
            continue
        out_path = cs_dir / "monthly" / month / f"{upper}-CrowdStrike-Activity-{month}.docx"
        try:
            res = render_report(data, out_path)
        except Exception as e:
            import traceback
            print(f"  [ERR ] {upper:<6s} {e}")
            traceback.print_exc()
            continue
        print(
            f"  [ok  ] {upper:<6s} alerts={res['alerts']:>4d}"
            f" crit/hi={res['critical_high']:>3d}"
            f" hosts={res['unique_hosts']:>3d}"
            f" snap={'Y' if res['has_host_snapshot'] else 'N'}"
            f"  -> {out_path.relative_to(REPO)}"
        )
        results.append(res)
    return results


def main() -> int:
    args = parse_args()
    skip_codes = {s.strip().upper() for raw in args.skip for s in raw.split(",") if s.strip()}
    only_codes: set[str] | None = None
    if args.only:
        only_codes = {s.strip().upper() for s in args.only.split(",") if s.strip()}

    months = [args.month] if args.month else discover_months()
    if not months:
        print("No CrowdStrike monthly data found. Run backfill_crowdstrike.py first.")
        return 1

    all_results: list[dict] = []
    for month in months:
        all_results.extend(run_month(month, only_codes, skip_codes))

    print(f"\n[{datetime.now():%H:%M:%S}] DONE — {len(all_results)} report(s) generated")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
