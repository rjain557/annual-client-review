"""Build a Technijian-branded weekly Word document for each tech with at least
one flagged entry this cycle.

Reuses the styling helpers in technijian/tech-training/scripts/_build-docx-report.py
but rebuilds the document targeting a single tech (rather than a single client).

Output:
    technijian/weekly-audit/<cycle>/by-tech/<slug>/<slug>-Weekly-Training.docx

Usage:
    python 3_build_weekly_docs.py
    python 3_build_weekly_docs.py --cycle 2026-W18 --only S-Kumar
"""
from __future__ import annotations

import argparse
import csv
import sys
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))
from _shared import cycle_dir, cycle_id_for, now_pacific

LOGO = Path(r"C:\VSCode\tech-branding\tech-branding\assets\logos\png\technijian-logo-full-color-600x125.png")

# Brand colors
CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0xCC, 0x00, 0x00)
GREEN = RGBColor(0x28, 0xA7, 0x45)
OFF_WHITE = "F8F9FA"
LIGHT_GREY = "E9ECEF"
FONT = "Open Sans"


def shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def remove_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for b in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{b}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tc_pr.append(borders)


def set_cell_border_color(cell, hex_color=LIGHT_GREY, sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for b in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{b}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:color"), hex_color)
        borders.append(e)
    tc_pr.append(borders)


def add_run(paragraph, text, *, bold=False, size=11, color=BRAND_GREY, italic=False):
    run = paragraph.add_run(text)
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def add_body(doc, text, *, bold=False, size=11, color=BRAND_GREY, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bold=bold, size=size, color=color)
    return p


def add_bullet(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=11, color=DARK_CHARCOAL)
        add_run(p, text, size=11, color=BRAND_GREY)
    else:
        add_run(p, text, size=11, color=BRAND_GREY)
    return p


def add_color_bar(doc, hex_color, height_pt=4):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade(cell, hex_color)
    remove_borders(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("")
    r.font.size = Pt(height_pt)
    return t


def add_section_header(doc, title, accent_color=CORE_BLUE):
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Emu(60000)
    t.columns[1].width = Emu(5700000)
    bar_cell, title_cell = t.rows[0].cells
    bar_cell.width = Emu(60000)
    title_cell.width = Emu(5700000)
    shade(bar_cell, "006DB6" if accent_color == CORE_BLUE else "F67D4B")
    remove_borders(bar_cell)
    remove_borders(title_cell)
    p = title_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "  " + title, bold=True, size=14, color=accent_color)
    return t


def set_col_widths(table, widths_in_inches):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_in_inches):
                cell.width = Inches(widths_in_inches[i])


def styled_table(doc, headers, rows, *, col_widths=None,
                  hrs_col=None, cap_col=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = False
    if col_widths:
        set_col_widths(t, col_widths)
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        shade(cell, "006DB6")
        set_cell_border_color(cell, LIGHT_GREY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, h, bold=True, size=10, color=WHITE)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            if i % 2 == 1:
                shade(cell, OFF_WHITE)
            set_cell_border_color(cell, LIGHT_GREY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            text = str(val)
            color = BRAND_GREY
            if isinstance(val, (int, float)):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if hrs_col is not None and j == hrs_col and cap_col is not None:
                try:
                    h_val = float(val)
                    c_val = float(row[cap_col])
                    if h_val > c_val * 1.5:
                        color = RED
                    elif h_val > c_val:
                        color = CORE_ORANGE
                except (ValueError, TypeError):
                    pass
            add_run(p, text, size=9.5, color=color)
    return t


def add_metric_card_row(doc, cards):
    t = doc.add_table(rows=1, cols=len(cards))
    t.autofit = False
    for i, (value, label, color) in enumerate(cards):
        cell = t.rows[0].cells[i]
        shade(cell, "F8F9FA")
        remove_borders(cell)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(12)
        p1.paragraph_format.space_after = Pt(0)
        add_run(p1, value, bold=True, size=24, color=color)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(12)
        add_run(p2, label, size=10, color=BRAND_GREY)
    return t


def set_default_style(doc):
    s = doc.styles["Normal"]
    s.font.name = FONT
    s.font.size = Pt(11)
    s.font.color.rgb = BRAND_GREY


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def add_footer(doc):
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Technijian  |  Internal Use - Weekly Tech Training  |  18 Technology Dr., Ste 141, Irvine, CA 92618",
            size=9, color=BRAND_GREY)


def add_header_logo(doc):
    p = doc.sections[0].header.paragraphs[0]
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(1.7))
    else:
        add_run(p, "TECHNIJIAN", bold=True, size=14, color=CORE_BLUE)


def slug_to_display(slug: str) -> str:
    parts = slug.split("-")
    if len(parts) >= 2:
        return f"{parts[0]}. {' '.join(parts[1:])}"
    return slug


def build_doc_for_tech(cycle: str, cycle_root: Path, slug: str) -> Path | None:
    folder = cycle_root / "by-tech" / slug
    csv_path = folder / "flagged-entries.csv"
    if not csv_path.exists():
        return None

    with csv_path.open("r", encoding="utf-8", newline="") as f:
        rows = list(csv.DictReader(f))
    if not rows:
        return None

    out = folder / f"{slug}-Weekly-Training.docx"
    display = slug_to_display(slug)

    flagged_n = len(rows)
    flagged_h = sum(float(r["Hours"]) for r in rows)
    by_client: dict = {}
    by_flag: dict = {"H1": 0, "H2": 0, "H3": 0, "H4": 0, "H5": 0}
    for r in rows:
        c = r["Client"].upper()
        by_client.setdefault(c, {"n": 0, "h": 0.0})
        by_client[c]["n"] += 1
        by_client[c]["h"] += float(r["Hours"])
        for code in (r["Flags"] or "").split(";"):
            if code in by_flag:
                by_flag[code] += 1

    today = now_pacific().strftime("%A, %B %d, %Y")
    invoice_evening = now_pacific().strftime("tonight (%B %d) at end of day")

    doc = Document()
    set_default_style(doc)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    add_header_logo(doc)
    add_footer(doc)

    # ---- Cover ----
    add_color_bar(doc, "006DB6", height_pt=6)
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(3.0))
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Weekly Tech Time-Entry Review", bold=True, size=26, color=DARK_CHARCOAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, display, bold=True, size=22, color=CORE_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"Cycle {cycle}  |  {today}", size=12, color=BRAND_GREY)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Adjust your flagged entries before tonight's invoice run",
             bold=True, size=14, color=CORE_ORANGE)

    for _ in range(6):
        doc.add_paragraph()
    add_color_bar(doc, "F67D4B", height_pt=6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "INTERNAL USE - WEEKLY TECH TRAINING", bold=True, size=10, color=BRAND_GREY)
    add_page_break(doc)

    # ---- What this is ----
    add_section_header(doc, "What this is", CORE_BLUE)
    doc.add_paragraph()
    add_body(doc,
        f"Every Friday morning we scan the prior week of time entries logged across "
        f"all clients and flag any that look like they would not pass a client review. "
        f"This is your personalized list for cycle {cycle}.")
    add_body(doc,
        f"Weekly in-contract invoices go out {invoice_evening}. Anything you adjust "
        f"in the Client Portal before then will appear on the invoice the way you "
        f"want it. Anything left as-is will appear exactly as logged.",
        color=DARK_CHARCOAL)

    doc.add_paragraph()
    add_metric_card_row(doc, [
        (f"{flagged_n}", "Flagged Entries", CORE_ORANGE),
        (f"{flagged_h:.1f}", "Flagged Hours", RED),
        (f"{len(by_client)}", "Clients Affected", CORE_BLUE),
    ])

    # ---- What to do ----
    add_page_break(doc)
    add_section_header(doc, "What to do for each flagged entry", CORE_ORANGE)
    doc.add_paragraph()
    add_body(doc, "For every entry on your flagged list, choose one of two actions:",
             color=DARK_CHARCOAL)
    add_bullet(doc,
        "if the work really did take that long, rewrite the title so a client reading "
        "the invoice can see why. The CSV column \"Suggested_Title_If_Hours_Stay\" gives "
        "a model rewrite for each entry.",
        bold_prefix="Option A - keep hours, rewrite title: ")
    add_bullet(doc,
        "drop the hours to the suggested amount in the \"Suggested_Adjusted_Hours\" column. "
        "You will not be paid for the difference, but the entry will appear cleanly on the invoice.",
        bold_prefix="Option B - reduce hours to the cap: ")
    doc.add_paragraph()
    add_body(doc,
        "If you do nothing, the entry goes onto tonight's invoice as written and may be "
        "flagged by the client. Repeated patterns will be reviewed with your team lead.",
        color=DARK_CHARCOAL)

    # ---- Flag breakdown ----
    add_page_break(doc)
    add_section_header(doc, "Why your entries were flagged", CORE_BLUE)
    doc.add_paragraph()
    legend = {
        "H1": ("Routine work over the category cap",
               "patch tickets, agent updates, monitoring alerts logged at hours far above what the work usually takes."),
        "H2": ("Vague title with too many hours",
               "titles like \"Help\", \"Fix\", \"Issue\", \"Test\" claimed at more than 30 minutes."),
        "H3": ("Single entry over 8 hours",
               "one block of time bigger than 8 hours - usually a whole-day dump that should be split."),
        "H4": ("Daily total over 12 hours",
               "your sum across all tickets in a single day exceeded 12 hours - check date assignments."),
        "H5": ("Same ticket / same day stacked",
               "multiple entries on the same ticket on the same day totalling more than 2x the cap."),
    }
    rows_legend = []
    for code in ("H1", "H2", "H3", "H4", "H5"):
        if by_flag[code]:
            label, desc = legend[code]
            rows_legend.append([code, by_flag[code], label, desc])
    if rows_legend:
        styled_table(doc,
            headers=["Code", "Count", "Label", "What it means"],
            rows=rows_legend,
            col_widths=[0.5, 0.6, 1.8, 3.6])

    # ---- Per-client breakdown ----
    if by_client:
        doc.add_paragraph()
        add_section_header(doc, "Flagged entries by client", CORE_ORANGE)
        doc.add_paragraph()
        rows_client = sorted(
            ([c, v["n"], round(v["h"], 2)] for c, v in by_client.items()),
            key=lambda r: -r[2])
        styled_table(doc,
            headers=["Client", "Entries", "Hours"],
            rows=rows_client,
            col_widths=[2.0, 1.0, 1.0])

    # ---- All flagged entries with suggestions ----
    add_page_break(doc)
    add_section_header(doc, "Your flagged entries with suggested fixes", CORE_BLUE)
    doc.add_paragraph()
    rows.sort(key=lambda r: -float(r["Hours"]))
    table_rows = []
    for r in rows:
        title = r["Title"][:60] + ("..." if len(r["Title"]) > 60 else "")
        suggested_t = r.get("Suggested_Title_If_Hours_Stay") or ""
        suggested_t = suggested_t[:60] + ("..." if len(suggested_t) > 60 else "")
        table_rows.append([
            r["Client"],
            r["Date"],
            f"{float(r['Hours']):.2f}",
            r.get("CategoryCap", ""),
            r.get("Suggested_Adjusted_Hours", ""),
            title,
            suggested_t,
        ])
    styled_table(doc,
        headers=["Client", "Date", "Hours", "Cap", "Suggest hrs",
                  "Your title", "Suggested rewrite"],
        rows=table_rows,
        col_widths=[0.7, 0.85, 0.55, 0.45, 0.7, 1.85, 1.9],
        hrs_col=2, cap_col=3)

    # ---- General rules reminder ----
    add_page_break(doc)
    add_section_header(doc, "Six rules that prevent flags", CORE_ORANGE)
    doc.add_paragraph()
    add_bullet(doc, "no standalone \"Help\", \"Test\", \"Fix\", \"Issue\".",
               bold_prefix="Descriptive titles - ")
    add_bullet(doc, "consolidate all work on the same ticket into one entry.",
               bold_prefix="One entry per ticket per day - ")
    add_bullet(doc, "if it takes longer, the title must explain why.",
               bold_prefix="Cap routine alerts at 1 hour - ")
    add_bullet(doc, "split it across the covered clients, not wholesale to one.",
               bold_prefix="Weekly Maintenance Window time - ")
    add_bullet(doc, "if it goes long, describe the problem in the title.",
               bold_prefix="Agent updates are ~0.25 hr/machine - ")
    add_bullet(doc, "if a title would confuse a client, rewrite it before submitting.",
               bold_prefix="Spot-check your own week - ")

    doc.save(str(out))
    return out


def parse_args() -> argparse.Namespace:
    ap = argparse.ArgumentParser()
    ap.add_argument("--cycle", help="cycle ID (default = current ISO week)")
    ap.add_argument("--only", help="comma-separated tech slugs to limit to")
    return ap.parse_args()


def main() -> int:
    args = parse_args()
    cycle = args.cycle or cycle_id_for()
    cycle_root = cycle_dir(cycle)
    by_tech = cycle_root / "by-tech"
    if not by_tech.exists():
        print(f"  no by-tech folder at {by_tech}; run 2_audit_weekly.py first.")
        return 1

    only = None
    if args.only:
        only = {s.strip() for s in args.only.split(",") if s.strip()}

    print(f"[{datetime.now():%H:%M:%S}] building weekly Word docs for cycle {cycle}")
    built = 0
    for d in sorted(by_tech.iterdir()):
        if not d.is_dir():
            continue
        if only is not None and d.name not in only:
            continue
        out = build_doc_for_tech(cycle, cycle_root, d.name)
        if out:
            built += 1
            print(f"  built {d.name:<22} -> {out.name}")
    print(f"\n[{datetime.now():%H:%M:%S}] DONE: {built} docs")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
