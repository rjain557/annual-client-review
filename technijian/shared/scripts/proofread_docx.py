#!/usr/bin/env python3
"""
proofread_docx.py

Structural + content proofreader for Technijian-branded DOCX reports.

Checks performed
----------------
  1. File exists and is readable
  2. File size >= min_kb (default 10 KB)
  3. Document opens without error
  4. Cover page: first non-empty paragraph contains a title (not placeholder)
  5. Section headers: expected sections are present (case-insensitive)
  6. Tables: each table has at least a header row
  7. Table widths: no table exceeds the usable page width (6.5" for letter/1" margins)
  8. Placeholder guard: no TODO / TBD / [placeholder] / [Your Name] text
  9. Mojibake guard: no common cp1252 encoding artifacts (A--- A,,"  etc.)
 10. Callout boxes: at least one callout table present (single-cell table)
 11. Metric cards row: at least one multi-cell table in first half of doc

Usage
-----
  python technijian/shared/scripts/proofread_docx.py path/to/report.docx

  # Specify expected sections to verify
  python ... --sections "Executive Summary,DLP Policy Status,Recommendations"

  # Adjust minimum file size
  python ... --min-kb 20

  # Fail hard on warnings (for CI / automated gates)
  python ... --strict

Exit codes
----------
  0  all checks passed
  1  one or more checks failed
  2  file not found or not openable
"""

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print(json.dumps({"error": "python-docx not installed", "passed": False}))
    sys.exit(2)

# ── layout constants ──────────────────────────────────────────────────────────

# 1 inch = 914 400 EMU (English Metric Units, the python-docx internal unit).
# Standard US Letter (8.5") with 1" left + 1" right margins → 6.5" usable.
_ONE_INCH_EMU = 914_400
_USABLE_PAGE_INCHES = 6.5
_USABLE_PAGE_EMU = int(_USABLE_PAGE_INCHES * _ONE_INCH_EMU)

# ── placeholder patterns ───────────────────────────────────────────────────────

PLACEHOLDER_PATTERNS = [
    r"\[Your Name\]", r"\[Recipient\]", r"\bTODO\b", r"\bTBD\b",
    r"\bPLACEHOLDER\b", r"\[placeholder\]", r"\[insert\b",
    r"\[client name\]", r"\[date\]",
]

# ── mojibake sentinel characters ───────────────────────────────────────────────

MOJIBAKE_PATTERNS = [
    r"A,+\"\"",         # A,,"
    r"A---",            # A—— rendered wrong
    r"\xc3\xa2",        # UTF-8 double-encoded
    r"â€[œ\x9c\x94\x99\x9d\x93]",  # common UTF-8->cp1252 artifacts
]

# ── helpers ────────────────────────────────────────────────────────────────────

def _iter_text(doc):
    """Yield (paragraph_or_cell, text) for all text in the document."""
    for para in doc.paragraphs:
        yield para, para.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para, para.text


def _is_heading(para):
    """True if the paragraph looks like a section heading."""
    if para.style and "heading" in para.style.name.lower():
        return True
    # Bold + size >= 13 heuristic (matches _brand.add_section_header)
    for run in para.runs:
        if run.bold and run.font.size and run.font.size.pt >= 12:
            return True
    return False


def _paragraph_text(doc):
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def _all_text_in_doc(doc):
    """Yield all non-empty text from paragraphs AND table cells."""
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            yield t
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        yield t


def _table_cell_count(table):
    if not table.rows:
        return 0
    return len(table.rows[0].cells)


def _is_color_bar_table(table):
    """True if this is a single-cell, single-row decorative color bar (no text)."""
    if len(table.rows) != 1 or len(table.rows[0].cells) != 1:
        return False
    return not table.rows[0].cells[0].text.strip()


# ── checks ─────────────────────────────────────────────────────────────────────

def check_file(path: Path, min_kb: float):
    issues = []
    warnings = []
    size_kb = path.stat().st_size / 1024
    if size_kb < min_kb:
        issues.append(f"File too small: {size_kb:.1f} KB (expected >= {min_kb} KB)")
    return issues, warnings, size_kb


def check_cover(doc):
    """First non-empty paragraph should look like a title."""
    issues = []
    warnings = []
    texts = _paragraph_text(doc)
    if not texts:
        issues.append("Document appears to be blank (no paragraph text)")
        return issues, warnings
    first = texts[0]
    if not first or len(first) < 4:
        issues.append(f"Cover page: first paragraph is too short or empty: {repr(first)}")
    for pat in PLACEHOLDER_PATTERNS:
        if re.search(pat, first, re.IGNORECASE):
            issues.append(f"Cover page contains placeholder text: {repr(first)}")
            break
    return issues, warnings


def check_sections(doc, expected_sections):
    """Verify each expected section header appears in the document.
    Searches all text including table cells (brand section headers are table-based)."""
    issues = []
    warnings = []
    if not expected_sections:
        return issues, warnings
    all_text = [t.lower() for t in _all_text_in_doc(doc)]
    for section in expected_sections:
        needle = section.lower().strip()
        if not any(needle in t for t in all_text):
            issues.append(f"Missing section: '{section}'")
    return issues, warnings


def check_tables(doc):
    """Every data table should have content. Decorative color bars are excluded."""
    issues = []
    warnings = []
    for i, table in enumerate(doc.tables):
        if len(table.rows) == 0:
            issues.append(f"Table {i+1}: completely empty (no rows)")
            continue
        # Skip decorative color bar tables (single-cell, no text)
        if _is_color_bar_table(table):
            continue
        # Skip section header tables (2-cell row: empty bar + title text)
        if len(table.rows) == 1 and len(table.rows[0].cells) == 2:
            bar_text = table.rows[0].cells[0].text.strip()
            title_text = table.rows[0].cells[1].text.strip()
            if not bar_text and title_text:
                continue  # legitimate section header bar+title
        # All text across all cells
        all_text = " ".join(
            c.text.strip() for row in table.rows for c in row.cells
        )
        if not all_text:
            issues.append(f"Table {i+1}: all cells are blank")
    return issues, warnings


def check_table_widths(doc):
    """Fail if any data table's column widths sum to more than the usable page
    width (6.5"). Overflow causes content to be cut off when printed or viewed
    at page width.

    Skips:
      - Decorative color-bar tables (single cell, no text)
      - Section-header tables (2 cells: empty bar + title text)
      - Tables where no cell has an explicit width set
    """
    issues = []
    warnings = []
    for i, table in enumerate(doc.tables):
        if not table.rows:
            continue
        if _is_color_bar_table(table):
            continue
        row0 = table.rows[0]
        # Skip section-header bar+title tables
        if len(row0.cells) == 2 and not row0.cells[0].text.strip():
            continue
        cell_widths = [c.width for c in row0.cells if c.width is not None]
        if not cell_widths:
            continue
        total_emu = sum(cell_widths)
        if total_emu > _USABLE_PAGE_EMU:
            total_in = total_emu / _ONE_INCH_EMU
            issues.append(
                f"Table {i+1} ({len(row0.cells)} columns, {total_in:.2f}\" wide) "
                f"overflows the {_USABLE_PAGE_INCHES:.1f}\" usable page width — "
                "content will be cut off when printed or viewed at full-page width"
            )
    return issues, warnings


def check_placeholders(doc):
    issues = []
    warnings = []
    combined_pattern = "|".join(PLACEHOLDER_PATTERNS)
    for _, text in _iter_text(doc):
        if re.search(combined_pattern, text, re.IGNORECASE):
            issues.append(f"Placeholder text found: {repr(text[:80])}")
    return issues, warnings


def check_mojibake(doc):
    issues = []
    warnings = []
    for _, text in _iter_text(doc):
        for pat in MOJIBAKE_PATTERNS:
            if re.search(pat, text):
                issues.append(f"Possible mojibake encoding artifact: {repr(text[:80])}")
                break
    return issues, warnings


def check_callout_boxes(doc):
    """At least one single-cell table should be present (callout / summary box)."""
    issues = []
    warnings = []
    single_cell = [t for t in doc.tables if _table_cell_count(t) == 1]
    if not single_cell:
        warnings.append("No callout box tables found (single-cell table). "
                        "Expected at least one status/callout section.")
    return issues, warnings


def check_metric_cards(doc):
    """At least one multi-column table in the first half of the document."""
    issues = []
    warnings = []
    tables = doc.tables
    first_half = tables[: max(1, len(tables) // 2)]
    multi_col = [t for t in first_half if _table_cell_count(t) >= 3]
    if not multi_col:
        warnings.append("No metric card table found in first half of document "
                        "(expected multi-column KPI row near Executive Summary).")
    return issues, warnings


# ── main ──────────────────────────────────────────────────────────────────────

DEFAULT_SECTIONS = [
    "Executive Summary",
    "Endpoint Monitoring Coverage",
    "DLP Policy Status",
    "Activity Summary",
    "Insider-Threat Risk Assessment",
    "What Technijian Did For You",
    "Recommendations",
    "About This Report",
]


def proofread(path: Path, expected_sections, min_kb: float, strict: bool):
    result = {
        "file": str(path),
        "passed": True,
        "issues": [],
        "warnings": [],
        "checks_passed": 0,
        "checks_total": 0,
        "size_kb": 0.0,
        "paragraphs": 0,
        "tables": 0,
    }

    # File existence
    if not path.exists():
        result["passed"] = False
        result["issues"].append(f"File not found: {path}")
        return result

    # File size
    fi, fw, size_kb = check_file(path, min_kb)
    result["size_kb"] = round(size_kb, 1)
    result["issues"].extend(fi)
    result["warnings"].extend(fw)

    # Open document
    try:
        doc = Document(str(path))
    except Exception as e:
        result["passed"] = False
        result["issues"].append(f"Failed to open DOCX: {e}")
        return result

    result["paragraphs"] = len([p for p in doc.paragraphs if p.text.strip()])
    result["tables"] = len(doc.tables)

    checks = [
        ("cover page",        check_cover(doc)),
        ("section headers",   check_sections(doc, expected_sections)),
        ("table structure",   check_tables(doc)),
        ("table widths",      check_table_widths(doc)),
        ("placeholder text",  check_placeholders(doc)),
        ("mojibake",          check_mojibake(doc)),
        ("callout boxes",     check_callout_boxes(doc)),
        ("metric cards",      check_metric_cards(doc)),
    ]

    for name, (issues, warnings) in checks:
        result["checks_total"] += 1
        if issues:
            result["issues"].extend(issues)
        else:
            result["checks_passed"] += 1
        result["warnings"].extend(warnings)

    if result["issues"]:
        result["passed"] = False
    elif strict and result["warnings"]:
        result["passed"] = False

    return result


def _print_result(result, verbose):
    status = "PASS" if result["passed"] else "FAIL"
    print(f"\n{'='*60}")
    print(f"  {status}  {Path(result['file']).name}")
    print(f"  {result['checks_passed']}/{result['checks_total']} checks passed  "
          f"| {result['size_kb']} KB | "
          f"{result['paragraphs']} paragraphs | {result['tables']} tables")

    if result["issues"]:
        print("\n  Issues:")
        for issue in result["issues"]:
            print(f"    [FAIL] {issue}")

    if result["warnings"] and verbose:
        print("\n  Warnings:")
        for warn in result["warnings"]:
            print(f"    [WARN] {warn}")

    print(f"{'='*60}")


def main():
    p = argparse.ArgumentParser(description="Proofread a Technijian DOCX report")
    p.add_argument("files", nargs="+", help="DOCX file(s) to check")
    p.add_argument("--sections", default=",".join(DEFAULT_SECTIONS),
                   help="Comma-separated expected section headers")
    p.add_argument("--min-kb", type=float, default=10.0,
                   help="Minimum file size in KB (default: 10)")
    p.add_argument("--strict", action="store_true",
                   help="Treat warnings as failures")
    p.add_argument("--json", action="store_true",
                   help="Output JSON results only")
    p.add_argument("--quiet", action="store_true",
                   help="Suppress warnings in output")
    args = p.parse_args()

    expected = [s.strip() for s in args.sections.split(",") if s.strip()]
    results = []
    overall_pass = True

    for filepath in args.files:
        r = proofread(Path(filepath), expected, args.min_kb, args.strict)
        results.append(r)
        if not r["passed"]:
            overall_pass = False
        if not args.json:
            _print_result(r, verbose=not args.quiet)

    if args.json:
        print(json.dumps(results if len(results) > 1 else results[0], indent=2))
    else:
        total = len(results)
        passed = sum(1 for r in results if r["passed"])
        print(f"\nSummary: {passed}/{total} reports passed proofreading.")

    sys.exit(0 if overall_pass else 1)


if __name__ == "__main__":
    main()
