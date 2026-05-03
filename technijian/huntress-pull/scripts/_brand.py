"""Reusable Technijian brand helpers for python-docx reports.

Distilled from `technijian/tech-training/scripts/_build-docx-report.py` and
the JS brand kit at `c:/vscode/tech-branding/tech-branding/scripts/brand-helpers.js`.
Keep this small and dependency-light so any pipeline (huntress-pull, weekly-audit,
monthly-pull) can `import _brand` and produce on-brand DOCX output without
duplicating low-level OXML plumbing.
"""
from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.document import Document as _Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---- Brand palette -------------------------------------------------------

CORE_BLUE = RGBColor(0x00, 0x6D, 0xB6)
CORE_ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
TEAL = RGBColor(0x1E, 0xAA, 0xC8)
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GREY = RGBColor(0x59, 0x59, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x28, 0xA7, 0x45)
RED = RGBColor(0xCC, 0x00, 0x00)

OFF_WHITE_HEX = "F8F9FA"
LIGHT_GREY_HEX = "E9ECEF"
CORE_BLUE_HEX = "006DB6"
CORE_ORANGE_HEX = "F67D4B"
TEAL_HEX = "1EAAC8"
GREEN_HEX = "28A745"
RED_HEX = "CC0000"

FONT = "Open Sans"

LOGO_PATH = Path(r"C:\VSCode\tech-branding\tech-branding\assets\logos\png\technijian-logo-full-color-600x125.png")

CONTACT_LINE = ("Technijian  |  18 Technology Dr., Ste 141, Irvine, CA 92618"
                "  |  949.379.8500  |  technijian.com")


# ---- Low-level OXML helpers ---------------------------------------------

def shade(cell, hex_color: str) -> None:
    """Set the background fill of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def remove_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for b in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{b}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tc_pr.append(borders)


def set_cell_border_color(cell, hex_color: str = LIGHT_GREY_HEX, sz: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for b in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{b}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:color"), hex_color)
        borders.append(e)
    tc_pr.append(borders)


def add_run(paragraph, text: str, *, bold: bool = False, size: float = 11,
             color: RGBColor | None = BRAND_GREY, font: str = FONT,
             italic: bool = False):
    run = paragraph.add_run(text)
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


# ---- Mid-level paragraph + table helpers --------------------------------

def add_body(doc: _Document, text: str, *, bold: bool = False, size: float = 11,
              color: RGBColor = BRAND_GREY, align: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bold=bold, size=size, color=color)
    return p


def add_bullet(doc: _Document, text: str, *, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=11, color=DARK_CHARCOAL)
        add_run(p, text, size=11, color=BRAND_GREY)
    else:
        add_run(p, text, size=11, color=BRAND_GREY)
    return p


def add_color_bar(doc: _Document, hex_color: str, height_pt: float = 4):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade(cell, hex_color)
    remove_borders(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("")
    r.font.size = Pt(height_pt)
    return t


def add_section_header(doc: _Document, title: str, accent_color: RGBColor = CORE_BLUE):
    """Colored left bar + bold title."""
    accent_hex = CORE_BLUE_HEX if accent_color == CORE_BLUE else CORE_ORANGE_HEX
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Emu(60000)
    t.columns[1].width = Emu(5700000)
    bar_cell, title_cell = t.rows[0].cells
    bar_cell.width = Emu(60000)
    title_cell.width = Emu(5700000)
    shade(bar_cell, accent_hex)
    remove_borders(bar_cell)
    remove_borders(title_cell)
    p = title_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "  " + title, bold=True, size=14, color=accent_color)
    return t


def set_col_widths(table, widths_in_inches: Sequence[float]) -> None:
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_in_inches):
                cell.width = Inches(widths_in_inches[i])


def styled_table(doc: _Document, headers: Sequence[str], rows: Sequence[Sequence],
                  *, col_widths: Sequence[float] | None = None,
                  status_col: int | None = None,
                  bold_last_row: bool = False):
    """Branded data table. Header = blue bg + white bold. Rows alternate
    white / off-white. If `status_col` is set, the value in that column is
    color-coded (red for critical/fail, orange for high/partial,
    teal for medium, green for pass/healthy)."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = False
    if col_widths:
        set_col_widths(t, col_widths)
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        shade(cell, CORE_BLUE_HEX)
        set_cell_border_color(cell, LIGHT_GREY_HEX)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, h, bold=True, size=10, color=WHITE)
    for i, row in enumerate(rows):
        is_total = bold_last_row and i == len(rows) - 1
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            if is_total:
                shade(cell, LIGHT_GREY_HEX)
            elif i % 2 == 1:
                shade(cell, OFF_WHITE_HEX)
            set_cell_border_color(cell, LIGHT_GREY_HEX)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            text = str(val) if val is not None else ""
            color = DARK_CHARCOAL if is_total else BRAND_GREY
            bold = is_total
            if status_col is not None and j == status_col and isinstance(val, str):
                lower = val.lower()
                if any(x in lower for x in ("critical", "fail", "high risk", "unhealthy", "disabled", "inactive (>30d)")):
                    color = RED
                    bold = True
                elif any(x in lower for x in ("high", "partial", "stale", "non compliant", "unmanaged", "incompatible")):
                    color = CORE_ORANGE
                    bold = True
                elif any(x in lower for x in ("medium", "recent")):
                    color = TEAL
                    bold = True
                elif any(x in lower for x in ("pass", "compliant", "complete", "healthy", "protected", "enabled", "active", "fresh")):
                    color = GREEN
                    bold = True
            add_run(p, text, size=9.5, color=color, bold=bold)
    return t


def add_metric_card_row(doc: _Document,
                          cards: Sequence[tuple[str, str, RGBColor]]):
    """Row of large-number KPI cards. Each card is (value, label, color)."""
    t = doc.add_table(rows=1, cols=len(cards))
    t.autofit = False
    for i, (value, label, color) in enumerate(cards):
        cell = t.rows[0].cells[i]
        shade(cell, OFF_WHITE_HEX)
        remove_borders(cell)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(12)
        p1.paragraph_format.space_after = Pt(0)
        add_run(p1, value, bold=True, size=24, color=color)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(12)
        add_run(p2, label, size=10, color=BRAND_GREY)
    return t


def add_callout_box(doc: _Document, text: str, *, accent_hex: str = CORE_ORANGE_HEX,
                     bg_hex: str = "FEF3EE"):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade(cell, bg_hex)
    set_cell_border_color(cell, accent_hex, sz="12")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, size=11, color=DARK_CHARCOAL)
    return t


def add_page_break(doc: _Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def add_footer(doc: _Document, line: str = CONTACT_LINE) -> None:
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, line, size=9, color=BRAND_GREY)


def add_header_logo(doc: _Document) -> None:
    section = doc.sections[0]
    p = section.header.paragraphs[0]
    if LOGO_PATH.exists():
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.7))
    else:
        add_run(p, "TECHNIJIAN", bold=True, size=14, color=CORE_BLUE)


def set_default_style(doc: _Document) -> None:
    s = doc.styles["Normal"]
    s.font.name = FONT
    s.font.size = Pt(11)
    s.font.color.rgb = BRAND_GREY


def standard_margins(doc: _Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


# ---- Cover page ----------------------------------------------------------

def render_cover(doc: _Document, *, title: str, subtitle: str, footer_note: str,
                  date_text: str | None = None) -> None:
    """Cover page in the Technijian style: blue accent bar -> centered logo
    -> orange divider -> title -> subtitle -> date -> blank space -> orange
    accent bar -> confidential note."""
    add_color_bar(doc, CORE_BLUE_HEX, height_pt=6)
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(3.0))
    else:
        add_run(p, "TECHNIJIAN", bold=True, size=28, color=CORE_BLUE)

    for _ in range(2):
        doc.add_paragraph()

    # Centered orange divider line
    div_t = doc.add_table(rows=1, cols=3)
    div_t.autofit = False
    for i, cell in enumerate(div_t.rows[0].cells):
        remove_borders(cell)
        if i == 1:
            shade(cell, CORE_ORANGE_HEX)
            cell.width = Inches(1.5)
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after = Pt(0)
            r = cp.add_run("")
            r.font.size = Pt(3)
        else:
            cell.width = Inches(2.45)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title, bold=True, size=28, color=DARK_CHARCOAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_run(p, subtitle, size=16, color=BRAND_GREY)

    if date_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, date_text, size=12, color=BRAND_GREY)

    for _ in range(6):
        doc.add_paragraph()
    add_color_bar(doc, CORE_ORANGE_HEX, height_pt=6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    add_run(p, footer_note, size=9, color=BRAND_GREY, italic=True)


def new_branded_document() -> _Document:
    """Convenience: return a Document with default style, margins, header logo,
    and footer line wired up."""
    doc = Document()
    set_default_style(doc)
    standard_margins(doc)
    add_header_logo(doc)
    add_footer(doc)
    return doc


__all__ = [
    "CORE_BLUE", "CORE_ORANGE", "TEAL", "DARK_CHARCOAL", "BRAND_GREY",
    "WHITE", "GREEN", "RED",
    "OFF_WHITE_HEX", "LIGHT_GREY_HEX", "CORE_BLUE_HEX", "CORE_ORANGE_HEX",
    "TEAL_HEX", "GREEN_HEX", "RED_HEX",
    "FONT", "LOGO_PATH", "CONTACT_LINE",
    "shade", "remove_borders", "set_cell_border_color",
    "add_run", "add_body", "add_bullet", "add_color_bar",
    "add_section_header", "set_col_widths", "styled_table",
    "add_metric_card_row", "add_callout_box", "add_page_break",
    "add_footer", "add_header_logo", "set_default_style", "standard_margins",
    "render_cover", "new_branded_document",
]
