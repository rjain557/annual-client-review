"""Build a client-facing Technijian-branded monthly cybersecurity activity
report from the Huntress per-client snapshots.

Inputs (per client, per target month YYYY-MM):
    clients/<code>/huntress/<latest>/agents.json + pull_summary.json
        (point-in-time agent inventory; the API has no historical filter)
    clients/<code>/huntress/monthly/<YYYY-MM>/incident_reports.json
        (incidents updated within the month - { "window": [...] })
    clients/<code>/huntress/monthly/<YYYY-MM>/signals.json
        (signals investigated within the month)
    clients/<code>/huntress/monthly/<YYYY-MM>/reports.json
        (Huntress reports overlapping the month)
    clients/<code>/huntress/monthly/<YYYY-MM>/pull_summary.json

Output:
    clients/<code>/huntress/monthly/<YYYY-MM>/<CODE>-Cybersecurity-Activity-<YYYY-MM>.docx

Usage:
    python build_monthly_report.py --month 2026-03
    python build_monthly_report.py --month 2026-03 --only BWH,ORX
    python build_monthly_report.py --month 2026-03 --skip ANI

Each report opens with an executive summary built for a non-technical
business reader: "here is how you are protected, here is what we caught,
here is what we recommend." Detail tables follow.
"""
from __future__ import annotations

import argparse
import json
import sys
from datetime import datetime, timezone, timedelta
from pathlib import Path

HERE = Path(__file__).resolve().parent
PIPELINE_ROOT = HERE.parent
REPO = PIPELINE_ROOT.parent.parent
CLIENTS_ROOT = REPO / "clients"

sys.path.insert(0, str(HERE))
import _brand as brand  # noqa: E402

from docx.shared import Inches, Pt, RGBColor  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402


# ---------------------------------------------------------------------------
# Inputs
# ---------------------------------------------------------------------------

def find_latest_daily_dir(client_dir: Path) -> Path | None:
    """Most recent clients/<code>/huntress/YYYY-MM-DD/ folder (skipping monthly/)."""
    if not client_dir.exists():
        return None
    candidates = [p for p in client_dir.iterdir()
                  if p.is_dir() and p.name != "monthly"
                  and len(p.name) == 10 and p.name[4] == "-" and p.name[7] == "-"]
    if not candidates:
        return None
    candidates.sort(reverse=True)
    return candidates[0]


def load_json(path: Path, default):
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError:
        return default
    except Exception as e:
        print(f"  WARN: could not parse {path}: {e}")
        return default


def parse_iso(s: str | None) -> datetime | None:
    if not s:
        return None
    try:
        if s.endswith("Z"):
            s = s[:-1] + "+00:00"
        return datetime.fromisoformat(s)
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Per-client data shape
# ---------------------------------------------------------------------------

class ClientMonthData:
    def __init__(self, code: str, month: str, daily_dir: Path | None,
                 monthly_dir: Path):
        self.code = code.upper()
        self.month = month
        self.daily_dir = daily_dir
        self.monthly_dir = monthly_dir

        self.daily_summary = load_json(daily_dir / "pull_summary.json", {}) if daily_dir else {}
        self.agents = load_json(daily_dir / "agents.json", []) if daily_dir else []

        inc_payload = load_json(monthly_dir / "incident_reports.json", {})
        if isinstance(inc_payload, dict):
            self.incidents = list(inc_payload.get("window") or [])
        else:
            self.incidents = list(inc_payload or [])
        self.signals = load_json(monthly_dir / "signals.json", []) or []
        self.reports = load_json(monthly_dir / "reports.json", []) or []
        self.month_summary = load_json(monthly_dir / "pull_summary.json", {})

    @property
    def location_name(self) -> str:
        return (self.daily_summary.get("Location_Name")
                or self.month_summary.get("Location_Name")
                or self.code)

    @property
    def has_data(self) -> bool:
        return bool(self.agents or self.incidents or self.signals or self.reports)

    # Agent rollups
    def agents_by_age_now(self, ref: datetime) -> dict:
        buckets = {"fresh": 0, "recent": 0, "stale": 0, "inactive": 0, "never": 0}
        for a in self.agents:
            last = parse_iso(a.get("last_callback_at") or a.get("last_survey_at"))
            if not last:
                buckets["never"] += 1
                continue
            delta = ref - last
            if delta <= timedelta(hours=24):
                buckets["fresh"] += 1
            elif delta <= timedelta(days=7):
                buckets["recent"] += 1
            elif delta <= timedelta(days=30):
                buckets["stale"] += 1
            else:
                buckets["inactive"] += 1
        return buckets

    def os_breakdown(self) -> dict:
        out: dict[str, int] = {}
        for a in self.agents:
            label = a.get("os") or "Unknown"
            out[label] = out.get(label, 0) + 1
        return dict(sorted(out.items(), key=lambda kv: kv[1], reverse=True))

    def defender_breakdown(self) -> dict:
        out: dict[str, int] = {}
        for a in self.agents:
            v = a.get("defender_status") or "Unknown"
            out[v] = out.get(v, 0) + 1
        return dict(sorted(out.items(), key=lambda kv: kv[1], reverse=True))

    def firewall_breakdown(self) -> dict:
        out: dict[str, int] = {}
        for a in self.agents:
            v = a.get("firewall_status") or "Unknown"
            out[v] = out.get(v, 0) + 1
        return dict(sorted(out.items(), key=lambda kv: kv[1], reverse=True))

    def stale_or_inactive_agents(self, ref: datetime) -> list[dict]:
        out = []
        for a in self.agents:
            last = parse_iso(a.get("last_callback_at") or a.get("last_survey_at"))
            if not last or (ref - last) > timedelta(days=7):
                out.append(a)
        return sorted(out, key=lambda a: parse_iso(a.get("last_callback_at"))
                                          or datetime(1970, 1, 1, tzinfo=timezone.utc))

    def attention_agents(self) -> list[dict]:
        """Agents with a Defender state that wants action."""
        bad_states = {"Unhealthy", "Disabled", "Incompatible"}
        out = [a for a in self.agents
               if (a.get("defender_status") in bad_states)
               or (a.get("firewall_status") == "Disabled")]
        return out


# ---------------------------------------------------------------------------
# Render
# ---------------------------------------------------------------------------

MONTH_NAMES = ["", "January", "February", "March", "April", "May", "June",
                "July", "August", "September", "October", "November", "December"]


def month_label(month: str) -> str:
    y, m = month.split("-")
    return f"{MONTH_NAMES[int(m)]} {y}"


def severity_color(sev: str) -> RGBColor:
    s = (sev or "").lower()
    if s in ("critical", "high"):
        return brand.RED
    if s == "medium":
        return brand.CORE_ORANGE
    if s == "low":
        return brand.TEAL
    return brand.BRAND_GREY


def fmt_dt(s: str | None) -> str:
    dt = parse_iso(s)
    return dt.strftime("%Y-%m-%d %H:%M UTC") if dt else "—"


def fmt_date(s: str | None) -> str:
    dt = parse_iso(s)
    return dt.strftime("%Y-%m-%d") if dt else "—"


def render_report(data: ClientMonthData, out_path: Path) -> dict:
    ref_now = datetime.now(timezone.utc)
    by_age = data.agents_by_age_now(ref_now)
    os_break = data.os_breakdown()
    defender_break = data.defender_breakdown()
    firewall_break = data.firewall_breakdown()
    attention = data.attention_agents()
    stale = [a for a in data.agents
             if (parse_iso(a.get("last_callback_at") or a.get("last_survey_at"))
                  is None
                  or (ref_now - parse_iso(a.get("last_callback_at")
                                          or a.get("last_survey_at")))
                     > timedelta(days=30))]

    inc_total = len(data.incidents)
    sev_counts: dict[str, int] = {}
    for inc in data.incidents:
        sev = (inc.get("severity") or "unknown").title()
        sev_counts[sev] = sev_counts.get(sev, 0) + 1

    sig_total = len(data.signals)
    rep_total = len(data.reports)

    doc = brand.new_branded_document()

    # ---- COVER ----
    brand.render_cover(
        doc,
        title="Cybersecurity Activity Report",
        subtitle=f"{data.location_name}",
        date_text=month_label(data.month),
        footer_note="Confidential — prepared by Technijian for the named client only.",
    )
    brand.add_page_break(doc)

    # ---- EXECUTIVE SUMMARY ----
    brand.add_section_header(doc, "Executive Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    brand.add_run(
        p,
        f"This report summarizes the cybersecurity activity Technijian delivered "
        f"for {data.location_name} during {month_label(data.month)}. The Huntress "
        f"managed detection and response (MDR) platform monitored your endpoints "
        f"24x7 and Technijian's security operations team reviewed every alert "
        f"that fired. Below is a plain-English summary of what we saw, what we "
        f"did about it, and where you stand today.",
        size=11,
    )

    # KPI cards
    cards = [
        (str(len(data.agents)), "Endpoints Protected", brand.CORE_BLUE),
        (str(by_age["fresh"] + by_age["recent"]), "Active in Last 7 Days", brand.GREEN),
        (str(inc_total), "Incidents Investigated", brand.CORE_ORANGE),
        (str(sig_total), "Signals Reviewed", brand.TEAL),
    ]
    brand.add_metric_card_row(doc, cards)

    if attention or stale:
        msg_bits = []
        if attention:
            msg_bits.append(f"{len(attention)} endpoint(s) with Microsoft Defender or firewall in a non-protected state")
        if stale:
            msg_bits.append(f"{len(stale)} endpoint(s) that have not phoned home in 30+ days")
        brand.add_callout_box(
            doc,
            "Attention items: " + "; ".join(msg_bits) +
            ". See Recommendations section below for the remediation list.",
            accent_hex=brand.CORE_ORANGE_HEX, bg_hex="FEF3EE")
    else:
        brand.add_callout_box(
            doc,
            "All monitored endpoints checked in on schedule and Microsoft "
            "Defender / firewall posture is healthy across your fleet. "
            "No action items this month.",
            accent_hex=brand.GREEN_HEX, bg_hex="EAF7EE")

    brand.add_page_break(doc)

    # ---- ENDPOINT PROTECTION ----
    brand.add_section_header(doc, "Endpoint Protection")
    brand.add_body(
        doc,
        "Every protected computer reports back to Huntress on a regular "
        "interval. The breakdown below shows which endpoints have been "
        "phoning home, how recently, and the state of their built-in "
        "defenses (Microsoft Defender + Windows Firewall).",
    )

    bullet_rows = []
    bullet_rows.append(("Total endpoints under management", len(data.agents)))
    bullet_rows.append(("Phoned home in the last 24 hours", by_age["fresh"]))
    bullet_rows.append(("Phoned home in the last 7 days", by_age["fresh"] + by_age["recent"]))
    bullet_rows.append(("Stale (8 - 30 days since last check-in)", by_age["stale"]))
    bullet_rows.append(("Inactive (more than 30 days since last check-in)", by_age["inactive"]))
    bullet_rows.append(("Never seen by the platform", by_age["never"]))
    brand.styled_table(
        doc,
        ["Metric", "Count"],
        bullet_rows,
        col_widths=[4.6, 1.4],
    )

    # OS breakdown
    if os_break:
        doc.add_paragraph()
        brand.add_body(doc, "Operating system distribution", bold=True, size=12, color=brand.DARK_CHARCOAL)
        brand.styled_table(
            doc,
            ["Operating System", "Endpoints"],
            list(os_break.items()),
            col_widths=[4.6, 1.4],
        )

    # Defender + Firewall
    if defender_break:
        doc.add_paragraph()
        brand.add_body(doc, "Microsoft Defender posture", bold=True, size=12, color=brand.DARK_CHARCOAL)
        brand.styled_table(
            doc,
            ["Defender Status", "Endpoints"],
            list(defender_break.items()),
            col_widths=[4.6, 1.4],
            status_col=0,
        )

    if firewall_break:
        doc.add_paragraph()
        brand.add_body(doc, "Windows Firewall posture", bold=True, size=12, color=brand.DARK_CHARCOAL)
        brand.styled_table(
            doc,
            ["Firewall Status", "Endpoints"],
            list(firewall_break.items()),
            col_widths=[4.6, 1.4],
            status_col=0,
        )

    brand.add_page_break(doc)

    # ---- THREAT ACTIVITY ----
    brand.add_section_header(doc, "Threat Activity This Month")

    if inc_total == 0 and sig_total == 0:
        brand.add_body(
            doc,
            f"No incident reports or active signals were raised against your "
            f"environment during {month_label(data.month)}. Continuous "
            f"monitoring stayed in place the entire period — a quiet month "
            f"means our platform did not detect actor activity that needed a "
            f"response.",
        )
    else:
        brand.add_body(
            doc,
            f"During {month_label(data.month)}, Huntress raised {inc_total} "
            f"incident report(s) and surfaced {sig_total} signal(s) on your "
            f"environment. Each one was reviewed by Technijian's security "
            f"operations team and routed to the appropriate response track. "
            f"The breakdown is below.",
        )

        if sev_counts:
            doc.add_paragraph()
            brand.add_body(doc, "Incidents by severity", bold=True, size=12, color=brand.DARK_CHARCOAL)
            brand.styled_table(
                doc,
                ["Severity", "Count"],
                [(k, v) for k, v in sev_counts.items()],
                col_widths=[3.5, 2.5],
                status_col=0,
            )

        if data.incidents:
            doc.add_paragraph()
            brand.add_body(doc, "Incident detail", bold=True, size=12, color=brand.DARK_CHARCOAL)
            inc_rows = []
            for inc in data.incidents[:50]:
                inc_rows.append((
                    fmt_date(inc.get("sent_at") or inc.get("updated_at") or inc.get("created_at")),
                    (inc.get("severity") or "").title() or "—",
                    (inc.get("status") or "").replace("_", " ").title() or "—",
                    (inc.get("subject") or inc.get("summary") or "")[:100] or "—",
                ))
            brand.styled_table(
                doc,
                ["Date", "Severity", "Status", "Summary"],
                inc_rows,
                col_widths=[1.0, 1.0, 1.0, 3.6],
                status_col=1,
            )
            if len(data.incidents) > 50:
                brand.add_body(
                    doc,
                    f"Showing 50 of {len(data.incidents)} incidents. Full list "
                    f"is available on request from your Technijian contact.",
                    size=10,
                )

    brand.add_page_break(doc)

    # ---- WHAT WE DID ----
    brand.add_section_header(doc, "What Technijian Did For You")
    bullets = [
        ("Continuous monitoring. ",
         f"Every endpoint was monitored 24x7 by the Huntress sensor. "
         f"{by_age['fresh'] + by_age['recent']} of {len(data.agents)} agent(s) "
         f"reported in within the last 7 days."),
        ("Alert review. ",
         f"Our security operations team triaged {sig_total} signal(s) and "
         f"actioned {inc_total} incident report(s) raised this month."),
        ("Reports & posture review. ",
         f"{rep_total} platform-generated security report(s) overlapped this "
         f"month and were reviewed for trends and emerging risks."),
        ("Posture data captured. ",
         "Microsoft Defender state, Windows Firewall state, agent version, "
         "and last-callback timestamp were captured for every endpoint to "
         "drive remediation discussions."),
    ]
    for prefix, body in bullets:
        brand.add_bullet(doc, body, bold_prefix=prefix)

    # ---- RECOMMENDATIONS ----
    doc.add_paragraph()
    brand.add_section_header(doc, "Recommendations")
    if not (attention or stale):
        brand.add_body(
            doc,
            "No action is required from your team this month. Coverage is "
            "healthy across the board.",
        )
    else:
        if stale:
            brand.add_body(doc, "Endpoints not phoning home (30+ days)",
                           bold=True, size=12, color=brand.DARK_CHARCOAL)
            stale_rows = []
            for a in stale[:20]:
                last = a.get("last_callback_at") or a.get("last_survey_at")
                stale_rows.append((
                    a.get("hostname") or "—",
                    a.get("os") or "—",
                    fmt_date(last),
                    "Inactive (>30d)",
                ))
            brand.styled_table(
                doc,
                ["Hostname", "OS", "Last Seen", "State"],
                stale_rows,
                col_widths=[2.0, 2.0, 1.5, 1.5],
                status_col=3,
            )
            if len(stale) > 20:
                brand.add_body(doc, f"... plus {len(stale) - 20} more. Full list available on request.",
                                size=10)
            brand.add_body(
                doc,
                "Recommended action: confirm whether each computer is still in "
                "service. If retired or reimaged, please notify Technijian to "
                "decommission the agent. If still in service, the agent likely "
                "needs to be reinstalled or the device powered on.",
            )

        if attention:
            doc.add_paragraph()
            brand.add_body(doc, "Endpoints with Defender or firewall issues",
                           bold=True, size=12, color=brand.DARK_CHARCOAL)
            att_rows = []
            for a in attention[:30]:
                att_rows.append((
                    a.get("hostname") or "—",
                    a.get("defender_status") or "—",
                    a.get("defender_policy_status") or "—",
                    a.get("firewall_status") or "—",
                ))
            brand.styled_table(
                doc,
                ["Hostname", "Defender Status", "Defender Policy", "Firewall"],
                att_rows,
                col_widths=[2.0, 1.5, 1.7, 1.3],
                status_col=1,
            )
            if len(attention) > 30:
                brand.add_body(doc, f"... plus {len(attention) - 30} more. Full list available on request.",
                                size=10)
            brand.add_body(
                doc,
                "Recommended action: review the listed endpoints with your "
                "Technijian contact to schedule remediation. Common fixes "
                "include re-enabling Microsoft Defender, applying a missing "
                "Group Policy / Intune setting, or re-enabling Windows Firewall.",
            )

    # ---- ABOUT ----
    doc.add_paragraph()
    brand.add_section_header(doc, "About This Report", accent_color=brand.CORE_ORANGE)
    brand.add_body(
        doc,
        "Technijian provides 24x7 managed detection and response (MDR) on top "
        "of the Huntress security platform. Activity captured here reflects "
        "data pulled directly from the Huntress API for your tenant. "
        "Endpoint inventory is a point-in-time snapshot taken at report "
        "generation; incident, signal, and report data are filtered to the "
        "calendar month above.",
        size=10,
    )
    brand.add_body(doc, brand.CONTACT_LINE, size=10, color=brand.BRAND_GREY)

    # Save
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    return {
        "code": data.code,
        "month": data.month,
        "out_path": str(out_path),
        "agents": len(data.agents),
        "incidents": inc_total,
        "signals": sig_total,
        "reports": rep_total,
        "attention": len(attention),
        "stale": len(stale),
    }


# ---------------------------------------------------------------------------
# Driver
# ---------------------------------------------------------------------------

def parse_args() -> argparse.Namespace:
    ap = argparse.ArgumentParser(description="Build branded Huntress monthly activity report per client.")
    ap.add_argument("--month", required=True, help="YYYY-MM target month")
    ap.add_argument("--only", help="comma-separated LocationCodes")
    ap.add_argument("--skip", action="append", default=[],
                    help="LocationCode to skip (repeatable)")
    return ap.parse_args()


def discover_clients_with_month(month: str) -> list[str]:
    """Return LocationCodes (lowercase) that have a clients/<code>/huntress/monthly/<month>/ folder."""
    out: list[str] = []
    for child in CLIENTS_ROOT.iterdir():
        if not child.is_dir():
            continue
        candidate = child / "huntress" / "monthly" / month
        if candidate.exists():
            out.append(child.name.lower())
    return sorted(out)


def main() -> int:
    args = parse_args()
    month = args.month
    skip_codes = {s.upper() for s in args.skip for s in s.split(",") if s.strip()}
    only_codes = None
    if args.only:
        only_codes = {s.strip().upper() for s in args.only.split(",") if s.strip()}

    codes = discover_clients_with_month(month)
    if not codes:
        print(f"No clients have data for month {month}. Run pull_huntress_daily.py "
              f"and/or backfill_huntress.py first.")
        return 1

    print(f"[{datetime.now():%H:%M:%S}] Building monthly reports for {month}")
    print(f"  candidates: {len(codes)} client(s) with data")

    results: list[dict] = []
    for code in codes:
        upper = code.upper()
        if upper in skip_codes:
            continue
        if only_codes is not None and upper not in only_codes:
            continue
        client_dir = CLIENTS_ROOT / code / "huntress"
        latest_daily = find_latest_daily_dir(client_dir)
        monthly_dir = client_dir / "monthly" / month
        data = ClientMonthData(upper, month, latest_daily, monthly_dir)
        if not data.has_data:
            print(f"  [skip] {upper:<6s} no data in monthly or daily folders")
            continue
        out_path = monthly_dir / f"{upper}-Cybersecurity-Activity-{month}.docx"
        try:
            res = render_report(data, out_path)
        except Exception as e:
            print(f"  [ERR ] {upper:<6s} {e}")
            continue
        print(f"  [ok  ] {upper:<6s} agents={res['agents']:>3d} inc={res['incidents']:>3d}"
              f" sig={res['signals']:>4d} rep={res['reports']:>2d}"
              f" attn={res['attention']:>2d} stale={res['stale']:>2d}"
              f"  -> {out_path.relative_to(REPO)}")
        results.append(res)

    print()
    print(f"[{datetime.now():%H:%M:%S}] DONE — generated {len(results)} report(s)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
